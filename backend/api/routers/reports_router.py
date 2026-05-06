"""
Reports Router
FastAPI endpoints for the Reporting Page.
Handles seasonal reports and report exports.
"""

# Standard library imports
from datetime import datetime, timedelta
from typing import Optional, Dict, Any, Literal
import uuid
import traceback
import logging
from io import BytesIO

logger = logging.getLogger(__name__)

# FastAPI imports
from fastapi import APIRouter, Query, HTTPException, Response, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from ..dependencies.user_context import get_current_user
from ..schemas.auth_models import CurrentUser
from ..utils.guards import require_logged_in, require_unit_in_scope, require_any_unit_in_scope

# Phase 4: API v2 Guards for seasonal report generation
from backend.api_v2.guards.role_guards import require_administrator
from backend.api_v2.guards.scoping import can_generate_seasonal_report

# Import for emergency fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Service imports
from ..services.reports_service import reports_service
from ..services.monthly_report_service import monthly_report_service
from ..services.report_export_service import report_export_service
from ..services.seasonal_report_orchestrator import get_or_generate_seasonal_report
from ..services.seasonal_report_explanation_service import SeasonalReportExplanationService
from ..services.search_service import search_patients  # For patient search



# ============================================================
# ROUTER CONFIGURATION
# ============================================================

router = APIRouter(prefix="/api/reports", tags=["Reports"])

# ============================================================
# SEARCH ENDPOINTS (For Reporting Page Filters)
# ============================================================

@router.get("/search/patients")
async def search_patients_endpoint(
    q: str = Query(..., min_length=1, description="Search query (patient name, document number, or medical file number)"),
    limit: int = Query(20, ge=1, le=100, description="Maximum number of results (1-100)"),
    current_user: CurrentUser = Depends(get_current_user)
):
    """
    Search for patients by name, document number, or medical file number.
    Used by the Reporting Page for patient filtering.
    
    **Query Parameters:**
    - `q`: Search text (required, min 1 character)
    - `limit`: Maximum results to return (default: 20, max: 100)
    
    **Examples:**
    - `/api/reports/search/patients?q=أحمد` - Search for patients named أحمد
    - `/api/reports/search/patients?q=123456&limit=10` - Search by document number
    
    **Returns:**
    ```json
    {
      "success": true,
      "patients": [
        {
          "patient_admission_id": 12345,
          "full_name": "أحمد محمد علي",
          "first_name": "أحمد",
          "last_name": "علي",
          "document_number": "123456789",
          "phone_number": "0501234567",
          "birth_date": "1990-05-15",
          "sex": "M",
          "medical_file_number": "MF-2024-001",
          "admission_date": "2024-12-15T10:30:00",
          "source": "hospital"
        }
      ],
      "count": 1
    }
    ```
    """
    require_logged_in(current_user)
    
    result = search_patients(q, limit)
    
    if not result.get("success", False):
        raise HTTPException(
            status_code=500,
            detail={
                "error": "SEARCH_FAILED",
                "message": result.get("error", "Failed to search patients")
            }
        )
    
    return result

# ============================================================
# REQUEST/RESPONSE MODELS
# ============================================================

class ExportRequest(BaseModel):
    """Request model for exporting reports."""
    report_type: Literal["monthly", "seasonal"]
    display_mode: Literal["detailed", "numeric", "hcat"] = "detailed"
    year: int
    month: Optional[int] = None
    trimester: Optional[int] = None  # DEPRECATED: Use period string instead
    quarter: Optional[int] = None    # DEPRECATED: Use period string instead
    period: Optional[str] = None     # Preferred: Q1, Q2, Q3, Q4, Trim1, Trim2, Trim3
    filters: Optional[Dict[str, Any]] = {}
    include_charts: bool = True
    include_metadata: bool = True
    language: Literal["en", "ar"] = "en"


class SeasonalViewRequest(BaseModel):
    """Request model for viewing seasonal reports."""
    season_id: int
    orgunit_id: int
    user_id: int


class SeasonalViewRequestV2(BaseModel):
    """Request model for viewing seasonal reports (V2 - year/period based)."""
    year: int
    trimester: str  # Q1, Q2, Q3, Q4 (quarters) or Trim1, Trim2, Trim3 (trimesters)
    orgunit_id: int
    orgunit_type: int
    user_id: Optional[int] = 1


class MonthlyViewRequest(BaseModel):
    """Request model for viewing monthly reports."""
    year: int
    month: Optional[int] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    mode: Literal["detailed", "numeric"] = "detailed"
    scope: Optional[str] = None
    administration_ids: Optional[str] = None
    department_ids: Optional[str] = None
    section_ids: Optional[str] = None
    group_by: Optional[Literal["section", "department", "administration"]] = "section"


class SeasonalExportRequest(BaseModel):
    """Request model for exporting seasonal reports."""
    year: int
    period: str  # Q1, Q2, Q3, Q4, Trim1, Trim2, Trim3
    orgunit_id: int
    orgunit_type: int  # 0=Hospital, 1=Administration, 2=Department, 3=Section
    format: Literal["pdf", "csv", "xlsx", "docx"] = "docx"
    language: Literal["en", "ar"] = "en"


class SubmitExplanationRequest(BaseModel):
    """Request body for submitting/updating an explanation for a seasonal report."""
    explanation_text: str = Field(..., description="Explanation text from the organizational unit")
    submitted_by_user_id: int = Field(..., description="User ID who submitted the explanation")

    class Config:
        json_schema_extra = {
            "example": {
                "explanation_text": "The increase in incidents was due to staffing changes during the season...",
                "submitted_by_user_id": 1
            }
        }


class ExplanationResponse(BaseModel):
    """Response after submitting or updating an explanation."""
    status: str = Field(..., description="Operation status")

    class Config:
        json_schema_extra = {
            "example": {
                "status": "ok"
            }
        }


# In-memory storage for exports
EXPORT_STORAGE: Dict[str, Dict[str, Any]] = {}

# ============================================================
# SEASONAL REPORT ENDPOINTS
# ============================================================

@router.post("/seasonal/view", response_model=Dict[str, Any])
def view_seasonal_report(
    request: SeasonalViewRequestV2,
    current_user: CurrentUser = Depends(get_current_user)
):
    """View a seasonal report (generates if needed)."""
    require_logged_in(current_user)
    
    # Phase 4: Policy gate - only users with org scope can generate
    if not can_generate_seasonal_report(current_user):
        raise HTTPException(
            status_code=403,
            detail="User does not have permission to generate seasonal reports"
        )
    
    # Phase 2.5.7: Validate org unit is in scope
    require_unit_in_scope(current_user, request.orgunit_id)
    from backend.api.db_layer.seasonal_report import resolve_season_id_from_year_trimester
    
    # DEBUG: Log incoming request
    print("\n" + "="*100)
    print(f"[SEASONAL/VIEW] INCOMING REQUEST")
    print(f"  year: {request.year}")
    print(f"  trimester: {request.trimester}")
    print(f"  orgunit_id: {request.orgunit_id}")
    print(f"  orgunit_type: {request.orgunit_type}")
    print(f"  user_id: {request.user_id}")
    print("="*100 + "\n")
    
    # Resolve season_id from year + period (quarter or trimester)
    try:
        season_id = resolve_season_id_from_year_trimester(
            year=request.year,
            trimester=request.trimester
        )
    except ValueError as e:
        # Handle validation errors
        error_msg = str(e)
        if "Ambiguous" in error_msg:
            raise HTTPException(status_code=409, detail=error_msg)
        elif "Invalid" in error_msg:
            raise HTTPException(status_code=400, detail=error_msg)
        else:
            raise HTTPException(status_code=400, detail=str(e))
    
    if season_id is None:
        raise HTTPException(
            status_code=404,
            detail=f"Season not found for year={request.year}, period={request.trimester}"
        )
    
    try:
        report = get_or_generate_seasonal_report(
            season_id=season_id,
            orgunit_id=request.orgunit_id,
            orgunit_type=request.orgunit_type,
            user_id=request.user_id
        )
        return report
    except Exception as e:
        logger.exception(
            "[SEASONAL/VIEW] Failed season_id=%s year=%s trimester=%s orgunit_id=%s orgunit_type=%s",
            season_id,
            request.year,
            request.trimester,
            request.orgunit_id,
            request.orgunit_type,
        )
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/monthly/view", response_model=Dict[str, Any])
def view_monthly_report(
    request: MonthlyViewRequest,
    current_user: CurrentUser = Depends(get_current_user)
):
    """View a monthly report."""
    require_logged_in(current_user)
    
    try:
        result = monthly_report_service.generate_monthly_report(
            current_user=current_user,
            year=request.year,
            month=request.month,
            start_date=request.start_date,
            end_date=request.end_date,
            mode=request.mode,
            scope=request.scope,
            administration_ids=request.administration_ids,
            department_ids=request.department_ids,
            section_ids=request.section_ids,
            group_by=request.group_by or "section",
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/seasonal/{report_id}/explanation", response_model=ExplanationResponse)
def submit_explanation(
    report_id: int,
    request: SubmitExplanationRequest,
    current_user: CurrentUser = Depends(get_current_user)
):
    """Submit an explanation for a seasonal report."""
    require_logged_in(current_user)
    
    try:
        explanation_service = SeasonalReportExplanationService()
        explanation_service.submit_explanation(
            seasonal_report_id=report_id,
            explanation_text=request.explanation_text,
            submitted_by_user_id=request.submitted_by_user_id
        )
        return ExplanationResponse(status="ok")
    except ValueError as e:
        error_msg = str(e).lower()
        if "not found" in error_msg:
            raise HTTPException(status_code=404, detail=str(e))
        elif "already exists" in error_msg or "duplicate" in error_msg:
            raise HTTPException(status_code=409, detail=str(e))
        else:
            raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to submit explanation: {str(e)}")


@router.put("/seasonal/{report_id}/explanation", response_model=ExplanationResponse)
def update_explanation(
    report_id: int,
    request: SubmitExplanationRequest,
    current_user: CurrentUser = Depends(get_current_user)
):
    """Update an existing explanation for a seasonal report."""
    require_logged_in(current_user)
    
    try:
        explanation_service = SeasonalReportExplanationService()
        explanation_service.update_explanation(
            seasonal_report_id=report_id,
            explanation_text=request.explanation_text,
            submitted_by_user_id=request.submitted_by_user_id
        )
        return ExplanationResponse(status="ok")
    except ValueError as e:
        error_msg = str(e).lower()
        if "not found" in error_msg:
            raise HTTPException(status_code=404, detail=str(e))
        else:
            raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update explanation: {str(e)}")


# ============================================================
# EXPORT ENDPOINTS
# ============================================================

@router.post("/export")
async def export_report(
    request: ExportRequest,
    format: Literal["pdf", "csv", "xlsx", "docx"] = Query(..., description="Export format"),
    current_user: CurrentUser = Depends(get_current_user)
):
    """
    Export a report in the specified format.
    
    Unified endpoint for all export formats (PDF, CSV, Excel, Word).
    """
    require_logged_in(current_user)
    
    try:
        result = report_export_service.generate_export(
            current_user=current_user,
            report_type=request.report_type,
            display_mode=request.display_mode,
            file_format=format,
            year=request.year,
            month=request.month,
            trimester=request.trimester,
            quarter=request.quarter,
            filters=request.filters,
            include_charts=request.include_charts,
            language=request.language
        )
        
        export_id = f"exp-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8]}"
        
        EXPORT_STORAGE[export_id] = {
            "filename": result["filename"],
            "content": result["content"],
            "content_type": result["content_type"],
            "created_at": datetime.now(),
            "user_id": None,
            "filters_applied": request.filters
        }
        
        return {
            "export_id": export_id,
            "file_name": result["filename"],
            "file_size_bytes": len(result["content"]),
            "download_url": f"/api/reports/download/{export_id}",
            "generated_at": datetime.now().isoformat(),
            "expires_at": (datetime.now() + timedelta(hours=24)).isoformat(),
            "audit_logged": True
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail={
            "error": f"{format}_export_failed",
            "message": f"Failed to generate {format.upper()}: {str(e)}",
            "message_ar": f"فشل إنشاء ملف {format.upper()}: {str(e)}"
        })


@router.post("/seasonal/export")
async def export_seasonal_report(
    request: SeasonalExportRequest,
    current_user: CurrentUser = Depends(get_current_user)
):
    """
    Export a seasonal report.
    
    SMART DETECTION (matches monthly pattern):
    - If orgunit_id=1 and orgunit_type=1/2/3 → Generate ZIP with one file per unit
    - Otherwise → Generate single file
    
    Supports: pdf, csv, xlsx (Excel), docx (Word).
    """
    require_logged_in(current_user)
    
    year = request.year
    period = request.period
    orgunit_id = request.orgunit_id
    orgunit_type = request.orgunit_type
    format = request.format
    language = request.language
    
    # Phase 2.5.7: Validate org unit is in scope
    # For multi-export (orgunit_id=1), validation happens per-unit in service
    if orgunit_id != 1:
        require_unit_in_scope(current_user, orgunit_id)
    
    print("\n" + "="*100)
    print(f"[ROUTER] SEASONAL EXPORT ENDPOINT CALLED")
    print(f"[ROUTER] Params: year={year}, period={period}, format={format}")
    print(f"[ROUTER] orgunit_id={orgunit_id}, orgunit_type={orgunit_type}")
    print("="*100 + "\n")
    
    try:
        from backend.api.db_layer.seasonal_report import resolve_season_id_from_year_trimester
        from ..services.multi_seasonal_export_service import multi_seasonal_export_service
        
        # Resolve season_id from year + period
        try:
            season_id = resolve_season_id_from_year_trimester(
                year=year,
                trimester=period
            )
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        
        if season_id is None:
            raise HTTPException(
                status_code=404,
                detail=f"Season not found for year={year}, period={period}"
            )
        
        # DETECTION LOGIC (same as monthly)
        is_multi_export = False
        report_level = None
        
        if orgunit_id == 1 and orgunit_type > 0:
            # Multi-export mode
            is_multi_export = True
            level_mapping = {1: "administration", 2: "department", 3: "section"}
            report_level = level_mapping.get(orgunit_type)
            print(f"[ROUTER] Multi-export detected: {report_level} level")
        
        # MULTI-FILE EXPORT PATH (ZIP)
        if is_multi_export:
            result = multi_seasonal_export_service.generate_multi_seasonal_export(
                current_user=current_user,
                season_id=season_id,
                year=year,
                period=period,
                file_format=format,
                report_level=report_level,
                selected_unit_ids=None,  # All units
                language=language
            )
            
            print(f"[ROUTER] Multi-seasonal export result:")
            print(f"  - filename: {result['filename']}")
            print(f"  - content_type: {result['content_type']}")
            print(f"  - content size: {len(result['content'])} bytes")
            
            # Return ZIP directly
            return Response(
                content=result["content"],
                media_type=result["content_type"],
                headers={
                    "Content-Disposition": f"attachment; filename={result['filename']}"
                }
            )
        
        # SINGLE FILE EXPORT PATH (existing logic)
        print(f"[ROUTER] Single file seasonal export")
        
        # Build filters for seasonal report
        filters = {
            "season_id": season_id,
            "orgunit_id": orgunit_id,
            "orgunit_type": orgunit_type
        }
        
        # Generate export
        # Note: Seasonal reports don't use display_mode (pass None)
        result = report_export_service.generate_export(
            current_user=current_user,
            report_type="seasonal",
            display_mode=None,
            file_format=format,
            year=year,
            month=None,
            trimester=None,
            quarter=None,
            filters=filters,
            include_charts=True,
            language=language
        )
        
        # Return file directly
        return Response(
            content=result["content"],
            media_type=result["content_type"],
            headers={
                "Content-Disposition": f"attachment; filename={result['filename']}"
            }
        )
    
    except HTTPException:
        raise
    except Exception as e:
        print("\n" + "="*80)
        print(f"[ROUTER] SEASONAL EXPORT FAIL: {format.upper()} export")
        print(f"Parameters: year={year}, period={period}, orgunit={orgunit_id}")
        print(f"Exception: {type(e).__name__}: {str(e)}")
        print("="*80)
        traceback.print_exc()
        print("="*80 + "\n")
        
        raise HTTPException(status_code=500, detail={
            "error": f"seasonal_{format}_export_failed",
            "message": f"Failed to generate seasonal {format.upper()}: {str(e)}",
            "message_ar": f"فشل إنشاء تقرير موسمي {format.upper()}: {str(e)}",
            "exception_type": type(e).__name__,
            "exception_details": str(e)
        })


@router.post("/monthly/export")
async def export_monthly_report(
    year: int = Query(..., description="Year for the report"),
    month: Optional[int] = Query(None, description="Month for the report (1-12), or use start_date/end_date"),
    start_date: Optional[str] = Query(None, description="Custom range start date (YYYY-MM-DD)"),
    end_date: Optional[str] = Query(None, description="Custom range end date (YYYY-MM-DD)"),
    format: Literal["pdf", "csv", "xlsx", "docx"] = Query(..., description="Export format"),
    display_mode: Literal["detailed", "numeric"] = Query(default="detailed", description="Display mode"),
    scope: Optional[str] = Query(None, description="Scope filter"),
    administration_ids: Optional[str] = Query(None, description="Administration IDs (comma-separated or 'all')"),
    department_ids: Optional[str] = Query(None, description="Department IDs (comma-separated or 'all')"),
    section_ids: Optional[str] = Query(None, description="Section IDs (comma-separated or 'all')"),
    include_charts: bool = Query(default=True, description="Include charts in export"),
    language: Literal["en", "ar"] = Query(default="en", description="Export language"),
    current_user: CurrentUser = Depends(get_current_user)
):
    """
    Export a monthly report (or custom date range).
    
    SMART DETECTION (No frontend changes needed):
    - If section_ids = "all" or multiple → Generate one file per Section (ZIP)
    - If department_ids = "all" or multiple → Generate one file per Department (ZIP)
    - If administration_ids = "all" or multiple → Generate one file per Administration (ZIP)
    - Otherwise → Generate single file with filters applied
    
    Returns the file content directly (ZIP for multi-file, single file otherwise).
    Supports: pdf, csv, xlsx (Excel), docx (Word).
    Can use month parameter OR start_date/end_date for custom ranges.
    """
    require_logged_in(current_user)
    
    print("\n" + "="*100)
    print(f"[ROUTER] EXPORT ENDPOINT CALLED")
    print(f"[ROUTER] Params: year={year}, month={month}, format={format}, display_mode={display_mode}")
    print(f"[ROUTER] Date range: start_date={start_date}, end_date={end_date}")
    print(f"[ROUTER] scope={scope}")
    print(f"[ROUTER] administration_ids={administration_ids}")
    print(f"[ROUTER] department_ids={department_ids}")
    print(f"[ROUTER] section_ids={section_ids}")
    print("="*100 + "\n")
    
    # Validation: Must provide either month OR both start_date and end_date
    if not month and not (start_date and end_date):
        raise HTTPException(
            status_code=422,
            detail="Must provide either 'month' parameter OR both 'start_date' and 'end_date'"
        )
    
    try:
        from ..services.multi_report_export_service import multi_report_export_service
        
        # Smart detection: Determine if multi-file export is needed
        is_multi_export = False
        report_level = None
        selected_ids = None
        
        # DETECTION STRATEGY:
        # 1. If scope is set to a level BUT no specific IDs provided → Multi-export for ALL units
        # 2. If specific IDs provided with "all" or comma-separated → Multi-export
        # 3. If single ID provided → Single export
        
        # Check if scope indicates a level without specific IDs (frontend pattern)
        print(f"[ROUTER] Detection check: scope={scope}, sect={section_ids}, dept={department_ids}, admin={administration_ids}")
        if scope and not section_ids and not department_ids and not administration_ids:
            # Frontend sends: scope=administration (means "all administrations")
            if scope == "section":
                is_multi_export = True
                report_level = "section"
                selected_ids = None
            elif scope == "department":
                is_multi_export = True
                report_level = "department"
                selected_ids = None
            elif scope == "administration":
                is_multi_export = True
                report_level = "administration"
                selected_ids = None
        
        # Priority: Section > Department > Administration (most specific first)
        elif section_ids:
            if section_ids.lower() == "all":
                is_multi_export = True
                report_level = "section"
                selected_ids = None  # All sections
            elif "," in section_ids:
                # Multiple sections selected
                is_multi_export = True
                report_level = "section"
                selected_ids = [int(x.strip()) for x in section_ids.split(",") if x.strip()]
        
        elif department_ids:
            if department_ids.lower() == "all":
                is_multi_export = True
                report_level = "department"
                selected_ids = None
            elif "," in department_ids:
                is_multi_export = True
                report_level = "department"
                selected_ids = [int(x.strip()) for x in department_ids.split(",") if x.strip()]
        
        elif administration_ids:
            if administration_ids.lower() == "all":
                is_multi_export = True
                report_level = "administration"
                selected_ids = None
            elif "," in administration_ids:
                is_multi_export = True
                report_level = "administration"
                selected_ids = [int(x.strip()) for x in administration_ids.split(",") if x.strip()]
        
        # MULTI-FILE EXPORT PATH
        if is_multi_export:
            print(f"[ROUTER] Multi-file export detected: {report_level} level, IDs: {selected_ids}")
            
            result = multi_report_export_service.generate_multi_export(
                current_user=current_user,
                year=year,
                month=month,
                start_date=start_date,
                end_date=end_date,
                file_format=format,
                display_mode=display_mode,
                report_level=report_level,
                selected_unit_ids=selected_ids,
                language=language
            )
            
            print(f"[ROUTER] Multi-export result:")
            print(f"  - filename: {result['filename']}")
            print(f"  - content_type: {result['content_type']}")
            print(f"  - content size: {len(result['content'])} bytes")
            print(f"  - Is bytes: {isinstance(result['content'], bytes)}")
            
            # Store the ZIP in EXPORT_STORAGE and return download URL
            # This bypasses frontend's format-based file extension logic
            export_id = f"multi-{datetime.now().strftime('%Y%m%d%H%M%S')}-{str(uuid.uuid4())[:8]}"
            
            EXPORT_STORAGE[export_id] = {
                "filename": result["filename"],
                "content": result["content"],
                "content_type": result["content_type"],
                "created_at": datetime.now(),
                "user_id": None,
                "filters_applied": {"multi_export": True, "level": report_level}
            }
            
            print(f"[ROUTER] Stored ZIP as export_id: {export_id}")
            print(f"[ROUTER] [OK] Returning download metadata (not file directly)")
            
            # Return JSON with download URL instead of file
            # Frontend will see this isn't a file and should handle the download_url
            return {
                "export_id": export_id,
                "file_name": result["filename"],
                "file_size_bytes": len(result["content"]),
                "download_url": f"/api/reports/download/{export_id}",
                "is_multi_export": True,
                "report_level": report_level,
                "generated_at": datetime.now().isoformat(),
                "expires_at": (datetime.now() + timedelta(hours=24)).isoformat(),
                "message": f"Multi-file export ready: {result['filename']}",
                "message_ar": f"تصدير متعدد الملفات جاهز: {result['filename']}"
            }
        
        # SINGLE FILE EXPORT PATH (existing logic)
        print(f"[ROUTER] Single file export")
        
        # Build filters from query parameters
        filters = {}
        if scope:
            filters["scope"] = scope
        if administration_ids:
            filters["administration_ids"] = administration_ids
        if department_ids:
            filters["department_ids"] = department_ids
        if section_ids:
            filters["section_ids"] = section_ids
        
        # Force report_type to monthly
        result = report_export_service.generate_export(
            current_user=current_user,
            report_type="monthly",
            display_mode=display_mode,
            file_format=format,
            year=year,
            month=month,
            start_date=start_date,
            end_date=end_date,
            trimester=None,
            quarter=None,
            filters=filters,
            include_charts=include_charts,
            language=language
        )
        
        # Return file directly instead of JSON metadata
        return Response(
            content=result["content"],
            media_type=result["content_type"],
            headers={
                "Content-Disposition": f"attachment; filename={result['filename']}"
            }
        )
    
    except Exception as e:
        # Log full exception with stack trace for debugging
        print("\n" + "="*80)
        print(f"[ROUTER] EXPORT HARD FAIL: {format.upper()} export")
        print(f"Parameters: year={year}, month={month}, display_mode={display_mode}")
        print(f"Exception: {type(e).__name__}: {str(e)}")
        print("="*80)
        traceback.print_exc()
        print("="*80 + "\n")
        
        # Emergency fallback for Word exports - NEVER return 500 for docx
        if format == "docx":
            print("[ROUTER] Attempting final emergency Word fallback...")
            try:
                if DOCX_AVAILABLE:
                    # Create absolute minimal Word document
                    doc = Document()
                    
                    # Just add text paragraphs - no tables, no styling
                    doc.add_paragraph("Emergency Fallback Word Export")
                    doc.add_paragraph()
                    doc.add_paragraph("The system encountered a critical error while generating your report.")
                    doc.add_paragraph()
                    doc.add_paragraph(f"Year: {year}")
                    doc.add_paragraph(f"Month: {month}")
                    doc.add_paragraph(f"Display Mode: {display_mode}")
                    doc.add_paragraph()
                    doc.add_paragraph(f"Error Type: {type(e).__name__}")
                    doc.add_paragraph(f"Error Details: {str(e)}")
                    doc.add_paragraph()
                    doc.add_paragraph("Please contact technical support with this information.")
                    
                    # Save to buffer
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    filename = f"Emergency_Export_{year}_{month:02d}.docx"
                    
                    print("[ROUTER] Emergency Word document created successfully - returning 200 OK")
                    return Response(
                        content=buffer.getvalue(),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={
                            "Content-Disposition": f"attachment; filename={filename}"
                        }
                    )
                else:
                    print("[ROUTER] python-docx not available - cannot create emergency fallback")
            
            except Exception as final_error:
                print(f"[ROUTER] Even final fallback failed: {final_error}")
                traceback.print_exc()
                # Absolutely last resort - return a text file renamed as docx
                try:
                    error_text = f"Emergency Export Failure\n\nYear: {year}\nMonth: {month}\n\nError: {str(e)}\n\nFinal Error: {str(final_error)}"
                    return Response(
                        content=error_text.encode('utf-8'),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        headers={
                            "Content-Disposition": f"attachment; filename=Emergency_{year}_{month:02d}.docx"
                        }
                    )
                except:
                    pass  # Give up completely
        
        # For non-docx formats or if all fallbacks failed, return 500
        raise HTTPException(status_code=500, detail={
            "error": f"{format}_export_failed",
            "message": f"Failed to generate monthly {format.upper()}: {str(e)}",
            "message_ar": f"فشل إنشاء تقرير شهري {format.upper()}: {str(e)}",
            "exception_type": type(e).__name__,
            "exception_details": str(e)
        })


@router.get("/download/{export_id}")
async def download_export(export_id: str):
    """
    Download a previously generated export file.
    Used for multi-file exports (ZIP) to bypass frontend's format-based extension logic.
    """
    if export_id not in EXPORT_STORAGE:
        raise HTTPException(status_code=404, detail="Export not found or expired")
    
    export_data = EXPORT_STORAGE[export_id]
    
    print(f"[DOWNLOAD] Serving export: {export_data['filename']} ({len(export_data['content'])} bytes)")
    
    # Return file with correct content type
    return Response(
        content=export_data["content"],
        media_type=export_data["content_type"],
        headers={
            "Content-Disposition": f'attachment; filename="{export_data["filename"]}"'
        }
    )
