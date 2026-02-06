"""  
Multi-Seasonal Export Service
Generates multiple seasonal report files (one per organizational unit) and packages them in a ZIP.
"""

from typing import Dict, Any, List, Literal, Optional
from io import BytesIO
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt

from ..schemas.auth_models import CurrentUser
from ..utils.guards import require_unit_in_scope
from .seasonal_report_orchestrator import (
    get_or_generate_seasonal_report,
    get_or_generate_comparative_seasonal_reports
)
from .seasonal_report_formatter import (
    generate_seasonal_word_report,
    generate_comparative_seasonal_word_report
)
from .reports_service import reports_service
from ..db_layer.admin_units import get_units_by_type
from ..db_layer.seasonal_report import get_previous_season


class MultiSeasonalExportService:
    """
    Service for generating multiple seasonal reports (one per organizational unit).
    Used when user selects "All" at a specific level (Administration/Department/Section).
    """
    
    def generate_multi_seasonal_export(
        self,
        *,
        current_user: CurrentUser,
        season_id: int,
        year: int,
        period: str,
        file_format: Literal["pdf", "csv", "xlsx", "docx"],
        report_level: Literal["administration", "department", "section"],
        selected_unit_ids: Optional[List[int]] = None,
        language: Literal["en", "ar"] = "en"
    ) -> Dict[str, Any]:
        """
        Generate multiple seasonal reports (one per unit) and package in ZIP.
        
        Phase 2.5.7: CRITICAL SECURITY - Each unit validated against current_user.allowed_unit_ids.
        If ANY unit is out of scope, entire request fails with 403.
        
        Args:
            current_user: Authenticated user with allowed org units
            season_id: Season ID for reports
            year: Year for reports
            period: Period string (Q1, Q2, Q3, Q4, Trim1, Trim2, Trim3)
            file_format: Output format (docx, pdf, xlsx, csv)
            report_level: Which level to generate reports for
            selected_unit_ids: Specific unit IDs, or None for "all"
            language: Export language
            
        Returns:
            Dict with filename (ZIP) and content (bytes)
            
        Raises:
            HTTPException: 403 if any requested unit is outside user's scope
        """
        # Map report level to Type values in database
        # Based on actual database values: 323=Administration, 324=Section, 325=Department
        type_mapping = {
            "administration": 323,
            "department": 325,
            "section": 324
        }
        
        # Map report level to orgunit_type for seasonal reports
        # 0=Hospital, 1=Administration, 2=Department, 3=Section
        orgunit_type_mapping = {
            "administration": 1,
            "department": 2,
            "section": 3
        }
        
        unit_type = type_mapping[report_level]
        orgunit_type = orgunit_type_mapping[report_level]
        
        # Phase 2.5.8: CRITICAL - Validate requested IDs BEFORE filtering
        # This ensures fail-fast behavior: if user requests [29, 6] and 6 is out of scope,
        # the entire request fails with 403, not just silently omitting unit 6
        if selected_unit_ids:
            # Validate ALL requested unit IDs before processing
            for unit_id in selected_unit_ids:
                require_unit_in_scope(current_user, unit_id)
        
        # Get units to process
        if selected_unit_ids:
            # Get specific units by ID
            all_units = get_units_by_type(unit_type)
            units = [u for u in all_units if u["id"] in selected_unit_ids]
        else:
            # Get all units of this type
            units = get_units_by_type(unit_type)
            # Validate all units when processing "all"
            for unit in units:
                require_unit_in_scope(current_user, unit["id"])
        
        print(f"[MULTI-SEASONAL] Generating {len(units)} seasonal reports for {report_level} level")
        print(f"[MULTI-SEASONAL] All {len(units)} units validated against user scope")
        print(f"[MULTI-SEASONAL] Season: {period} {year} (season_id={season_id})")
        print(f"[MULTI-SEASONAL] Mode: PHASE 2 - Generating 2 files per unit (Regular + Comparison with Charts)")
        
        # Get previous season for comparisons
        previous_season_id = get_previous_season(season_id)
        if previous_season_id:
            print(f"[MULTI-SEASONAL] Previous season ID: {previous_season_id}")
        else:
            print(f"[MULTI-SEASONAL] No previous season found - comparisons will show zero data")
        
        # Track results for summary
        successful_units = []
        empty_units = []
        failed_units = []
        
        # Generate ZIP file
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # Generate report for each unit
            for unit in units:
                unit_id = unit["id"]
                unit_name = unit["name"]
                
                try:
                    # PHASE 2: Generate BOTH current and previous seasonal reports
                    comparative_data = self._generate_unit_seasonal_reports(
                        season_id=season_id,
                        orgunit_id=unit_id,
                        orgunit_type=orgunit_type
                    )
                    
                    current_report = comparative_data['current_report']
                    previous_report = comparative_data['previous_report']
                    has_previous = comparative_data['has_previous']
                    
                    # Check if unit has data in current season
                    has_data = False
                    incidents_count = 0
                    
                    if isinstance(current_report, dict):
                        # Check classification stats
                        classification_stats = current_report.get("classification_stats", [])
                        if classification_stats and len(classification_stats) > 0:
                            has_data = True
                            # Sum total incidents from classification stats
                            incidents_count = sum(stat.get("total_count", 0) for stat in classification_stats)
                    
                    if not has_data or incidents_count == 0:
                        print(f"[MULTI-SEASONAL] {unit_name} (ID {unit_id}): No data - skipping files")
                        empty_units.append({"id": unit_id, "name": unit_name})
                        continue
                    
                    safe_name = self._sanitize_filename(unit_name)
                    files_generated = []
                    
                    # FILE 1: Generate regular seasonal report
                    if file_format == "docx":
                        regular_content = generate_seasonal_word_report(
                            seasonal_data=current_report,
                            language=language
                        )
                    else:
                        # For other formats, use existing converter
                        regular_content = self._generate_file_content(
                            report_data=current_report,
                            file_format=file_format,
                            language=language,
                            unit_name=unit_name,
                            unit_type=report_level,
                            period=period,
                            year=year
                        )
                    
                    regular_filename = f"Seasonal_Report_{safe_name}_{period}{year}.{file_format}"
                    zip_file.writestr(regular_filename, regular_content)
                    files_generated.append(regular_filename)
                    print(f"[MULTI-SEASONAL] ✓ {unit_name}: Regular report → {regular_filename}")
                    
                    # FILE 2: Generate comparison report (if previous data exists or for DOCX)
                    if file_format == "docx":
                        # Always generate comparison for DOCX (shows zero data gracefully)
                        comparison_content = self._generate_comparison_content(
                            current_data=current_report,
                            previous_data=previous_report,
                            file_format=file_format,
                            language=language,
                            unit_name=unit_name,
                            current_period=period,
                            previous_period=previous_report.get('header', {}).get('period', 'N/A')
                        )
                        
                        comparison_filename = f"Comparison_{safe_name}_{period}_vs_{previous_report.get('header', {}).get('period', 'Previous')}.{file_format}"
                        zip_file.writestr(comparison_filename, comparison_content)
                        files_generated.append(comparison_filename)
                        print(f"[MULTI-SEASONAL] ✓ {unit_name}: Comparison report → {comparison_filename}")
                    
                    successful_units.append({
                        "id": unit_id,
                        "name": unit_name,
                        "filenames": files_generated,
                        "incidents_count": incidents_count,
                        "has_comparison": file_format == "docx",
                        "has_previous_data": has_previous
                    })
                    
                    print(f"[MULTI-SEASONAL] ✅ {unit_name}: {incidents_count} incidents, {len(files_generated)} files generated")
                    
                except Exception as e:
                    print(f"[MULTI-SEASONAL] ✗ {unit_name} (ID {unit_id}): Failed - {str(e)}")
                    import traceback
                    traceback.print_exc()
                    failed_units.append({
                        "id": unit_id,
                        "name": unit_name,
                        "error": str(e)
                    })
            
            # Generate summary file
            summary_content = self._generate_summary_file(
                year=year,
                period=period,
                report_level=report_level,
                successful_units=successful_units,
                empty_units=empty_units,
                failed_units=failed_units,
                language=language
            )
            
            zip_file.writestr("_SUMMARY_Report.docx", summary_content)
            print(f"[MULTI-SEASONAL] Summary file added")
        
        # Prepare ZIP for download
        zip_buffer.seek(0)
        
        # Create ZIP filename (Phase 2: includes "With_Comparison")
        zip_filename = f"Seasonal_Reports_With_Comparison_{report_level.capitalize()}_{period}{year}.zip"
        
        total_files = sum(len(u.get("filenames", [])) for u in successful_units)
        print(f"[MULTI-SEASONAL] Complete: {len(successful_units)} units, {total_files} files, {len(empty_units)} empty, {len(failed_units)} failed")
        
        return {
            "filename": zip_filename,
            "content": zip_buffer.getvalue(),
            "content_type": "application/zip"
        }
    
    def _generate_unit_seasonal_report(
        self,
        season_id: int,
        orgunit_id: int,
        orgunit_type: int
    ) -> Dict[str, Any]:
        """Generate seasonal report data for a specific organizational unit."""
        
        # Use the same orchestrator as the view endpoint
        result = get_or_generate_seasonal_report(
            season_id=season_id,
            orgunit_id=orgunit_id,
            orgunit_type=orgunit_type,
            user_id=1  # System user for batch exports
        )
        
        return result
    
    def _generate_unit_seasonal_reports(
        self,
        season_id: int,
        orgunit_id: int,
        orgunit_type: int
    ) -> Dict[str, Any]:
        """
        Generate BOTH current and previous seasonal reports for a unit.
        
        Returns:
            Dict with:
                - current_report: Current season report data
                - previous_report: Previous season report data (or empty dict)
                - has_previous: Boolean indicating if previous season has data
        """
        # Get current season report
        current_report = get_or_generate_seasonal_report(
            season_id=season_id,
            orgunit_id=orgunit_id,
            orgunit_type=orgunit_type,
            user_id=1
        )
        
        # Get previous season
        previous_season_id = get_previous_season(season_id)
        
        if previous_season_id:
            try:
                previous_report = get_or_generate_seasonal_report(
                    season_id=previous_season_id,
                    orgunit_id=orgunit_id,
                    orgunit_type=orgunit_type,
                    user_id=1
                )
                
                # Check if previous has data
                prev_stats = previous_report.get('classification_stats', [])
                has_previous = len(prev_stats) > 0
                
            except Exception as e:
                print(f"[MULTI-SEASONAL] Could not fetch previous season: {e}")
                previous_report = self._create_empty_seasonal_report()
                has_previous = False
        else:
            previous_report = self._create_empty_seasonal_report()
            has_previous = False
        
        return {
            'current_report': current_report,
            'previous_report': previous_report,
            'has_previous': has_previous
        }
    
    def _create_empty_seasonal_report(self) -> Dict[str, Any]:
        """Create an empty seasonal report structure for cases with no previous data."""
        return {
            'header': {
                'season_id': None,
                'period': 'N/A',
                'total_cases': 0,
                'clinical_domain_count': 0,
                'management_domain_count': 0,
                'relational_domain_count': 0,
                'low_severity_count': 0,
                'medium_severity_count': 0,
                'high_severity_count': 0,
                'is_compliant': True
            },
            'classification_stats': [],
            'domain_totals': {}
        }
    
    def _generate_comparison_content(
        self,
        current_data: Dict[str, Any],
        previous_data: Dict[str, Any],
        file_format: str,
        language: str,
        unit_name: str,
        current_period: str,
        previous_period: str
    ) -> bytes:
        """
        Generate comparison report content with charts.
        
        This generates the beautiful comparison report with 9 charts
        (3 levels × 3 chart types) using the comparative formatter.
        """
        if file_format == "docx":
            # Use the Phase 1 comparative formatter with all 9 charts!
            return generate_comparative_seasonal_word_report(
                current_data=current_data,
                previous_data=previous_data,
                language=language
            )
        elif file_format == "pdf":
            # Generate Word first, then convert to PDF (future implementation)
            word_bytes = generate_comparative_seasonal_word_report(
                current_data=current_data,
                previous_data=previous_data,
                language=language
            )
            # TODO: Convert to PDF using docx2pdf or similar
            return word_bytes
        else:
            # For CSV/XLSX: Generate side-by-side comparison table
            # Use existing export logic
            export_data = self._convert_seasonal_to_export_format(current_data)
            return reports_service.generate_csv_export(
                report_data=export_data,
                filename="temp.csv",
                language=language
            )
    
    def _generate_file_content(
        self,
        report_data: Dict[str, Any],
        file_format: str,
        language: str,
        unit_name: str = None,
        unit_type: str = None,
        period: str = None,
        year: int = None
    ) -> bytes:
        """Generate file content in specified format for seasonal report."""
        
        # Seasonal reports have a different structure than monthly
        # They contain: classification_stats, domain_totals, policy_snapshot, etc.
        
        # For now, we'll convert seasonal report format to a flat list for export
        # This matches what the frontend expects and what the export services can handle
        
        export_data = self._convert_seasonal_to_export_format(report_data)
        
        # Determine entity parameters for header
        report_administration = None
        report_department = None
        report_section = None
        
        # Helper function to get unique values from data
        def get_unique_values(data, field_name):
            """Get unique non-null values from data"""
            if not data or not isinstance(data, list):
                return []
            values = set()
            for row in data:
                val = row.get(field_name)
                if val and val != "—":
                    values.add(val)
            return list(values)
        
        # Extract hierarchy information
        unique_admins = get_unique_values(export_data, "administration_name")
        unique_depts = get_unique_values(export_data, "department_name")
        unique_sections = get_unique_values(export_data, "section_name")
        
        print(f"[MULTI SEASONAL EXPORT] Unit: {unit_name} ({unit_type})")
        print(f"[MULTI SEASONAL EXPORT] Data has {len(export_data) if isinstance(export_data, list) else 0} records")
        
        # Set header values based on report level
        if unit_type == "administration":
            report_administration = unit_name
            report_department = unique_depts[0] if len(unique_depts) == 1 else ("متعدد" if len(unique_depts) > 1 else "—")
            report_section = unique_sections[0] if len(unique_sections) == 1 else ("متعدد" if len(unique_sections) > 1 else "—")
        elif unit_type == "department":
            report_department = unit_name
            report_administration = unique_admins[0] if unique_admins else "—"
            report_section = unique_sections[0] if len(unique_sections) == 1 else ("متعدد" if len(unique_sections) > 1 else "—")
        elif unit_type == "section":
            report_section = unit_name
            report_department = unique_depts[0] if unique_depts else "—"
            report_administration = unique_admins[0] if unique_admins else "—"
        
        # Generate file based on format
        # Use the same export services as monthly reports
        if file_format == "docx":
            return reports_service.generate_docx_export(
                report_data=export_data,
                filename="temp.docx",
                language=language,
                report_entity_name=unit_name,
                report_entity_type=unit_type,
                report_administration=report_administration,
                report_department=report_department,
                report_section=report_section
            )
        elif file_format == "xlsx":
            return reports_service.generate_xlsx_export(
                report_data=export_data,
                filename="temp.xlsx",
                language=language
            )
        elif file_format == "pdf":
            return reports_service.generate_pdf_export(
                report_data=export_data,
                filename="temp.pdf",
                language=language,
                include_charts=True,
                report_entity_name=unit_name,
                report_entity_type=unit_type,
                report_administration=report_administration,
                report_department=report_department,
                report_section=report_section
            )
        else:  # csv
            return reports_service.generate_csv_export(
                report_data=export_data,
                filename="temp.csv",
                language=language
            )
    
    def _convert_seasonal_to_export_format(self, report_data: Dict[str, Any]) -> List[Dict]:
        """
        Convert seasonal report structure to flat export format.
        
        Seasonal report has: classification_stats, domain_totals, policy_snapshot
        Export format needs: list of rows with standard fields
        """
        # For seasonal exports, we'll create summary rows from classification_stats
        export_rows = []
        
        classification_stats = report_data.get("classification_stats", [])
        
        for stat in classification_stats:
            row = {
                "classification_name": stat.get("classification_name", "—"),
                "total_count": stat.get("total_count", 0),
                "preventive_measures_count": stat.get("preventive_measures_count", 0),
                "preventive_percentage": stat.get("preventive_percentage", 0),
                "administration_name": "—",  # Seasonal reports aggregate across units
                "department_name": "—",
                "section_name": "—"
            }
            export_rows.append(row)
        
        return export_rows
    
    def _generate_summary_file(
        self,
        year: int,
        period: str,
        report_level: str,
        successful_units: List[Dict],
        empty_units: List[Dict],
        failed_units: List[Dict],
        language: str
    ) -> bytes:
        """Generate summary Word document."""
        
        doc = Document()
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(f"Seasonal Reports Summary - {report_level.capitalize()} Level")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = 1  # Center
        
        # Period
        doc.add_paragraph(f"Period: {period} {year}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Statistics
        doc.add_heading("Summary Statistics", level=2)
        total_units = len(successful_units) + len(empty_units) + len(failed_units)
        total_incidents = sum(u["incidents_count"] for u in successful_units)
        total_files = sum(len(u.get("filenames", [])) for u in successful_units)
        units_with_comparison = sum(1 for u in successful_units if u.get("has_comparison"))
        
        doc.add_paragraph(f"Total Units Processed: {total_units}")
        doc.add_paragraph(f"Units with Data: {len(successful_units)}")
        doc.add_paragraph(f"Units with No Incidents: {len(empty_units)}")
        doc.add_paragraph(f"Failed Units: {len(failed_units)}")
        doc.add_paragraph(f"Total Incidents: {total_incidents}")
        doc.add_paragraph(f"Total Files Generated: {total_files}")
        doc.add_paragraph(f"Units with Comparison Reports: {units_with_comparison}")
        doc.add_paragraph()
        
        # Successful units
        if successful_units:
            doc.add_heading("Units with Data (Files Generated)", level=2)
            for unit in successful_units:
                filenames_str = ", ".join(unit.get('filenames', []))
                comparison_status = "✓ With Comparison" if unit.get("has_comparison") else "○ No Comparison"
                doc.add_paragraph(
                    f"✓ {unit['name']} - {unit['incidents_count']} incidents - {comparison_status}",
                    style='List Bullet'
                )
                for filename in unit.get('filenames', []):
                    doc.add_paragraph(f"    → {filename}", style='List Bullet 2')
            doc.add_paragraph()
        
        # Empty units
        if empty_units:
            doc.add_heading("Units with No Incidents (No Files)", level=2)
            for unit in empty_units:
                doc.add_paragraph(f"○ {unit['name']} - No incidents in this period", style='List Bullet')
            doc.add_paragraph()
        
        # Failed units
        if failed_units:
            doc.add_heading("Failed Units (Errors)", level=2)
            for unit in failed_units:
                doc.add_paragraph(f"✗ {unit['name']} - Error: {unit['error']}", style='List Bullet')
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _sanitize_filename(self, name: str) -> str:
        """Sanitize unit name for use in filename."""
        # Remove or replace characters that are invalid in filenames
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '_')
        # Limit length
        return name[:50]


# Singleton instance
multi_seasonal_export_service = MultiSeasonalExportService()
