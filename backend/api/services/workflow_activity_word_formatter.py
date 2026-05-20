"""
Workflow Activity Report — Word Formatter  (Compact Narrative Design)

Philosophy:
  - Workflow expressed as inline colored runs inside a single paragraph, NOT tables
  - Case header = shaded paragraph with bottom border (no table)
  - Subcase = one line: identifier + workflow states inline
  - Responses = plain labeled paragraphs (answered levels only)
  - Action items = compact inline lines (<=3) or minimal 3-col table (4+)
  - Target density: 6-10 cases per page

No large workflow tables. No empty placeholder rows. No decorative grid noise.
"""

import os
from typing import Dict, Any, Tuple, Optional
from datetime import date, datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell, _Row


# ============================================================================
# WORKFLOW STATUS MAPS
# ============================================================================

WORKFLOW_LANE_MAP = {
    "SUBMITTED_TO_SECTION":             ("waiting",  "not_reached", "not_reached", "بانتظار رد القسم"),
    "SECTION_DENIED":                   ("rejected", "not_reached", "not_reached", "القسم رفض الإحالة"),
    "SECTION_ACCEPTED_PENDING_DEPT":    ("done",     "waiting",     "not_reached", "بانتظار رد الدائرة"),
    "RETURNED_TO_SECTION_FOR_REVISION": ("revision", "not_reached", "not_reached", "أُعيدت للقسم للمراجعة"),
    "DEPT_REJECTED":                    ("done",     "rejected",    "not_reached", "الدائرة رفضت الإحالة"),
    "DEPT_ACCEPTED_PENDING_ADMIN":      ("done",     "done",        "waiting",     "بانتظار موافقة الإدارة"),
    "RETURNED_TO_DEPT_FOR_REVISION":    ("done",     "revision",    "not_reached", "أُعيدت للدائرة للمراجعة"),
    "ADMIN_REJECTED":                   ("done",     "done",        "rejected",    "الإدارة رفضت"),
    "ADMIN_APPROVED":                   ("done",     "done",        "done",        "مكتملة"),
    "CLOSED":                           ("done",     "done",        "done",        "مغلقة"),
    "FORCE_CLOSED":                     ("forced",   "forced",      "forced",      "مغلقة قسراً"),
    "FORCE_CLOSED_DRAFT":               ("forced",   "forced",      "forced",      "مغلقة قسراً (مسودة)"),
    "FORCE_CLOSED_COMPLETE":            ("forced",   "forced",      "forced",      "مغلقة قسراً — مكتملة"),
}

# Visual config for each state as inline run
STATE_RUN = {
    "done":        {"sym": "✔", "sz": 9,  "bold": False, "color": "2E7D32"},
    "waiting":     {"sym": "●", "sz": 11, "bold": True,  "color": "6B4200"},
    "not_reached": {"sym": "",  "sz": 8,  "bold": False, "color": "CCCCCC"},
    "rejected":    {"sym": "✖", "sz": 10, "bold": True,  "color": "8B0000"},
    "revision":    {"sym": "↩", "sz": 10, "bold": True,  "color": "7B3800"},
    "forced":      {"sym": "■", "sz": 8,  "bold": False, "color": "888888"},
}

LANE_NAMES = ["القسم", "الدائرة", "الإدارة"]  # logical order = RTL visual order

AI_STATUS_AR = {
    "DRAFT":              "مسودة",
    "SUBMITTED_TO_DEPT":  "مُقدَّم للدائرة",
    "DEPT_REJECTED":      "رفض الدائرة",
    "SUBMITTED_TO_ADMIN": "مُقدَّم للإدارة",
    "ADMIN_REJECTED":     "رفض الإدارة",
    "ADMIN_APPROVED":     "موافق عليه",
    "IN_PROGRESS":        "قيد التنفيذ",
    "DONE":               "مُنجز",
    "VERIFIED":           "مُتحقق منه",
    "CANCELLED":          "ملغى",
}

TOTAL_W = 267.0
ACTIVE_STATES = {"waiting", "rejected", "revision"}


# ============================================================================
# WORKFLOW INTELLIGENCE
# ============================================================================

def _derive_workflow_lane(status: str) -> Tuple[str, str, str, Optional[str]]:
    key = (status or "").strip().upper()
    return WORKFLOW_LANE_MAP.get(
        key, ("waiting", "not_reached", "not_reached", f"حالة غير مصنَّفة: {status}")
    )


def _derive_case_health(subcases: list) -> Tuple[str, str]:
    """Return (symbol, phrase) for inline health display in the case header."""
    if not subcases:
        return ("—", "لا توجد قضايا فرعية")

    statuses = [sc.get("status", "") for sc in subcases]
    has_overdue = any(
        ai.get("is_overdue")
        for sc in subcases
        for ai in sc.get("action_items", [])
    )
    has_rejected  = any("REJECTED" in s or "DENIED" in s for s in statuses)
    force_closed  = any("FORCE_CLOSED" in s for s in statuses)
    all_complete  = all(s in ("ADMIN_APPROVED", "CLOSED", "FORCE_CLOSED_COMPLETE")
                        for s in statuses)
    any_revision  = any("REVISION" in s for s in statuses)

    if has_overdue:   return ("!", "إجراءات متأخرة")
    if has_rejected:  return ("✖", "مرفوضة — يلزم مراجعة")
    if force_closed and all_complete: return ("■", "مغلقة قسراً")
    if all_complete:  return ("✔", "مكتملة")
    if any_revision:  return ("↩", "مراجعة مطلوبة")

    for sc in subcases:
        _, _, _, phrase = _derive_workflow_lane(sc.get("status", ""))
        if phrase and phrase not in ("مكتملة", "مغلقة"):
            return ("●", phrase)

    return ("●", "قيد المتابعة")


# ============================================================================
# LOW-LEVEL HELPERS
# ============================================================================

def _mm_to_dxa(mm: float) -> int:
    return int(mm * 1440 / 25.4)


def _fmt_date(d) -> str:
    if d is None: return "—"
    if isinstance(d, datetime): return d.strftime("%Y-%m-%d")
    if isinstance(d, date):     return d.strftime("%Y-%m-%d")
    return str(d)


def _ar_run(para, text: str, size: int = 10, bold: bool = False,
            italic: bool = False, color: str = None) -> object:
    run = para.add_run(text or "")
    run.font.name = 'Traditional Arabic'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _new_para(doc: Document, align: str = "right",
              space_before: float = 0, space_after: float = 3,
              rtl: bool = True) -> object:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if rtl:
        p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    return p


def _set_para_shading(para, fill: str):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _set_para_bottom_border(para, color: str = '2E4999', sz: int = 6):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_row_col_width(row: _Row, col_idx: int, mm: float):
    dxa = str(_mm_to_dxa(mm))
    tc = row.cells[col_idx]._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), dxa)
    tcW.set(qn('w:type'), 'dxa')


def _set_rtl_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblPr.append(OxmlElement('w:bidiVisual'))


def _set_cell_shading(cell: _Cell, fill: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)


def _cell_para(cell: _Cell, align: str = "center") -> object:
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                   else WD_ALIGN_PARAGRAPH.RIGHT)
    p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return p


def _apply_minimal_table_borders(table,
                                  outer: str = 'BBBBBB', outer_sz: int = 4,
                                  inner: str = 'DDDDDD', inner_sz: int = 4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    bdr = OxmlElement('w:tblBorders')
    for name, color, sz in [
        ('top',     outer, outer_sz), ('left',    outer, outer_sz),
        ('bottom',  outer, outer_sz), ('right',   outer, outer_sz),
        ('insideH', inner, inner_sz), ('insideV', inner, inner_sz),
    ]:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        bdr.append(b)
    tblPr.append(bdr)


# ============================================================================
# CASE HEADER  (shaded paragraph, no table)
# ============================================================================

def _make_case_header(doc: Document, case_id, patient_name,
                      feedback_date, health_sym: str, health_phrase: str):
    p = _new_para(doc, space_before=10, space_after=0)
    _set_para_shading(p, 'EEF2FF')
    _set_para_bottom_border(p, color='1C3A7A', sz=8)
    p.paragraph_format.left_indent  = Mm(3)
    p.paragraph_format.right_indent = Mm(3)

    name_text = patient_name or "غير محدد"
    _ar_run(p, f"الحالة #{case_id}   |   {name_text}   |   {_fmt_date(feedback_date)}",
            size=11, bold=True, color='1C3A7A')
    _ar_run(p, f"      {health_sym}  {health_phrase}",
            size=9, italic=True, color='667799')


# ============================================================================
# COMPLAINT  (plain paragraph)
# ============================================================================

def _make_complaint_para(doc: Document, complaint_text: str):
    if not (complaint_text and complaint_text.strip()):
        return
    p = _new_para(doc, space_before=5, space_after=6)
    p.paragraph_format.right_indent = Mm(4)
    clipped = complaint_text.strip()
    if len(clipped) > 350:
        clipped = clipped[:350] + "..."
    _ar_run(p, "الشكوى:   ", size=9, bold=False, color='999999')
    _ar_run(p, clipped, size=10, italic=True, color='333333')


# ============================================================================
# SUBCASE LINE  (identity + workflow states — all in one paragraph)
# ============================================================================

def _make_subcase_line(doc: Document, subcase_id, org_name: str,
                        sec_state: str, dept_state: str, admin_state: str):
    """
    Single paragraph:
      قضية #N — [unit]      القسم ✔   ·   الدائرة ●   ·   الإدارة
    Owner state is bold + colored. Done is green. Not-reached is near-invisible.
    """
    p = _new_para(doc, space_before=8, space_after=2)
    p.paragraph_format.right_indent = Mm(4)

    # Subcase identifier
    org_display = (org_name[:40] + "...") if len(org_name) > 40 else org_name
    _ar_run(p, f"قضية #{subcase_id}  —  {org_display}",
            size=10, bold=True, color='1C3A7A')

    # Spacer then workflow states
    _ar_run(p, "          ", size=9)   # visual gap between identifier and states

    states = [sec_state, dept_state, admin_state]
    for idx, (state, name) in enumerate(zip(states, LANE_NAMES)):
        cfg = STATE_RUN[state]

        if state == "not_reached":
            _ar_run(p, name, size=8, color='CCCCCC')
        elif state == "forced":
            _ar_run(p, f"{name} ■", size=8, italic=True, color='999999')
        else:
            sym = cfg["sym"]
            text = f"{name} {sym}" if sym else name
            _ar_run(p, text, size=cfg["sz"], bold=cfg["bold"], color=cfg["color"])

        if idx < 2:
            _ar_run(p, "   ·   ", size=8, color='CCCCCC')


# ============================================================================
# EXPLANATION PARAGRAPHS  (answered levels only)
# ============================================================================

RESPONSE_LABELS = {
    "section": "رد القسم",
    "dept":    "رد الدائرة",
    "admin":   "رد الإدارة",
}


def _make_explanation_paras(doc: Document, section_exp, dept_exp, admin_exp):
    entries = [
        ("section", section_exp),
        ("dept",    dept_exp),
        ("admin",   admin_exp),
    ]
    for key, text in entries:
        if not (text and text.strip()):
            continue
        p = _new_para(doc, space_before=3, space_after=2)
        p.paragraph_format.right_indent = Mm(6)
        clipped = text.strip()
        if len(clipped) > 400:
            clipped = clipped[:400] + "..."
        _ar_run(p, f"{RESPONSE_LABELS[key]}:   ", size=9, bold=True, color='555555')
        _ar_run(p, clipped, size=10, color='222222')


# ============================================================================
# ACTION ITEMS  (compact lines for <=3, minimal table for 4+)
# ============================================================================

def _make_action_items_section(doc: Document, action_items: list):
    if not action_items:
        return

    if len(action_items) <= 3:
        # Compact inline paragraphs — one per item
        for ai in action_items:
            is_overdue   = ai.get("is_overdue", False)
            completed_at = ai.get("completed_at")
            title = (ai.get("title") or "—")[:80]
            due   = _fmt_date(ai.get("due_date"))

            status_raw = (ai.get("status") or "").upper()
            status_ar  = AI_STATUS_AR.get(status_raw, ai.get("status") or "—")

            p = _new_para(doc, space_before=2, space_after=1)
            p.paragraph_format.right_indent = Mm(6)

            if is_overdue:
                _ar_run(p, "! ", size=9, bold=True, color='C62828')
                _ar_run(p, title, size=9, bold=True, color='333333')
                _ar_run(p, f"   ·   {due}   ·   ", size=8, color='AAAAAA')
                _ar_run(p, f"متأخر {ai.get('days_overdue', '?')} يوم",
                        size=9, bold=True, color='C62828')
            elif completed_at:
                _ar_run(p, "✔ ", size=9, color='2E7D32')
                _ar_run(p, title, size=9, color='444444')
                _ar_run(p, f"   ·   مُنجز {_fmt_date(completed_at)}",
                        size=8, color='2E7D32')
            else:
                _ar_run(p, "· ", size=9, color='AAAAAA')
                _ar_run(p, title, size=9, color='333333')
                _ar_run(p, f"   ·   {due}   ·   {status_ar}",
                        size=8, color='888888')
    else:
        # Minimal 3-col table for larger lists
        _make_action_items_table(doc, action_items)


def _make_action_items_table(doc: Document, action_items: list):
    headers    = ["العنوان", "تاريخ الاستحقاق", "الحالة"]
    col_widths = [155.0, 47.0, 65.0]

    tbl = doc.add_table(rows=1 + len(action_items), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_rtl_table(tbl)
    _apply_minimal_table_borders(tbl, outer='AAAAAA', outer_sz=4,
                                      inner='DDDDDD', inner_sz=4)

    for ci, hdr in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        _set_cell_shading(c, '3A5280')
        ph = _cell_para(c, "center")
        _ar_run(ph, hdr, size=9, bold=True, color='FFFFFF')

    for ri, ai in enumerate(action_items, start=1):
        row = tbl.rows[ri]
        is_overdue   = ai.get("is_overdue", False)
        completed_at = ai.get("completed_at")
        row_bg = 'E8F5E9' if completed_at else ('FFF0F0' if is_overdue else None)

        status_raw = (ai.get("status") or "").upper()
        status_ar  = AI_STATUS_AR.get(status_raw, ai.get("status") or "—")

        def _tc(ci, text, bold=False, color=None):
            cell = row.cells[ci]
            if row_bg:
                _set_cell_shading(cell, row_bg)
            p = _cell_para(cell, "right" if ci == 0 else "center")
            _ar_run(p, text, size=9, bold=bold, color=color)

        _tc(0, (ai.get("title") or "—")[:80])
        _tc(1, _fmt_date(ai.get("due_date")))
        if is_overdue:
            _tc(2, f"! متأخر {ai.get('days_overdue', '?')} يوم",
                bold=True, color='C62828')
        elif completed_at:
            _tc(2, f"✔ مُنجز — {_fmt_date(completed_at)}", color='2E7D32')
        else:
            _tc(2, status_ar)

    for ci, w in enumerate(col_widths):
        for ri in range(len(action_items) + 1):
            _set_row_col_width(tbl.rows[ri], ci, w)


# ============================================================================
# CASE SEPARATOR
# ============================================================================

def _add_case_separator(doc: Document):
    p = _new_para(doc, space_before=10, space_after=2)
    _set_para_bottom_border(p, color='8899BB', sz=4)


# ============================================================================
# PER-SUBCASE / PER-CASE RENDERERS
# ============================================================================

def _render_subcase(doc: Document, subcase: Dict[str, Any]):
    org_name = (subcase.get("target_org_unit_name")
                or f"ID:{subcase.get('target_org_unit_id', '?')}")
    status = subcase.get("status", "")
    sec_state, dept_state, admin_state, _ = _derive_workflow_lane(status)

    _make_subcase_line(doc, subcase["subcase_id"], org_name,
                       sec_state, dept_state, admin_state)

    _make_explanation_paras(
        doc,
        subcase.get("section_explanation"),
        subcase.get("department_explanation"),
        subcase.get("administration_explanation"),
    )

    _make_action_items_section(doc, subcase.get("action_items", []))


def _render_case(doc: Document, case: Dict[str, Any]):
    subcases = case.get("subcases", [])
    health_sym, health_phrase = _derive_case_health(subcases)

    _make_case_header(doc, case["case_id"], case.get("patient_name"),
                      case.get("feedback_date"), health_sym, health_phrase)

    _make_complaint_para(doc, case.get("complaint_text") or "")

    if not subcases:
        p = _new_para(doc, space_before=4, space_after=4)
        p.paragraph_format.right_indent = Mm(4)
        _ar_run(p, "لا توجد قضايا فرعية مسجَّلة لهذه الحالة.",
                size=9, italic=True, color='AAAAAA')
    else:
        for subcase in subcases:
            _render_subcase(doc, subcase)

    _add_case_separator(doc)


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def generate_workflow_activity_word(report_data: Dict[str, Any]) -> bytes:
    doc = Document()

    sec = doc.sections[0]
    sec.page_height    = int(Mm(210))
    sec.page_width     = int(Mm(297))
    sec.orientation    = WD_ORIENT.LANDSCAPE
    sec.left_margin    = int(Mm(15))
    sec.right_margin   = int(Mm(15))
    sec.top_margin     = int(Mm(15))
    sec.bottom_margin  = int(Mm(15))

    doc.styles['Normal'].font.name = 'Traditional Arabic'
    doc.styles['Normal'].font.size = Pt(10)

    meta  = report_data.get("meta", {})
    cases = report_data.get("cases", [])

    # Logo
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            sec.header_distance = int(Inches(0.1))
            hdr_p = sec.header.paragraphs[0]
            hdr_p.clear()
            hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_p.add_run().add_picture(logo_path, width=int(Inches(0.9)))
    except Exception:
        pass

    # Report title
    title_p = _new_para(doc, align="center", space_before=0, space_after=2)
    _ar_run(title_p, "تقرير نشاط سير العمل القسمي", size=20, bold=True, color='1C3A7A')

    sub_p = _new_para(doc, align="center", space_before=0, space_after=6, rtl=False)
    r = sub_p.add_run("Section Workflow Activity Report")
    r.font.size = Pt(13)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

    # Summary info table
    info = doc.add_table(rows=2, cols=3)
    info.style = 'Table Grid'
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_rtl_table(info)
    _apply_minimal_table_borders(info, outer='AAAAAA', outer_sz=4,
                                       inner='CCCCCC', inner_sz=4)

    for ci, lbl in enumerate(["إجمالي الإجراءات", "النطاق", "الفترة الزمنية"]):
        _set_cell_shading(info.rows[0].cells[ci], '1C3A7A')
        _ar_run(_cell_para(info.rows[0].cells[ci], "center"),
                lbl, size=10, bold=True, color='FFFFFF')

    period_str = f"{_fmt_date(meta.get('start_date'))}  —  {_fmt_date(meta.get('end_date'))}"
    counts_str = (
        f"الحالات: {meta.get('total_cases', 0)}  |  "
        f"الفرعية: {meta.get('total_subcases', 0)}  |  "
        f"الإجراءات: {meta.get('total_action_items', 0)}"
    )
    for ci, val in enumerate([counts_str, meta.get('scope', '—'), period_str]):
        _set_cell_shading(info.rows[1].cells[ci], 'F0F4FF')
        _ar_run(_cell_para(info.rows[1].cells[ci], "center"), val, size=10)

    for ci, w in enumerate([95.0, 65.0, 107.0]):
        for ri in range(2):
            _set_row_col_width(info.rows[ri], ci, w)

    _new_para(doc, space_before=0, space_after=10)

    # Cases
    if not cases:
        p = _new_para(doc, space_before=0, space_after=0)
        _ar_run(p, "لا توجد حالات في الفترة المحددة — No cases found for this period.",
                size=11, italic=True, color='888888')
    else:
        for case in cases:
            _render_case(doc, case)

    # Footer
    footer_p = _new_para(doc, align="center", space_before=8, space_after=0)
    _ar_run(footer_p,
            f"تاريخ الإنشاء: {_fmt_date(meta.get('generated_at'))}"
            f"  |  بواسطة: {meta.get('generated_by', 'System')}",
            size=9, color='888888')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
