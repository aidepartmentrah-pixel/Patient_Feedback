"""
Reports Service Layer
Handles business logic, data aggregation, filtering, and export preparation.
"""

from datetime import datetime, date, timedelta
from typing import Dict, List, Any, Optional, Tuple
from io import BytesIO, StringIO
import csv
import json
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE


try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from ..db_layer.reports_db import (
    get_filtered_complaints,
    get_monthly_statistics,
    get_seasonal_hcat,
    get_bulk_summary
)
from ..db_layer.admin_units import get_admin_unit_by_id, get_units_by_type


class ReportsService:
    """Service class for report operations."""
    
    @staticmethod
    def get_period_dates(
        report_type: str,
        year: int,
        month: Optional[int] = None,
        trimester: Optional[int] = None,
        quarter: Optional[int] = None,
        start_date: Optional[date] = None,
        end_date: Optional[date] = None
    ) -> Tuple[date, date, str, str]:
        """
        Calculate period start/end dates and labels.
        
        Returns:
            Tuple of (start_date, end_date, label, label_ar)
        """
        if start_date and end_date:
            label = f"Custom Range {start_date} to {end_date}"
            label_ar = f"نطاق مخصص {start_date} إلى {end_date}"
            return start_date, end_date, label, label_ar
        
        if report_type == "monthly":
            if not month or month < 1 or month > 12:
                raise ValueError("Month required and must be 1-12")
            
            # First day of month
            start = date(year, month, 1)
            # Last day of month
            if month == 12:
                end = date(year + 1, 1, 1) - timedelta(days=1)
            else:
                end = date(year, month + 1, 1) - timedelta(days=1)
            
            months = {
                1: ("January", "يناير"),
                2: ("February", "فبراير"),
                3: ("March", "مارس"),
                4: ("April", "أبريل"),
                5: ("May", "مايو"),
                6: ("June", "يونيو"),
                7: ("July", "يوليو"),
                8: ("August", "أغسطس"),
                9: ("September", "سبتمبر"),
                10: ("October", "أكتوبر"),
                11: ("November", "نوفمبر"),
                12: ("December", "ديسمبر")
            }
            
            en_label = f"{months[month][0]} {year}"
            ar_label = f"{months[month][1]} {year}"
            return start, end, en_label, ar_label
        
        elif report_type == "seasonal":
            if trimester:
                # Trimester: 4-month periods
                if trimester == 1:
                    start = date(year, 1, 1)
                    end = date(year, 4, 30)
                    label = f"Trimester 1 - {year}"
                    label_ar = f"الفصل الأول - {year}"
                elif trimester == 2:
                    start = date(year, 5, 1)
                    end = date(year, 8, 31)
                    label = f"Trimester 2 - {year}"
                    label_ar = f"الفصل الثاني - {year}"
                elif trimester == 3:
                    start = date(year, 9, 1)
                    end = date(year, 12, 31)
                    label = f"Trimester 3 - {year}"
                    label_ar = f"الفصل الثالث - {year}"
                else:
                    raise ValueError("Trimester must be 1, 2, or 3")
                return start, end, label, label_ar
            
            elif quarter:
                # Quarter: 3-month periods
                if quarter == 1:
                    start = date(year, 1, 1)
                    end = date(year, 3, 31)
                    label = f"Q1 - {year}"
                    label_ar = f"الربع الأول - {year}"
                elif quarter == 2:
                    start = date(year, 4, 1)
                    end = date(year, 6, 30)
                    label = f"Q2 - {year}"
                    label_ar = f"الربع الثاني - {year}"
                elif quarter == 3:
                    start = date(year, 7, 1)
                    end = date(year, 9, 30)
                    label = f"Q3 - {year}"
                    label_ar = f"الربع الثالث - {year}"
                elif quarter == 4:
                    start = date(year, 10, 1)
                    end = date(year, 12, 31)
                    label = f"Q4 - {year}"
                    label_ar = f"الربع الرابع - {year}"
                else:
                    raise ValueError("Quarter must be 1, 2, 3, or 4")
                return start, end, label, label_ar
            else:
                raise ValueError("Trimester or Quarter required for seasonal report")
        
        else:
            raise ValueError("report_type must be 'monthly' or 'seasonal'")
    
    @staticmethod
    def get_filtered_complaints_with_pagination(
        report_type: str,
        year: int,
        month: Optional[int] = None,
        trimester: Optional[int] = None,
        quarter: Optional[int] = None,
        start_date: Optional[date] = None,
        end_date: Optional[date] = None,
        building_id: Optional[int] = None,
        idara_id: Optional[int] = None,
        dayra_id: Optional[int] = None,
        qism_id: Optional[int] = None,
        domain_id: Optional[int] = None,
        category_id: Optional[int] = None,
        severity_id: Optional[int] = None,
        status: Optional[str] = None,
        page: int = 1,
        page_size: int = 50
    ) -> Dict[str, Any]:
        """Fetch complaints with pagination for detailed monthly view."""
        
        # Get period dates
        period_start, period_end, label, label_ar = ReportsService.get_period_dates(
            report_type, year, month, trimester, quarter, start_date, end_date
        )
        
        # Fetch complaints
        complaints, total_records = get_filtered_complaints(
            year=year,
            month=month,
            start_date=period_start,
            end_date=period_end,
            building_id=building_id,
            idara_id=idara_id,
            dayra_id=dayra_id,
            qism_id=qism_id,
            domain_id=domain_id,
            category_id=category_id,
            severity_id=severity_id,
            status=status,
            page=page,
            page_size=page_size
        )
        
        total_pages = (total_records + page_size - 1) // page_size
        
        return {
            "complaints": complaints,
            "pagination": {
                "page": page,
                "page_size": page_size,
                "total_records": total_records,
                "total_pages": total_pages
            },
            "period": {
                "label": label,
                "label_ar": label_ar,
                "start_date": period_start.isoformat(),
                "end_date": period_end.isoformat()
            }
        }
    
    @staticmethod
    def get_monthly_statistics_report(
        year: int,
        month: Optional[int] = None,
        start_date: Optional[date] = None,
        end_date: Optional[date] = None,
        building_id: Optional[int] = None,
        idara_id: Optional[int] = None,
        dayra_id: Optional[int] = None,
        qism_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """Fetch aggregated monthly statistics."""
        
        # Get period dates
        period_start, period_end, label, label_ar = ReportsService.get_period_dates(
            "monthly", year, month, None, None, start_date, end_date
        )
        
        # Fetch statistics
        stats = get_monthly_statistics(
            year=year,
            month=month,
            start_date=period_start,
            end_date=period_end,
            building_id=building_id,
            idara_id=idara_id,
            dayra_id=dayra_id,
            qism_id=qism_id
        )
        
        return {
            "period": {
                "year": year,
                "month": month,
                "label": label,
                "label_ar": label_ar,
                "start_date": period_start.isoformat(),
                "end_date": period_end.isoformat()
            },
            "summary": stats["summary"],
            "by_domain": stats["by_domain"],
            "by_category": stats["by_category"],
            "by_severity": stats["by_severity"],
            "by_department": stats["by_department"]
        }
    
    @staticmethod
    def get_seasonal_hcat_report(
        year: int,
        trimester: Optional[int] = None,
        quarter: Optional[int] = None,
        start_date: Optional[date] = None,
        end_date: Optional[date] = None,
        threshold: Optional[int] = None,
        building_id: Optional[int] = None,
        idara_id: Optional[int] = None,
        dayra_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """Fetch seasonal HCAT analysis with threshold evaluation."""
        
        # Get period dates
        period_start, period_end, label, label_ar = ReportsService.get_period_dates(
            "seasonal", year, None, trimester, quarter, start_date, end_date
        )
        
        # Use default threshold if not provided
        if threshold is None:
            threshold = 50
        
        # Fetch HCAT analysis
        hcat_data = get_seasonal_hcat(
            year=year,
            start_date=period_start,
            end_date=period_end,
            threshold=threshold,
            building_id=building_id,
            idara_id=idara_id,
            dayra_id=dayra_id
        )
        
        return {
            "period": {
                "year": year,
                "trimester": trimester,
                "quarter": quarter,
                "label": label,
                "label_ar": label_ar,
                "start_date": period_start.isoformat(),
                "end_date": period_end.isoformat()
            },
            "threshold": {
                "value": threshold,
                "source": "user_input"
            },
            "total_complaints": hcat_data["total_complaints"],
            "domains": hcat_data["domains"],
            "exceeding_count": hcat_data["exceeding_count"],
            "within_threshold_count": hcat_data["within_threshold_count"]
        }
    
    @staticmethod
    def get_bulk_summary_report(
        report_type: str,
        year: int,
        month: Optional[int] = None,
        trimester: Optional[int] = None,
        quarter: Optional[int] = None,
        building_id: Optional[int] = None,
        idara_id: Optional[int] = None
    ) -> Dict[str, Any]:
        """Fetch bulk export summary per department."""
        
        # Get period dates
        period_start, period_end, label, label_ar = ReportsService.get_period_dates(
            report_type, year, month, trimester, quarter
        )
        
        # Fetch summaries
        departments = get_bulk_summary(
            year=year,
            month=month,
            start_date=period_start,
            end_date=period_end,
            building_id=building_id,
            idara_id=idara_id
        )
        
        return {
            "period": {
                "year": year,
                "month": month,
                "label": label,
                "label_ar": label_ar,
                "start_date": period_start.isoformat(),
                "end_date": period_end.isoformat()
            },
            "departments": departments,
            "total_departments": len(departments),
            "grand_total_complaints": sum(d["total_complaints"] for d in departments)
        }
    
    @staticmethod
    def _get_org_unit_name(unit_id: int) -> str:
        """Helper to get organizational unit name by ID."""
        try:
            unit = get_admin_unit_by_id(unit_id)
            if unit and hasattr(unit, 'Name'):
                return unit.Name
            return "—"
        except:
            return "—"
    
    @staticmethod
    def generate_monthly_numeric_word_report(
        report_data: Dict[str, Any],
        filename: str,
        language: str = "ar",
        report_entity_name: str = None,
        report_entity_type: str = None
    ) -> bytes:
        """
        Generate professional Word document for monthly numeric (aggregated) reports.
        
        Creates a comprehensive Arabic report with:
        - Summary statistics table
        - Domain breakdown with counts and percentages
        - Severity breakdown (Low/Medium/High)
        - Department breakdown (top departments)
        
        Args:
            report_data: Numeric report dict with period, summary, by_domain, by_severity, by_department
            filename: Target filename
            language: Language code (defaults to ar for Arabic)
            report_entity_name: Name of the organizational unit (e.g., "الادارة الطبية")
            report_entity_type: Type of entity (hospital/administration/department/section)
        
        Returns:
            bytes: Valid Word .docx file content
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word export. "
                "Install with: pip install python-docx"
            )
        
        def _safe(v):
            """Convert dimension values to int (python-docx requirement)"""
            return int(v)
        
        def center_cell_content(cell):
            """Center-align cell content"""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        def apply_cell_borders(cell):
            """Apply borders to cell"""
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()

            # Add borders
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

        def add_section_divider(text_ar, text_en):
            """
            Major section heading (Session 4): Executive Summary / Complaint
            Statistics / Notice Statistics / Combined Totals. More prominent
            than the existing sub-headings (e.g. "التوزيع حسب المجال") so the
            four-section structure reads clearly without redesigning the
            report's existing look.
            """
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text_ar)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Traditional Arabic'
            run.font.color.rgb = RGBColor(0x1C, 0x3A, 0x7A)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(1)

            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sub_run = sub.add_run(text_en)
            sub_run.font.size = Pt(11)
            sub_run.font.italic = True
            sub.paragraph_format.space_after = Pt(6)

            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            bottom.set(qn('w:space'), '2')
            bottom.set(qn('w:color'), '1C3A7A')
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_label_value_table(rows_data):
            """2-column (value | label) table, same visual style as the
            existing Summary Statistics table - used by Executive Summary."""
            table = doc.add_table(rows=len(rows_data), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = _safe(Cm(4))
            table.columns[1].width = _safe(Cm(10))

            for idx, (label, value) in enumerate(rows_data):
                row_cells = table.rows[idx].cells
                row_cells[0].text = str(value)
                for paragraph in row_cells[0].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

                row_cells[1].text = label
                for paragraph in row_cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

                apply_cell_borders(row_cells[0])
                apply_cell_borders(row_cells[1])
            return table

        def add_unit_count_table(units, count_field):
            """2-column (count | unit name) table - same style as the
            existing Severity breakdown table. Used by Notice Statistics'
            per-level distribution tables."""
            table = doc.add_table(rows=len(units) + 1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = _safe(Cm(3))
            table.columns[1].width = _safe(Cm(11))

            header_cells = table.rows[0].cells
            for idx, header_text in enumerate(["العدد", "الوحدة"]):
                header_cells[idx].text = header_text
                for paragraph in header_cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                apply_cell_borders(header_cells[idx])
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D3D3D3')
                header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)

            for idx, unit in enumerate(units, start=1):
                row_cells = table.rows[idx].cells
                row_cells[0].text = str(unit.get(count_field, 0))
                center_cell_content(row_cells[0])
                row_cells[1].text = unit.get("unit_name", "—")
                for paragraph in row_cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                for cell in row_cells:
                    apply_cell_borders(cell)
            return table

        def add_no_data_note(text_ar):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text_ar)
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.name = 'Traditional Arabic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

        # Extract data
        period = report_data.get("period", {})
        summary = report_data.get("summary", {})
        by_domain = report_data.get("by_domain", [])
        by_severity = report_data.get("by_severity", [])
        by_department = report_data.get("by_department", [])
        executive_summary = report_data.get("executive_summary")
        notice_summary = report_data.get("notice_summary")
        intent_counts = report_data.get("intent_counts")
        
        # Build period label
        year = period.get("year", "")
        month = period.get("month")
        start_date = period.get("start_date", "")
        end_date = period.get("end_date", "")
        
        if month:
            # Month-based label
            period_label = period.get("label", f"Month {month}, {year}")
            period_label_ar = period.get("label_ar", f"الشهر {month}، {year}")
        else:
            # Date range label
            period_label = f"{start_date} to {end_date}"
            period_label_ar = f"{start_date} إلى {end_date}"
        
        # Build scope label
        if report_entity_type == "hospital" or not report_entity_type:
            scope_label = "التقرير الشهري الإحصائي للمستشفى"
            scope_name = "مستشفى الرّسول الأعظم"
        elif report_entity_type == "administration":
            scope_label = "التقرير الشهري الإحصائي للإدارة"
            scope_name = report_entity_name or "—"
        elif report_entity_type == "department":
            scope_label = "التقرير الشهري الإحصائي للقسم"
            scope_name = report_entity_name or "—"
        elif report_entity_type == "section":
            scope_label = "التقرير الشهري الإحصائي للشعبة"
            scope_name = report_entity_name or "—"
        else:
            scope_label = "التقرير الشهري الإحصائي"
            scope_name = report_entity_name or "—"
        
        # Create Word document
        doc = Document()

        # Set page to A4 Landscape
        section = doc.sections[0]
        section.page_height = _safe(Mm(210))
        section.page_width = _safe(Mm(297))
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = _safe(Mm(15))
        section.right_margin = _safe(Mm(15))
        section.top_margin = _safe(Mm(15))
        section.bottom_margin = _safe(Mm(15))

        # Hospital logo in Word page header (same pattern as generate_docx_export
        # and generate_workflow_activity_word — all official hospital reports)
        import os
        try:
            logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
            if os.path.exists(logo_path):
                section.header_distance = _safe(Mm(5))
                hdr_p = section.header.paragraphs[0]
                hdr_p.clear()
                hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                hdr_p.add_run().add_picture(logo_path, width=_safe(Inches(0.9)))
        except Exception:
            pass  # logo is cosmetic — report must never fail because of it

        # ============================================================
        # HEADER SECTION
        # ============================================================

        # Main title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(scope_label)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.name = 'Traditional Arabic'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Entity name
        if scope_name != "—":
            entity_para = doc.add_paragraph()
            entity_run = entity_para.add_run(scope_name)
            entity_run.font.size = Pt(14)
            entity_run.font.bold = True
            entity_run.font.name = 'Traditional Arabic'
            entity_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            entity_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Period
        period_para = doc.add_paragraph()
        period_run = period_para.add_run(f"الفترة: {period_label_ar}")
        period_run.font.size = Pt(12)
        period_run.font.name = 'Traditional Arabic'
        period_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer

        # ============================================================
        # SECTION 1 — EXECUTIVE SUMMARY (Session 4)
        # ============================================================

        if executive_summary:
            add_section_divider("الملخص التنفيذي", "Section 1 — Executive Summary")
            exec_rows = [
                ("إجمالي الشكاوى (Total Complaints)", executive_summary.get("total_complaints", 0)),
                ("إجمالي الملاحظات (Total Notices)", executive_summary.get("total_notices", 0)),
                ("إجمالي السجلات (Total Records)", executive_summary.get("total_records", 0)),
                ("عدد الأقسام المعنية (Sections Involved)", executive_summary.get("sections_involved", 0)),
                ("عدد الدوائر المعنية (Departments Involved)", executive_summary.get("departments_involved", 0)),
                ("عدد الإدارات المعنية (Administrations Involved)", executive_summary.get("administrations_involved", 0)),
            ]
            add_label_value_table(exec_rows)
            doc.add_paragraph()

        # ============================================================
        # SECTION 2 — COMPLAINT STATISTICS (existing tables, unchanged)
        # ============================================================

        add_section_divider("إحصائيات الشكاوى", "Section 2 — Complaint Statistics")

        # ============================================================
        # SUMMARY STATISTICS TABLE (Complaint-only — retained as-is)
        # ============================================================

        summary_heading = doc.add_paragraph()
        summary_heading_run = summary_heading.add_run("الإحصائيات العامة للشكاوى")
        summary_heading_run.font.size = Pt(14)
        summary_heading_run.font.bold = True
        summary_heading_run.font.name = 'Traditional Arabic'
        summary_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Create summary table (2 columns: label, value)
        summary_table = doc.add_table(rows=7, cols=2)
        summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        summary_table.autofit = False
        summary_table.allow_autofit = False
        
        # Set column widths (RTL: value on right, label on left)
        summary_table.columns[0].width = _safe(Cm(4))  # Value column
        summary_table.columns[1].width = _safe(Cm(10))  # Label column
        
        # Define summary rows (RTL order)
        summary_rows = [
            ("إجمالي الشكاوى", summary.get("total_complaints", 0)),
            ("الشكاوى المفتوحة", summary.get("open_complaints", 0)),
            ("الشكاوى المغلقة", summary.get("closed_complaints", 0)),
            ("الحالات الحرجة (Red Flags)", summary.get("red_flags_count", 0)),
            ("الأحداث التي لا يجب أن تحدث (Never Events)", summary.get("never_events_count", 0)),
            ("متوسط أيام الإغلاق", f"{summary.get('avg_closure_days', 0):.1f}"),
            ("الوسيط لأيام الإغلاق", f"{summary.get('median_closure_days', 0):.1f}")
        ]
        
        # Populate summary table
        for idx, (label, value) in enumerate(summary_rows):
            row_cells = summary_table.rows[idx].cells
            
            # Value cell (right side in RTL)
            row_cells[0].text = str(value)
            for paragraph in row_cells[0].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Traditional Arabic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            # Label cell (left side in RTL)
            row_cells[1].text = label
            for paragraph in row_cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.name = 'Traditional Arabic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            # Apply borders
            apply_cell_borders(row_cells[0])
            apply_cell_borders(row_cells[1])
        
        doc.add_paragraph()  # Spacer
        
        # ============================================================
        # DOMAIN BREAKDOWN TABLE
        # ============================================================
        
        domain_heading = doc.add_paragraph()
        domain_heading_run = domain_heading.add_run("التوزيع حسب المجال (Domain)")
        domain_heading_run.font.size = Pt(14)
        domain_heading_run.font.bold = True
        domain_heading_run.font.name = 'Traditional Arabic'
        domain_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        domain_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Create domain table (3 columns: percentage, count, domain name)
        domain_table = doc.add_table(rows=len(by_domain) + 1, cols=3)
        domain_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        domain_table.autofit = False
        domain_table.allow_autofit = False
        
        # Set column widths (RTL order: percentage, count, name)
        domain_table.columns[0].width = _safe(Cm(3))   # Percentage
        domain_table.columns[1].width = _safe(Cm(3))   # Count
        domain_table.columns[2].width = _safe(Cm(8))   # Domain name
        
        # Header row
        header_cells = domain_table.rows[0].cells
        headers = ["النسبة المئوية", "العدد", "المجال"]
        
        for idx, header_text in enumerate(headers):
            header_cells[idx].text = header_text
            for paragraph in header_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.name = 'Traditional Arabic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            apply_cell_borders(header_cells[idx])
            
            # Shade header
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D3D3D3')
            header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
        
        # Data rows
        for idx, domain in enumerate(by_domain, start=1):
            row_cells = domain_table.rows[idx].cells
            
            # Percentage
            percentage = domain.get("percentage", 0)
            row_cells[0].text = f"{percentage:.1f}%"
            center_cell_content(row_cells[0])
            
            # Count
            count = domain.get("count", 0)
            row_cells[1].text = str(count)
            center_cell_content(row_cells[1])
            
            # Domain name (use Arabic if available)
            domain_name = domain.get("domain_name_ar", domain.get("domain_name", "—"))
            row_cells[2].text = domain_name
            for paragraph in row_cells[2].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Traditional Arabic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            # Apply borders
            for cell in row_cells:
                apply_cell_borders(cell)
        
        doc.add_paragraph()  # Spacer
        
        # ============================================================
        # SEVERITY BREAKDOWN TABLE
        # ============================================================
        
        if by_severity:
            severity_heading = doc.add_paragraph()
            severity_heading_run = severity_heading.add_run("التوزيع حسب درجة الخطورة")
            severity_heading_run.font.size = Pt(14)
            severity_heading_run.font.bold = True
            severity_heading_run.font.name = 'Traditional Arabic'
            severity_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            severity_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Create severity table (2 columns: count, severity name)
            severity_table = doc.add_table(rows=len(by_severity) + 1, cols=2)
            severity_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            severity_table.autofit = False
            severity_table.allow_autofit = False
            
            # Set column widths
            severity_table.columns[0].width = _safe(Cm(3))   # Count
            severity_table.columns[1].width = _safe(Cm(11))  # Severity name
            
            # Header row
            header_cells = severity_table.rows[0].cells
            headers = ["العدد", "درجة الخطورة"]
            
            for idx, header_text in enumerate(headers):
                header_cells[idx].text = header_text
                for paragraph in header_cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                apply_cell_borders(header_cells[idx])
                
                # Shade header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D3D3D3')
                header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
            
            # Data rows
            for idx, severity in enumerate(by_severity, start=1):
                row_cells = severity_table.rows[idx].cells
                
                # Count
                count = severity.get("count", 0)
                row_cells[0].text = str(count)
                center_cell_content(row_cells[0])
                
                # Severity name
                severity_name = severity.get("severity_name_ar", severity.get("severity_name", "—"))
                row_cells[1].text = severity_name
                for paragraph in row_cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                
                # Apply borders
                for cell in row_cells:
                    apply_cell_borders(cell)
            
            doc.add_paragraph()  # Spacer
        
        # ============================================================
        # DEPARTMENT BREAKDOWN TABLE (TOP DEPARTMENTS)
        # ============================================================
        
        if by_department:
            dept_heading = doc.add_paragraph()
            dept_heading_run = dept_heading.add_run("التوزيع حسب الأقسام")
            dept_heading_run.font.size = Pt(14)
            dept_heading_run.font.bold = True
            dept_heading_run.font.name = 'Traditional Arabic'
            dept_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            dept_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Create department table (2 columns: count, department name)
            # Show top 15 departments or all if less
            top_departments = by_department[:15]
            dept_table = doc.add_table(rows=len(top_departments) + 1, cols=2)
            dept_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            dept_table.autofit = False
            dept_table.allow_autofit = False
            
            # Set column widths
            dept_table.columns[0].width = _safe(Cm(3))   # Count
            dept_table.columns[1].width = _safe(Cm(11))  # Department name
            
            # Header row
            header_cells = dept_table.rows[0].cells
            headers = ["العدد", "القسم"]
            
            for idx, header_text in enumerate(headers):
                header_cells[idx].text = header_text
                for paragraph in header_cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                apply_cell_borders(header_cells[idx])
                
                # Shade header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D3D3D3')
                header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
            
            # Data rows
            for idx, dept in enumerate(top_departments, start=1):
                row_cells = dept_table.rows[idx].cells
                
                # Count
                count = dept.get("count", 0)
                row_cells[0].text = str(count)
                center_cell_content(row_cells[0])
                
                # Department name
                dept_name = dept.get("dayra_name_ar", dept.get("dayra_name", "—"))
                row_cells[1].text = dept_name
                for paragraph in row_cells[1].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                
                # Apply borders
                for cell in row_cells:
                    apply_cell_borders(cell)
            
            # Add note if there are more departments
            if len(by_department) > 15:
                note_para = doc.add_paragraph()
                note_run = note_para.add_run(f"ملاحظة: تم عرض أعلى 15 قسماً من أصل {len(by_department)} قسماً")
                note_run.font.size = Pt(10)
                note_run.font.italic = True
                note_run.font.name = 'Traditional Arabic'
                note_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                note_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # ============================================================
        # ORGANIZATIONAL BREAKDOWN (FOR "ALL" SCOPE)
        # ============================================================
        
        organizational_breakdown = report_data.get("organizational_breakdown", [])
        if organizational_breakdown:
            doc.add_paragraph()  # Spacer
            
            # Determine breakdown label based on entity type
            if report_entity_type == "all_administrations":
                breakdown_label = "مقارنة الإحصائيات بين الإدارات"
            elif report_entity_type == "all_departments":
                breakdown_label = "مقارنة الإحصائيات بين الأقسام"
            elif report_entity_type == "all_sections":
                breakdown_label = "مقارنة الإحصائيات بين الشعب"
            else:
                breakdown_label = "مقارنة الإحصائيات بين الوحدات التنظيمية"
            
            breakdown_heading = doc.add_paragraph()
            breakdown_heading_run = breakdown_heading.add_run(breakdown_label)
            breakdown_heading_run.font.size = Pt(14)
            breakdown_heading_run.font.bold = True
            breakdown_heading_run.font.name = 'Traditional Arabic'
            breakdown_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            breakdown_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Create comparison table
            breakdown_table = doc.add_table(rows=len(organizational_breakdown) + 1, cols=6)
            breakdown_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            breakdown_table.autofit = False
            breakdown_table.allow_autofit = False
            
            # Set column widths (RTL order)
            breakdown_table.columns[0].width = _safe(Cm(2.5))  # Average closure days
            breakdown_table.columns[1].width = _safe(Cm(2.5))  # Never events
            breakdown_table.columns[2].width = _safe(Cm(2.5))  # Red flags
            breakdown_table.columns[3].width = _safe(Cm(2.5))  # Closed
            breakdown_table.columns[4].width = _safe(Cm(2.5))  # Open
            breakdown_table.columns[5].width = _safe(Cm(7))    # Unit name
            
            # Header row
            header_cells = breakdown_table.rows[0].cells
            headers = ["متوسط أيام الإغلاق", "Never Events", "Red Flags", "مغلقة", "مفتوحة", "الوحدة التنظيمية"]
            
            for idx, header_text in enumerate(headers):
                header_cells[idx].text = header_text
                for paragraph in header_cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                apply_cell_borders(header_cells[idx])
                
                # Shade header
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D3D3D3')
                header_cells[idx]._element.get_or_add_tcPr().append(shading_elm)
            
            # Data rows
            for idx, unit in enumerate(organizational_breakdown, start=1):
                row_cells = breakdown_table.rows[idx].cells
                
                # Average closure days
                avg_days = unit.get("avg_closure_days", 0)
                row_cells[0].text = f"{avg_days:.1f}"
                center_cell_content(row_cells[0])
                
                # Never events
                never_events = unit.get("never_events_count", 0)
                row_cells[1].text = str(never_events)
                center_cell_content(row_cells[1])
                
                # Red flags
                red_flags = unit.get("red_flags_count", 0)
                row_cells[2].text = str(red_flags)
                center_cell_content(row_cells[2])
                
                # Closed complaints
                closed = unit.get("closed_complaints", 0)
                row_cells[3].text = str(closed)
                center_cell_content(row_cells[3])
                
                # Open complaints
                open_complaints = unit.get("open_complaints", 0)
                row_cells[4].text = str(open_complaints)
                center_cell_content(row_cells[4])
                
                # Unit name
                unit_name = unit.get("unit_name", "—")
                row_cells[5].text = unit_name
                for paragraph in row_cells[5].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                
                # Apply borders
                for cell in row_cells:
                    apply_cell_borders(cell)
            
            doc.add_paragraph()  # Spacer
            
            # Add note about total units
            breakdown_note = doc.add_paragraph()
            breakdown_note_run = breakdown_note.add_run(f"عدد الوحدات المعروضة: {len(organizational_breakdown)} (تم عرض الوحدات التي بها شكاوى فقط)")
            breakdown_note_run.font.size = Pt(10)
            breakdown_note_run.font.italic = True
            breakdown_note_run.font.name = 'Traditional Arabic'
            breakdown_note_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            breakdown_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # ============================================================
        # SECTION 3 — NOTICE STATISTICS (Session 4)
        # ============================================================

        if intent_counts is not None:
            add_section_divider("إحصائيات الملاحظات", "Section 3 — Notice Statistics")

            total_notices_val = (notice_summary or {}).get("total_notices", 0)
            add_label_value_table([("إجمالي الملاحظات (Total Notices)", total_notices_val)])
            doc.add_paragraph()

            for level_key, level_label_ar, level_label_en in [
                ("sections", "توزيع الملاحظات حسب القسم", "Notices by Section"),
                ("departments", "توزيع الملاحظات حسب الدائرة", "Notices by Department"),
                ("administrations", "توزيع الملاحظات حسب الإدارة", "Notices by Administration"),
            ]:
                units_with_notices = [u for u in intent_counts.get(level_key, []) if u.get("notice_count", 0) > 0]

                level_heading = doc.add_paragraph()
                level_run = level_heading.add_run(level_label_ar)
                level_run.font.size = Pt(13)
                level_run.font.bold = True
                level_run.font.name = 'Traditional Arabic'
                level_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                level_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                if units_with_notices:
                    add_unit_count_table(units_with_notices, "notice_count")
                else:
                    add_no_data_note("لا توجد ملاحظات لهذه الوحدات في الفترة المحددة — No notices for this level in the selected period")
                doc.add_paragraph()

        # ============================================================
        # SECTION 4 — COMBINED TOTALS (Session 4)
        # ============================================================

        if intent_counts is not None:
            add_section_divider("الإجماليات المدمجة", "Section 4 — Combined Totals")

            # Per-unit table (Section / Department / Administration rows with
            # Complaint + Notice + Total columns) — same visual style as the
            # Session 3 Detailed Report count summary for consistency.
            all_units = []
            type_labels = {
                "sections": "قسم (Section)",
                "departments": "دائرة (Department)",
                "administrations": "إدارة (Administration)",
            }
            for level_key, type_label in type_labels.items():
                for unit in intent_counts.get(level_key, []):
                    all_units.append((unit, type_label))

            if all_units:
                combined_table = doc.add_table(rows=1, cols=5)
                combined_table.style = 'Table Grid'
                combined_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                combined_table.autofit = False
                combined_table.allow_autofit = False

                col_widths = [_safe(Cm(9)), _safe(Cm(4)), _safe(Cm(3)), _safe(Cm(3)), _safe(Cm(3))]
                for idx, w in enumerate(col_widths):
                    combined_table.columns[idx].width = w

                header_cells = combined_table.rows[0].cells
                combined_headers = [
                    "اسم الوحدة / Unit Name",
                    "نوع الوحدة / Unit Type",
                    "عدد الشكاوى / Complaints",
                    "عدد الملاحظات / Notices",
                    "المجموع / Total",
                ]
                for idx, (cell, hdr) in enumerate(zip(header_cells, combined_headers)):
                    cell.text = hdr
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.bold = True
                            run.font.name = 'Traditional Arabic'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                    apply_cell_borders(cell)
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D3D3D3')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                for idx, w in enumerate(col_widths):
                    header_cells[idx].width = w

                for unit, type_label in all_units:
                    row = combined_table.add_row()
                    values = [
                        unit.get("unit_name", "—"),
                        type_label,
                        str(unit.get("complaint_count", 0)),
                        str(unit.get("notice_count", 0)),
                        str(unit.get("total_count", 0)),
                    ]
                    for idx, (cell, val) in enumerate(zip(row.cells, values)):
                        cell.text = val
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.RIGHT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.name = 'Traditional Arabic'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                        apply_cell_borders(cell)
                        cell.width = col_widths[idx]

                doc.add_paragraph()

            # Hospital / scope total row
            ex = executive_summary or {}
            total_row_data = [
                ("إجمالي الشكاوى (Total Complaints)", ex.get("total_complaints", 0)),
                ("إجمالي الملاحظات (Total Notices)", ex.get("total_notices", 0)),
                ("الإجمالي الكلي (Grand Total)", ex.get("total_records", 0)),
            ]
            hosp_label = doc.add_paragraph()
            hosp_run = hosp_label.add_run("إجمالي المستشفى / Hospital Total")
            hosp_run.font.size = Pt(13)
            hosp_run.font.bold = True
            hosp_run.font.name = 'Traditional Arabic'
            hosp_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            hosp_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_label_value_table(total_row_data)
            doc.add_paragraph()

        # ============================================================
        # FOOTER
        # ============================================================

        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(f"تم إنشاء التقرير في: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.name = 'Traditional Arabic'
        footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.read()
    
    @staticmethod
    def generate_csv_export(
        report_data: List[Dict[str, Any]],
        filename: str,
        language: str = "en"
    ) -> bytes:
        """Generate CSV from report data."""
        
        csv_buffer = StringIO()
        
        if not report_data:
            return b""
        
        # Get fieldnames from first record
        fieldnames = list(report_data[0].keys())
        
        writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
        writer.writeheader()
        
        for row in report_data:
            writer.writerow(row)
        
        return csv_buffer.getvalue().encode('utf-8')
    
    @staticmethod
    def generate_pdf_export(
        report_data: Dict[str, Any],
        filename: str,
        language: str = "en",
        include_charts: bool = True,
        report_entity_name: str = None,
        report_entity_type: str = None,
        report_administration: str = None,
        report_department: str = None,
        report_section: str = None
    ) -> bytes:
        """
        Generate professional PDF matching Word export exactly.
        A4 landscape with RTL layout, vertical headers, logo, footer.
        
        Args:
            report_data: List of dictionaries OR dict with "complaints" key
            filename: Target filename
            language: Language code (en or ar)
            include_charts: Not used (kept for compatibility)
            report_entity_name: Name of the entity being reported
            report_entity_type: Type of entity (administration/department/section)
            report_administration: Administration name for header
            report_department: Department name for header
            report_section: Section name for header
        
        Returns:
            bytes: Valid PDF file content
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "reportlab is required for PDF export. "
                "Install with: pip install reportlab"
            )
        
        # Import additional reportlab components
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_RIGHT, TA_CENTER
        from reportlab.lib.pagesizes import landscape, A4 as portrait_a4
        from reportlab.pdfgen import canvas
        from reportlab.platypus import Flowable
        import os
        
        # Try to import Arabic text reshaping libraries (optional)
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display
            ARABIC_RESHAPER_AVAILABLE = True
        except ImportError:
            ARABIC_RESHAPER_AVAILABLE = False
            print("[PDF] WARNING: arabic-reshaper and python-bidi not installed. Arabic text may not render correctly.")
            print("[PDF] Install with: pip install arabic-reshaper python-bidi")
        
        def sanitize_value(value):
            """Convert value to string, handling dates and None"""
            try:
                if value is None:
                    return ""
                if isinstance(value, (datetime, date)):
                    if isinstance(value, datetime):
                        return value.strftime("%Y-%m-%d")
                    return value.isoformat()
                return str(value)
            except:
                return ""
        
        def normalize_text(text: str) -> str:
            """Normalize text - remove manual line breaks"""
            text = str(text)
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            lines = [l.strip() for l in text.split("\n") if l.strip()]
            return " ".join(lines)
        
        def ar(text: str) -> str:
            """Apply Arabic text reshaping for proper rendering in PDFs.
            ONLY use when creating Paragraph text - never in rotated/transformed contexts.
            """
            if not text:
                return ""
            if not ARABIC_RESHAPER_AVAILABLE:
                return str(text)  # Return as-is if libraries not available
            try:
                reshaped = arabic_reshaper.reshape(str(text))
                return get_display(reshaped)
            except:
                return str(text)
        
        # Normalize data source
        try:
            if isinstance(report_data, dict) and "complaints" in report_data:
                rows = report_data["complaints"]
            elif isinstance(report_data, list):
                rows = report_data
            else:
                rows = []
            
            if not isinstance(rows, list):
                rows = []
        except:
            rows = []
        
        # Extract metadata for header
        start_date = "—"
        end_date = "—"
        
        # Use provided parameters for header info
        print(f"[PDF EXPORT] Parameters: admin={report_administration}, dept={report_department}, section={report_section}")
        
        if report_administration or report_department or report_section:
            Administration = report_administration or "—"
            Department = report_department or "—"
            Section = report_section or "—"
        else:
            # Fallback to first row
            Administration = "—"
            Department = "—"
            Section = "—"
            if rows:
                try:
                    first_record = rows[0]
                    Administration = first_record.get("administration_name", "—")
                    Department = first_record.get("department_name", "—")
                    Section = first_record.get("section_name", "—")
                except:
                    pass

        # Extract date range
        if rows:
            try:
                first_record = rows[0]
                last_record = rows[-1] if len(rows) > 1 else first_record
                start_date = sanitize_value(first_record.get("received_date", "—"))
                end_date = sanitize_value(last_record.get("received_date", "—"))
            except:
                pass
        
        # Register Arabic font with strict priority order
        try:
            # PRIORITY ORDER (strict):
            # 1. Amiri-Regular.ttf (best Arabic support)
            # 2. NotoNaskhArabic-Regular.ttf (good Arabic support)
            # 3. Tahoma (fallback with partial Arabic support)
            # 4. Arial (last resort)
            arabic_font_paths = [
                # Amiri fonts (HIGHEST PRIORITY)
                ("C:\\Windows\\Fonts\\Amiri-Regular.ttf", "Amiri"),
                ("/usr/share/fonts/truetype/amiri/Amiri-Regular.ttf", "Amiri"),
                ("assets/fonts/Amiri-Regular.ttf", "Amiri"),
                ("Amiri-Regular.ttf", "Amiri"),
                
                # NotoNaskh fonts (SECOND PRIORITY)
                ("C:\\Windows\\Fonts\\NotoNaskhArabic-Regular.ttf", "NotoNaskh"),
                ("/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf", "NotoNaskh"),
                
                # Tahoma (THIRD PRIORITY - Windows default with Arabic)
                ("C:\\Windows\\Fonts\\tahoma.ttf", "Tahoma"),
                
                # Arial (LAST RESORT - minimal Arabic support)
                ("C:\\Windows\\Fonts\\arial.ttf", "Arial"),
            ]
            
            font_registered = False
            font_name = None
            
            for font_path, font_label in arabic_font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ArabicFont', font_path))
                        font_registered = True
                        font_name = 'ArabicFont'
                        print(f"[PDF] ✅ SUCCESS: Registered Arabic font: {font_label} ({font_path})")
                        break
                    except Exception as e:
                        print(f"[PDF] ❌ Failed to register {font_label} ({font_path}): {e}")
                        continue
            
            if not font_registered:
                print("[PDF] ⚠️ WARNING: No Arabic font found! Arabic text will NOT render correctly.")
                print("[PDF] ⚠️ Please install Amiri or NotoNaskhArabic fonts for proper Arabic support.")
                # Use Helvetica as absolute last resort
                font_name = 'Helvetica'
            
        except Exception as e:
            print(f"[PDF] ❌ CRITICAL: Font registration error: {e}")
            font_name = 'Helvetica'
        
        # Create PDF buffer
        pdf_buffer = BytesIO()
        
        # A4 Landscape dimensions
        page_width, page_height = landscape(portrait_a4)
        
        # Custom canvas with header/footer
        class HeaderFooterCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                canvas.Canvas.__init__(self, *args, **kwargs)
                self.pages = []
                
            def showPage(self):
                self.pages.append(dict(self.__dict__))
                self._startPage()
                
            def save(self):
                for page in self.pages:
                    self.__dict__.update(page)
                    self.draw_header_footer()
                    canvas.Canvas.showPage(self)
                canvas.Canvas.save(self)
                
            def draw_header_footer(self):
                self.saveState()
                
                # Logo in header (top right)
                try:
                    logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
                    if os.path.exists(logo_path):
                        x = page_width - 100
                        y = page_height - 90
                        self.drawImage(logo_path, x, y, width=70, height=70, preserveAspectRatio=True)
                except Exception as e:
                    print(f"[PDF] Logo error: {e}")
                
                # Footer with quote
                try:
                    self.setFont(font_name, 9)
                    footer_text = "نؤمن أن الإبتكار لا يكون فقط في التقنيات، بل في أسلوب الخدمة والتواصل والتعاطف… فلنبتكر معًا تجربة ذات أثر طيب"
                    footer_text_shaped = ar(footer_text)
                    
                    # Border line
                    self.setStrokeColorRGB(0.8, 0.8, 0.8)
                    self.setLineWidth(0.5)
                    self.line(30, 40, page_width - 30, 40)
                    
                    # Centered text
                    text_width = self.stringWidth(footer_text_shaped, font_name, 9)
                    x = (page_width - text_width) / 2
                    self.drawString(x, 25, footer_text_shaped)
                except Exception as e:
                    print(f"[PDF] Footer error: {e}")
                
                self.restoreState()
        
        # Create document
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=landscape(portrait_a4),
            rightMargin=30,
            leftMargin=30,
            topMargin=90,
            bottomMargin=50
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            alignment=TA_CENTER,
            fontName=font_name,
            spaceAfter=6
        )
        
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_CENTER,
            fontName=font_name,
            spaceAfter=6
        )
        
        period_style = ParagraphStyle(
            'Period',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            fontName=font_name,
            spaceAfter=6
        )
        
        dept_style = ParagraphStyle(
            'Department',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            fontName=font_name
        )
        
        # Build elements
        elements = []
        
        # Title
        elements.append(Paragraph(
            ar("نموذج التقرير الشهري لفرص التحسين والإجراءات التصحيحية الواردة من المرضى وذويهم"),
            title_style
        ))
        
        # Subtitle
        elements.append(Paragraph(
            ar("(إصدار رسمي — للاستخدام الإداري والجودة)"),
            subtitle_style
        ))
        
        # Period
        elements.append(Paragraph(
            ar(f"الشهر المعني: من {start_date} إلى {end_date}"),
            period_style
        ))
        
        # Department info table
        dept_data = [[
            Paragraph(ar(f"الإدارة: {Administration}"), dept_style),
            Paragraph(ar(f"الدائرة: {Department}"), dept_style),
            Paragraph(ar(f"القسم المعني: {Section}"), dept_style)
        ]]
        dept_table = Table(dept_data, colWidths=[250, 250, 250])
        dept_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0, colors.white),
        ]))
        elements.append(dept_table)
        elements.append(Spacer(1, 6))
        
        # Handle empty data
        if not rows:
            elements.append(Paragraph(ar("No data available"), period_style))
            doc.build(elements, canvasmaker=HeaderFooterCanvas)
            pdf_buffer.seek(0)
            return pdf_buffer.getvalue()
        
        # Calculate usable width
        usable_width = page_width - 60
        
        # Define columns (19 columns - ALL HORIZONTAL, no vertical text)
        columns = [
            ("Date", "received_date", False, 0.6),
            ("ID", "id", False, 0.4),
            ("Name", "patient_name", False, 0.8),
            ("Section", "section_name", False, 0.8),
            ("Admin", "administration_name", False, 0.6),
            ("Dept", "department_name", False, 0.8),
            ("Source", "source_name", False, 0.6),
            ("Type", "feedback_intent_type_name", False, 0.6),
            ("Domain", "domain_name", False, 0.8),
            ("Category", "category_name", False, 0.8),
            ("Sub-Cat", "subcategory_name", False, 0.8),
            ("Target Depts", "target_departments_display", False, 1.2),
            ("Classification", "classification_name_en", False, 1.2),
            ("Complaint", "complaint_text", False, 3.0),
            ("Immediate Action", "immediate_action", False, 2.5),
            ("Actions Taken", "taken_action", False, 2.0),
            ("Severity", "severity_name", False, 0.6),
            ("Stage", "stage_name", False, 0.6),
            ("Harm", "harm_level", False, 0.5)
        ]
        
        # Reverse for RTL
        columns = list(reversed(columns))
        
        # Calculate column widths
        total_ratio = sum(col[3] for col in columns)
        col_widths = [(col[3] / total_ratio) * usable_width for col in columns]
        
        # Build table data
        table_data = []
        
        # Header row - ALL HORIZONTAL (simple and stable)
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontSize=7,
            alignment=TA_CENTER,
            fontName=font_name,
            wordWrap='LTR'
        )
        
        headers = []
        for col_idx, (header_name, _, _, _) in enumerate(columns):
            shaped_text = ar(header_name)
            headers.append(Paragraph(f"<b>{shaped_text}</b>", header_style))
        table_data.append(headers)
        
        # Data rows (limit to 30)
        row_count = min(len(rows), 30)
        for row_dict in rows[:row_count]:
            row_values = []
            
            # Check for red flag
            is_red_flag = False
            try:
                clinical_risk = row_dict.get("clinical_risk_type_name", "")
                if clinical_risk and clinical_risk != "Ordinary":
                    is_red_flag = True
            except:
                pass
            
            for col_idx, (header_name, field_name, is_vertical, _) in enumerate(columns):
                # Handle target departments
                if field_name == "target_departments_display":
                    target_depts = row_dict.get("target_departments", [])
                    if target_depts and isinstance(target_depts, list):
                        primary_and_matching = []
                        primary_only = []
                        matching_only = []
                        others = []
                        
                        for dept in target_depts:
                            if dept.get("section_name"):
                                display = dept["section_name"]
                            elif dept.get("department_name"):
                                display = dept["department_name"]
                            elif dept.get("administration_name"):
                                display = dept["administration_name"]
                            else:
                                continue
                            
                            is_primary = dept.get("is_primary", False)
                            matches_entity = report_entity_name and report_entity_name in display
                            
                            if is_primary and matches_entity:
                                primary_and_matching.append(display)
                            elif is_primary:
                                primary_only.append(display)
                            elif matches_entity:
                                matching_only.append(display)
                            else:
                                others.append(display)
                        
                        all_displays = primary_and_matching + primary_only + matching_only + others
                        
                        MAX_DISPLAY = 3
                        if len(all_displays) > MAX_DISPLAY:
                            displayed = all_displays[:MAX_DISPLAY]
                            remaining = len(all_displays) - MAX_DISPLAY
                            raw_value = ", ".join(displayed) + f" +{remaining}"
                        else:
                            raw_value = ", ".join(all_displays) if all_displays else "—"
                    else:
                        raw_value = "—"
                else:
                    raw_value = sanitize_value(row_dict.get(field_name, ""))
                
                # Normalize text
                value = normalize_text(raw_value)
                
                # Truncate (all fields treated the same)
                if len(value) > 200:
                    value = value[:200] + "..."
                
                # Create cell content - ALL HORIZONTAL (simple)
                shaped_value = ar(value)
                cell_style = ParagraphStyle(
                    'Cell',
                    parent=styles['Normal'],
                    fontSize=6,
                    alignment=TA_CENTER,
                    fontName=font_name,
                    wordWrap='LTR'
                )
                row_values.append(Paragraph(shaped_value, cell_style))
            
            table_data.append(row_values)
        
        # Create table
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        
        # Base table style - SIMPLE AND STABLE (no height constraints)
        table_style = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.706, 0.906, 0.808)),  # #B4E7CE
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ]
        
        # Alternating row colors
        for i in range(1, len(table_data)):
            if i % 2 == 0:
                table_style.append(('BACKGROUND', (0, i), (-1, i), colors.Color(0.97, 0.97, 0.97)))
            else:
                table_style.append(('BACKGROUND', (0, i), (-1, i), colors.white))
        
        # Semantic coloring
        for row_idx, row_dict in enumerate(rows[:row_count], start=1):
            is_red_flag = False
            try:
                clinical_risk = row_dict.get("clinical_risk_type_name", "")
                if clinical_risk and clinical_risk != "Ordinary":
                    is_red_flag = True
                    table_style.append(('BACKGROUND', (0, row_idx), (-1, row_idx), colors.Color(1.0, 0.898, 0.898)))
            except:
                pass
            
            # Column-specific coloring
            for col_idx, (_, field_name, _, _) in enumerate(columns):
                value = sanitize_value(row_dict.get(field_name, ""))
                value_lower = value.lower()
                cell_color = None
                
                if field_name == "severity_name":
                    if "high" in value_lower:
                        cell_color = colors.Color(1.0, 0.702, 0.729)
                    elif "medium" in value_lower:
                        cell_color = colors.Color(1.0, 0.875, 0.729)
                    elif "low" in value_lower:
                        cell_color = colors.Color(0.729, 1.0, 0.788)
                
                elif field_name == "harm_level":
                    if "death" in value_lower:
                        cell_color = colors.Color(1.0, 0.420, 0.420)
                    elif "severe" in value_lower:
                        cell_color = colors.Color(1.0, 0.647, 0.0)
                    elif "no harm" in value_lower or "none" in value_lower:
                        cell_color = colors.Color(0.729, 1.0, 0.788)
                    elif "minor" in value_lower or "temporary" in value_lower:
                        cell_color = colors.Color(1.0, 1.0, 0.729)
                
                elif field_name == "stage_name":
                    if "admission" in value_lower:
                        cell_color = colors.Color(0.729, 0.882, 1.0)
                    elif "discharge" in value_lower or "transfer" in value_lower:
                        cell_color = colors.Color(0.878, 0.733, 0.894)
                    elif "examination" in value_lower or "diagnosis" in value_lower:
                        cell_color = colors.Color(0.706, 0.973, 0.973)
                
                elif field_name == "domain_name":
                    if "clinical" in value_lower:
                        cell_color = colors.Color(0.729, 0.882, 1.0)
                    elif "management" in value_lower:
                        cell_color = colors.Color(0.878, 0.733, 0.894)
                    elif "relational" in value_lower:
                        cell_color = colors.Color(1.0, 0.875, 0.729)
                
                if cell_color and not is_red_flag:
                    table_style.append(('BACKGROUND', (col_idx, row_idx), (col_idx, row_idx), cell_color))
        
        table.setStyle(TableStyle(table_style))
        elements.append(table)
        
        # Truncation note
        if len(rows) > 30:
            elements.append(Spacer(1, 10))
            elements.append(Paragraph(
                ar(f"عرض أول 30 من {len(rows)} سجل. قم بتنزيل Excel للحصول على البيانات الكاملة."),
                period_style
            ))
        
        # Signature block
        elements.append(Spacer(1, 20))
        
        sig_data = [
            [
                Paragraph(ar("<b>التاريخ:</b>"), dept_style),
                Paragraph(ar("<b>التوقيع:</b>"), dept_style),
                Paragraph(ar("<b>إسم مسؤول العملية</b>"), dept_style),
                Paragraph(ar("<b>خاص خدمات المرضى<br/>الإسم:</b>"), dept_style)
            ],
            [
                Paragraph(ar("<b>التاريخ:</b>"), dept_style),
                Paragraph(ar("<b>التوقيع:</b>"), dept_style),
                Paragraph(ar("<b>إسم رئيس الدائرة</b>"), dept_style),
                Paragraph(ar("<b>تاريخ الإستلام:</b>"), dept_style)
            ],
            [
                Paragraph(ar("<b>التاريخ:</b>"), dept_style),
                Paragraph(ar("<b>التوقيع:</b>"), dept_style),
                Paragraph(ar("<b>إسم مدير الإدارة</b>"), dept_style),
                Paragraph(ar("<b>التوقيع:</b>"), dept_style)
            ]
        ]
        
        sig_table = Table(sig_data, colWidths=[100, 120, 150, 180])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
        ]))
        elements.append(sig_table)
        
        # Build PDF
        try:
            doc.build(elements, canvasmaker=HeaderFooterCanvas)
        except Exception as e:
            print(f"[PDF] Error building report: {e}")
            elements = []
            elements.append(Paragraph(ar("تقرير شهري - Monthly Report"), title_style))
            elements.append(Paragraph(ar(f"خطأ في إنشاء التقرير: {str(e)}"), period_style))
            doc.build(elements, canvasmaker=HeaderFooterCanvas)
        
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue()

    
    @staticmethod
    def generate_docx_export(
        report_data: Dict[str, Any],
        filename: str,
        language: str = "en",
        report_entity_name: str = None,
        report_entity_type: str = None,
        report_administration: str = None,
        report_department: str = None,
        report_section: str = None
    ) -> bytes:
        """
        Generate official hospital audit form Word document.
        A4 Landscape with RTL layout, vertical headers, signature block.
        
        Args:
            report_data: List of dictionaries OR dict with "complaints" key
            filename: Target filename
            language: Language code (en or ar)
            report_entity_name: Name of the entity being reported (for prioritization)
            report_entity_type: Type of entity (administration/department/section)
            report_administration: Administration name for header
            report_department: Department name for header
            report_section: Section name for header
        
        Returns:
            bytes: Valid Word .docx file content
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word export. "
                "Install with: pip install python-docx"
            )
        
        def _safe(v):
            """Convert dimension values to int (python-docx requirement)"""
            return int(v)

        def sanitize_value(value):
            """Convert value to string, handling dates and None"""
            try:
                if value is None:
                    return ""
                if isinstance(value, (datetime, date)):
                    if isinstance(value, datetime):
                        return value.strftime("%Y-%m-%d")
                    return value.isoformat()
                return str(value)
            except:
                return ""

        def normalize_text(text: str) -> str:
            """Normalize text for Word table cells - remove manual line breaks"""
            text = str(text)
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            lines = [l.strip() for l in text.split("\n") if l.strip()]
            return " ".join(lines)

        def is_arabic(text: str) -> bool:
            """Returns True if text contains Arabic characters."""
            import re
            return bool(re.search(r'[؀-ۿ]', str(text) if text else ''))

        def render_intent_count_summary(doc, intent_counts):
            """
            Render the Complaint/Notice Count Summary by Unit (Session 3).
            intent_counts is prepared by get_monthly_intent_counts_by_unit() —
            this only renders already-counted data, no counting logic here.
            """
            if not intent_counts:
                return

            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run("ملخص عدد الشكاوى والملاحظات بحسب الوحدة")
            title_run.font.size = Pt(13)
            title_run.font.bold = True
            title_run.font.name = 'Traditional Arabic'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            title_p.paragraph_format.space_after = int(Pt(1))

            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_p.add_run("Complaint and Notice Count Summary by Unit")
            sub_run.font.size = Pt(10)
            sub_run.font.italic = True
            sub_p.paragraph_format.space_after = int(Pt(6))

            level_labels = [
                ("sections", "قسم (Section)"),
                ("departments", "دائرة (Department)"),
                ("administrations", "إدارة (Administration)"),
            ]

            rows_to_render = []
            for key, type_label in level_labels:
                for unit in intent_counts.get(key, []) or []:
                    rows_to_render.append((unit, type_label))

            if not rows_to_render:
                empty_p = doc.add_paragraph()
                empty_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                empty_run = empty_p.add_run("لا توجد سجلات لهذا الشهر — No records for this month")
                empty_run.font.size = Pt(10)
                empty_run.font.italic = True
                empty_run.font.name = 'Traditional Arabic'
                empty_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                doc.add_paragraph()
                return

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            tbl = table._element
            tblPr = tbl.tblPr
            tblPr.append(OxmlElement('w:bidiVisual'))

            headers = [
                ("اسم الوحدة", "Unit Name"),
                ("نوع الوحدة", "Unit Type"),
                ("عدد الشكاوى", "Complaints"),
                ("عدد الملاحظات", "Notices"),
                ("المجموع", "Total"),
            ]
            header_cells = table.rows[0].cells
            for idx, (ar_label, en_label) in enumerate(headers):
                cell = header_cells[idx]
                cell.text = f"{ar_label}\n{en_label}"
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = int(Pt(9))
                        run.font.name = 'Traditional Arabic'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                try:
                    cell._element.get_or_add_tcPr().get_or_add_shd().fill = "B4E7CE"
                except Exception:
                    pass

            for unit, type_label in rows_to_render:
                row_cells = table.add_row().cells
                values = [
                    unit.get("unit_name", "—"),
                    type_label,
                    str(unit.get("complaint_count", 0)),
                    str(unit.get("notice_count", 0)),
                    str(unit.get("total_count", 0)),
                ]
                for idx, val in enumerate(values):
                    row_cells[idx].text = val
                    for paragraph in row_cells[idx].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = int(Pt(9))
                            run.font.name = 'Traditional Arabic'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

            doc.add_paragraph()

        # Normalize data source
        try:
            if isinstance(report_data, dict) and "complaints" in report_data:
                rows = report_data["complaints"]
            elif isinstance(report_data, list):
                rows = report_data
            else:
                rows = []

            if not isinstance(rows, list):
                rows = []
        except:
            rows = []

        # Complaint/Notice count summary (Session 3) — only present when the
        # caller passed the full monthly-detailed report_data dict (not a bare
        # list), e.g. the single-file monthly export. Multi-export (per-unit
        # ZIP) still passes a bare list and simply won't render this section.
        intent_counts = report_data.get("intent_counts") if isinstance(report_data, dict) else None

        # Load institutional config once (header title, subtitle, footer, report code)
        try:
            from ..db_layer.report_config_db import get_report_config
            _cfg = get_report_config()
        except Exception:
            _cfg = {}
        _header_title    = _cfg.get("header_title",    "نموذج التقرير الشهري لفرص التحسين والإجراءات التصحيحية الواردة من المرضى وذويهم")
        _header_subtitle = _cfg.get("header_subtitle", "(إصدار رسمي — للاستخدام الإداري والجودة)")
        _footer_text     = _cfg.get("footer_text",     "نؤمن أن الإبتكار لا يكون فقط في التقنيات، بل في أسلوب الخدمة والتواصل والتعاطف… فلنبتكر معًا تجربة ذات أثر طيب")
        _report_code     = _cfg.get("report_code",     "")

        # Create Word document
        doc = Document()
        
        # Set page to A4 Landscape
        section = doc.sections[0]
        section.page_height = _safe(Mm(210))  # A4 width becomes height in landscape
        section.page_width = _safe(Mm(297))   # A4 height becomes width in landscape
        section.orientation = WD_ORIENT.LANDSCAPE
        section.left_margin = _safe(Mm(15))
        section.right_margin = _safe(Mm(15))
        section.top_margin = _safe(Mm(15))
        section.bottom_margin = _safe(Mm(15))
        
        # Set document RTL
        try:
            section.start_type = 2  # New page
        except:
            pass
        
        # Extract metadata
        start_date = "—"
        end_date = "—"
        
        # Use provided parameters for header info, fallback to first row if not provided
        print(f"[DOCX EXPORT] Parameters received: admin={report_administration}, dept={report_department}, section={report_section}")
        
        if report_administration or report_department or report_section:
            Administration = report_administration or "—"
            Department = report_department or "—"
            Section = report_section or "—"
            print(f"[DOCX EXPORT] Using parameters: admin={Administration}, dept={Department}, section={Section}")
        elif report_entity_name:
            # Use the filter entity name — avoids showing the wrong issuing-section from row[0]
            Administration = "—"
            Department = "—"
            Section = "—"
            if report_entity_type == "section":
                Section = report_entity_name
            elif report_entity_type == "department":
                Department = report_entity_name
            elif report_entity_type in ("administration", "all_administrations"):
                Administration = report_entity_name
            else:
                Administration = report_entity_name
            print(f"[DOCX EXPORT] Using entity name: {report_entity_name} ({report_entity_type})")
        elif report_entity_type == "hospital":
            # Hospital-level (all-hospital) report — no specific unit filter was applied,
            # so do NOT fall back to row[0]'s issuing section (that would show whichever
            # unit filed the first complaint as if it were the report's scope).
            Administration = "—"
            Department = "—"
            Section = "—"
            print(f"[DOCX EXPORT] Hospital-level report — header left blank (all units)")
        else:
            # Fallback for unknown/unhandled entity types
            Administration = "—"
            Department = "—"
            Section = "—"
            print(f"[DOCX EXPORT] Unknown entity type '{report_entity_type}' — header left blank")

        # Extract date range from rows
        if rows:
            try:
                first_record = rows[0]
                last_record = rows[-1] if len(rows) > 1 else first_record
                start_date = sanitize_value(first_record.get("received_date", "—"))
                end_date = sanitize_value(last_record.get("received_date", "—"))
            except:
                pass
        
        # ========== REAL DOCUMENT HEADER (repeats on every page) ==========
        import os
        try:
            logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')

            # Increase top margin to accommodate compact header
            section.top_margin = _safe(Mm(50))
            section.header_distance = _safe(Mm(5))

            hdr = section.header

            # Logo — right-aligned in first paragraph
            logo_para = hdr.paragraphs[0]
            logo_para.clear()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_para.paragraph_format.space_after = int(Pt(2))
            if os.path.exists(logo_path):
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(0.7))

            # Title
            hdr_title_para = hdr.add_paragraph()
            hdr_title_run = hdr_title_para.add_run(_header_title)
            hdr_title_run.font.size = int(Pt(13))
            hdr_title_run.font.bold = True
            hdr_title_run.font.name = 'Traditional Arabic'
            hdr_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            hdr_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_title_para.paragraph_format.space_after = int(Pt(1))

            # Subtitle
            hdr_sub_para = hdr.add_paragraph()
            hdr_sub_run = hdr_sub_para.add_run(_header_subtitle)
            hdr_sub_run.font.size = int(Pt(10))
            hdr_sub_run.font.name = 'Traditional Arabic'
            hdr_sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            hdr_sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_sub_para.paragraph_format.space_after = int(Pt(1))

            # Period
            hdr_period_para = hdr.add_paragraph()
            hdr_period_run = hdr_period_para.add_run(f"الشهر المعني: من {start_date} إلى {end_date}")
            hdr_period_run.font.size = int(Pt(10))
            hdr_period_run.font.bold = True
            hdr_period_run.font.name = 'Traditional Arabic'
            hdr_period_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            hdr_period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_period_para.paragraph_format.space_after = int(Pt(1))

            # Admin / Dept / Section on one compact line
            hdr_dept_para = hdr.add_paragraph()
            hdr_dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_dept_para.paragraph_format.space_after = int(Pt(2))
            for i, (lbl, val) in enumerate([
                ("الإدارة: ", Administration),
                ("الدائرة: ", Department),
                ("القسم المعني: ", Section),
            ]):
                if i > 0:
                    sep_r = hdr_dept_para.add_run("   |   ")
                    sep_r.font.size = int(Pt(10))
                    sep_r.font.name = 'Traditional Arabic'
                    sep_r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                lbl_r = hdr_dept_para.add_run(lbl)
                lbl_r.font.size = int(Pt(10))
                lbl_r.font.bold = True
                lbl_r.font.name = 'Traditional Arabic'
                lbl_r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                val_r = hdr_dept_para.add_run(str(val))
                val_r.font.size = int(Pt(10))
                val_r.font.name = 'Traditional Arabic'
                val_r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

            # Report code — shown only when configured
            if _report_code:
                hdr_code_para = hdr.add_paragraph()
                hdr_code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_code_para.paragraph_format.space_after = int(Pt(2))
                code_lbl = hdr_code_para.add_run("رمز التقرير: ")
                code_lbl.font.size = int(Pt(10))
                code_lbl.font.bold = True
                code_lbl.font.name = 'Traditional Arabic'
                code_lbl._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                code_val = hdr_code_para.add_run(_report_code)
                code_val.font.size = int(Pt(10))
                code_val.font.name = 'Traditional Arabic'
                code_val._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

            # Blue separator line
            hdr_sep_para = hdr.add_paragraph()
            hdr_sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _pPr = hdr_sep_para._element.get_or_add_pPr()
            _pBdr = OxmlElement('w:pBdr')
            _bot = OxmlElement('w:bottom')
            _bot.set(qn('w:val'), 'single')
            _bot.set(qn('w:sz'), '12')
            _bot.set(qn('w:space'), '1')
            _bot.set(qn('w:color'), '4472C4')
            _pBdr.append(_bot)
            _pPr.append(_pBdr)

        except Exception as _hdr_err:
            print(f"[DOCX] Header setup error: {_hdr_err}")

        # ========== END OF HEADER ==========

        # Complaint/Notice Count Summary by Unit (Session 3) — rendered before
        # the detailed table so it survives even when `rows` is empty below
        # (e.g. a month with only Notices has zero complaint rows, but should
        # still show its Notice counts here).
        try:
            render_intent_count_summary(doc, intent_counts)
        except Exception as _summary_err:
            print(f"[DOCX] Intent count summary error: {_summary_err}")

        # Define columns (23 columns with behavior)
        # Format: (header_label, field_name, is_vertical, width_ratio)
        # is_vertical=True  → narrow column, header rotated 90°
        # is_vertical=False → wide content column, horizontal text
        # Column classes: S=0.360, M=0.750, L-narrow=0.660, L-equal=1.050, XL-large=3.380, XL-mid=2.000, XL-small=1.400
        columns = [
            ("تاريخ الاستلام",        "received_date",                 True,  0.360),  # S
            ("رقم الحادثة",           "incident_id",                   True,  0.360),  # S
            ("رقم الحالة",            "id",                            True,  0.360),  # S
            ("P. Full Name",         "patient_name",                  True,  0.360),  # S
            ("قسم الصّادر",           "section_name",                  True,  0.750),  # M — issuing section
            ("قسم معني",              "target_section_name",           True,  0.750),  # M — target section
            ("دائرة معنيّة",           "target_department_name",        True,  0.750),  # M — target department
            ("إدارة معنيّة",           "target_administration_name",    True,  0.750),  # M — target administration
            ("المصدر",                "source_name",                   True,  0.360),  # S
            ("النوع",                 "feedback_intent_type_name_ar",  True,  0.360),  # S
            ("Domain",               "domain_name",                   True,  0.360),  # S
            ("Category",             "category_name",                 True,  0.360),  # S
            ("Sub-Category",         "subcategory_name",              True,  0.750),  # M
            ("التصنيف (عربي)",        "classification_name",           True,  1.050),  # L-equal, AR first
            ("Classification (EN)",  "classification_name_en",        True,  1.703),  # L-EN-XL (+30%)
            ("محتوى الشكوى",          "complaint_text",                False, 3.380),  # XL-large
            ("Immediate Action",     "immediate_action",              False, 2.000),  # XL-mid
            ("الإجراءات المتخذة",      "taken_action",                  False, 1.400),  # XL-small
            ("Severity",             "severity_name",                 True,  0.360),  # S
            ("Stage",                "stage_name",                    True,  0.750),  # M
            ("Harm",                 "harm_level",                    True,  0.360),  # S
            ("Status",               "status_name",                   True,  0.360),  # S
            ("Field Type",           "clinical_risk_type_name",       True,  0.360),  # S
        ]
        
        # Handle empty data
        if not rows:
            doc.add_paragraph("No data available")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        
        # Create main data table
        table = doc.add_table(rows=1, cols=23)
        table.style = 'Table Grid'
        # Force fixed table layout so Word respects column widths
        tbl = table._element
        tblPr = tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)


        # Set table to RTL (right-to-left) direction
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Force RTL table
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)

        # Center the table on the page
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Also force center alignment at table level
        tblJc = OxmlElement('w:jc')
        tblJc.set(qn('w:val'), 'center')
        tblPr.append(tblJc)
        
        # Header row: taller minimum to accommodate rotated Arabic text
        header_row = table.rows[0]
        header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        header_row.height = Inches(1.4)  # ~35mm — fits full rotated text

        # Header row setup
        header_cells = table.rows[0].cells
        for idx, (header_name, _, is_vertical, _) in enumerate(columns):
            cell = header_cells[idx]
            cell.text = header_name

            tc = cell._element
            tcPr = tc.get_or_add_tcPr()

            # Vertical center alignment for header cell content
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)

            # Style header text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # RTL only on non-vertical columns — RTL+btLr conflicts and shifts centering
                if not is_vertical:
                    paragraph.paragraph_format.right_to_left = True
                paragraph.paragraph_format.space_before = int(Pt(0))
                paragraph.paragraph_format.space_after = int(Pt(0))
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = int(Pt(8))
                    run.font.name = 'Traditional Arabic'
                    run.italic = False
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

            # Light green background
            try:
                cell._element.get_or_add_tcPr().get_or_add_shd().fill = "B4E7CE"
            except:
                pass

            # Rotate all header columns 90° (btLr) for compact layout
            if is_vertical:
                try:
                    textDir = OxmlElement('w:textDirection')
                    textDir.set(qn('w:val'), 'btLr')
                    tcPr.append(textDir)
                except:
                    pass
        
        # Set column widths (force exact page width fitting)
        # Calculate usable width from actual page dimensions
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        
        # Disable autofit so Word respects our exact widths
        table.autofit = False
        
        # Calculate total ratio units and normalize
        total_ratio = sum(col[3] for col in columns)
        
        # Track actual widths to ensure exact sum
        assigned_widths = []
        total_assigned = 0
        
        # Calculate each column width as exact proportion of usable width
        for idx, (_, _, _, width_ratio) in enumerate(columns):
            if idx == len(columns) - 1:  # Last column gets remaining width
                col_width = int(usable_width - total_assigned)
            else:
                col_width = int((width_ratio / total_ratio) * usable_width)
                total_assigned += col_width
            
            assigned_widths.append(col_width)
            
            # Apply width to column
            table.columns[idx].width = col_width
            
            # Apply width to every cell in this column for enforcement
            for row in table.rows:
                row.cells[idx].width = col_width
        
        # Data rows — all records (no cap)
        row_count = len(rows)
        for row_dict in rows[:row_count]:
            new_row = table.add_row()
            row_cells = new_row.cells
            
            new_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            
            # Check if this is a red flag / never event
            is_red_flag = False
            try:
                clinical_risk = row_dict.get("clinical_risk_type_name", "")
                if clinical_risk and clinical_risk != "Ordinary":
                    is_red_flag = True
            except:
                pass
            
            for idx, (header_name, field_name, is_vertical, _) in enumerate(columns):
                # Special handling for target departments: concatenate all department names with hierarchy
                if field_name == "target_departments_display":
                    target_depts = row_dict.get("target_departments", [])
                    if target_depts and isinstance(target_depts, list):
                        # Categorize departments by priority
                        primary_and_matching = []  # Primary AND matches report entity
                        primary_only = []          # Primary but doesn't match
                        matching_only = []         # Matches entity but not primary
                        others = []                # Neither primary nor matching
                        
                        for dept in target_depts:
                            # For compact display, show only the most specific level (Section)
                            # If no section, show Department, if no department show Administration
                            if dept.get("section_name"):
                                display = dept["section_name"]
                            elif dept.get("department_name"):
                                display = dept["department_name"]
                            elif dept.get("administration_name"):
                                display = dept["administration_name"]
                            else:
                                continue  # Skip if no name available
                            
                            is_primary = dept.get("is_primary", False)
                            matches_entity = report_entity_name and report_entity_name in display
                            
                            # Categorize by priority
                            if is_primary and matches_entity:
                                primary_and_matching.append(display)
                            elif is_primary:
                                primary_only.append(display)
                            elif matches_entity:
                                matching_only.append(display)
                            else:
                                others.append(display)
                        
                        # Combine in priority order
                        all_displays = primary_and_matching + primary_only + matching_only + others
                        
                        # Limit to 3 departments for compact display
                        MAX_DISPLAY = 3
                        if len(all_displays) > MAX_DISPLAY:
                            displayed = all_displays[:MAX_DISPLAY]
                            remaining = len(all_displays) - MAX_DISPLAY
                            raw_value = ", ".join(displayed) + f" +{remaining}"
                        else:
                            raw_value = ", ".join(all_displays) if all_displays else "—"
                    else:
                        raw_value = "—"
                elif field_name in ("target_section_name", "target_department_name", "target_administration_name"):
                    target_depts = row_dict.get("target_departments", [])
                    primary = next((d for d in target_depts if d.get("is_primary")), target_depts[0] if target_depts else None)
                    key = {"target_section_name": "section_name", "target_department_name": "department_name", "target_administration_name": "administration_name"}[field_name]
                    raw_value = (primary.get(key) or "—") if primary else "—"
                elif field_name == "incident_id":
                    _inc = row_dict.get("incident_id")
                    raw_value = f"INC-{int(_inc):06d}" if _inc is not None else ""
                else:
                    raw_value = sanitize_value(row_dict.get(field_name, ""))

                # CRITICAL: Normalize text to remove manual line breaks from UI
                value = normalize_text(raw_value)
                
                # Truncate text appropriately (except for target departments)
                if field_name == "target_departments_display":
                    # Don't truncate target departments - show full hierarchy
                    pass
                elif not is_vertical:  # Wide horizontal columns
                    if len(value) > 400:
                        value = value[:400] + "..."
                else:  # Narrow vertical columns
                    if len(value) > 60:
                        value = value[:60] + "..."
                
                cell = row_cells[idx]
                cell.text = ""  # Clear cell

                p = cell.paragraphs[0]
                run = p.add_run(value)

                # Font size 9; Calibri for Arabic content, Traditional Arabic for English/numeric
                run.font.size = int(Pt(9))
                run.italic = False
                if is_arabic(value):
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
                else:
                    run.font.name = 'Traditional Arabic'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                    run._element.rPr.rFonts.set(qn('w:cs'), 'Traditional Arabic')

                # Line spacing 1.15 + optional 6pt before/after for horizontal cells
                _ppPr = p._element.get_or_add_pPr()
                for _s in _ppPr.findall(qn('w:spacing')):
                    _ppPr.remove(_s)
                _sp = OxmlElement('w:spacing')
                _sp.set(qn('w:line'), '276')   # 240 * 1.15 ≈ 276 twips
                _sp.set(qn('w:lineRule'), 'auto')
                if not is_vertical:
                    _sp.set(qn('w:before'), '120')  # 6pt = 120 twips
                    _sp.set(qn('w:after'),  '120')
                _ppPr.append(_sp)

                # Apply semantic coloring based on column and value (applies to ALL columns)
                cell_color = None
                    
                # Severity coloring
                if field_name == "severity_name":
                    value_lower = value.lower()
                    if "high" in value_lower:
                        cell_color = "FFB3BA"  # Light red
                    elif "medium" in value_lower:
                        cell_color = "FFDFBA"  # Light orange
                    elif "low" in value_lower:
                        cell_color = "BAFFC9"  # Light green
                    
                # Harm coloring
                elif field_name == "harm_level":
                    value_lower = value.lower()
                    if "death" in value_lower:
                        cell_color = "FF6B6B"  # Red
                    elif "severe" in value_lower:
                        cell_color = "FFA500"  # Orange
                    elif "no harm" in value_lower or "none" in value_lower:
                        cell_color = "BAFFC9"  # Light green
                    elif "minor" in value_lower or "temporary" in value_lower:
                        cell_color = "FFFFBA"  # Light yellow
                
                # Stage coloring
                elif field_name == "stage_name":
                    value_lower = value.lower()
                    if "admission" in value_lower:
                        cell_color = "BAE1FF"  # Light blue
                    elif "discharge" in value_lower or "transfer" in value_lower:
                        cell_color = "E0BBE4"  # Light purple
                    elif "examination" in value_lower or "diagnosis" in value_lower:
                        cell_color = "B4F8F8"  # Light cyan
                
                # Domain coloring
                elif field_name == "domain_name":
                    value_lower = value.lower()
                    if "clinical" in value_lower:
                        cell_color = "BAE1FF"  # Light blue
                    elif "management" in value_lower:
                        cell_color = "E0BBE4"  # Light purple
                    elif "relational" in value_lower:
                        cell_color = "FFDFBA"  # Light orange
                
                # Red flag row highlighting (very light red background for entire row)
                if is_red_flag and cell_color is None:
                    cell_color = "FFE5E5"  # Very light red
                
                # Apply cell background color if determined
                if cell_color:
                    try:
                        shading_elm = cell._element.get_or_add_tcPr()
                        shading = shading_elm.get_or_add_shd()
                        shading.fill = cell_color
                    except:
                        pass

                # Apply layout based on column type - ALL TEXT CENTERED
                if is_vertical:
                    # Vertical columns: center horizontally and vertically
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Remove spacing but allow natural wrapping
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    # Do NOT set noWrap - let Word handle text flow

                    # Set vertical text direction
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()

                    # Remove previous textDirection if exists
                    for el in tcPr.findall(qn('w:textDirection')):
                        tcPr.remove(el)

                    textDirection = OxmlElement('w:textDirection')
                    textDirection.set(qn('w:val'), 'btLr')
                    tcPr.append(textDirection)

                    # Vertical center alignment for cell
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)

                else:
                    # Horizontal columns: center horizontally and vertically
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # (6pt before/after spacing is set in the w:spacing block above)

                    # Set vertical center alignment for cell
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()

                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'center')
                    tcPr.append(vAlign)
                    # Allow Word to wrap text naturally for proper row expansion
                
        # Record count footer
        doc.add_paragraph()
        count_para = doc.add_paragraph()
        count_run = count_para.add_run(f"إجمالي السجلات: {len(rows)}")
        count_run.font.size = int(Pt(9))
        count_run.font.name = 'Traditional Arabic'
        count_run.italic = True
        count_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add signature block
        doc.add_paragraph()  # Spacer
        
        # Simple signature table (3 rows × 4 columns)
        sig_table = doc.add_table(rows=3, cols=4)
        sig_table.style = 'Table Grid'
        
        # Set column widths
        sig_table.columns[0].width = Cm(4)   # التاريخ
        sig_table.columns[1].width = Cm(5)   # التوقيع
        sig_table.columns[2].width = Cm(6)   # الاسم
        sig_table.columns[3].width = Cm(7)   # خاص خدمات المرضى
        
        # Define approval role names
        names = [
            "إسم مسؤول العملية",
            "إسم رئيس الدائرة", 
            "إسم مدير الإدارة"
        ]
        
        # Define patient services content for each row
        patient_services_content = [
            "خاص خدمات المرضى\nالإسم:",
            "تاريخ الإستلام:",
            "التوقيع:"
        ]
        
        # Fill all cells
        for row_idx in range(3):
            row_cells = sig_table.rows[row_idx].cells
            
            # Column 0: التاريخ
            cell = row_cells[0]
            para = cell.paragraphs[0]
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.right_to_left = True
            
            run = para.add_run("التاريخ:")
            run.font.bold = True
            run.font.size = int(Pt(10))
            run.font.name = 'Traditional Arabic'
            run.italic = False
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            # Column 1: التوقيع
            cell = row_cells[1]
            para = cell.paragraphs[0]
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.right_to_left = True
            
            run = para.add_run("التوقيع:")
            run.font.bold = True
            run.font.size = int(Pt(10))
            run.font.name = 'Traditional Arabic'
            run.italic = False
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            # Column 2: الاسم (role names)
            cell = row_cells[2]
            para = cell.paragraphs[0]
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.right_to_left = True
            
            run = para.add_run(names[row_idx])
            run.font.bold = True
            run.font.size = int(Pt(10))
            run.font.name = 'Traditional Arabic'
            run.italic = False
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            # Column 3: خاص خدمات المرضى (patient services)
            cell = row_cells[3]
            para = cell.paragraphs[0]
            para.clear()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.right_to_left = True
            
            run = para.add_run(patient_services_content[row_idx])
            run.font.bold = True
            run.font.size = int(Pt(10))
            run.font.name = 'Traditional Arabic'
            run.italic = False
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        
        # ========== ADD MOTIVATIONAL QUOTE TO REAL FOOTER ==========
        try:
            section = doc.sections[0]
            footer = section.footer
            footer.is_linked_to_previous = False  # Activate footer so Word writes footer1.xml
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            
            # Add the Arabic quote
            run = footer_para.add_run(_footer_text)
            run.font.size = int(Pt(10))
            run.font.name = 'Traditional Arabic'
            run.italic = False
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add top border to footer for professional look
            pPr = footer_para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '6')
            top.set(qn('w:color'), 'DDDDDD')
            pBdr.append(top)
            pPr.append(pBdr)
        except:
            pass  # If footer setup fails, continue without it
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


# Export service instance
reports_service = ReportsService()

