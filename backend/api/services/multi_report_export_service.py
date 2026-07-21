"""  
Multi-Report Export Service
Generates multiple report files (one per organizational unit) and packages them in a ZIP.
"""

from typing import Dict, Any, List, Literal, Optional
from io import BytesIO
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt

from ..schemas.auth_models import CurrentUser
from ..utils.guards import require_unit_in_scope
from ..constants.org_unit_types import (
    ORG_TYPE_ADMINISTRATION,
    ORG_TYPE_DEPARTMENT,
    ORG_TYPE_SECTION
)
from .monthly_report_service import monthly_report_service
from .reports_service import reports_service
from ..db_layer.admin_units import get_units_by_type, get_unit_hierarchy


class MultiReportExportService:
    """
    Service for generating multiple reports (one per organizational unit).
    Used when user selects "All" at a specific level (Administration/Department/Section).
    """
    
    def generate_multi_export(
        self,
        *,
        current_user: CurrentUser,
        year: int,
        month: int = None,
        start_date: str = None,
        end_date: str = None,
        file_format: Literal["pdf", "csv", "xlsx", "docx"],
        display_mode: Literal["detailed", "numeric"],
        report_level: Literal["administration", "department", "section"],
        selected_unit_ids: Optional[List[int]] = None,
        language: Literal["en", "ar"] = "en"
    ) -> Dict[str, Any]:
        """
        Generate multiple reports (one per unit) and package in ZIP.
        
        Phase 2.5.7: CRITICAL SECURITY - Each unit validated against current_user.allowed_unit_ids.
        If ANY unit is out of scope, entire request fails with 403.
        
        Args:
            current_user: Authenticated user with allowed org units
            year: Year for reports
            month: Month for reports (or None if using date range)
            start_date: Custom range start date (YYYY-MM-DD) (optional)
            end_date: Custom range end date (YYYY-MM-DD) (optional)
            file_format: Output format (docx, pdf, xlsx, csv)
            display_mode: detailed or numeric
            report_level: Which level to generate reports for
            selected_unit_ids: Specific unit IDs, or None for "all"
            language: Export language
            
        Returns:
            Dict with filename (ZIP) and content (bytes)
            
        Raises:
            HTTPException: 403 if any requested unit is outside user's scope
        """
        # Map report level to Type values in database
        type_mapping = {
            "administration": ORG_TYPE_ADMINISTRATION,
            "department": ORG_TYPE_DEPARTMENT,
            "section": ORG_TYPE_SECTION
        }
        
        unit_type = type_mapping[report_level]
        
        # Phase 2.5.8: CRITICAL - Validate requested IDs BEFORE filtering
        # This ensures fail-fast behavior: if user requests [29, 6] and 6 is out of scope,
        # the entire request fails with 403, not just silently omitting unit 6
        if selected_unit_ids:
            # Validate ALL requested unit IDs before processing
            for unit_id in selected_unit_ids:
                require_unit_in_scope(current_user, unit_id)
        
        # Get units to process
        if selected_unit_ids:
            # Get specific units by ID
            all_units = get_units_by_type(unit_type)
            units = [u for u in all_units if u["id"] in selected_unit_ids]
        else:
            # Get all units of this type
            units = get_units_by_type(unit_type)
            # Validate all units when processing "all"
            for unit in units:
                require_unit_in_scope(current_user, unit["id"])
        
        print(f"[MULTI-EXPORT] Generating {len(units)} reports for {report_level} level")
        print(f"[MULTI-EXPORT] All {len(units)} units validated against user scope")
        
        # Track results for summary
        successful_units = []
        empty_units = []
        failed_units = []
        
        # Generate ZIP file
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # Generate report for each unit
            for unit in units:
                unit_id = unit["id"]
                unit_name = unit["name"]
                
                try:
                    # Generate report for this specific unit
                    report_data = self._generate_unit_report(
                        current_user=current_user,
                        year=year,
                        month=month,
                        start_date=start_date,
                        end_date=end_date,
                        display_mode=display_mode,
                        report_level=report_level,
                        unit_id=unit_id
                    )
                    
                    # Check if unit has data
                    complaints_count = 0
                    if display_mode == "detailed" and isinstance(report_data, dict):
                        complaints = report_data.get("complaints", [])
                        complaints_count = len(complaints)
                    elif display_mode == "numeric" and isinstance(report_data, dict):
                        # For numeric mode, check summary total_complaints
                        summary = report_data.get("summary", {})
                        complaints_count = summary.get("total_complaints", 0)
                    
                    if complaints_count == 0:
                        print(f"[MULTI-EXPORT] {unit_name} (ID {unit_id}): No data - skipping file")
                        empty_units.append({"id": unit_id, "name": unit_name})
                        continue
                    
                    # Get proper hierarchy for this unit from database
                    unit_hierarchy = get_unit_hierarchy(unit_id)
                    
                    # Generate file content
                    file_content = self._generate_file_content(
                        report_data=report_data,
                        file_format=file_format,
                        language=language,
                        unit_name=unit_name,
                        unit_type=report_level,
                        display_mode=display_mode,
                        unit_id=unit_id,
                        unit_hierarchy=unit_hierarchy
                    )
                    
                    # Create filename
                    mode_prefix = "Numeric" if display_mode == "numeric" else "Detailed"
                    safe_name = self._sanitize_filename(unit_name)
                    
                    if month is not None:
                        month_name = self._get_month_name(month, language)
                        filename = f"{mode_prefix}_Report_{safe_name}_{month_name}{year}.{file_format}"
                    else:
                        # For date range
                        filename = f"{mode_prefix}_Report_{safe_name}_{start_date}_to_{end_date}.{file_format}"
                    
                    # Add to ZIP
                    zip_file.writestr(filename, file_content)
                    
                    successful_units.append({
                        "id": unit_id,
                        "name": unit_name,
                        "filename": filename,
                        "complaints_count": complaints_count
                    })
                    
                    print(f"[MULTI-EXPORT] [OK] {unit_name}: {complaints_count} complaints -> {filename}")
                    
                except Exception as e:
                    print(f"[MULTI-EXPORT] [FAIL] {unit_name} (ID {unit_id}): Failed - {str(e)}")
                    failed_units.append({
                        "id": unit_id,
                        "name": unit_name,
                        "error": str(e)
                    })
            
            # Generate summary file
            summary_content = self._generate_summary_file(
                year=year,
                month=month,
                report_level=report_level,
                successful_units=successful_units,
                empty_units=empty_units,
                failed_units=failed_units,
                language=language
            )
            
            zip_file.writestr("_SUMMARY_Report.docx", summary_content)
            print(f"[MULTI-EXPORT] Summary file added")
        
        # Prepare ZIP for download
        zip_buffer.seek(0)
        
        # Create ZIP filename
        month_name = self._get_month_name(month, language)
        zip_filename = f"Monthly_Reports_{report_level.capitalize()}_{month_name}{year}.zip"
        
        print(f"[MULTI-EXPORT] Complete: {len(successful_units)} files, {len(empty_units)} empty, {len(failed_units)} failed")
        
        return {
            "filename": zip_filename,
            "content": zip_buffer.getvalue(),
            "content_type": "application/zip"
        }
    
    def _generate_unit_report(
        self,
        current_user: CurrentUser,
        year: int,
        month: Optional[int],
        start_date: Optional[str],
        end_date: Optional[str],
        display_mode: str,
        report_level: str,
        unit_id: int
    ) -> Dict[str, Any]:
        """
        Generate report data for a specific organizational unit.
        
        Phase 2.5.7: Passes current_user for scope enforcement at service layer.
        """
        
        # Map report level to filter parameter
        filter_param_map = {
            "administration": "administration_ids",
            "department": "department_ids",
            "section": "section_ids"
        }
        
        # Build filter (single unit ID as string)
        filters = {
            filter_param_map[report_level]: str(unit_id)
        }
        
        # Use the same service as normal export
        result = monthly_report_service.generate_monthly_report(
            current_user=current_user,
            year=year,
            month=month,
            start_date=start_date,
            end_date=end_date,
            mode=display_mode,
            scope=filters.get("scope"),
            administration_ids=filters.get("administration_ids"),
            department_ids=filters.get("department_ids"),
            section_ids=filters.get("section_ids"),
            page=1,
            page_size=9999  # Get all records
        )
        
        return result
    
    def _generate_file_content(
        self,
        report_data: Dict[str, Any],
        file_format: str,
        language: str,
        unit_name: str = None,
        unit_type: str = None,
        display_mode: str = "detailed",
        unit_id: int = None,
        unit_hierarchy: dict = None
    ) -> bytes:
        """Generate file content in specified format."""
        
        # Extract complaints data
        if isinstance(report_data, dict) and "complaints" in report_data:
            export_data = report_data["complaints"]
        else:
            export_data = report_data
        
        # Determine entity parameters for header
        # We need to populate ALL THREE fields (Administration, Department, Section)
        # based on what level we're reporting at and the ACTUAL HIERARCHY from the database
        report_administration = None
        report_department = None
        report_section = None
        
        # Helper function to get unique values from data (fallback only)
        def get_unique_values(data, field_name):
            """Get unique non-null values from complaints data"""
            if not data or not isinstance(data, list):
                return []
            values = set()
            for row in data:
                val = row.get(field_name)
                if val and val != "—":
                    values.add(val)
            return list(values)
        
        # Extract hierarchy information from the actual data (for fallback and multi-value detection)
        unique_admins = get_unique_values(export_data, "administration_name")
        unique_depts = get_unique_values(export_data, "department_name")
        unique_sections = get_unique_values(export_data, "section_name")
        
        print(f"[MULTI EXPORT] Unit: {unit_name} ({unit_type}) ID={unit_id}")
        print(f"[MULTI EXPORT] Data has {len(export_data) if isinstance(export_data, list) else 0} records")
        print(f"[MULTI EXPORT] Unique admins from data: {unique_admins}")
        print(f"[MULTI EXPORT] Unique depts from data: {unique_depts}")
        print(f"[MULTI EXPORT] Unique sections from data: {unique_sections}")
        
        if unit_hierarchy:
            print(f"[MULTI EXPORT] Hierarchy from DB: parent={unit_hierarchy.get('parent_name')}, grandparent={unit_hierarchy.get('grandparent_name')}")
        
        # Set header values based on report level using ACTUAL HIERARCHY from database
        if unit_type == "administration":
            # Administration report: show admin name, aggregate dept/section info
            report_administration = unit_name
            report_department = unique_depts[0] if len(unique_depts) == 1 else ("متعدد" if len(unique_depts) > 1 else "—")
            report_section = unique_sections[0] if len(unique_sections) == 1 else ("متعدد" if len(unique_sections) > 1 else "—")
        elif unit_type == "department":
            # Department report: use hierarchy from database for parent administration
            report_department = unit_name
            # Use DB hierarchy for administration (parent of department)
            if unit_hierarchy and unit_hierarchy.get('parent_name'):
                report_administration = unit_hierarchy['parent_name']
                print(f"[MULTI EXPORT] Using DB hierarchy for admin: {report_administration}")
            else:
                # Fallback to data extraction (may be wrong if data has mixed parents)
                report_administration = unique_admins[0] if unique_admins else "—"
                print(f"[MULTI EXPORT] Fallback to data for admin: {report_administration}")
            report_section = unique_sections[0] if len(unique_sections) == 1 else ("متعدد" if len(unique_sections) > 1 else "—")
        elif unit_type == "section":
            # Section report: use hierarchy from database for parent (department) and grandparent (administration)
            report_section = unit_name
            # Use DB hierarchy for department (parent) and administration (grandparent)
            if unit_hierarchy:
                if unit_hierarchy.get('parent_name'):
                    report_department = unit_hierarchy['parent_name']
                    print(f"[MULTI EXPORT] Using DB hierarchy for dept: {report_department}")
                else:
                    report_department = unique_depts[0] if unique_depts else "—"
                    print(f"[MULTI EXPORT] Fallback to data for dept: {report_department}")
                    
                if unit_hierarchy.get('grandparent_name'):
                    report_administration = unit_hierarchy['grandparent_name']
                    print(f"[MULTI EXPORT] Using DB hierarchy for admin: {report_administration}")
                else:
                    report_administration = unique_admins[0] if unique_admins else "—"
                    print(f"[MULTI EXPORT] Fallback to data for admin: {report_administration}")
            else:
                # Full fallback
                report_department = unique_depts[0] if unique_depts else "—"
                report_administration = unique_admins[0] if unique_admins else "—"
        
        # Generate file based on format
        if file_format == "docx":
            # Check if numeric or detailed mode
            if display_mode == "numeric":
                # Use numeric report generator
                return reports_service.generate_monthly_numeric_word_report(
                    report_data=report_data,  # Pass full dict (not just export_data)
                    filename="temp.docx",
                    language=language,
                    report_entity_name=unit_name,
                    report_entity_type=unit_type
                )
            else:
                # Formatter routing: read monthly_report_format from
                # APP_ReportConfig and route to Classical or Stylish, exactly
                # like the single-file export path in report_export_service.py.
                # This was previously missing here — the ZIP/multi-export path
                # always called the Classical formatter regardless of the
                # configured format, so switching the setting to "stylish" had
                # no effect on ZIP downloads (only single-file exports).
                try:
                    from ..db_layer.report_config_db import get_report_config as _get_cfg
                    _monthly_format = (_get_cfg().get("monthly_report_format") or "classical").strip().lower()
                    if _monthly_format not in ("classical", "stylish"):
                        _monthly_format = "classical"
                except Exception as _cfg_err:
                    print(f"[MULTI EXPORT] Warning: could not read monthly_report_format ({_cfg_err}), defaulting to classical")
                    _monthly_format = "classical"

                if _monthly_format == "stylish":
                    from .monthly_stylish_word_formatter import generate_monthly_stylish_docx
                    print(f"[MULTI EXPORT] Stylish formatter -> generate_monthly_stylish_docx (unit={unit_name})")
                    return generate_monthly_stylish_docx(
                        report_data=report_data,  # full dict (complaints/notices/period/intent_counts)
                        filename="temp.docx",
                        language=language,
                        report_entity_name=unit_name,
                        report_entity_type=unit_type,
                        report_administration=report_administration,
                        report_department=report_department,
                        report_section=report_section
                    )
                else:
                    # Classical formatter (default, unchanged behaviour)
                    return reports_service.generate_docx_export(
                        report_data=export_data,
                        filename="temp.docx",
                        language=language,
                        report_entity_name=unit_name,
                        report_entity_type=unit_type,
                        report_administration=report_administration,
                        report_department=report_department,
                        report_section=report_section
                    )
        elif file_format == "xlsx":
            return reports_service.generate_xlsx_export(
                report_data=export_data,
                filename="temp.xlsx",
                language=language
            )
        elif file_format == "pdf":
            return reports_service.generate_pdf_export(
                report_data=export_data,
                filename="temp.pdf",
                language=language,
                include_charts=True,
                report_entity_name=unit_name,
                report_entity_type=unit_type,
                report_administration=report_administration,
                report_department=report_department,
                report_section=report_section
            )
        else:  # csv
            return reports_service.generate_csv_export(
                report_data=export_data,
                filename="temp.csv",
                language=language
            )
    
    def _generate_summary_file(
        self,
        year: int,
        month: int,
        report_level: str,
        successful_units: List[Dict],
        empty_units: List[Dict],
        failed_units: List[Dict],
        language: str
    ) -> bytes:
        """Generate summary Word document."""
        
        doc = Document()
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(f"Monthly Reports Summary - {report_level.capitalize()} Level")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = 1  # Center
        
        # Period
        doc.add_paragraph(f"Period: {self._get_month_name(month, 'en')} {year}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Statistics
        doc.add_heading("Summary Statistics", level=2)
        total_units = len(successful_units) + len(empty_units) + len(failed_units)
        total_complaints = sum(u["complaints_count"] for u in successful_units)
        
        doc.add_paragraph(f"Total Units Processed: {total_units}")
        doc.add_paragraph(f"Units with Data: {len(successful_units)}")
        doc.add_paragraph(f"Units with No Complaints: {len(empty_units)}")
        doc.add_paragraph(f"Failed Units: {len(failed_units)}")
        doc.add_paragraph(f"Total Complaints: {total_complaints}")
        doc.add_paragraph()
        
        # Successful units
        if successful_units:
            doc.add_heading("Units with Data (Files Generated)", level=2)
            for unit in successful_units:
                doc.add_paragraph(
                    f"✓ {unit['name']} - {unit['complaints_count']} complaints → {unit['filename']}",
                    style='List Bullet'
                )
            doc.add_paragraph()
        
        # Empty units
        if empty_units:
            doc.add_heading("Units with No Complaints (No Files)", level=2)
            for unit in empty_units:
                doc.add_paragraph(f"○ {unit['name']} - No complaints in this period", style='List Bullet')
            doc.add_paragraph()
        
        # Failed units
        if failed_units:
            doc.add_heading("Failed Units (Errors)", level=2)
            for unit in failed_units:
                doc.add_paragraph(f"✗ {unit['name']} - Error: {unit['error']}", style='List Bullet')
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _get_month_name(self, month: int, language: str) -> str:
        """Get month name in specified language."""
        months_en = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
        months_ar = ["", "يناير", "فبراير", "مارس", "أبريل", "مايو", "يونيو", "يوليو", "أغسطس", "سبتمبر", "أكتوبر", "نوفمبر", "ديسمبر"]
        
        if language == "ar":
            return months_ar[month]
        return months_en[month]
    
    def _sanitize_filename(self, name: str) -> str:
        """Sanitize unit name for use in filename."""
        # Remove or replace characters that are invalid in filenames
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '_')
        # Limit length
        return name[:50]


# Singleton instance
multi_report_export_service = MultiReportExportService()
