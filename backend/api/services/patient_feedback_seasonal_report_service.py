"""
Patient Feedback Seasonal Report Word Generator

Generates comprehensive Word documents with RCA (Root Cause Analysis) 
and Satisfaction statistics for a given seasonal period.

This report aggregates:
- RCA breakdowns by cause type (Staff, Process, Equipment, Environment)
- Detailed sub-cause analysis with individual cause factors
- Satisfaction status distribution (Satisfied, Not Satisfied, Not Present)
- Preventability analysis with specific measures breakdown
- Feedback coverage metrics
- Visual charts and graphs for data visualization
"""

from typing import List, Dict, Any
from datetime import date, datetime
from io import BytesIO
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

# For chart generation
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server use
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import rcParams
import numpy as np

from core.database import get_connection

# Configure matplotlib for Arabic text support
rcParams['font.family'] = ['Arial', 'DejaVu Sans', 'sans-serif']


def _fetch_rca_statistics(season_start: date, season_end: date) -> Dict[str, Any]:
    """
    Fetch structured RCA statistics from the new APP_SubcaseRCASuggestionSelection model.

    Returns per-category breakdowns of selected causes and corrective action suggestions,
    grouped and counted for the seasonal period.
    """
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor()

        query = """
            SELECT
                cat.CategoryID,
                cat.CategoryNameAr,
                cat.CategoryNameEn,
                cat.SortOrder AS CatSort,
                sug.SuggestionID,
                sug.SuggestionType,
                sug.SuggestionTextAr,
                sug.SuggestionTextEn,
                sug.SortOrder AS SugSort,
                COUNT(DISTINCT sel.SubcaseID) AS SelectionCount
            FROM dbo.APP_SubcaseRCASuggestionSelection sel
            JOIN dbo.APP_RCASuggestion sug
                ON sel.SuggestionID = sug.SuggestionID
            JOIN dbo.APP_RCAFactorCategory cat
                ON sug.CategoryID = cat.CategoryID
            JOIN dbo.APP_AdministrativeSubcase sub
                ON sel.SubcaseID = sub.SubcaseID
            JOIN dbo.APP_IncidentCase ic
                ON sub.IncidentRequestCaseID = ic.IncidentRequestCaseID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
            GROUP BY
                cat.CategoryID, cat.CategoryNameAr, cat.CategoryNameEn, cat.SortOrder,
                sug.SuggestionID, sug.SuggestionType,
                sug.SuggestionTextAr, sug.SuggestionTextEn, sug.SortOrder
            ORDER BY cat.SortOrder, cat.CategoryID, sug.SortOrder, sug.SuggestionID
        """
        cursor.execute(query, (season_start.isoformat(), season_end.isoformat()))
        rows = cursor.fetchall()

        # Count distinct subcases with any selection
        count_query = """
            SELECT COUNT(DISTINCT sel.SubcaseID)
            FROM dbo.APP_SubcaseRCASuggestionSelection sel
            JOIN dbo.APP_AdministrativeSubcase sub ON sel.SubcaseID = sub.SubcaseID
            JOIN dbo.APP_IncidentCase ic ON sub.IncidentRequestCaseID = ic.IncidentRequestCaseID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
        """
        cursor.execute(count_query, (season_start.isoformat(), season_end.isoformat()))
        total_subcases_row = cursor.fetchone()
        total_subcases_with_selections = total_subcases_row[0] if total_subcases_row else 0

        # Build per-category structure
        from collections import OrderedDict
        categories: dict = OrderedDict()
        for row in rows:
            (cat_id, cat_name_ar, cat_name_en, cat_sort,
             sug_id, sug_type, sug_text_ar, sug_text_en,
             sug_sort, count) = row

            if cat_id not in categories:
                categories[cat_id] = {
                    "category_id": cat_id,
                    "category_name_ar": cat_name_ar or cat_name_en or "",
                    "category_name_en": cat_name_en or "",
                    "causes": [],
                    "action_items": [],
                }
            item = {
                "text_ar": sug_text_ar or "",
                "text_en": sug_text_en or "",
                "count": count or 0,
            }
            if sug_type == "CAUSE":
                categories[cat_id]["causes"].append(item)
            else:
                categories[cat_id]["action_items"].append(item)

        return {
            "total_subcases_with_selections": total_subcases_with_selections,
            "by_category": list(categories.values()),
        }
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


def _fetch_satisfaction_statistics(season_start: date, season_end: date) -> Dict[str, Any]:
    """
    Fetch Satisfaction statistics for the seasonal period.
    
    Returns breakdown by status and feedback metrics.
    """
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # Total cases in period
        total_query = """
            SELECT COUNT(DISTINCT ic.IncidentRequestCaseID)
            FROM dbo.APP_IncidentCase ic
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
        """
        cursor.execute(total_query, (season_start.isoformat(), season_end.isoformat()))
        total_cases = cursor.fetchone()[0] or 0
        
        # Satisfaction by status
        status_query = """
            SELECT
                ss.SatisfactionStatusID,
                ss.StatusNameEn,
                ss.StatusNameAr,
                COUNT(*) as count
            FROM dbo.APP_IncidentCaseSatisfaction s
            INNER JOIN dbo.APP_Lookup_SatisfactionStatus ss ON s.SatisfactionStatusID = ss.SatisfactionStatusID
            INNER JOIN dbo.APP_IncidentCase ic ON s.IncidentRequestCaseID = ic.IncidentRequestCaseID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
            GROUP BY ss.SatisfactionStatusID, ss.StatusNameEn, ss.StatusNameAr
            ORDER BY ss.SatisfactionStatusID
        """
        cursor.execute(status_query, (season_start.isoformat(), season_end.isoformat()))
        
        by_status = []
        total_with_satisfaction = 0
        for row in cursor.fetchall():
            count = row[3]
            total_with_satisfaction += count
            by_status.append({
                "status_id": row[0],
                "status_en": row[1],
                "status_ar": row[2],
                "count": count
            })
        
        # Feedback needed/given stats
        feedback_query = """
            SELECT
                SUM(CASE WHEN s.FeedbackNeeded = 1 THEN 1 ELSE 0 END) as needed,
                SUM(CASE WHEN s.FeedbackGiven = 1 THEN 1 ELSE 0 END) as given
            FROM dbo.APP_IncidentCaseSatisfaction s
            INNER JOIN dbo.APP_IncidentCase ic ON s.IncidentRequestCaseID = ic.IncidentRequestCaseID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
        """
        cursor.execute(feedback_query, (season_start.isoformat(), season_end.isoformat()))
        fb_row = cursor.fetchone()
        
        return {
            "total_cases": total_cases,
            "cases_with_satisfaction": total_with_satisfaction,
            "cases_without_satisfaction": total_cases - total_with_satisfaction,
            "coverage_percentage": round(total_with_satisfaction / total_cases * 100, 1) if total_cases > 0 else 0,
            "by_status": by_status,
            "feedback_stats": {
                "feedback_needed": fb_row[0] or 0 if fb_row else 0,
                "feedback_given": fb_row[1] or 0 if fb_row else 0
            }
        }
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


def _fetch_subcase_coverage(season_start: date, season_end: date) -> Dict[str, Any]:
    """
    Fetch subcase RCA coverage statistics.
    """
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        query = """
            SELECT
                COUNT(DISTINCT sub.SubcaseID) as total_subcases,
                COUNT(DISTINCT CASE WHEN f.AdministrativeSubcaseID IS NOT NULL THEN sub.SubcaseID END) as subcases_with_rca
            FROM dbo.APP_AdministrativeSubcase sub
            INNER JOIN dbo.APP_IncidentCase ic ON sub.IncidentRequestCaseID = ic.IncidentRequestCaseID
            LEFT JOIN dbo.APP_IncidentCaseFeedback f ON sub.SubcaseID = f.AdministrativeSubcaseID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
        """
        cursor.execute(query, (season_start.isoformat(), season_end.isoformat()))
        row = cursor.fetchone()
        
        total = row[0] or 0
        with_rca = row[1] or 0
        
        return {
            "total_subcases": total,
            "subcases_with_rca": with_rca,
            "subcases_without_rca": total - with_rca,
            "rca_coverage_percentage": round(with_rca / total * 100, 1) if total > 0 else 0
        }
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _format_table_header(table, headers: List[str], color_hex: str = "667eea"):
    """Format table header row with styling."""
    header_row = table.rows[0]
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header_text
        _set_cell_shading(cell, color_hex)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(255, 255, 255)
                run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')


def _generate_rca_pie_chart(rca_stats: Dict[str, Any]) -> BytesIO:
    """
    Generate a pie chart showing RCA cause type distribution.
    
    Returns the chart as a BytesIO object for embedding in Word document.
    """
    cause_types = rca_stats.get("by_cause_type", {})
    
    labels = [
        'الكوادر البشرية\n(Staff)',
        'العمليات\n(Process)', 
        'المعدات\n(Equipment)',
        'البيئة\n(Environment)'
    ]
    sizes = [
        cause_types.get("staff", 0),
        cause_types.get("process", 0),
        cause_types.get("equipment", 0),
        cause_types.get("environment", 0)
    ]
    
    # Filter out zero values
    non_zero_data = [(l, s) for l, s in zip(labels, sizes) if s > 0]
    if not non_zero_data:
        return None
    
    labels, sizes = zip(*non_zero_data)
    
    # Professional color palette
    colors = ['#667eea', '#4CAF50', '#FF9800', '#E91E63'][:len(sizes)]
    explode = [0.02] * len(sizes)  # Slight separation
    
    fig, ax = plt.subplots(figsize=(8, 6), facecolor='white')
    
    wedges, texts, autotexts = ax.pie(
        sizes, 
        labels=labels,
        autopct=lambda pct: f'{pct:.1f}%\n({int(pct/100*sum(sizes))})',
        explode=explode,
        colors=colors,
        shadow=False,
        startangle=90,
        textprops={'fontsize': 10}
    )
    
    # Style the percentage text
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')
        autotext.set_fontsize(9)
    
    ax.set_title('RCA Cause Type Distribution\nتوزيع أنواع الأسباب الجذرية', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.axis('equal')
    
    plt.tight_layout()
    
    # Save to BytesIO
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    plt.close(fig)
    img_buffer.seek(0)
    
    return img_buffer


def _generate_subcause_bar_chart(subcauses: Dict[str, int], title: str, color: str) -> BytesIO:
    """
    Generate a horizontal bar chart for sub-cause analysis.
    
    Returns the chart as a BytesIO object for embedding in Word document.
    """
    # Filter out zero values
    filtered = {k: v for k, v in subcauses.items() if v > 0}
    if not filtered:
        return None
    
    # Sort by value descending
    sorted_items = sorted(filtered.items(), key=lambda x: x[1], reverse=True)
    labels, values = zip(*sorted_items)
    
    fig, ax = plt.subplots(figsize=(8, max(3, len(labels) * 0.5)), facecolor='white')
    
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, values, color=color, edgecolor='white', height=0.6)
    
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()  # Highest value at top
    
    ax.set_xlabel('عدد الحالات (Count)', fontsize=10)
    ax.set_title(title, fontsize=12, fontweight='bold', pad=15)
    
    # Add value labels on bars
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2,
                str(int(val)), va='center', fontsize=9, fontweight='bold')
    
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    img_buffer.seek(0)
    
    return img_buffer


def _generate_preventive_measures_chart(measures: Dict[str, int]) -> BytesIO:
    """
    Generate a bar chart for preventive measures breakdown.
    """
    # Filter out zero values
    filtered = {k: v for k, v in measures.items() if v > 0}
    if not filtered:
        return None
    
    sorted_items = sorted(filtered.items(), key=lambda x: x[1], reverse=True)
    labels, values = zip(*sorted_items)
    
    fig, ax = plt.subplots(figsize=(8, 4), facecolor='white')
    
    x_pos = np.arange(len(labels))
    bars = ax.bar(x_pos, values, color='#4CAF50', edgecolor='white', width=0.6)
    
    ax.set_xticks(x_pos)
    ax.set_xticklabels(labels, fontsize=8, rotation=15, ha='right')
    
    ax.set_ylabel('عدد الحالات (Count)', fontsize=10)
    ax.set_title('Preventive Measures Distribution\nتوزيع الإجراءات الوقائية', 
                 fontsize=12, fontweight='bold', pad=15)
    
    # Add value labels on bars
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2,
                str(int(val)), ha='center', fontsize=9, fontweight='bold')
    
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    img_buffer.seek(0)
    
    return img_buffer


def _get_rca_insights(rca_stats: Dict[str, Any]) -> List[str]:
    """
    Generate analytical insights based on RCA data.
    
    Returns a list of insight strings for the report.
    """
    insights = []
    
    total = rca_stats.get("total_rca_records", 0)
    if total == 0:
        return ["لا توجد بيانات كافية لتحليل الأسباب الجذرية في هذه الفترة."]
    
    cause_types = rca_stats.get("by_cause_type", {})
    
    # Find dominant cause type
    max_cause = max(cause_types.items(), key=lambda x: x[1])
    cause_names = {
        "staff": "الكوادر البشرية (Staff)",
        "process": "العمليات (Process)",
        "equipment": "المعدات (Equipment)",
        "environment": "البيئة (Environment)"
    }
    
    if max_cause[1] > 0:
        pct = round(max_cause[1] / total * 100, 1)
        insights.append(f"• السبب الأكثر شيوعاً هو {cause_names[max_cause[0]]} بنسبة {pct}% من الحالات المحللة.")
    
    # Staff analysis
    staff_subcauses = rca_stats.get("staff_subcauses", {})
    if staff_subcauses:
        top_staff = max(staff_subcauses.items(), key=lambda x: x[1])
        if top_staff[1] > 0:
            insights.append(f"• أبرز مشكلة في الكوادر البشرية: {top_staff[0]} ({top_staff[1]} حالة).")
    
    # Process analysis
    process_subcauses = rca_stats.get("process_subcauses", {})
    if process_subcauses:
        top_process = max(process_subcauses.items(), key=lambda x: x[1])
        if top_process[1] > 0:
            insights.append(f"• أبرز مشكلة في العمليات: {top_process[0]} ({top_process[1]} حالة).")
    
    # Preventive measures analysis
    preventability = rca_stats.get("preventability", {})
    has_prev = preventability.get("has_preventive_measures", 0)
    no_prev = preventability.get("no_preventive_measures", 0)
    
    if has_prev + no_prev > 0:
        prev_pct = round(has_prev / (has_prev + no_prev) * 100, 1)
        if prev_pct >= 70:
            insights.append(f"• نسبة عالية ({prev_pct}%) من الحالات لديها إجراءات وقائية مقترحة - مؤشر إيجابي.")
        elif prev_pct < 50:
            insights.append(f"• تحذير: نسبة منخفضة ({prev_pct}%) من الحالات لديها إجراءات وقائية - يُنصح بتعزيز التحليل.")
    
    # Recommendations based on data
    insights.append("")
    insights.append("التوصيات (Recommendations):")
    
    if cause_types.get("staff", 0) > total * 0.3:
        insights.append("• التركيز على برامج التدريب وتطوير الكفاءات للموظفين.")
    if cause_types.get("process", 0) > total * 0.2:
        insights.append("• مراجعة وتحديث البروتوكولات والإجراءات التشغيلية.")
    if cause_types.get("equipment", 0) > total * 0.15:
        insights.append("• تقييم احتياجات المعدات والأنظمة التقنية.")
    if cause_types.get("environment", 0) > total * 0.1:
        insights.append("• تحسين بيئة العمل وظروفها.")
    
    return insights


def generate_patient_feedback_seasonal_word(
    season_start: date,
    season_end: date
) -> bytes:
    """
    Generate Patient Feedback Seasonal Report as Word document.
    
    Contains:
    - Executive summary with coverage metrics
    - RCA (Root Cause Analysis) breakdown
    - Satisfaction status distribution
    - Preventability analysis
    - Department-wise breakdown
    
    Args:
        season_start: Start date of reporting period
        season_end: End date of reporting period
    
    Returns:
        Word document as bytes
    """
    # Fetch all statistics
    rca_stats = _fetch_rca_statistics(season_start, season_end)
    satisfaction_stats = _fetch_satisfaction_statistics(season_start, season_end)
    subcase_coverage = _fetch_subcase_coverage(season_start, season_end)
    
    # Create document
    doc = Document()

    # Set default font to Traditional Arabic
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Page setup
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ========== LOGO HEADER ==========
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            section.header_distance = Inches(0.1)
            header_section = section.header
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[PATIENT_FEEDBACK_SEASONAL] Could not add logo: {e}")

    # ========== TITLE PAGE ==========
    doc.add_paragraph()

    title_ar = doc.add_heading('تقرير ملاحظات المرضى الموسمي', 0)
    title_ar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_ar.runs:
        run.font.size = Pt(22)
        run.font.name = 'Calibri'
        run.font.bold = True
        run.font.color.rgb = RGBColor(102, 126, 234)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

    title_en = doc.add_heading('Patient Feedback Seasonal Report', 1)
    title_en.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_en.runs:
        run.font.size = Pt(16)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(118, 75, 162)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run = subtitle.add_run('تحليل الأسباب الجذرية ورضا المرضى')
    sub_run.font.size = Pt(14)
    sub_run.font.name = 'Calibri'
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    sub_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_run2 = subtitle2.add_run('Root Cause Analysis & Patient Satisfaction')
    sub_run2.font.size = Pt(12)
    sub_run2.font.name = 'Calibri'
    sub_run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Period info — keep as a single Arabic line to avoid LTR/RTL number collision
    period = doc.add_paragraph()
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    period_run = period.add_run(
        f'{season_start.strftime("%Y-%m-%d")} — {season_end.strftime("%Y-%m-%d")} :الفترة'
    )
    period_run.font.size = Pt(12)
    period_run.font.name = 'Calibri'
    period_run.font.bold = True
    period_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

    # Generation timestamp
    gen_time = doc.add_paragraph()
    gen_time.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen_run = gen_time.add_run(f'{datetime.now().strftime("%Y-%m-%d %H:%M")} :تم الإنشاء')
    gen_run.font.size = Pt(9)
    gen_run.font.name = 'Calibri'
    gen_run.font.color.rgb = RGBColor(128, 128, 128)
    gen_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
    
    doc.add_page_break()
    
    # ========== EXECUTIVE SUMMARY ==========
    summary_heading = doc.add_heading('الملخص التنفيذي / Executive Summary', 1)
    for run in summary_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(102, 126, 234)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
    
    # Summary metrics table
    summary_table = doc.add_table(rows=6, cols=3)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    _format_table_header(summary_table, ['المقياس / Metric', 'القيمة / Value', 'النسبة / Percentage'])
    
    summary_data = [
        ('إجمالي الحالات / Total Cases', str(satisfaction_stats['total_cases']), '-'),
        ('إجمالي الحالات الفرعية / Total Subcases', str(subcase_coverage['total_subcases']), '-'),
        ('الحالات الفرعية مع RCA / Subcases with RCA', 
         str(subcase_coverage['subcases_with_rca']), 
         f"{subcase_coverage['rca_coverage_percentage']}%"),
        ('حالات مع رضا المريض / Cases with Satisfaction', 
         str(satisfaction_stats['cases_with_satisfaction']), 
         f"{satisfaction_stats['coverage_percentage']}%"),
        ('الحالات مع اختيارات RCA / Subcases with RCA Selections', str(rca_stats['total_subcases_with_selections']), '-'),
    ]
    
    for idx, (metric, value, pct) in enumerate(summary_data):
        row = summary_table.rows[idx + 1]
        row.cells[0].text = metric
        row.cells[1].text = value
        row.cells[2].text = pct
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    # ========== RCA CAUSES SECTION (DB-driven) ==========
    rca_heading = doc.add_heading('تحليل الأسباب الجذرية / Root Cause Analysis', 1)
    for run in rca_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(102, 126, 234)
        run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

    rca_intro = doc.add_paragraph()
    rca_intro_run = rca_intro.add_run(
        f"يوضح هذا القسم الأسباب الجذرية المختارة لعدد "
        f"{rca_stats['total_subcases_with_selections']} "
        f"حالة فرعية خلال الفترة من {season_start.strftime('%Y-%m-%d')} "
        f"إلى {season_end.strftime('%Y-%m-%d')}."
    )
    rca_intro_run.font.size = Pt(11)
    rca_intro_run.font.name = 'Calibri'
    rca_intro.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    by_category = rca_stats.get('by_category', [])
    if not by_category:
        no_data = doc.add_paragraph()
        no_data_run = no_data.add_run('لا توجد بيانات RCA لهذه الفترة / No RCA data for this period')
        no_data_run.font.size = Pt(11)
        no_data_run.font.name = 'Calibri'
        no_data_run.font.color.rgb = RGBColor(150, 150, 150)
        no_data_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
    else:
        def _add_suggestion_table(items: list, color_hex: str, col_label: str):
            if not items:
                return
            tbl = doc.add_table(rows=len(items) + 1, cols=3)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _format_table_header(tbl, [col_label, 'النص / Text', 'عدد الاختيارات / Count'], color_hex)
            total = sum(i['count'] for i in items) or 1
            for idx, item in enumerate(items):
                row = tbl.rows[idx + 1]
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = item['text_ar'] or item['text_en'] or ''
                row.cells[2].text = str(item['count'])
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Calibri'
                            run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

        for cat in by_category:
            cat_heading = doc.add_heading(
                f"{cat['category_name_ar']} / {cat['category_name_en']}", 2
            )
            for run in cat_heading.runs:
                run.font.name = 'Calibri'
                run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

            if cat['causes']:
                causes_label = doc.add_paragraph()
                causes_lbl_run = causes_label.add_run('العوامل المسبّبة المختارة')
                causes_lbl_run.font.bold = True
                causes_lbl_run.font.size = Pt(11)
                causes_lbl_run.font.name = 'Calibri'
                causes_lbl_run.font.color.rgb = RGBColor(180, 90, 0)
                causes_lbl_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
                _add_suggestion_table(cat['causes'], 'F5A623', '#')

            if cat['action_items']:
                actions_label = doc.add_paragraph()
                actions_lbl_run = actions_label.add_run('الإجراءات التصحيحية المقترحة المختارة')
                actions_lbl_run.font.bold = True
                actions_lbl_run.font.size = Pt(11)
                actions_lbl_run.font.name = 'Calibri'
                actions_lbl_run.font.color.rgb = RGBColor(0, 90, 180)
                actions_lbl_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
                _add_suggestion_table(cat['action_items'], '667eea', '#')

            doc.add_paragraph()

    doc.add_page_break()
    
    # ========== SATISFACTION SECTION ==========
    sat_heading = doc.add_heading('رضا المرضى / Patient Satisfaction', 1)
    for run in sat_heading.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(102, 126, 234)
    
    # Satisfaction by Status
    status_subheading = doc.add_heading('توزيع حسب حالة الرضا / Distribution by Satisfaction Status', 2)
    for run in status_subheading.runs:
        run.font.name = 'Calibri'
    
    status_table = doc.add_table(rows=len(satisfaction_stats['by_status']) + 2, cols=4)
    status_table.style = 'Table Grid'
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    _format_table_header(status_table, ['الحالة (EN)', 'الحالة (AR)', 'العدد / Count', 'النسبة / Percentage'])
    
    total_satisfaction = satisfaction_stats['cases_with_satisfaction'] or 1
    
    status_colors = {
        1: 'F0F0F0',  # Not Present - gray
        2: 'D4EDDA',  # Satisfied - green
        3: 'F8D7DA',  # Not Satisfied - red
    }
    
    for idx, status in enumerate(satisfaction_stats['by_status']):
        row = status_table.rows[idx + 1]
        row.cells[0].text = status['status_en']
        row.cells[1].text = status['status_ar']
        row.cells[2].text = str(status['count'])
        row.cells[3].text = f"{round(status['count'] / total_satisfaction * 100, 1)}%"
        
        # Color code based on status
        color = status_colors.get(status['status_id'], 'FFFFFF')
        _set_cell_shading(row.cells[0], color)
        _set_cell_shading(row.cells[1], color)
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    # Add row for cases without satisfaction
    no_sat_row = status_table.rows[-1]
    no_sat_row.cells[0].text = "No Record"
    no_sat_row.cells[1].text = "بدون سجل"
    no_sat_row.cells[2].text = str(satisfaction_stats['cases_without_satisfaction'])
    no_sat_row.cells[3].text = f"{round(satisfaction_stats['cases_without_satisfaction'] / satisfaction_stats['total_cases'] * 100, 1)}%" if satisfaction_stats['total_cases'] > 0 else "0%"
    _set_cell_shading(no_sat_row.cells[0], 'FFFACD')
    _set_cell_shading(no_sat_row.cells[1], 'FFFACD')
    
    for cell in no_sat_row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    # Feedback Follow-up Stats
    followup_subheading = doc.add_heading('متابعة ملاحظات المرضى / Patient Feedback Follow-up', 2)
    for run in followup_subheading.runs:
        run.font.name = 'Calibri'
    
    followup_table = doc.add_table(rows=3, cols=2)
    followup_table.style = 'Table Grid'
    followup_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    _format_table_header(followup_table, ['المقياس / Metric', 'القيمة / Value'])
    
    followup_data = [
        ('حالات تحتاج متابعة / Feedback Needed', str(satisfaction_stats['feedback_stats']['feedback_needed'])),
        ('حالات تمت المتابعة / Feedback Given', str(satisfaction_stats['feedback_stats']['feedback_given'])),
    ]
    
    for idx, (metric, value) in enumerate(followup_data):
        row = followup_table.rows[idx + 1]
        row.cells[0].text = metric
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    # ========== FOOTER ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer.add_run('— نهاية التقرير / End of Report —')
    footer_run.font.size = Pt(10)
    footer_run.font.name = 'Calibri'
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
