"""
Aggregate Person Seasonal Report Word Generator

Generates comprehensive Word documents for ALL doctors or ALL workers
in the system for a given seasonal period.

Uses direct SQL queries for efficiency — fetches all doctors/workers
with their incident counts and severity breakdowns in single queries.
"""

from typing import List, Dict, Any
from datetime import date, datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.database import get_connection
from core.table_config import HR_EMPLOYEES_TABLE


def _fetch_doctors_with_incident_counts(
    season_start: date,
    season_end: date
) -> List[Dict[str, Any]]:
    """
    Fetch all doctors who have incidents in the given date range,
    with counts by severity. Single efficient SQL query.
    """
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        query = """
            SELECT
                icd.DoctorID,
                COALESCE(d.DoctorName, vw.Name, rd.DoctorName, CONCAT('Doctor ', icd.DoctorID)) as DoctorName,
                COALESCE(d.Specialty, vw.SpecialityName, rd.Specialty, '') as Specialty,
                COUNT(*) as TotalIncidents,
                SUM(CASE WHEN sev.SeverityName = 'High' THEN 1 ELSE 0 END) as HighSeverity,
                SUM(CASE WHEN sev.SeverityName = 'Medium' THEN 1 ELSE 0 END) as MediumSeverity,
                SUM(CASE WHEN sev.SeverityName = 'Low' THEN 1 ELSE 0 END) as LowSeverity,
                SUM(CASE WHEN crt.Code IN ('RED_FLAG', 'NEVER_EVENT') THEN 1 ELSE 0 END) as RedFlags
            FROM dbo.APP_IncidentCaseDoctor icd
            INNER JOIN dbo.APP_IncidentCase ic 
                ON ic.IncidentRequestCaseID = icd.IncidentRequestCaseID
            LEFT JOIN dbo.APP_LOOKUP_SEVERITY sev 
                ON ic.SeverityID = sev.SeverityID
            LEFT JOIN dbo.APP_LOOKUP_CLINICAL_RISK_TYPE crt 
                ON ic.ClinicalRiskTypeID = crt.ClinicalRiskTypeID
            LEFT JOIN dbo.APP_LOOKUP_DOCTOR d 
                ON icd.DoctorID = d.DoctorID
            LEFT JOIN dbo.APP_VIEWTABLE_VW_DOCTORS vw 
                ON icd.DoctorID = vw.DoctorID
            LEFT JOIN dbo.APP_RESERVE_DOCTOR rd 
                ON icd.DoctorID = rd.DoctorID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
            GROUP BY icd.DoctorID,
                     COALESCE(d.DoctorName, vw.Name, rd.DoctorName, CONCAT('Doctor ', icd.DoctorID)),
                     COALESCE(d.Specialty, vw.SpecialityName, rd.Specialty, '')
            ORDER BY COUNT(*) DESC
        """
        
        cursor.execute(query, (season_start.isoformat(), season_end.isoformat()))
        
        results = []
        for row in cursor.fetchall():
            results.append({
                "doctor_id": row.DoctorID,
                "name": row.DoctorName or f"Doctor {row.DoctorID}",
                "specialty": row.Specialty or "Unknown",
                "total_incidents": row.TotalIncidents,
                "high_severity": row.HighSeverity,
                "medium_severity": row.MediumSeverity,
                "low_severity": row.LowSeverity,
                "red_flags": row.RedFlags
            })
        
        return results
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


def _fetch_workers_with_incident_counts(
    season_start: date,
    season_end: date
) -> List[Dict[str, Any]]:
    """
    Fetch all workers (employees) who have incidents in the given date range,
    with counts by severity. Single efficient SQL query.
    """
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        query = f"""
            SELECT
                ice.EmployeeID,
                COALESCE(e.FullName, CONCAT('Employee ', ice.EmployeeID)) as FullName,
                COALESCE(e.JobTitle, '') as JobTitle,
                COUNT(*) as TotalIncidents,
                SUM(CASE WHEN sev.SeverityName = 'High' THEN 1 ELSE 0 END) as HighSeverity,
                SUM(CASE WHEN sev.SeverityName = 'Medium' THEN 1 ELSE 0 END) as MediumSeverity,
                SUM(CASE WHEN sev.SeverityName = 'Low' THEN 1 ELSE 0 END) as LowSeverity,
                SUM(CASE WHEN crt.Code IN ('RED_FLAG', 'NEVER_EVENT') THEN 1 ELSE 0 END) as RedFlags
            FROM dbo.APP_IncidentCaseEmployee ice
            INNER JOIN dbo.APP_IncidentCase ic 
                ON ic.IncidentRequestCaseID = ice.IncidentRequestCaseID
            LEFT JOIN dbo.APP_LOOKUP_SEVERITY sev 
                ON ic.SeverityID = sev.SeverityID
            LEFT JOIN dbo.APP_LOOKUP_CLINICAL_RISK_TYPE crt 
                ON ic.ClinicalRiskTypeID = crt.ClinicalRiskTypeID
            LEFT JOIN dbo.{HR_EMPLOYEES_TABLE} e 
                ON ice.EmployeeID = e.EmployeeID
            WHERE CONVERT(DATE, ic.FeedbackRecievedDate) >= ?
              AND CONVERT(DATE, ic.FeedbackRecievedDate) <= ?
            GROUP BY ice.EmployeeID,
                     COALESCE(e.FullName, CONCAT('Employee ', ice.EmployeeID)),
                     COALESCE(e.JobTitle, '')
            ORDER BY COUNT(*) DESC
        """
        
        cursor.execute(query, (season_start.isoformat(), season_end.isoformat()))
        
        results = []
        for row in cursor.fetchall():
            results.append({
                "employee_id": row.EmployeeID,
                "name": row.FullName or f"Employee {row.EmployeeID}",
                "job_title": row.JobTitle or "Unknown",
                "total_incidents": row.TotalIncidents,
                "high_severity": row.HighSeverity,
                "medium_severity": row.MediumSeverity,
                "low_severity": row.LowSeverity,
                "red_flags": row.RedFlags
            })
        
        return results
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


def generate_all_doctors_seasonal_word(
    season_start: date,
    season_end: date
) -> bytes:
    """
    Generate comprehensive Word document for ALL doctors with incidents.
    Shows comparison table of doctors sorted by incident count (most cases first).
    
    Uses a single SQL query for efficiency instead of per-doctor service calls.
    """
    if season_start >= season_end:
        raise ValueError("season_start must be before season_end")
    
    # Single efficient query: all doctors with incidents + severity breakdown
    doctor_reports = _fetch_doctors_with_incident_counts(season_start, season_end)
    
    if not doctor_reports:
        raise ValueError(
            f"No doctors with incidents found in the period {season_start} to {season_end}"
        )
    
    # Generate Word document
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    
    # ========== TITLE ==========
    title = doc.add_heading('تقرير مقارنة الأطباء الموسمي', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.name = 'Arial'
        run.font.bold = True
    
    subtitle = doc.add_heading('Seasonal Doctor Comparison Report', 1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.name = 'Arial'
    
    # Period info
    period = doc.add_paragraph()
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    period_run = period.add_run(
        f'الفترة: {season_start.strftime("%Y-%m-%d")} إلى {season_end.strftime("%Y-%m-%d")}'
    )
    period_run.font.size = Pt(12)
    period_run.font.name = 'Arial'
    
    # Total count
    count = doc.add_paragraph()
    count.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count_run = count.add_run(f'عدد الأطباء الذين لديهم حالات: {len(doctor_reports)}')
    count_run.font.size = Pt(11)
    count_run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # ========== SUMMARY STATISTICS ==========
    summary_heading = doc.add_heading('الإحصائيات الإجمالية', 2)
    for run in summary_heading.runs:
        run.font.name = 'Arial'
    
    total_all = sum(dr["total_incidents"] for dr in doctor_reports)
    total_high = sum(dr["high_severity"] for dr in doctor_reports)
    total_medium = sum(dr["medium_severity"] for dr in doctor_reports)
    total_low = sum(dr["low_severity"] for dr in doctor_reports)
    total_red = sum(dr["red_flags"] for dr in doctor_reports)
    avg_per = total_all / len(doctor_reports) if doctor_reports else 0
    
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Light List Accent 1'
    
    summary_data = [
        ('إجمالي الحالات / Total Incidents', str(total_all)),
        ('متوسط الحالات لكل طبيب / Avg per Doctor', f'{avg_per:.1f}'),
        ('حالات شدة عالية / High Severity', str(total_high)),
        ('حالات شدة متوسطة / Medium Severity', str(total_medium)),
        ('حالات شدة منخفضة / Low Severity', str(total_low)),
        ('علامات حمراء / Red Flags', str(total_red)),
    ]
    
    for idx, (label, value) in enumerate(summary_data):
        row = summary_table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # ========== COMPARISON TABLE ==========
    comparison_heading = doc.add_heading('جدول المقارنة - مرتب حسب عدد الحالات', 2)
    for run in comparison_heading.runs:
        run.font.name = 'Arial'
    
    note = doc.add_paragraph()
    note_run = note.add_run('(الأطباء ذوي العدد الأكبر من الحالات يظهرون أولاً)')
    note_run.font.size = Pt(9)
    note_run.font.name = 'Arial'
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(100, 100, 100)
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = [
        '#', 'اسم الطبيب\nDoctor Name', 'التخصص\nSpecialty',
        'إجمالي\nTotal', 'عالية\nHigh', 'متوسطة\nMedium',
        'منخفضة\nLow', 'علامات حمراء\nRed Flags'
    ]
    
    for idx, header_text in enumerate(headers):
        hdr_cells[idx].text = header_text
        for paragraph in hdr_cells[idx].paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Arial'
    
    for idx, doctor in enumerate(doctor_reports, 1):
        row = table.add_row()
        cells = row.cells
        
        row_data = [
            str(idx),
            doctor["name"],
            doctor["specialty"] or "غير محدد",
            str(doctor["total_incidents"]),
            str(doctor["high_severity"]),
            str(doctor["medium_severity"]),
            str(doctor["low_severity"]),
            str(doctor["red_flags"])
        ]
        
        for cell_idx, value in enumerate(row_data):
            cells[cell_idx].text = value
            for paragraph in cells[cell_idx].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # ========== FOOTER ==========
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer.add_run(f'تم الإنشاء في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    footer_run.font.size = Pt(9)
    footer_run.font.name = 'Arial'
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_all_workers_seasonal_word(
    season_start: date,
    season_end: date
) -> bytes:
    """
    Generate comprehensive Word document for ALL workers with incidents.
    Shows comparison table of workers sorted by incident count (most cases first).
    
    Uses a single SQL query for efficiency instead of per-worker service calls.
    """
    if season_start >= season_end:
        raise ValueError("season_start must be before season_end")
    
    # Single efficient query: all workers with incidents + severity breakdown
    worker_reports = _fetch_workers_with_incident_counts(season_start, season_end)
    
    if not worker_reports:
        raise ValueError(
            f"No workers with incidents found in the period {season_start} to {season_end}"
        )
    
    # Generate Word document
    doc = Document()
    
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    
    # ========== TITLE ==========
    title = doc.add_heading('تقرير مقارنة الموظفين الموسمي', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.name = 'Arial'
        run.font.bold = True
    
    subtitle = doc.add_heading('Seasonal Worker Comparison Report', 1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.name = 'Arial'
    
    # Period info
    period = doc.add_paragraph()
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    period_run = period.add_run(
        f'الفترة: {season_start.strftime("%Y-%m-%d")} إلى {season_end.strftime("%Y-%m-%d")}'
    )
    period_run.font.size = Pt(12)
    period_run.font.name = 'Arial'
    
    # Total count
    count = doc.add_paragraph()
    count.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    count_run = count.add_run(f'عدد الموظفين الذين لديهم حالات: {len(worker_reports)}')
    count_run.font.size = Pt(11)
    count_run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # ========== SUMMARY STATISTICS ==========
    summary_heading = doc.add_heading('الإحصائيات الإجمالية', 2)
    for run in summary_heading.runs:
        run.font.name = 'Arial'
    
    total_all = sum(wr["total_incidents"] for wr in worker_reports)
    total_high = sum(wr["high_severity"] for wr in worker_reports)
    total_medium = sum(wr["medium_severity"] for wr in worker_reports)
    total_low = sum(wr["low_severity"] for wr in worker_reports)
    total_red = sum(wr["red_flags"] for wr in worker_reports)
    avg_per = total_all / len(worker_reports) if worker_reports else 0
    
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Light List Accent 1'
    
    summary_data = [
        ('إجمالي الحالات / Total Incidents', str(total_all)),
        ('متوسط الحالات لكل موظف / Avg per Worker', f'{avg_per:.1f}'),
        ('حالات شدة عالية / High Severity', str(total_high)),
        ('حالات شدة متوسطة / Medium Severity', str(total_medium)),
        ('حالات شدة منخفضة / Low Severity', str(total_low)),
        ('علامات حمراء / Red Flags', str(total_red)),
    ]
    
    for idx, (label, value) in enumerate(summary_data):
        row = summary_table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # ========== COMPARISON TABLE ==========
    comparison_heading = doc.add_heading('جدول المقارنة - مرتب حسب عدد الحالات', 2)
    for run in comparison_heading.runs:
        run.font.name = 'Arial'
    
    note = doc.add_paragraph()
    note_run = note.add_run('(الموظفون ذوي العدد الأكبر من الحالات يظهرون أولاً)')
    note_run.font.size = Pt(9)
    note_run.font.name = 'Arial'
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(100, 100, 100)
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = [
        '#', 'اسم الموظف\nWorker Name', 'المسمى الوظيفي\nJob Title',
        'إجمالي\nTotal', 'عالية\nHigh', 'متوسطة\nMedium',
        'منخفضة\nLow', 'علامات حمراء\nRed Flags'
    ]
    
    for idx, header_text in enumerate(headers):
        hdr_cells[idx].text = header_text
        for paragraph in hdr_cells[idx].paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Arial'
    
    for idx, worker in enumerate(worker_reports, 1):
        row = table.add_row()
        cells = row.cells
        
        row_data = [
            str(idx),
            worker["name"],
            worker["job_title"] or "غير محدد",
            str(worker["total_incidents"]),
            str(worker["high_severity"]),
            str(worker["medium_severity"]),
            str(worker["low_severity"]),
            str(worker["red_flags"])
        ]
        
        for cell_idx, value in enumerate(row_data):
            cells[cell_idx].text = value
            for paragraph in cells[cell_idx].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # ========== FOOTER ==========
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer.add_run(f'تم الإنشاء في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    footer_run.font.size = Pt(9)
    footer_run.font.name = 'Arial'
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
