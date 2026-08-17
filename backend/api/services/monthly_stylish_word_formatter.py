"""
Monthly Report - Stylish Word Formatter (Formal Hospital Form Rewrite)

Philosophy
----------
This formatter reproduces the hospital's official paper form
("Stylish Reporting Template") as closely as practical: a restrained,
table-based, formal document — not a modern dashboard export.

Two independent page templates:
  - Complaint / Improvement Opportunity form (one full landscape page
    per complaint): scope strip, classification table, complaint-data
    block (narrative gets the majority of the width), immediate-action /
    actions-taken block, RCA instruction note.
  - Notice / Commendation form (one flowing table per unit): scope strip,
    then a 7-column table of that unit's notices.

Records are grouped by primary target org unit (see _group_by_unit), and
each unit's batch of complaint/notice pages is followed by its own
standalone signature page (_render_signature_page) — the paper form's
approval grid is physically handed to one person per batch, so it can
appear at most once per unit, not once per record. See Complaint Details /
Immediate Action / Actions Taken sizing via the 3 registered narrative
styles (_register_narrative_styles) for the one piece of user-adjustable
formatting in the document.

Pure renderer: zero DB queries, zero business-logic calculations.
All data arrives pre-computed in report_data.

Imports low-level OOXML helpers from workflow_activity_word_formatter —
no parallel utilities; other reports depend on that shared module, so
it is never modified from here. Local wrappers below cover two proven
gotchas: complex-script Arabic font (_ar_run) and tblGrid sync
(_set_row_col_width) — see their docstrings.
"""

import os
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .workflow_activity_word_formatter import (
    _ar_run as _ar_run_base,
    _new_para,
    _set_rtl_table,
    _apply_minimal_table_borders,
    _set_cell_shading,
    _cell_para,
    _set_row_col_width as _set_row_col_width_base,
    _set_para_shading,
    _set_para_bottom_border,
    _fmt_date,
    _mm_to_dxa,
)


def _ar_run(para, text: str, size: int = 10, bold: bool = False,
            italic: bool = False, color: str = None):
    """
    Local wrapper around the shared _ar_run helper that additionally sets
    the w:cs (complex-script) font. The shared helper only sets w:eastAsia,
    which leaves RTL Arabic glyph rendering to the theme's default complex-
    script font — can show as missing-glyph boxes in some renderers. Not
    touching the shared helper itself since other reports depend on it.
    """
    run = _ar_run_base(para, text, size=size, bold=bold, italic=italic, color=color)
    run._element.rPr.rFonts.set(qn('w:cs'), 'Traditional Arabic')
    return run


def _ltr_run(para, text: str, size: int = 10, bold: bool = False, color: str = None):
    """
    Like _ar_run, but forces this run to render left-to-right via an
    explicit w:rPr/w:rtl w:val="0" override, instead of relying on Unicode
    bidi-isolate characters (LRI U+2066 / PDI U+2069) to keep an LTR chunk
    (dates, report codes) from reordering inside surrounding RTL Arabic.
    Those isolate marks are a newer Unicode 6.3 feature that some Word
    font/rendering combinations don't treat as invisible — they render as
    visible placeholder boxes labeled "LRI"/"PDI" instead. The run-level
    w:rtl override is the standard OOXML mechanism for this and has no
    such fallback-glyph failure mode.
    """
    run = _ar_run(para, text, size=size, bold=bold, color=color)
    rtl_el = OxmlElement('w:rtl')
    rtl_el.set(qn('w:val'), '0')
    run._element.rPr.append(rtl_el)
    return run


# ---------------------------------------------------------------------------
# NARRATIVE PARAGRAPH STYLES  (user-adjustable text size in Word)
# ---------------------------------------------------------------------------
#
# Complaint Details / Immediate Action / Actions Taken body text used to
# carry hardcoded per-run font sizes, which Word's Styles pane can't touch —
# direct run formatting always wins over style formatting. These three real
# Word paragraph styles let a user select the text, bump the size once in
# Word, and "Update Style to Match Selection" to resize every occurrence of
# that block type across the whole document in one step.

STYLE_COMPLAINT_NARRATIVE = 'نص محتوى الشكوى'
STYLE_IMMEDIATE_ACTION    = 'نص الإجراءات الفورية'
STYLE_ACTIONS_TAKEN       = 'نص الإجراءات المتخذة'
STYLE_COMPLAINT_DATA      = 'نص بيانات الشكوى'  # قسم الصادر / المصدر / P.Name values

# Bumped up from the previous hardcoded run sizes (11 / 10 / 10) per the
# "make it bigger" request — these are now just the starting point; the
# whole reason they're styles is so the user can go further from here.
_NARRATIVE_STYLE_SIZES = {
    STYLE_COMPLAINT_NARRATIVE: 13,
    STYLE_IMMEDIATE_ACTION: 12,
    STYLE_ACTIONS_TAKEN: 12,
    STYLE_COMPLAINT_DATA: 12,
}


def _register_narrative_styles(doc: Document) -> None:
    """
    Registers the narrative + complaint-data styles on doc.styles. Must run
    once per document, before any paragraph tries to use them.
    """
    for style_name, size in _NARRATIVE_STYLE_SIZES.items():
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles['Normal']
        style.font.name = 'Traditional Arabic'
        style.font.size = Pt(size)
        # Complex-script (Arabic) font override, same w:cs trick _ar_run
        # applies per-run — python-docx's font.name only sets w:ascii/w:hAnsi.
        rPr = style.element.find(qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts'))
        rFonts.set(qn('w:cs'), 'Traditional Arabic')
        if style_name == STYLE_COMPLAINT_DATA:
            # Matches the bold=True the direct-formatting version of these
            # 3 fields always used (_labeled_cell's value_bold default).
            style.font.bold = True
        # Without qFormat, Word's on-ribbon Quick Style Gallery doesn't show
        # a newly added custom style at all — it's only reachable via the
        # full Styles pane. quick_style (w:qFormat) + a low priority number
        # is what actually makes "one click in the ribbon" true, matching
        # the whole point of using a style instead of direct formatting.
        style.quick_style = True
        style.priority = 2


def _styled_ar_run(para, text: str, color: str = None):
    """
    Adds a run with NO direct font name/size — it inherits both entirely
    from the paragraph's assigned narrative style. Setting size here would
    silently defeat the point of using a style at all, since direct
    run-level formatting overrides style formatting in Word.
    """
    run = para.add_run(text or '')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _set_row_col_width(row, col_idx: int, mm: float):
    """
    Local wrapper around the shared _set_row_col_width helper that ALSO
    updates the table's tblGrid column width via the column object.

    The shared helper only sets w:tcW on the specific cell — it never
    touches w:tblGrid, which is what Word's tblLayout="fixed" rendering
    actually uses to determine column boundaries. Every prior "width fix"
    in this file changed cell tcW values that Word was silently ignoring,
    because the table's real grid stayed frozen at its auto-generated
    default. Not touching the shared helper since other reports use it.
    """
    _set_row_col_width_base(row, col_idx, mm)
    row.table.columns[col_idx].width = Mm(mm)


# ---------------------------------------------------------------------------
# CONSTANTS
# ---------------------------------------------------------------------------

NAVY        = '1C3A7A'   # used sparingly — headings/accents only, not fills
NAVY_LIGHT  = 'EBF3FB'   # scope-strip shading
HEADER_CYAN = 'D6EAF8'   # classification / notice table header shading
ACTION_CREAM = 'FFF9DB'  # immediate-action / actions-taken header shading
WHITE       = 'FFFFFF'
GREY_LINE   = 'D0D7E4'
GREY_TEXT   = '6B7A99'
DARK_TEXT   = '1A1A2E'

BORDER_OUTER = '2B2B2B'  # near-black, thin — formal paper-form look
BORDER_INNER = '999999'

# Usable width: A4 landscape 297mm - 12mm left - 12mm right = 273mm.
# Every table below targets <=270mm (3mm safety margin).
USABLE_WIDTH_MM = 273.0
MAX_TABLE_WIDTH_MM = 270.0

# Last-resort truncation ceilings (character counts) — reactivated only as
# a safety net for pathological outliers, not a normal-case page budget.
# Normal complaints (~1,200 chars) fit page 1 comfortably within these caps.
NARRATIVE_MAX_CHARS = 6000
ACTION_MAX_CHARS    = 3000
NOTICE_MAX_CHARS    = 3000


# ---------------------------------------------------------------------------
# LOW-LEVEL DOCX UTILITIES (supplement to imported helpers)
# ---------------------------------------------------------------------------

def _truncate_for_fit(text: str, max_chars: int,
                       marker: str = '… (النص الكامل في النظام)') -> str:
    """
    Last-resort safety net: cuts text at the last word boundary before
    max_chars and appends a marker pointing back to the system for the
    full record. Normal-length content never reaches this ceiling —
    truncation should only fire for pathological outliers (e.g. pasted
    logs), never for an ordinary long complaint.
    """
    text = (text or '').strip()
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars]
    last_space = cut.rfind(' ')
    if last_space > max_chars * 0.7:
        cut = cut[:last_space]
    return cut.rstrip(' ,.;:،؛') + ' ' + marker


def _cell_v_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _labeled_cell(cell, label: str, value: str,
                  label_color: str = GREY_TEXT,
                  value_color: str = DARK_TEXT,
                  value_size: int = 10,
                  label_size: int = 7,
                  value_bold: bool = True,
                  bg: str = None,
                  align: str = 'right',
                  doc: Document = None,
                  style_name: str = None):
    """
    Two-paragraph cell: small bilingual label on top (always centered —
    label text is "Arabic / English"), larger value below (right-aligned —
    pure Arabic/data content, ragged-left like classic RTL print layouts).

    When doc + style_name are both given, the value paragraph uses that
    real Word style (no direct size) instead of value_size/value_bold, so
    it becomes one-click resizable in Word — same mechanism as the
    narrative styles. Existing callers that don't pass these two keep the
    prior direct-formatting behavior unchanged.
    """
    if bg:
        _set_cell_shading(cell, bg)
    _cell_v_center(cell)
    cell.text = ''

    lp = cell.paragraphs[0]
    lp.clear()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(3)
    lp.paragraph_format.space_after  = Pt(1)
    lp._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    _ar_run(lp, label, size=label_size, color=label_color)

    vp = cell.add_paragraph()
    if doc is not None and style_name:
        vp.style = doc.styles[style_name]
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == 'right' else WD_ALIGN_PARAGRAPH.CENTER
    vp.paragraph_format.space_before = Pt(2)
    vp.paragraph_format.space_after  = Pt(3)
    vp._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    if doc is not None and style_name:
        _styled_ar_run(vp, value or '—', color=value_color)
    else:
        _ar_run(vp, value or '—', size=value_size, bold=value_bold, color=value_color)


def _cell_para0(cell, align: str = 'center'):
    """
    Local wrapper around the shared _cell_para that also zeroes paragraph
    spacing. The document's default template (w:docDefaults/w:pPrDefault)
    sets 10pt space-after + 1.15 line spacing on every paragraph that
    doesn't explicitly override it — for compact single-line table cells
    that silently added ~3.5mm of trailing space per cell, on top of
    whatever AT_LEAST row-height floor was set. This is why repeatedly
    shrinking that floor had no visible effect: the paragraph's own default
    spacing was the actual floor, never touched. Not modifying the shared
    _cell_para helper itself since other reports use it.
    """
    p = _cell_para(cell, align)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _bold_underline_value_cell(cell, text: str):
    """Classification table's data-row cells: bold + underlined value, centered."""
    p = _cell_para0(cell, 'center')
    run = _ar_run(p, text or '—', size=9, bold=True, color=DARK_TEXT)
    run.font.underline = True


def _rotate_cell(cell):
    """Sets bottom-to-top text direction (w:textDirection btLr) on a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    direction = OxmlElement('w:textDirection')
    direction.set(qn('w:val'), 'btLr')
    tcPr.append(direction)


def _vertical_label_cell(cell, text_ar: str):
    """Rotated (bottom-to-top) side label, e.g. 'بيانات الشكوى' / 'المتابعة'."""
    _set_cell_shading(cell, NAVY)
    _rotate_cell(cell)
    p = _cell_para0(cell, 'center')
    _ar_run(p, text_ar, size=9, bold=True, color=WHITE)


def _vertical_header_cell(cell, label: str):
    """Rotated classification-table header cell (short label, cyan bg)."""
    _set_cell_shading(cell, HEADER_CYAN)
    _rotate_cell(cell)
    p = _cell_para0(cell, 'center')
    _ar_run(p, label, size=7, bold=True, color=NAVY)


def _vertical_bold_underline_value_cell(cell, text: str):
    """Rotated classification-table data cell: bold + underlined value."""
    _rotate_cell(cell)
    p = _cell_para0(cell, 'center')
    run = _ar_run(p, text or '—', size=8, bold=True, color=DARK_TEXT)
    run.font.underline = True


def _spacer_para(doc: Document, pt: float = 2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)


def _gap(doc: Document, mm: float):
    """
    Precise vertical gap between blocks, in mm. Uses a near-zero-size run
    so the empty paragraph's own line-height doesn't add unaccounted
    height — the gap is controlled almost entirely by space_after.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(mm * 2.6)
    r = p.add_run('')
    r.font.size = Pt(1)


def _page_break(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _set_row_cant_split(row):
    """Keeps a row from splitting across a page break — whole row jumps instead."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def _set_table_header_repeat(row):
    """Marks a row as a repeating header on multi-page tables."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))


def _load_report_config() -> Dict[str, Any]:
    try:
        from ..db_layer.report_config_db import get_report_config
        return get_report_config() or {}
    except Exception:
        return {}


# ---------------------------------------------------------------------------
# TARGET DEPARTMENT HELPERS
# ---------------------------------------------------------------------------

def _primary_target(record: Dict) -> Dict:
    tds = record.get('target_departments') or []
    primary = next((d for d in tds if d.get('is_primary')), tds[0] if tds else {})
    return primary


def _target_names(record: Dict):
    p = _primary_target(record)
    return (
        p.get('section_name') or '—',
        p.get('department_name') or '—',
        p.get('administration_name') or '—',
    )


def _target_display(record: Dict) -> str:
    """Most specific named target unit as a single string."""
    p = _primary_target(record)
    return (p.get('section_name') or p.get('department_name') or
            p.get('administration_name') or '—')


def _scope_labels(report_entity_name: Optional[str], report_entity_type: Optional[str]):
    """Maps the report-level scope (single name+type) onto the 3-way strip."""
    name = report_entity_name or '—'
    t = (report_entity_type or '').lower()
    if 'administration' in t:
        return name, '—', '—'
    if 'department' in t:
        return '—', name, '—'
    if 'section' in t:
        return '—', '—', name
    return '—', '—', '—'


def _unit_key(record: Dict):
    """
    Identity key for a record's primary target unit — prefers IDs over
    names (two different sections can share a display name) so grouping
    for the per-batch signature pages below is exact, not string-based.
    """
    p = _primary_target(record)
    return (
        p.get('section_id') or p.get('section_name'),
        p.get('department_id') or p.get('department_name'),
        p.get('administration_id') or p.get('administration_name'),
    )


def _group_by_unit(records: List[Dict]):
    """
    Stable group-by primary target unit: preserves each group's first-seen
    order and keeps every record belonging to that unit contiguous, even if
    the incoming list interleaves units — needed so a single signature page
    can immediately follow each unit's complete, uninterrupted batch of
    records. Returns [(unit_display_label, [records...]), ...].
    """
    groups: Dict[Any, List[Dict]] = {}
    for record in records:
        groups.setdefault(_unit_key(record), []).append(record)
    return [(_target_display(recs[0]), recs) for recs in groups.values()]


# ---------------------------------------------------------------------------
# SCOPE STRIP  (الإدارة / الدائرة / الوحدة الإدارية-القسم / الشهر)
# ---------------------------------------------------------------------------

_STRIP_WIDTHS = [68, 68, 68, 66]  # sum = 270mm


def _four_cell_strip(doc: Document, admin_name: str, dept_name: str,
                     sec_name: str, period_label: str):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=6, inner=BORDER_INNER, inner_sz=3)
    _set_rtl_table(tbl)

    row = tbl.rows[0]
    boxes = [
        ('الإدارة / Administration', admin_name),
        ('الدائرة / Circle', dept_name),
        ('الوحدة الإدارية/القسم / Section', sec_name),
        ('الشهر / Month', period_label),
    ]
    for ci, (label, value) in enumerate(boxes):
        _labeled_cell(row.cells[ci], label, value, bg=NAVY_LIGHT,
                      value_size=9.5, label_size=7.5, align='center')
        _set_row_col_width(row, ci, _STRIP_WIDTHS[ci])

    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = _mm_to_dxa(13)
    _set_row_cant_split(row)
    return tbl


# ---------------------------------------------------------------------------
# CLASSIFICATION TABLE  (11 columns, label-only header + bold/underline data)
# ---------------------------------------------------------------------------

def _case_number(complaint: Dict) -> str:
    inc_id = complaint.get('incident_id')
    if inc_id is not None:
        try:
            return f"INC-{int(inc_id):06d}"
        except (TypeError, ValueError):
            return str(inc_id)
    return str(complaint.get('id', '—'))


# RTL reading order (index 0 = rightmost column once bidiVisual is applied).
# Three distinct dates are shown side by side:
#   - received_date (تاريخ تلقي الملاحظة): the date the feedback reached
#     Patient Services. This is the field that actually determines which
#     month a case is reported under — see the date_filter WHERE-clause in
#     reports_db.get_filtered_complaints, which filters exclusively on
#     FeedbackRecievedDate. Never IncidentDate.
#   - incident_date (تاريخ وقوع الحادثة): the date the underlying event
#     occurred. Display-only — plays no role in month/period classification.
#   - publication_date (تاريخ النشر): the date the case first entered
#     workflow (earliest APP_AdministrativeSubcase.CreatedAt) — same
#     definition already used by Table View and Inbox. Display-only.
#
# The 3 dates + case number are rendered ROTATED (header and value both) so
# the Classification/Classification columns can be doubled and Severity/Harm
# widened without exceeding the page budget — see _ROTATED_CLASS_COLS below.
# Rotating the value (not just the header) needs real vertical room for a
# ~10-char string, so the data row is taller here (22mm vs the previous
# 11mm) than a purely horizontal table would need.
# Problem Domain widened 30% (20mm -> 26mm, prior round). Sub-Category +15%
# (20mm -> 23mm), Complaint Field Type +15% (16mm -> 18mm), offset by
# trimming Category/Stage/Status (prior round). This round: Severity and
# Harm each cut 15% (19->16.15mm, 21->17.85mm — 6mm saved combined) and that
# 6mm moved entirely onto Stage (16mm -> 22mm).
# Sum of widths = 269mm (<= 270mm target ceiling).
_CLASS_COLS = [
    ('الاستلام', 8, lambda c: _fmt_date(c.get('received_date'))),
    ('الحادثة', 8, lambda c: _fmt_date(c.get('incident_date'))),
    ('النشر', 8, lambda c: _fmt_date(c.get('publication_date'))),
    ('الرقم', 8, _case_number),
    ('Problem Domain\nالمجال', 26, lambda c: c.get('domain_name') or '—'),
    ('Problem Category\nفئة المشكلة', 20, lambda c: c.get('category_name') or '—'),
    ('Sub-Category\nالفئة الفرعية', 23, lambda c: c.get('subcategory_name') or '—'),
    ('Classification (Arb.)', 40, lambda c: c.get('classification_name') or '—'),
    ('Classification (Eng.)', 40, lambda c: c.get('classification_name_en') or '—'),
    ('Severity\nالخطورة', 16.15, lambda c: c.get('severity_name') or '—'),
    ('Stage\nالمرحلة', 22, lambda c: c.get('stage_name') or '—'),
    ('Harm\nالضرر', 17.85, lambda c: c.get('harm_level') or '—'),
    ('Status\nالحالة', 14, lambda c: c.get('status_name') or '—'),
    ('Complaint Field Type\nنوع السجل', 18, lambda c: c.get('clinical_risk_type_name') or 'Ordinary'),
]

# Indices of the 4 rotated columns (received/incident/publication date, case number).
_ROTATED_CLASS_COLS = {0, 1, 2, 3}


def _classification_table(doc: Document, complaint: Dict):
    n = len(_CLASS_COLS)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=6, inner=BORDER_INNER, inner_sz=3)
    _set_rtl_table(tbl)

    hdr = tbl.rows[0]
    for ci, (label, w, _getter) in enumerate(_CLASS_COLS):
        c = hdr.cells[ci]
        if ci in _ROTATED_CLASS_COLS:
            _vertical_header_cell(c, label)
        else:
            _set_cell_shading(c, HEADER_CYAN)
            cp = _cell_para0(c, 'center')
            _ar_run(cp, label, size=6.5, bold=True, color=NAVY)
        _set_row_col_width(hdr, ci, w)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr.height = _mm_to_dxa(12)
    _set_row_cant_split(hdr)

    data = tbl.rows[1]
    for ci, (_label, w, getter) in enumerate(_CLASS_COLS):
        if ci in _ROTATED_CLASS_COLS:
            _vertical_bold_underline_value_cell(data.cells[ci], getter(complaint))
        else:
            _bold_underline_value_cell(data.cells[ci], getter(complaint))
        _set_row_col_width(data, ci, w)
    # Floor bumped 22mm -> 27mm: the rotated case-number text (~11 unbroken
    # characters, e.g. "INC-000187") only looked clean when something else in
    # the row (typically a long Classification (Eng.) wrap) happened to push
    # the row taller than 22mm — at the bare 22mm floor it looked cramped.
    # Raising the floor itself makes that a fixed, guaranteed minimum instead
    # of something the number's readability accidentally depended on. This
    # does cost ~5mm of the vertical space reclaimed earlier this session.
    data.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    data.height = _mm_to_dxa(27)
    _set_row_cant_split(data)
    return tbl


# ---------------------------------------------------------------------------
# COMPLAINT DATA BLOCK  (بيانات الشكوى)
# ---------------------------------------------------------------------------

# Label cell first (index 0) so it renders at the RIGHT edge under RTL
# bidiVisual — Arabic reads right-to-left, so the "بيانات الشكوى" side label
# belongs where reading starts, not at the left.
# The 3 value boxes (قسم الصادر/المصدر/P.Name) are equal width on purpose —
# previously 22/18/26mm, an inconsistent spread that left المصدر cramped;
# now uniform so the row reads as one consistent group.
_COMPLAINT_DATA_WIDTHS = [8, 26, 26, 26, 184]  # sum = 270mm; narrative = 184mm (68.1%)


def _complaint_data_block(doc: Document, complaint: Dict):
    tbl = doc.add_table(rows=1, cols=5)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=6, inner=BORDER_INNER, inner_sz=3)
    _set_rtl_table(tbl)

    row = tbl.rows[0]
    label_cell, issuing_cell, source_cell, patient_cell, narrative_cell = row.cells

    _labeled_cell(issuing_cell, 'قسم الصادر', complaint.get('section_name') or '—',
                  label_size=7.5, align='center', doc=doc, style_name=STYLE_COMPLAINT_DATA)
    _labeled_cell(source_cell, 'المصدر', complaint.get('source_name') or '—',
                  label_size=7.5, align='center', doc=doc, style_name=STYLE_COMPLAINT_DATA)
    _labeled_cell(patient_cell, 'P.Name', complaint.get('patient_name') or '—',
                  label_size=7.5, align='center', doc=doc, style_name=STYLE_COMPLAINT_DATA)

    _cell_v_center(narrative_cell)
    narrative_cell.text = ''
    nh = narrative_cell.paragraphs[0]
    nh.clear()
    nh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nh.paragraph_format.space_before = Pt(2)
    nh.paragraph_format.space_after  = Pt(2)
    nh._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    _ar_run(nh, 'محتوى الشكوى  /  Complaint Details', size=8, bold=True, color=NAVY)

    text = _truncate_for_fit(complaint.get('complaint_text') or '', NARRATIVE_MAX_CHARS)
    nb = narrative_cell.add_paragraph()
    nb.style = doc.styles[STYLE_COMPLAINT_NARRATIVE]
    nb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nb.paragraph_format.space_before = Pt(3)
    nb.paragraph_format.space_after  = Pt(3)
    nb.paragraph_format.right_indent = Mm(3)
    nb.paragraph_format.left_indent  = Mm(3)
    nb._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    pPr = nb._p.get_or_add_pPr()
    sp_el = OxmlElement('w:spacing')
    sp_el.set(qn('w:line'), '265')
    sp_el.set(qn('w:lineRule'), 'auto')
    pPr.append(sp_el)
    _styled_ar_run(nb, text or 'لا يوجد نص للشكوى', color=DARK_TEXT)

    _vertical_label_cell(label_cell, 'بيانات الشكوى')

    for ci, w in enumerate(_COMPLAINT_DATA_WIDTHS):
        _set_row_col_width(row, ci, w)

    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = _mm_to_dxa(20)
    # Deliberately NOT cantSplit — a long narrative should flow onto a
    # continuation page rather than jump whole and leave a blank gap.
    return tbl


# ---------------------------------------------------------------------------
# ACTION BLOCK  (Immediate Action | Actions Taken) — المتابعة
# ---------------------------------------------------------------------------

# Label cell first (index 0) so it renders at the RIGHT edge under RTL
# bidiVisual, same rationale as _COMPLAINT_DATA_WIDTHS above.
_ACTION_WIDTHS = [8, 131, 131]  # sum = 270mm


def _action_block(doc: Document, complaint: Dict):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=6, inner=BORDER_INNER, inner_sz=3)
    _set_rtl_table(tbl)

    row = tbl.rows[0]
    label_cell, immediate_cell, taken_cell = row.cells

    def _fill(cell, header_ar, header_en, text, style_name):
        _set_cell_shading(cell, WHITE)
        cell.text = ''
        hp = cell.paragraphs[0]
        hp.clear()
        _set_para_shading(hp, ACTION_CREAM)
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(2)
        hp.paragraph_format.space_after  = Pt(2)
        hp._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
        _ar_run(hp, f'{header_ar}  /  {header_en}', size=8, bold=True, color=NAVY)

        body_text = _truncate_for_fit(text or '', ACTION_MAX_CHARS)
        bp = cell.add_paragraph()
        bp.style = doc.styles[style_name]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(3)
        bp.paragraph_format.space_after  = Pt(3)
        bp.paragraph_format.right_indent = Mm(2)
        bp.paragraph_format.left_indent  = Mm(2)
        bp._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
        _styled_ar_run(bp, body_text or '—', color=DARK_TEXT)

    _fill(immediate_cell, 'الإجراءات الفورية', 'Immediate Action',
          complaint.get('immediate_action'), STYLE_IMMEDIATE_ACTION)
    _fill(taken_cell, 'الإجراءات المتخذة', 'Actions Taken',
          complaint.get('taken_action'), STYLE_ACTIONS_TAKEN)
    _vertical_label_cell(label_cell, 'المتابعة')

    for ci, w in enumerate(_ACTION_WIDTHS):
        _set_row_col_width(row, ci, w)

    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = _mm_to_dxa(20)
    _set_row_cant_split(row)
    return tbl


# ---------------------------------------------------------------------------
# APPROVAL / SIGNATURE GRID
# ---------------------------------------------------------------------------

# English removed — this was the actual cause of the table staying tall
# despite repeatedly shrinking the AT_LEAST floor: AT_LEAST is a minimum,
# it can never compress a row below what its content physically needs, and
# a 2-line "Arabic\nEnglish" cell needs roughly double a 1-line cell's
# height no matter how small the floor is set. Single-line Arabic-only
# labels are what actually let the row get shorter.
_APPROVAL_ROLES = ['مسؤول العملية', 'رئيس الدائرة', 'مدير الإدارة', 'خاص خدمات المرضى']
_APPROVAL_COL_WIDTHS = [30, 60, 60, 60, 60]  # sum = 270mm


def _tighten_cell_margins(cell, top_mm=0.3, bottom_mm=0.3, left_mm=1.0, right_mm=1.0):
    """
    Cuts a cell's internal top/bottom padding to near-zero. The second,
    quieter contributor to this table's height: python-docx's default table
    style carries its own cell margins that add to whatever the row's
    AT_LEAST floor specifies, regardless of that floor's value.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, mm in (('top', top_mm), ('bottom', bottom_mm), ('left', left_mm), ('right', right_mm)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(_mm_to_dxa(mm)))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def _set_para_mark_size(p, pt: float):
    """
    Sets the font size an EMPTY paragraph uses for its own line-height
    calculation (the paragraph-mark run properties, w:pPr/w:rPr). A cell
    with no text/run in it — the blank signature cells here — silently
    falls back to the Normal style's font size (10pt in this document) for
    that calculation, which is LARGER than every other font size used in
    this table (7-8pt). That makes an invisible, empty cell the tallest
    thing in its row and the actual row-height driver, with nothing in the
    visible content explaining why the row wouldn't shrink further.
    """
    pPr = p._p.get_or_add_pPr()
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(pt * 2)))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(pt * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    pPr.append(rPr)


def _approval_grid(doc: Document):
    """
    Row-height floors here are deliberately generous (unlike most tables in
    this file): _approval_grid is only ever called from
    _render_signature_page, a dedicated, otherwise-empty page — not squeezed
    alongside the rest of a complaint's content anymore, so there's no
    reason to keep it as compact as the old shared-page version needed.
    """
    tbl = doc.add_table(rows=4, cols=5)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=6, inner=BORDER_INNER, inner_sz=3)
    _set_rtl_table(tbl)

    hdr = tbl.rows[0]
    hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    # NOTE: row.height expects a Length (EMU) value, NOT the raw twips int
    # _mm_to_dxa() returns — python-docx silently treats a bare int as
    # already-EMU and divides it by 635 when serializing to w:trHeight/@w:val,
    # so _mm_to_dxa(12) here would have written an effectively-zero explicit
    # height (AT_LEAST + near-zero floor = pure content-driven autosize,
    # not the deliberate floor this file's comments describe). Mm() is the
    # correct Length type — same pattern _set_row_col_width already uses
    # correctly for column widths.
    hdr.height = Mm(12)
    for ci, label in enumerate([''] + _APPROVAL_ROLES):
        c = hdr.cells[ci]
        _set_cell_shading(c, 'F2F2F2')
        _tighten_cell_margins(c, top_mm=3, bottom_mm=3)
        cp = _cell_para0(c, 'center')
        _ar_run(cp, label, size=11, bold=True, color=DARK_TEXT)
        _set_row_col_width(hdr, ci, _APPROVAL_COL_WIDTHS[ci])
    _set_row_cant_split(hdr)

    # Field labels: Arabic only (English dropped), centered instead of
    # right-aligned now that there's no bilingual "/" split to anchor.
    # التوقيع (signature) row is tallest of the three — needs real room for
    # an actual pen signature, not just a label-height line.
    field_rows = [('الاسم', 20), ('التاريخ', 20), ('التوقيع', 45)]
    for ri, (field_lbl, h) in enumerate(field_rows):
        row = tbl.rows[ri + 1]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Mm(h)  # see Mm(12) note above — same EMU-vs-twips fix
        c0 = row.cells[0]
        _set_cell_shading(c0, 'F8F8F8')
        _tighten_cell_margins(c0, top_mm=3, bottom_mm=3)
        p0 = _cell_para0(c0, 'center')
        _ar_run(p0, field_lbl, size=11, bold=True, color=NAVY)
        _set_row_col_width(row, 0, _APPROVAL_COL_WIDTHS[0])
        for ci in range(1, 5):
            _set_cell_shading(row.cells[ci], WHITE)
            _tighten_cell_margins(row.cells[ci], top_mm=3, bottom_mm=3)
            # Blank signature cell: no visible text, but its empty
            # paragraph mark still needs an explicit font size or it
            # silently falls back to Normal's 10pt for line-height purposes.
            # See _set_para_mark_size's docstring.
            bp = _cell_para0(row.cells[ci], 'center')
            _set_para_mark_size(bp, 11)
            _set_row_col_width(row, ci, _APPROVAL_COL_WIDTHS[ci])
        _set_row_cant_split(row)
    return tbl


def _render_signature_page(doc: Document, unit_label: str, count: int, noun_label: str):
    """
    Standalone page carrying just one batch's approval/signature grid.
    Physically, complaints/notices for one org unit get sent by hand to one
    person for one signature — so the grid can't repeat per record (that
    forces multiple signatures for a single batch) and can't be shared
    across units in the same file either. Captioned with the unit name and
    record count since this page is meant to be separated from the rest and
    handed off on its own — without the caption there'd be no way to tell
    whose batch it belongs to once detached.
    """
    cap = _new_para(doc, align='center', space_before=6, space_after=2)
    _ar_run(cap, 'توقيع الدفعة  /  Batch Signature', size=12, bold=True, color=NAVY)

    sub = _new_para(doc, align='center', space_before=0, space_after=10)
    _ar_run(sub, f'{unit_label}   —   عدد {noun_label}: {count}', size=9.5, color=GREY_TEXT)

    _approval_grid(doc)


# ---------------------------------------------------------------------------
# RCA / QUARTERLY INSTRUCTION NOTE  (complaint pages only)
# ---------------------------------------------------------------------------

_RCA_NOTE_LINE1 = (
    'ملاحظة: 1- التقرير الشهري: الشكاوى المصنفة High يلزم ملء استمارة تحليل السبب الجذري '
    'RCA (Root Cause Analysis) إذا لم يتم ملؤها خلال المتابعة، أما المصنفة Medium أو Low '
    'فملؤها يكون تبعاً للحاجة بناءً على قرار مسؤول العملية.'
)
_RCA_NOTE_LINE2 = '2- التقرير الفصلي: ترفع استمارة تحسين تلقائياً تبعاً لـ Target الشكاوى.'


def _instruction_note(doc: Document):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=4, inner=BORDER_INNER, inner_sz=4)

    row = tbl.rows[0]
    _set_row_col_width(row, 0, sum(_APPROVAL_COL_WIDTHS))
    cell = row.cells[0]
    _set_cell_shading(cell, 'FFFDE7')
    cell.text = ''

    p1 = cell.paragraphs[0]
    p1.clear()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after  = Pt(2)
    p1.paragraph_format.right_indent = Mm(2)
    p1.paragraph_format.left_indent  = Mm(2)
    p1._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    _ar_run(p1, _RCA_NOTE_LINE1, size=7.5, color='5D4037')

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.right_indent = Mm(2)
    p2.paragraph_format.left_indent  = Mm(2)
    p2._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
    _ar_run(p2, _RCA_NOTE_LINE2, size=7.5, color='5D4037')

    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = _mm_to_dxa(14)
    _set_row_cant_split(row)
    return tbl


# ---------------------------------------------------------------------------
# FULL COMPLAINT PAGE
# ---------------------------------------------------------------------------

def _render_complaint_page(doc: Document, complaint: Dict, index: int, total: int, period: Dict):
    # Inter-block gaps: 2mm -> 1mm -> 0.6mm -> 0.3mm -> 0.15mm (this round).
    # Change log only, not a reason to stop reducing further.
    sec_name, dept_name, admin_name = _target_names(complaint)
    period_label = period.get('label_ar') or period.get('label') or '—'
    _four_cell_strip(doc, admin_name, dept_name, sec_name, period_label)
    _gap(doc, 0.15)

    _classification_table(doc, complaint)
    _gap(doc, 0.15)

    _complaint_data_block(doc, complaint)
    _gap(doc, 0.15)

    _action_block(doc, complaint)
    _gap(doc, 0.15)

    _instruction_note(doc)

    pg_para = _new_para(doc, align='center', space_before=2, space_after=0)
    _ar_run(pg_para, f'شكوى {index} من {total}  •  {_fmt_date(complaint.get("received_date"))}',
            size=7, color=GREY_TEXT)

    # No page break here — the caller inserts one only BETWEEN complaints.
    # Adding one unconditionally after every complaint (including the last)
    # produced a trailing blank page: either a dangling empty page at the
    # very end of the document, or a blank page sandwiched between the last
    # complaint and the next section (notices/appendix), since add_section
    # (WD_SECTION.NEW_PAGE) already forces its own page break.


# ---------------------------------------------------------------------------
# NOTICES SECTION  (flowing table, multiple per page)
# ---------------------------------------------------------------------------

# RTL reading order (index 0 = rightmost). Sum of widths = 270mm.
_NOTICE_COLS = [
    ('تاريخ تلقي الملاحظة', 20),
    ('الرقم', 20),
    ('قسم الصادر', 26),
    ('المصدر', 20),
    ('P.Name', 26),
    ('تفصيل الملاحظة', 128),
    ('الوحدة المنوّه بها', 30),
]


def _notice_case_number(notice: Dict) -> str:
    inc_id = notice.get('incident_id')
    if inc_id is not None:
        try:
            return f"RTG-{int(inc_id):06d}"
        except (TypeError, ValueError):
            return str(inc_id)
    return str(notice.get('id', '—'))


def _render_notices_table(doc: Document, notices: List[Dict],
                           admin_name: str, dept_name: str, sec_name: str,
                           period_label: str):
    """
    Scope strip + notices table for ONE unit's batch (or, when notices is
    empty, a single empty-state table for the whole report). No signature
    grid here — see _render_notices_section, which appends one standalone
    _render_signature_page per batch instead.
    """
    _four_cell_strip(doc, admin_name, dept_name, sec_name, period_label)
    _gap(doc, 3)

    n = len(_NOTICE_COLS)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_minimal_table_borders(tbl, outer=BORDER_OUTER, outer_sz=6, inner=BORDER_INNER, inner_sz=3)
    _set_rtl_table(tbl)

    hdr = tbl.rows[0]
    for ci, (label, w) in enumerate(_NOTICE_COLS):
        c = hdr.cells[ci]
        _set_cell_shading(c, HEADER_CYAN)
        cp = _cell_para0(c, 'center')
        _ar_run(cp, label, size=8, bold=True, color=NAVY)
        _set_row_col_width(hdr, ci, w)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr.height = _mm_to_dxa(8)
    _set_table_header_repeat(hdr)
    _set_row_cant_split(hdr)

    for notice in notices:
        row = tbl.add_row()
        vals = [
            _fmt_date(notice.get('received_date')),
            _notice_case_number(notice),
            notice.get('section_name') or '—',
            notice.get('source_name') or '—',
            notice.get('patient_name') or '—',
            _truncate_for_fit(notice.get('notice_text') or '', NOTICE_MAX_CHARS) or '—',
            _target_display(notice),
        ]
        for ci, (val, (_label, w)) in enumerate(zip(vals, _NOTICE_COLS)):
            c = row.cells[ci]
            _cell_v_center(c)
            c.text = ''
            p = c.paragraphs[0]
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci == 5 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
            _ar_run(p, val, size=8.5, color=DARK_TEXT)
            _set_row_col_width(row, ci, w)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _mm_to_dxa(9)
        _set_row_cant_split(row)

    if not notices:
        row = tbl.add_row()
        merged = row.cells[0].merge(row.cells[n - 1])
        cp = _cell_para0(merged, 'center')
        _ar_run(cp, 'لا توجد تنويهات لهذه الفترة', size=9, italic=True, color=GREY_TEXT)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = _mm_to_dxa(9)


def _render_notices_section(doc: Document, notices: List[Dict],
                             report_entity_name: Optional[str],
                             report_entity_type: Optional[str],
                             period: Dict,
                             skip_signature: bool = False):
    """
    Groups notices by primary target unit — same rationale as the
    complaints path (see the 'complaints' branch in
    generate_monthly_stylish_docx): one physical batch, one signature page,
    on its own, immediately after that batch's table.

    skip_signature=True (whole-hospital/unfiltered exports — see the
    _NO_SINGLE_UNIT_TYPES check in generate_monthly_stylish_docx) omits every
    signature page: that export isn't routed to one physical unit to sign.
    """
    period_label = period.get('label_ar') or period.get('label') or '—'

    if not notices:
        admin_name, dept_name, sec_name = _scope_labels(report_entity_name, report_entity_type)
        _render_notices_table(doc, [], admin_name, dept_name, sec_name, period_label)
        if not skip_signature:
            _page_break(doc)
            _render_signature_page(doc, report_entity_name or '—', 0, 'التنويهات')
        return

    groups = _group_by_unit(notices)
    items: List[tuple] = []
    for unit_label, group_notices in groups:
        items.append(('table', unit_label, group_notices))
        if not skip_signature:
            items.append(('signature', unit_label, len(group_notices)))

    for i, item in enumerate(items):
        tag, unit_label, payload = item
        try:
            if tag == 'table':
                p = _primary_target(payload[0])
                admin_name = p.get('administration_name') or '—'
                dept_name = p.get('department_name') or '—'
                sec_name = p.get('section_name') or '—'
                _render_notices_table(doc, payload, admin_name, dept_name, sec_name, period_label)
            else:
                _render_signature_page(doc, unit_label, payload, 'التنويهات')
        except Exception as e:
            print(f'[STYLISH] Warning: failed to render notices item #{i} ({unit_label}): {e}')
        if i < len(items) - 1:
            _page_break(doc)


# ---------------------------------------------------------------------------
# SECTION SETUP + REPEATING HEADER/FOOTER
# ---------------------------------------------------------------------------

def _setup_section(sec, title_ar: str, subtitle: str, footer_text: str,
                    report_code: str, period_str: str):
    """
    Configures page geometry and a compact repeating header/footer for one
    document section. python-docx section-property inheritance across
    doc.add_section() is not fully reliable, so every geometry value is
    re-applied explicitly on every call rather than assumed inherited.
    """
    sec.page_width    = int(Mm(297))
    sec.page_height   = int(Mm(210))
    sec.orientation   = WD_ORIENT.LANDSCAPE
    sec.left_margin   = int(Mm(12))
    sec.right_margin  = int(Mm(12))
    sec.top_margin    = int(Mm(13))
    sec.bottom_margin = int(Mm(3))
    sec.header_distance = int(Mm(4))
    sec.footer_distance = int(Mm(3))

    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False

    logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')

    # Repeating header — compact single-row table (logo beside the title
    # block, not stacked above it) so total header height fits a 13mm top
    # margin. Plain 2-column table, no borders, no merged cells, AT_LEAST
    # height — a proven-safe pattern from the previous implementation.
    hdr = sec.header
    hdr.paragraphs[0].clear()

    hdr_tbl = hdr.add_table(1, 2, int(Mm(273)))
    hdr_tbl.autofit = False
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_rtl_table(hdr_tbl)

    logo_cell  = hdr_tbl.rows[0].cells[0]
    title_cell = hdr_tbl.rows[0].cells[1]
    _set_row_col_width(hdr_tbl.rows[0], 0, 32)
    _set_row_col_width(hdr_tbl.rows[0], 1, 241)

    _cell_v_center(logo_cell)
    logo_cell.text = ''
    lp = logo_cell.paragraphs[0]
    lp.clear()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        if os.path.exists(logo_path):
            lp.add_run().add_picture(logo_path, width=int(Inches(0.35)))
    except Exception:
        pass

    _cell_v_center(title_cell)
    title_cell.text = ''
    tp = title_cell.paragraphs[0]
    tp.clear()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = int(Pt(0))
    tp.paragraph_format.space_after  = int(Pt(1))
    _ar_run(tp, title_ar, size=10, bold=True, color=NAVY)

    # Info line — two bidi fixes:
    # 1. Strip ASCII parens from config text (bidi-mirrors backwards in RTL).
    # 2. Insert RLM (U+200F) after each colon preceding an LTR token, so the
    #    colon stays anchored to the Arabic RTL context, then render the LTR
    #    chunk itself (date range, report code) as its own run with an
    #    explicit w:rtl="0" override via _ltr_run — see that helper's
    #    docstring for why this replaced the old Unicode-isolate approach.
    hdr_info_para = title_cell.add_paragraph()
    hdr_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_info_para.paragraph_format.space_before = int(Pt(0))
    hdr_info_para.paragraph_format.space_after = int(Pt(1))
    hdr_info_para._p.get_or_add_pPr().append(OxmlElement('w:bidi'))

    clean_subtitle = (subtitle or '').strip('() ')
    _ar_run(hdr_info_para, clean_subtitle, size=7, italic=True, color=GREY_TEXT)
    _ar_run(hdr_info_para, '   |   ', size=7, color=GREY_TEXT)
    RLM = '‏'
    _ar_run(hdr_info_para, f'{RLM}الفترة:{RLM} ', size=7, bold=True, color=GREY_TEXT)
    _ltr_run(hdr_info_para, period_str, size=7, bold=True, color=GREY_TEXT)
    if report_code:
        _ar_run(hdr_info_para, '   |   ', size=7, color=GREY_TEXT)
        _ar_run(hdr_info_para, f'{RLM}رمز التقرير:{RLM} ', size=7, color=GREY_TEXT)
        _ltr_run(hdr_info_para, report_code, size=7, color=GREY_TEXT)

    hdr_tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr_tbl.rows[0].height = _mm_to_dxa(9)

    # Thin separator line under the header
    hdr_sep = hdr.add_paragraph()
    hdr_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_sep.paragraph_format.space_before = int(Pt(0))
    hdr_sep.paragraph_format.space_after  = int(Pt(0))
    hdr_sep_run = hdr_sep.add_run('')
    hdr_sep_run.font.size = int(Pt(2))
    pPr = hdr_sep._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '4472C4')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer — compact single line
    ftr = sec.footer
    fp = ftr.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    _set_para_bottom_border(fp, color=GREY_LINE, sz=4)
    _ar_run(fp, footer_text, size=6, italic=True, color=GREY_TEXT)


# ---------------------------------------------------------------------------
# PUBLIC ENTRY POINT
# ---------------------------------------------------------------------------

_COMPLAINT_TITLE_DEFAULT = 'التقرير الشهري لفرص التحسين من المرضى ومرافقيهم'
_NOTICE_TITLE = 'التقرير الشهري للتنويهات من المرضى ومرافقيهم'


def generate_monthly_stylish_docx(
    report_data: Dict[str, Any],
    filename: str = '',
    language: str = 'ar',
    report_entity_name: str = None,
    report_entity_type: str = None,
    report_administration: str = None,
    report_department: str = None,
    report_section: str = None,
) -> bytes:
    """
    Generate the Stylish Monthly Report DOCX — a formal, table-based
    document closely modeled on the hospital's official paper form.

    Args:
        report_data: Prepared report model from get_detailed_monthly_report()
                     Keys used: complaints, notices, period, intent_counts
        report_entity_name: Org unit name for scope display
        report_entity_type: 'section' | 'department' | 'administration' | 'hospital'

    Returns:
        bytes: Valid DOCX file content.
    """
    # Normalise inputs
    complaints: List[Dict] = []
    notices: List[Dict] = []
    try:
        if isinstance(report_data, dict):
            raw = report_data.get('complaints', [])
            complaints = raw if isinstance(raw, list) else []
            notices = report_data.get('notices', []) or []
        elif isinstance(report_data, list):
            complaints = report_data
    except Exception:
        complaints = []

    if not report_entity_name:
        if report_administration:    report_entity_name = report_administration
        elif report_department:      report_entity_name = report_department
        elif report_section:         report_entity_name = report_section
        else:                        report_entity_name = 'مستوى المستشفى'

    period: Dict = {}
    if isinstance(report_data, dict):
        period = report_data.get('period', {}) or {}

    cfg = _load_report_config()
    subtitle = cfg.get('header_subtitle', 'Health Care Analysis Tool - HCAT')
    footer_text = cfg.get('footer_text',
                          'نؤمن أن الابتكار لا يكون فقط في التقنيات، بل في أسلوب الخدمة والتواصل والتعاطف')
    report_code = cfg.get('report_code', '')
    complaint_title = cfg.get('header_title', _COMPLAINT_TITLE_DEFAULT)
    period_str = f"{period.get('start_date', '—')}  —  {period.get('end_date', '—')}"

    doc = Document()
    doc.styles['Normal'].font.name = 'Traditional Arabic'
    doc.styles['Normal'].font.size = Pt(10)
    _register_narrative_styles(doc)

    # A whole-hospital/unfiltered export isn't routed to one physical unit
    # to sign, so it gets no signature page at all — see
    # report_export_service.py:130-176 for where each of these values comes
    # from. Multi-export ZIP entries always pass a real unit_type
    # ('administration'/'department'/'section'), never one of these, so
    # they're unaffected (see multi_report_export_service.py:395-396).
    _NO_SINGLE_UNIT_TYPES = {None, 'hospital', 'all_administrations', 'all_departments', 'all_sections'}
    skip_signature = report_entity_type in _NO_SINGLE_UNIT_TYPES

    plan: List[str] = []
    if complaints: plan.append('complaints')
    if notices:    plan.append('notices')
    if not plan: plan = ['empty']

    titles = {
        'complaints': complaint_title,
        'notices': _NOTICE_TITLE,
        'empty': complaint_title,
    }

    first = True
    for kind in plan:
        sec = doc.sections[0] if first else doc.add_section(WD_SECTION.NEW_PAGE)
        first = False
        _setup_section(sec, titles[kind], subtitle, footer_text, report_code, period_str)

        if kind == 'complaints':
            total_c = len(complaints)
            # Group complaints by primary target unit so each unit's batch
            # gets exactly one signature page, immediately after its own
            # complaint pages (see _group_by_unit / _render_signature_page).
            # Flattened into one (tag, payload) item list so the existing
            # "break BETWEEN items only, never after the last one" pattern
            # — which avoids a blank trailing page — applies uniformly to
            # both complaint pages and the signature pages interleaved
            # between batches.
            groups = _group_by_unit(complaints)
            items: List[tuple] = []
            idx = 0
            for unit_label, group_complaints in groups:
                for complaint in group_complaints:
                    idx += 1
                    items.append(('complaint', complaint, idx))
                if not skip_signature:
                    items.append(('signature', unit_label, len(group_complaints)))

            for i, item in enumerate(items):
                try:
                    if item[0] == 'complaint':
                        _, complaint, cidx = item
                        _render_complaint_page(doc, complaint, cidx, total_c, period)
                    else:
                        _, unit_label, count = item
                        _render_signature_page(doc, unit_label, count, 'الشكاوى')
                except Exception as e:
                    print(f'[STYLISH] Warning: failed to render complaints item #{i}: {e}')
                if i < len(items) - 1:
                    _page_break(doc)
        elif kind == 'notices':
            try:
                _render_notices_section(doc, notices, report_entity_name, report_entity_type, period,
                                         skip_signature=skip_signature)
            except Exception as e:
                print(f'[STYLISH] Warning: failed to render notices section: {e}')
        elif kind == 'empty':
            ep = _new_para(doc, align='center', space_before=20, space_after=0)
            _ar_run(ep, 'لا توجد سجلات لهذه الفترة — No records for this period.',
                    size=13, italic=True, color=GREY_TEXT)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
