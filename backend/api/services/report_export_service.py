"""  
Report Export Service
Handles export logic for monthly and seasonal reports.
"""

from typing import Dict, Any, Literal, Optional, List
import traceback
from io import BytesIO
from ..schemas.auth_models import CurrentUser
from .monthly_report_service import monthly_report_service
from .reports_service import reports_service

# Import for emergency fallback
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ReportExportService:
    """
    Service responsible for generating report exports (PDF/CSV).
    Orchestrates data fetching and file generation.
    """

    def generate_export(
        self,
        *,
        current_user: CurrentUser,
        report_type: Literal["monthly", "seasonal"],
        display_mode: Optional[Literal["detailed", "numeric", "hcat"]],
        file_format: Literal["pdf", "csv", "xlsx", "docx"],
        year: int,
        month: Optional[int] = None,
        start_date: Optional[str] = None,
        end_date: Optional[str] = None,
        trimester: int = None,
        quarter: int = None,
        filters: Dict[str, Any] = None,
        include_charts: bool = True,
        language: Literal["en", "ar"] = "en"
    ) -> Dict[str, Any]:
        """
        Generate an export file for a report.
        
        Phase 2.5.7: Accepts current_user to enforce organizational scope.
        
        Args:
            current_user: Authenticated user with allowed org units
            report_type: Type of report (monthly or seasonal)
            display_mode: Display mode (detailed, numeric, hcat)
            file_format: Output format (pdf, csv, xlsx, docx)
            year: Year for the report
            month: Month for monthly reports (or use start_date/end_date)
            start_date: Custom range start date (YYYY-MM-DD) for monthly reports
            end_date: Custom range end date (YYYY-MM-DD) for monthly reports
            trimester: Trimester for seasonal reports
            quarter: Quarter for seasonal reports
            filters: Additional filters
            include_charts: Include charts in PDF
            language: Language for export (en or ar)
        
        Returns:
            Dictionary containing:
                - filename: Generated filename
                - content: File content as bytes
                - content_type: MIME type
        
        Note:
            - xlsx uses CSV generator internally (temporary)
            - docx uses PDF generator internally (temporary)
        """
        try:
            filters = filters or {}
            
            # Map file formats to MIME types
            content_type_mapping = {
                "pdf": "application/pdf",
                "csv": "text/csv",
                "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            
            content_type = content_type_mapping.get(file_format, "application/pdf")
            
            # Step 1: Fetch data based on report type and display mode
            if report_type == "monthly":
                report_data = self._fetch_monthly_data(
                    current_user=current_user,
                    display_mode=display_mode,
                    year=year,
                    month=month,
                    start_date=start_date,
                    end_date=end_date,
                    filters=filters
                )
                # Normalize monthly detailed data: extract complaints list
                if display_mode == "detailed" and isinstance(report_data, dict) and "complaints" in report_data:
                    export_data = report_data["complaints"]
                else:
                    export_data = report_data
            else:  # seasonal
                # For seasonal reports with display_mode=None, use orchestrator
                if display_mode is None:
                    from ..services.seasonal_report_orchestrator import get_or_generate_seasonal_report
                    season_id = filters.get("season_id")
                    orgunit_id = filters.get("orgunit_id")
                    orgunit_type = filters.get("orgunit_type")
                    
                    # PHASE 6: Single season only (no auto-comparison)
                    # Comparisons should be done via /api/seasonal-comparison endpoints
                    report_data = get_or_generate_seasonal_report(
                        season_id=season_id,
                        orgunit_id=orgunit_id,
                        orgunit_type=orgunit_type,
                        user_id=1  # System user for exports
                    )
                else:
                    report_data = self._fetch_seasonal_data(
                        display_mode=display_mode,
                        year=year,
                        trimester=trimester,
                        quarter=quarter,
                        filters=filters
                    )
                export_data = report_data
            
            # Extract entity info from filters for meaningful filenames and headers
            report_entity_name = None
            report_entity_type = None
            report_administration = None
            report_department = None
            report_section = None
            scope_is_all = False  # Track if this is an "all" scope report
            
            if filters:
                # Determine which entity is being reported on
                administration_ids = filters.get("administration_ids")
                department_ids = filters.get("department_ids")
                section_ids = filters.get("section_ids")
                
                if administration_ids:
                    if administration_ids == "all":
                        report_entity_type = "all_administrations"
                        scope_is_all = True
                        report_entity_name = "جميع الإدارات"  # All Administrations
                    else:
                        report_entity_type = "administration"
                        # Get specific administration name
                        try:
                            from ..db_layer.admin_units import get_admin_unit_by_id
                            unit = get_admin_unit_by_id(int(administration_ids.split(',')[0]))
                            report_entity_name = unit.Name if unit else None
                        except:
                            pass
                elif department_ids:
                    if department_ids == "all":
                        report_entity_type = "all_departments"
                        scope_is_all = True
                        report_entity_name = "جميع الأقسام"  # All Departments
                    else:
                        report_entity_type = "department"
                        try:
                            from ..db_layer.admin_units import get_admin_unit_by_id
                            unit = get_admin_unit_by_id(int(department_ids.split(',')[0]))
                            report_entity_name = unit.Name if unit else None
                        except:
                            pass
                elif section_ids:
                    if section_ids == "all":
                        report_entity_type = "all_sections"
                        scope_is_all = True
                        report_entity_name = "جميع الشعب"  # All Sections
                    else:
                        report_entity_type = "section"
                        try:
                            from ..db_layer.admin_units import get_admin_unit_by_id
                            unit = get_admin_unit_by_id(int(section_ids.split(',')[0]))
                            report_entity_name = unit.Name if unit else None
                        except:
                            pass
                else:
                    # Hospital level (no specific filter)
                    report_entity_type = "hospital"
            
            # Step 2: Generate file based on format
            if file_format == "pdf":
                # Use dedicated seasonal formatter for seasonal reports
                if report_type == "seasonal" and display_mode is None:
                    from .seasonal_report_formatter import (
                        generate_seasonal_word_report,
                        generate_seasonal_pdf_report
                    )
                    
                    print(f"\n[EXPORT SERVICE] Generating single season PDF report")
                    print(f"  - report_data type: {type(report_data)}")
                    
                    # PHASE 6: Generate single season report only
                    # Use Word format (better quality than PDF direct generation)
                    content = generate_seasonal_word_report(
                        seasonal_data=report_data,
                        language=language
                    )
                    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    
                    print(f"[EXPORT SERVICE] [OK] Single season report generated")
                else:
                    content = reports_service.generate_pdf_export(
                        report_data=export_data,
                        filename=f"report_{year}.pdf",
                        language=language,
                        include_charts=include_charts,
                        report_entity_name=report_entity_name,
                        report_entity_type=report_entity_type,
                        report_administration=report_administration,
                        report_department=report_department,
                        report_section=report_section
                    )
            elif file_format == "xlsx":
                content = reports_service.generate_xlsx_export(
                    report_data=export_data,
                    filename=f"report_{year}.xlsx",
                    language=language
                )
            elif file_format == "docx":
                # Use dedicated seasonal formatter for seasonal reports
                if report_type == "seasonal" and display_mode is None:
                    from .seasonal_report_formatter import generate_seasonal_word_report
                    
                    print(f"\n[EXPORT SERVICE] Generating single season DOCX report")
                    print(f"  - report_data type: {type(report_data)}")
                    
                    # PHASE 6: Generate single season report only
                    content = generate_seasonal_word_report(
                        seasonal_data=report_data,
                        language=language
                    )
                    
                    print(f"[EXPORT SERVICE] [OK] Single season report generated")
                else:
                    # Check if this is a numeric (aggregated) report or detailed report
                    if display_mode == "numeric":
                        # Use numeric report generator
                        print(f"[EXPORT SERVICE] Generating numeric monthly Word report")
                        print(f"[EXPORT SERVICE] Scope is all: {scope_is_all}, Entity type: {report_entity_type}")
                        
                        # If scope is "all", fetch organizational breakdown data
                        if scope_is_all and report_entity_type in ["all_administrations", "all_departments", "all_sections"]:
                            print(f"[EXPORT SERVICE] Fetching organizational breakdown for {report_entity_type}")
                            org_breakdown = self._fetch_organizational_breakdown(
                                report_entity_type=report_entity_type,
                                year=year,
                                month=month,
                                start_date=start_date,
                                end_date=end_date
                            )
                            # Add organizational breakdown to report data
                            report_data["organizational_breakdown"] = org_breakdown
                        
                        content = reports_service.generate_monthly_numeric_word_report(
                            report_data=report_data,  # Pass full dict (not export_data)
                            filename=f"report_{year}.docx",
                            language=language,
                            report_entity_name=report_entity_name,
                            report_entity_type=report_entity_type
                        )
                    else:
                        # Formatter routing (Session 5): read monthly_report_format from
                        # APP_ReportConfig and route to Classical or Stylish formatter.
                        # Fallback to "classical" on any read error or unrecognised value.
                        # The selector only governs the Monthly Detailed Report;
                        # the Statistical Report uses its own fixed formatter.
                        try:
                            from ..db_layer.report_config_db import get_report_config as _get_cfg
                            _monthly_format = (_get_cfg().get("monthly_report_format") or "classical").strip().lower()
                            if _monthly_format not in ("classical", "stylish"):
                                _monthly_format = "classical"
                        except Exception as _cfg_err:
                            print(f"[EXPORT SERVICE] Warning: could not read monthly_report_format ({_cfg_err}), defaulting to classical")
                            _monthly_format = "classical"

                        print(f"[EXPORT SERVICE] Generating detailed monthly Word report [format={_monthly_format}]")

                        if _monthly_format == "stylish":
                            # Session 6: Real Stylish Monthly formatter.
                            from .monthly_stylish_word_formatter import generate_monthly_stylish_docx
                            print(f"[EXPORT SERVICE] Stylish formatter -> generate_monthly_stylish_docx")
                            content = generate_monthly_stylish_docx(
                                report_data=report_data,
                                filename=f"report_{year}.docx",
                                language=language,
                                report_entity_name=report_entity_name,
                                report_entity_type=report_entity_type,
                                report_administration=report_administration,
                                report_department=report_department,
                                report_section=report_section
                            )
                        else:
                            # Classical formatter (default, unchanged behaviour)
                            content = reports_service.generate_docx_export(
                                report_data=report_data,
                                filename=f"report_{year}.docx",
                                language=language,
                                report_entity_name=report_entity_name,
                                report_entity_type=report_entity_type,
                                report_administration=report_administration,
                                report_department=report_department,
                                report_section=report_section
                            )
            else:  # csv
                content = reports_service.generate_csv_export(
                    report_data=export_data,
                    filename=f"report_{year}.csv",
                    language=language
                )
            
            # Step 3: Build filename
            if report_type == "monthly" and month:
                filename = f"Monthly_Report_{year}_{month:02d}.{file_format}"
            elif report_type == "monthly" and start_date and end_date:
                filename = f"Monthly_Report_{start_date}_to_{end_date}.{file_format}"
            elif report_type == "seasonal" and trimester and display_mode is None:
                # PHASE 6: Single season report (no ZIP)
                filename = f"Seasonal_Report_{year}_T{trimester}.{file_format}"
            elif report_type == "seasonal" and quarter and display_mode is None:
                # PHASE 6: Single season report (no ZIP)
                filename = f"Seasonal_Report_{year}_Q{quarter}.{file_format}"
            elif report_type == "seasonal" and trimester:
                filename = f"Seasonal_Report_{year}_T{trimester}.{file_format}"
            elif report_type == "seasonal" and quarter:
                filename = f"Seasonal_Report_{year}_Q{quarter}.{file_format}"
            else:
                filename = f"Report_{year}.{file_format}"
            
            return {
                "filename": filename,
                "content": content,
                "content_type": content_type
            }
        
        except Exception as e:
            print("\n" + "="*80)
            print(f"[EXPORT DISPATCHER] HARD FAIL: {file_format.upper()} export")
            print(f"Report Type: {report_type}, Year: {year}, Month: {month}")
            print(f"Exception: {type(e).__name__}: {str(e)}")
            print("="*80)
            traceback.print_exc()
            print("="*80 + "\n")
            
            # Emergency fallback for Word exports
            if file_format == "docx":
                print("[EXPORT DISPATCHER] Generating emergency fallback Word document...")
                try:
                    if DOCX_AVAILABLE:
                        # Create minimal emergency Word document
                        doc = Document()
                        
                        # Title
                        title = doc.add_paragraph()
                        title_run = title.add_run("Word Export Emergency Fallback")
                        title_run.bold = True
                        title_run.font.size = 16
                        
                        # Error message
                        doc.add_paragraph()
                        doc.add_paragraph("The system failed to generate the real report.")
                        doc.add_paragraph()
                        doc.add_paragraph(f"Error Type: {type(e).__name__}")
                        doc.add_paragraph(f"Error Message: {str(e)}")
                        doc.add_paragraph()
                        doc.add_paragraph(f"Report Type: {report_type}")
                        doc.add_paragraph(f"Year: {year}")
                        if month:
                            doc.add_paragraph(f"Month: {month}")
                        
                        # Save to buffer
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        filename = f"Emergency_Fallback_{year}_{month or 0:02d}.docx"
                        
                        print("[EXPORT DISPATCHER] Emergency Word document created successfully")
                        return {
                            "filename": filename,
                            "content": buffer.getvalue(),
                            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        }
                    else:
                        print("[EXPORT DISPATCHER] python-docx not available, cannot create fallback")
                except Exception as fallback_error:
                    print(f"[EXPORT DISPATCHER] Even fallback failed: {fallback_error}")
                    traceback.print_exc()
            
            # Re-raise if not Word or if fallback failed
            raise
    
    def _fetch_organizational_breakdown(
        self,
        report_entity_type: str,
        year: int,
        month: Optional[int],
        start_date: Optional[str],
        end_date: Optional[str]
    ) -> List[Dict[str, Any]]:
        """
        Fetch organizational breakdown for 'all' scope reports.
        Gets statistics for each organizational unit (administration/department/section).
        
        Args:
            report_entity_type: 'all_administrations', 'all_departments', or 'all_sections'
            year: Report year
            month: Optional month
            start_date: Optional custom start date
            end_date: Optional custom end date
            
        Returns:
            List of dicts with organizational unit name and statistics
        """
        from ..db_layer.admin_units import get_units_by_type
        from ..constants.org_unit_types import (
            ORG_TYPE_ADMINISTRATION,
            ORG_TYPE_DEPARTMENT,
            ORG_TYPE_SECTION
        )
        
        # Map entity type to organizational unit type
        type_mapping = {
            "all_administrations": ORG_TYPE_ADMINISTRATION,
            "all_departments": ORG_TYPE_DEPARTMENT,
            "all_sections": ORG_TYPE_SECTION
        }
        
        unit_type = type_mapping.get(report_entity_type)
        if not unit_type:
            return []
        
        # Get all units of this type
        units = get_units_by_type(unit_type)
        
        breakdown = []
        for unit in units:
            unit_id = unit["id"]
            unit_name = unit["name"]
            
            # Build filters for this specific unit
            filters = {}
            if report_entity_type == "all_administrations":
                filters["administration_ids"] = str(unit_id)
            elif report_entity_type == "all_departments":
                filters["department_ids"] = str(unit_id)
            elif report_entity_type == "all_sections":
                filters["section_ids"] = str(unit_id)
            
            try:
                # Fetch statistics for this unit
                unit_data = self._fetch_monthly_data(
                    display_mode="numeric",
                    year=year,
                    month=month,
                    start_date=start_date,
                    end_date=end_date,
                    filters=filters
                )
                
                # Extract summary statistics
                summary = unit_data.get("summary", {})
                total_complaints = summary.get("total_complaints", 0)
                
                # Only include units with data
                if total_complaints > 0:
                    breakdown.append({
                        "unit_id": unit_id,
                        "unit_name": unit_name,
                        "total_complaints": total_complaints,
                        "open_complaints": summary.get("open_complaints", 0),
                        "closed_complaints": summary.get("closed_complaints", 0),
                        "red_flags_count": summary.get("red_flags_count", 0),
                        "never_events_count": summary.get("never_events_count", 0),
                        "avg_closure_days": summary.get("avg_closure_days", 0),
                        "by_domain": unit_data.get("by_domain", []),
                        "by_severity": unit_data.get("by_severity", [])
                    })
            except Exception as e:
                print(f"[EXPORT SERVICE] Error fetching data for unit {unit_name} (ID {unit_id}): {e}")
                continue
        
        # Sort by total complaints descending
        breakdown.sort(key=lambda x: x["total_complaints"], reverse=True)
        
        return breakdown
    
    def _fetch_monthly_data(
        self,
        current_user: CurrentUser,
        display_mode: str,
        year: int,
        month: Optional[int],
        start_date: Optional[str],
        end_date: Optional[str],
        filters: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Fetch monthly report data based on display mode.
        
        Phase 2.5.7: Passes current_user for scope enforcement.
        
        IMPORTANT: Uses the EXACT same logic as the view endpoint to ensure
        export data matches what the user sees in the Generate button.
        
        The only difference is page_size=9999 to get ALL records for export.
        """
        # Use the same unified entry point as the view endpoint
        # This ensures 100% consistency between view and export data
        result = monthly_report_service.generate_monthly_report(
            current_user=current_user,
            year=year,
            month=month,
            start_date=start_date,
            end_date=end_date,
            mode=display_mode,
            scope=filters.get("scope"),
            administration_ids=filters.get("administration_ids"),
            department_ids=filters.get("department_ids"),
            section_ids=filters.get("section_ids"),
            page=1,
            page_size=9999  # Get ALL records for export (not paginated like view)
        )
        
        return result
    
    def _fetch_seasonal_data(
        self,
        display_mode: str,
        year: int,
        trimester: int,
        quarter: int,
        filters: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Fetch seasonal report data based on display mode."""
        if display_mode == "hcat":
            return reports_service.get_seasonal_hcat_report(
                year=year,
                trimester=trimester,
                quarter=quarter
            )
        else:
            raise ValueError(f"Invalid display_mode for seasonal report: {display_mode}")


# Singleton instance
report_export_service = ReportExportService()
