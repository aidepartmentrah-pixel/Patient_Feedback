"""
D-B8: Word Generator Reuse Adapter - Enhanced with Detailed Reports

Generates professional Word documents for doctor and worker seasonal reports.
Includes detailed incident tables and comprehensive metrics.
"""

from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
from datetime import datetime


def generate_person_seasonal_word_report(
    person_type: str,
    payload: Dict[str, Any]
) -> bytes:
    """
    Generate comprehensive seasonal report for doctor or worker with detailed incident table.
    
    Creates a professional Word document containing:
    - Person identity and information
    - Summary metrics and counts
    - Detailed table of ALL incidents/cases in the period
    - Professional Arabic formatting
    
    Args:
        person_type: "doctor" or "worker"
        payload: Seasonal data dict containing:
            - identity section (doctor_identity or worker_identity)
            - period section (start, end)
            - metrics section (totals and breakdowns)
            - incidents_details (list of incident dicts)
    
    Returns:
        bytes: Word document (.docx) as bytes
    """
    # Validate person_type
    if person_type not in ["doctor", "worker"]:
        raise ValueError(f"Invalid person_type: {person_type}. Must be 'doctor' or 'worker'")
    
    # Create document
    doc = Document()
    
    # Document setup - A4 Portrait
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    
    # Set default font for Arabic
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Extract data
    identity_key = 'doctor_identity' if person_type == "doctor" else 'worker_identity'
    identity = payload.get(identity_key, {})
    period = payload.get('period', {})
    metrics = payload.get('metrics', {})
    incidents = payload.get('incidents_details', [])
    
    # Title
    title_text = "التقرير الموسمي للطبيب" if person_type == "doctor" else "التقرير الموسمي للموظف"
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph('نظام ملاحظات المرضى')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Spacing
    
    # ============================================================
    # PERSON INFORMATION
    # ============================================================
    doc.add_heading('المعلومات الشخصية', level=2)
    
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.cell(0, 0).text = 'الرقم الوظيفي'
    info_table.cell(0, 1).text = str(identity.get('id', 'N/A'))
    
    info_table.cell(1, 0).text = 'الاسم'
    info_table.cell(1, 1).text = str(identity.get('name', 'Unknown'))
    
    info_table.cell(2, 0).text = 'التخصص' if person_type == "doctor" else 'المسمى الوظيفي'
    info_table.cell(2, 1).text = str(identity.get('specialty') or identity.get('job_title', 'Unknown'))
    
    doc.add_paragraph()
    
    # ============================================================
    # REPORTING PERIOD
    # ============================================================
    doc.add_heading('الفترة المشمولة بالتقرير', level=2)
    
    period_para = doc.add_paragraph()
    period_para.add_run(f"من: {period.get('start', 'N/A')} إلى: {period.get('end', 'N/A')}")
    period_para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # ============================================================
    # SUMMARY METRICS
    # ============================================================
    doc.add_heading('ملخص الإحصائيات', level=2)
    
    metrics_table = doc.add_table(rows=7, cols=2)
    metrics_table.style = 'Light Grid Accent 1'
    
    metrics_table.cell(0, 0).text = 'إجمالي الحالات'
    metrics_table.cell(0, 1).text = str(metrics.get('total_incidents', 0))
    
    metrics_table.cell(1, 0).text = 'حالات خطورة عالية'
    metrics_table.cell(1, 1).text = str(metrics.get('high_severity', 0))
    
    metrics_table.cell(2, 0).text = 'حالات خطورة متوسطة'
    metrics_table.cell(2, 1).text = str(metrics.get('medium_severity', 0))
    
    metrics_table.cell(3, 0).text = 'حالات خطورة منخفضة'
    metrics_table.cell(3, 1).text = str(metrics.get('low_severity', 0))
    
    metrics_table.cell(4, 0).text = 'ملاحظات إيجابية (تنويه)'
    metrics_table.cell(4, 1).text = str(metrics.get('good_feedback_count', 0))
    
    metrics_table.cell(5, 0).text = 'ملاحظات سلبية (نقد/اقتراح)'
    metrics_table.cell(5, 1).text = str(metrics.get('bad_feedback_count', 0))
    
    metrics_table.cell(6, 0).text = 'ملاحظات محايدة'
    metrics_table.cell(6, 1).text = str(metrics.get('neutral_feedback_count', 0))
    
    doc.add_paragraph()
    
    # ============================================================
    # DETAILED INCIDENTS TABLE
    # ============================================================
    doc.add_heading('سجل الحالات التفصيلي', level=2)
    
    if incidents:
        # Create incidents table with classification column
        incidents_table = doc.add_table(rows=len(incidents) + 1, cols=7)
        incidents_table.style = 'Light List Accent 1'
        
        # Header row
        header_cells = incidents_table.rows[0].cells
        header_cells[0].text = 'التاريخ'
        header_cells[1].text = 'المريض'
        header_cells[2].text = 'التصنيف'
        header_cells[3].text = 'الخطورة'
        header_cells[4].text = 'نوع الملاحظة'
        header_cells[5].text = 'الحالة'
        header_cells[6].text = 'رقم الحالة'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Classification display mapping
        classification_labels = {
            'good': 'إيجابي ✓',
            'bad': 'سلبي ✗',
            'neutral': 'محايد ―'
        }
        
        # Data rows
        for idx, incident in enumerate(incidents, start=1):
            row_cells = incidents_table.rows[idx].cells
            
            row_cells[0].text = str(incident.get('date', 'N/A'))
            row_cells[1].text = str(incident.get('patient_name') or incident.get('patient_id', 'N/A'))[:30]
            row_cells[2].text = str(incident.get('category', 'غير محدد'))
            row_cells[3].text = str(incident.get('severity', 'N/A'))
            
            # Intent classification column
            classification = incident.get('classification', 'neutral')
            row_cells[4].text = classification_labels.get(classification, 'محايد ―')
            
            row_cells[5].text = str(incident.get('status', 'N/A'))
            row_cells[6].text = str(incident.get('id', 'N/A'))
            
            # Set font size for data cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    else:
        # Elegant zero-complaint message
        zero_para = doc.add_paragraph()
        zero_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        zero_run = zero_para.add_run('لا توجد حالات مسجلة في هذه الفترة')
        zero_run.font.size = Pt(12)
        zero_run.font.color.rgb = RGBColor(0, 128, 0)
        zero_run.bold = True
        
        praise_para = doc.add_paragraph()
        praise_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        praise_run = praise_para.add_run('سجل نظيف — أداء ممتاز')
        praise_run.font.size = Pt(11)
        praise_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # ============================================================
    # FOOTER
    # ============================================================
    export_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    metadata_para = doc.add_paragraph(f'تاريخ التصدير: {export_date}')
    metadata_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    metadata_para.runs[0].font.size = Pt(9)
    metadata_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    footer_note = doc.add_paragraph('تم إنشاؤه بواسطة نظام ملاحظات المرضى')
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_note.runs[0].font.size = Pt(9)
    footer_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()