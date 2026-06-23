"""
Workflow Performance Report — Word Formatter (HCAT Reporting Refactor - Session 1)

Builds the DOCX export for the unified Workflow Performance Report:
  Section 1 - أداء الردود / Response Performance
  Section 2 - التأخير الحالي / Current Delay

Same hospital report family as the Seasonal/Monthly reports and the Section
Workflow Activity Report: A4 landscape, Traditional Arabic font, hospital
logo, navy header tables. Low-level DOCX helpers are reused as-is from
workflow_activity_word_formatter.py - no parallel formatting utilities.

No business logic lives here - this module only renders the Unified Report
Model produced by workflow_performance_report_service.py.
"""

import os
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from .workflow_activity_word_formatter import (
    _ar_run,
    _new_para,
    _set_rtl_table,
    _apply_minimal_table_borders,
    _set_cell_shading,
    _cell_para,
    _set_row_col_width,
    _fmt_date,
)

NAVY = '1C3A7A'

_UNIT_TYPE_LABELS = {
    "Administration": "إدارة (Administration)",
    "Department": "دائرة (Department)",
    "Section": "قسم (Section)",
}

_PERFORMANCE_LABELS_AR = {
    "Good": "جيد",
    "Warning": "تحذير",
    "Poor": "ضعيف",
}

_LEVEL_LABELS_AR = {
    "All": "الكل (All)",
    "Administration": "إدارة (Administration)",
    "Department": "دائرة (Department)",
    "Section": "قسم (Section)",
}

# Section 1 - Response Performance table layout
_RP_HEADERS = [
    ("الوحدة\nUnit", 40.0),
    ("النوع\nType", 26.0),
    ("الإجمالي\nTotal", 17.0),
    ("في الوقت\nOn Time", 17.0),
    ("متأخر\nLate", 15.0),
    ("إغلاق قسري\nForce Closed", 19.0),
    ("وقت إضافي\nExtra Time", 17.0),
    ("مفتوحة\nStill Open", 16.0),
    ("متوسط الرد\nAvg Days", 18.0),
    ("معدل الاستجابة\nResp. Rate", 20.0),
    ("معدل التأخير\nLate Rate", 18.0),
    ("الأداء\nPerformance", 44.0),
]

# Section 2 - Current Delay table layout
_CD_HEADERS = [
    ("الوحدة\nUnit", 55.0),
    ("النوع\nType", 30.0),
    ("قيد الانتظار\nPending", 30.0),
    ("متأخرة حالياً\nOverdue", 30.0),
    ("إغلاق قسري\nForce Closed", 28.0),
    ("وقت إضافي\nExtra Time", 28.0),
    ("أقدم حالة (يوم)\nOldest (d)", 33.0),
    ("متوسط الانتظار\nAvg (d)", 33.0),
]


def _section_title(doc: Document, ar_title: str, en_title: str):
    p = _new_para(doc, align="right", space_before=14, space_after=0)
    _ar_run(p, ar_title, size=14, bold=True, color=NAVY)
    sub = _new_para(doc, align="right", space_before=0, space_after=6, rtl=False)
    r = sub.add_run(en_title)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    r.font.bold = True


def _empty_section_notice(doc: Document):
    p = _new_para(doc, align="center", space_before=4, space_after=10)
    _ar_run(
        p,
        "لا توجد بيانات للفلاتر المحددة — No data found for the selected filters.",
        size=10, italic=True, color='888888',
    )


def _build_table(doc: Document, headers: list, col_widths: list) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_rtl_table(table)
    _apply_minimal_table_borders(table, outer='888888', outer_sz=4, inner='CCCCCC', inner_sz=4)

    header_row = table.rows[0]
    for ci, label in enumerate(headers):
        _set_cell_shading(header_row.cells[ci], NAVY)
        _ar_run(_cell_para(header_row.cells[ci], "center"), label, size=8, bold=True, color='FFFFFF')

    for ci, w in enumerate(col_widths):
        _set_row_col_width(header_row, ci, w)

    return table


def _add_data_row(table, values: list, col_widths: list, shade: str = None):
    row = table.add_row()
    for ci, val in enumerate(values):
        if shade:
            _set_cell_shading(row.cells[ci], shade)
        _ar_run(_cell_para(row.cells[ci], "center"), str(val), size=8)
    for ci, w in enumerate(col_widths):
        _set_row_col_width(row, ci, w)


def _render_response_performance(doc: Document, rows: List[Dict[str, Any]]):
    _section_title(doc, "أداء الردود", "Response Performance")

    if not rows:
        _empty_section_notice(doc)
        return

    headers = [h for h, _ in _RP_HEADERS]
    widths = [w for _, w in _RP_HEADERS]
    table = _build_table(doc, headers, widths)

    for row in rows:
        unit_type = _UNIT_TYPE_LABELS.get(row.get("unit_type"), row.get("unit_type", "—"))
        perf_label = row.get("performance_label", "—")
        perf_ar = _PERFORMANCE_LABELS_AR.get(perf_label, perf_label)
        _add_data_row(table, [
            row.get("unit_name", "—"),
            unit_type,
            row.get("total_assigned", 0),
            row.get("answered_on_time", 0),
            row.get("answered_late", 0),
            row.get("force_closed", 0),
            row.get("extra_time_granted", 0),
            row.get("still_open", 0),
            row.get("average_response_days", 0.0),
            f"{row.get('response_rate', 0.0)}%",
            f"{row.get('late_rate', 0.0)}%",
            f"{perf_label} / {perf_ar}",
        ], widths)

    _new_para(doc, space_before=4, space_after=0)


def _render_current_delay(doc: Document, rows: List[Dict[str, Any]]):
    _section_title(doc, "التأخير الحالي", "Current Delay")

    if not rows:
        _empty_section_notice(doc)
        return

    headers = [h for h, _ in _CD_HEADERS]
    widths = [w for _, w in _CD_HEADERS]
    table = _build_table(doc, headers, widths)

    for row in rows:
        unit_type = _UNIT_TYPE_LABELS.get(row.get("unit_type"), row.get("unit_type", "—"))
        _add_data_row(table, [
            row.get("unit_name", "—"),
            unit_type,
            row.get("currently_pending", 0),
            row.get("currently_overdue", 0),
            row.get("force_closed_at_unit", 0),
            row.get("extra_time_granted", 0),
            row.get("oldest_pending_days", 0),
            row.get("average_pending_days", 0.0),
        ], widths)

    _new_para(doc, space_before=4, space_after=0)


def generate_workflow_performance_word(report_data: Dict[str, Any]) -> bytes:
    doc = Document()

    sec = doc.sections[0]
    sec.page_height = int(Mm(210))
    sec.page_width = int(Mm(297))
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.left_margin = int(Mm(15))
    sec.right_margin = int(Mm(15))
    sec.top_margin = int(Mm(15))
    sec.bottom_margin = int(Mm(15))

    doc.styles['Normal'].font.name = 'Traditional Arabic'
    doc.styles['Normal'].font.size = Pt(10)

    meta = report_data.get("meta", {})

    # Logo
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            sec.header_distance = int(Inches(0.1))
            hdr_p = sec.header.paragraphs[0]
            hdr_p.clear()
            hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_p.add_run().add_picture(logo_path, width=int(Inches(0.9)))
    except Exception:
        pass

    # Title
    title_p = _new_para(doc, align="center", space_before=0, space_after=2)
    _ar_run(title_p, "تقرير أداء سير العمل", size=20, bold=True, color=NAVY)

    sub_p = _new_para(doc, align="center", space_before=0, space_after=6, rtl=False)
    r = sub_p.add_run("Workflow Performance Report")
    r.font.size = Pt(13)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

    # Report information table
    info = doc.add_table(rows=2, cols=4)
    info.style = 'Table Grid'
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_rtl_table(info)
    _apply_minimal_table_borders(info, outer='AAAAAA', outer_sz=4, inner='CCCCCC', inner_sz=4)

    info_labels = ["الفترة\nPeriod", "تاريخ الإنشاء\nGenerated", "النطاق\nScope", "الوحدة المستهدفة\nTarget Unit"]
    for ci, lbl in enumerate(info_labels):
        _set_cell_shading(info.rows[0].cells[ci], NAVY)
        _ar_run(_cell_para(info.rows[0].cells[ci], "center"), lbl, size=9, bold=True, color='FFFFFF')

    period_str = f"{_fmt_date(meta.get('date_from'))}  —  {_fmt_date(meta.get('date_to'))}"
    target_unit_str = str(meta.get("target_unit_id")) if meta.get("target_unit_id") else "الكل (All units in scope)"
    level_str = _LEVEL_LABELS_AR.get(meta.get("level"), meta.get("level", "All"))

    info_values = [
        period_str,
        _fmt_date(meta.get("generated_at")),
        level_str,
        target_unit_str,
    ]
    for ci, val in enumerate(info_values):
        _set_cell_shading(info.rows[1].cells[ci], 'F0F4FF')
        _ar_run(_cell_para(info.rows[1].cells[ci], "center"), val, size=9)

    for ci, w in enumerate([70.0, 70.0, 60.0, 67.0]):
        for ri in range(2):
            _set_row_col_width(info.rows[ri], ci, w)

    _new_para(doc, space_before=0, space_after=8)

    # Section 1
    _render_response_performance(doc, report_data.get("response_performance", []))

    # Section 2
    _render_current_delay(doc, report_data.get("current_delay", []))

    # Footer
    footer_p = _new_para(doc, align="center", space_before=10, space_after=0)
    _ar_run(
        footer_p,
        f"نهاية التقرير — End of Report  |  بواسطة: {meta.get('generated_by', 'System')}",
        size=9, color='888888',
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
