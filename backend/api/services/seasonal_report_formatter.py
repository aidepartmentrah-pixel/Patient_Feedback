"""
Seasonal Report Formatter
Generates professional Word documents for seasonal reports with hierarchical classification tables.

Features:
- A4 Landscape orientation with Arabic header and logo
- RTL table layout (right-to-left for Arabic readers)
- Hierarchical table structure with merged cells (Domain > Category > Sub-Category > Classification)
- Bilingual classification names (Arabic + English)
- Domain-level grouping with subtotals
- Professional styling with colors, borders, and proper alignment
- Clear Severity section (Low/Medium/High)
- Clear Prevention Action section (Yes/No)
- Policy targets footer in Arabic
"""

from typing import Dict, Any, List, Tuple, Optional
import os
from ..db_layer.report_config_db import get_report_config
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
import io
from collections import defaultdict
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend for server environments
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib import patches
import numpy as np


def set_cell_shading(cell: _Cell, color: str):
    """
    Set background color for a table cell.
    
    Args:
        cell: The cell to shade
        color: Hex color code (e.g., 'D9E2F3')
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def center_cell_content(cell: _Cell):
    """
    Center cell content both horizontally and vertically.
    
    Args:
        cell: The cell to center
    """
    # Horizontal centering
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Vertical centering
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def merge_cells_vertically(table, col: int, start_row: int, end_row: int):
    """
    Merge cells vertically in a table column.
    
    Args:
        table: The table object
        col: Column index
        start_row: Starting row index
        end_row: Ending row index (inclusive)
    """
    if start_row >= end_row:
        return
    
    # Get the first cell
    first_cell = table.rows[start_row].cells[col]
    
    # Merge with subsequent cells
    for row_idx in range(start_row + 1, end_row + 1):
        cell_to_merge = table.rows[row_idx].cells[col]
        first_cell.merge(cell_to_merge)


def generate_seasonal_word_report(
    seasonal_data: Dict[str, Any],
    language: str = "en"
) -> bytes:
    """
    Generate a professional Word document for a seasonal report.
    
    Creates a hierarchical table structure matching the user's specification:
    - A4 Landscape with Arabic header and logo
    - RTL table layout for Arabic readers
    - Domain level (e.g., "Clinical" n=6) with merged cells
    - Category level (e.g., "Quality of Care" n=4)
    - Sub-Category level (e.g., "Examination & Monitoring" n=2)
    - Classification level with Arabic + English names
    - Clear Severity section (Low/Medium/High)
    - Clear Prevention Action section (Yes/No)
    - Policy targets footer in Arabic
    
    Args:
        seasonal_data: Seasonal report data from orchestrator
        language: Language for the report (en or ar)
    
    Returns:
        Bytes of the generated Word document
    """
    # Debug: Log what we received
    print(f"\n[FORMATTER] Received seasonal_data type: {type(seasonal_data)}")
    print(f"[FORMATTER] seasonal_data preview: {str(seasonal_data)[:200]}")
    
    # Validate input
    if not isinstance(seasonal_data, dict):
        raise TypeError(f"Expected dict, got {type(seasonal_data)}. Data: {seasonal_data}")

    # Pull institutional config from DB (same pattern as monthly report)
    _cfg              = get_report_config()
    _header_title     = _cfg.get("seasonal_header_title",    "نموذج التقرير الموسمي لفرص التحسين والإجراءات التصحيحية")
    _header_subtitle  = _cfg.get("seasonal_header_subtitle", "(إصدار رسمي — للاستخدام الإداري والجودة)")
    _footer_text      = _cfg.get("seasonal_footer_text",     "")
    _report_code      = _cfg.get("seasonal_report_code",     "")

    # Utility functions
    def _safe(v):
        """Convert dimension values to int (python-docx requirement)"""
        return int(v)
    
    doc = Document()
    
    # ============================================================
    # DOCUMENT SETUP - A4 LANDSCAPE
    # ============================================================
    section = doc.sections[0]
    section.page_height = _safe(Mm(210))  # A4 width becomes height in landscape
    section.page_width = _safe(Mm(297))   # A4 height becomes width in landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = _safe(Mm(15))
    section.right_margin = _safe(Mm(15))
    section.top_margin = _safe(Mm(15))
    section.bottom_margin = _safe(Mm(15))
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Traditional Arabic'
    font.size = Pt(11)
    
    # Extract data
    header = seasonal_data.get("header", {})
    classification_stats = seasonal_data.get("classification_stats", [])
    policy_snapshot = seasonal_data.get("policy_snapshot", {})
    
    # ============================================================
    # HEADER - LOGO (TOP RIGHT)
    # ============================================================
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            # Make header compact
            section.header_distance = Inches(0.1)
            header_section = section.header
            
            # Use only one paragraph and clear it
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[FORMATTER] Could not add logo: {e}")
        pass
    
    # ============================================================
    # TITLE SECTION (ARABIC)
    # ============================================================
    
    # Main Title (big, bold, centered)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(_header_title)
    title_run.font.size = int(Pt(21))
    title_run.font.bold = True
    title_run.font.name = 'Traditional Arabic'
    title_run.italic = False
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = int(Pt(3))
    
    # Subtitle (smaller, centered)
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(_header_subtitle)
    subtitle_run.font.size = int(Pt(14))
    subtitle_run.font.name = 'Traditional Arabic'
    subtitle_run.italic = False
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.space_after = int(Pt(6))

    # Report code line (only shown when configured)
    if _report_code:
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(f"رمز التقرير: {_report_code}")
        code_run.font.size = int(Pt(11))
        code_run.font.name = 'Traditional Arabic'
        code_run.italic = False
        code_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        code_para.space_after = int(Pt(3))

    # ============================================================
    # PERIOD & ORGANIZATION INFO
    # ============================================================
    
    # Period line (centered, bold)
    period_str = header.get('period', 'N/A')
    # Convert period to Arabic
    period_arabic = period_str
    if 'Q1' in period_str:
        period_arabic = period_str.replace('Q1', 'الربع الأول')
    elif 'Q2' in period_str:
        period_arabic = period_str.replace('Q2', 'الربع الثاني')
    elif 'Q3' in period_str:
        period_arabic = period_str.replace('Q3', 'الربع الثالث')
    elif 'Q4' in period_str:
        period_arabic = period_str.replace('Q4', 'الربع الرابع')
    elif 'Trim1' in period_str:
        period_arabic = period_str.replace('Trim1', 'الفصل الأول')
    elif 'Trim2' in period_str:
        period_arabic = period_str.replace('Trim2', 'الفصل الثاني')
    elif 'Trim3' in period_str:
        period_arabic = period_str.replace('Trim3', 'الفصل الثالث')
    
    period_para = doc.add_paragraph()
    period_run = period_para.add_run(f"الموسم المعني: {period_arabic}")
    period_run.font.size = int(Pt(12))
    period_run.font.bold = True
    period_run.font.name = 'Traditional Arabic'
    period_run.italic = False
    period_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_para.space_after = int(Pt(3))
    
    # Organization info table (3-column, borderless)
    org_table = doc.add_table(rows=1, cols=3)
    
    # Remove all table borders
    org_tbl = org_table._element
    org_tblPr = org_tbl.tblPr
    if org_tblPr is None:
        org_tblPr = OxmlElement('w:tblPr')
        org_tbl.insert(0, org_tblPr)
    
    org_tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_elem = OxmlElement(f'w:{border_name}')
        border_elem.set(qn('w:val'), 'nil')
        org_tblBorders.append(border_elem)
    org_tblPr.append(org_tblBorders)
    
    # Center the table
    org_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    org_tblJc = OxmlElement('w:jc')
    org_tblJc.set(qn('w:val'), 'center')
    org_tblPr.append(org_tblJc)
    
    # Set column widths
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    target_width = int(usable_width * 0.7)
    col_width = int(target_width / 3)
    
    for i in range(3):
        org_table.columns[i].width = col_width
    
    # Determine organization type name
    orgunit_type = header.get('orgunit_type', 0)
    type_names = {
        0: "المستشفى",
        1: "الإدارة",
        2: "الدائرة",
        3: "القسم"
    }
    type_name = type_names.get(orgunit_type, "الوحدة التنظيمية")
    
    # Fill cells with data
    org_cells = org_table.rows[0].cells
    org_data = [
        ("المستشفى: ", "مستشفى الرّسول الأعظم"),
        ("الوحدة التنظيمية: ", header.get('orgunit_name', 'N/A')),
        ("النوع: ", type_name)
    ]
    
    for i, (label, value) in enumerate(org_data):
        cell = org_cells[i]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.right_to_left = True
        paragraph.space_after = int(Pt(6))
        
        # Bold label
        label_run = paragraph.add_run(label)
        label_run.font.bold = True
        label_run.font.size = int(Pt(15))
        label_run.font.name = 'Traditional Arabic'
        label_run.italic = False
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        
        # Normal value
        value_run = paragraph.add_run(str(value))
        value_run.font.bold = False
        value_run.font.size = int(Pt(15))
        value_run.font.name = 'Traditional Arabic'
        value_run.italic = False
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    # Visual separator line
    separator_para = doc.add_paragraph()
    separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        pPr = separator_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4472C4')
        pBdr.append(bottom)
        pPr.append(pBdr)
    except:
        pass
    
    doc.add_paragraph()  # Spacer
    
    # ============================================================
    # HIERARCHICAL CLASSIFICATION TABLES (RTL LAYOUT - ONE PER DOMAIN)
    # ============================================================
    
    # Build hierarchical structure from flat classification_stats
    hierarchy = _build_hierarchy(classification_stats)
    
    # Create separate table for each domain
    _create_hierarchical_tables_by_domain_rtl(doc, hierarchy, language)
    
    # ============================================================
    # POLICY COMPLIANCE EVALUATION (ARABIC RTL)
    # ============================================================
    
    doc.add_paragraph()  # Spacer
    _add_policy_compliance_section(doc, header, policy_snapshot, classification_stats)
    doc.add_paragraph()  # Spacer

    # ============================================================
    # SUMMARY STATISTICS (AT THE END, IN ARABIC)
    # ============================================================
    
    # Domain totals
    clinical_count = header.get('clinical_domain_count', 0)
    management_count = header.get('management_domain_count', 0)
    relational_count = header.get('relational_domain_count', 0)
    total_cases = header.get('total_cases', 0)
    
    # Severity counts
    low_count = header.get('low_severity_count', 0)
    medium_count = header.get('medium_severity_count', 0)
    high_count = header.get('high_severity_count', 0)
    
    # Add heading in Arabic
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary_heading.paragraph_format.right_to_left = True
    
    sh_run = summary_heading.add_run("📊 إحصائيات ملخصة (Summary Statistics)")
    sh_run.font.bold = True
    sh_run.font.size = Pt(14)
    sh_run.font.name = 'Traditional Arabic'
    sh_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    doc.add_paragraph()
    
    summary_table = doc.add_table(rows=5, cols=3)
    summary_table.style = 'Table Grid'
    
    # Headers in Arabic (RTL order: النسبة المئوية, العدد, الفئة)
    summary_table.rows[0].cells[0].text = "النسبة المئوية"
    summary_table.rows[0].cells[1].text = "العدد"
    summary_table.rows[0].cells[2].text = "الفئة"
    for cell in summary_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        center_cell_content(cell)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Domain breakdown in Arabic
    domain_data = [
        ("Clinical (الجوانب السريرية)", clinical_count),
        ("Management (الجوانب الإدارية)", management_count),
        ("Relational (الجوانب العلائقية)", relational_count)
    ]
    
    for idx, (label, count) in enumerate(domain_data, start=1):
        row = summary_table.rows[idx]
        
        # RTL order: percentage (col 0), count (col 1), label (col 2)
        percentage = (count / total_cases * 100) if total_cases > 0 else 0
        row.cells[0].text = f"{percentage:.1f}%"
        center_cell_content(row.cells[0])
        
        row.cells[1].text = str(count)
        center_cell_content(row.cells[1])
        
        row.cells[2].text = label
        row.cells[2].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        row.cells[2].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Total row in Arabic (RTL order: percentage, count, label)
    total_row = summary_table.rows[4]
    
    total_row.cells[0].text = "100.0%"
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    center_cell_content(total_row.cells[0])
    
    total_row.cells[1].text = str(total_cases)
    total_row.cells[1].paragraphs[0].runs[0].bold = True
    center_cell_content(total_row.cells[1])
    
    total_row.cells[2].text = "المجموع TOTAL"
    total_row.cells[2].paragraphs[0].runs[0].bold = True
    total_row.cells[2].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
    total_row.cells[2].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    set_cell_shading(total_row.cells[0], 'D9E2F3')
    set_cell_shading(total_row.cells[1], 'D9E2F3')
    set_cell_shading(total_row.cells[2], 'D9E2F3')
    
    doc.add_paragraph()
    
    # Severity breakdown in Arabic
    severity_para = doc.add_paragraph()
    severity_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    severity_para.paragraph_format.right_to_left = True
    
    sev_bold = severity_para.add_run("تفصيل الشدة (Severity Breakdown): ")
    sev_bold.font.bold = True
    sev_bold.font.name = 'Traditional Arabic'
    sev_bold._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    sev_text = severity_para.add_run(f"منخفضة ({low_count}) • متوسطة ({medium_count}) • عالية ({high_count})")
    sev_text.font.name = 'Traditional Arabic'
    sev_text._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    # ============================================================
    # FOOTER TEXT (institutional quote / closing statement)
    # ============================================================
    if _footer_text:
        doc.add_paragraph()  # spacer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.right_to_left = True
        footer_run = footer_para.add_run(_footer_text)
        footer_run.font.size = int(Pt(10))
        footer_run.italic = True
        footer_run.font.name = 'Traditional Arabic'
        footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    # ============================================================
    # SAVE AND RETURN
    # ============================================================

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _tbl_rtl(table) -> None:
    """Apply bidiVisual RTL to a table (idempotent helper)."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblPr.append(OxmlElement('w:bidiVisual'))


def _hdr_run(cell, text: str, size: int = 10) -> None:
    """Write a bold white header cell (reused by both policy tables)."""
    cell.text = text
    if cell.paragraphs[0].runs:
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = 'Traditional Arabic'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        r.font.color.rgb = RGBColor(255, 255, 255)
    center_cell_content(cell)
    set_cell_shading(cell, '4472C4')


def _render_section_policy_table(
    doc: Document,
    policy_snapshot: Dict[str, Any],
    classification_stats: List[Dict[str, Any]],
) -> None:
    """
    Section policy: per-classification evaluation table.

    SEMANTICS:
      LowSeverityLimit  = Low-severity incidents threshold per classification
      MediumSeverityLimit = Medium incidents threshold per classification
      HighSeverityLimit   = High incidents threshold per classification

    A classification is violating if ANY enabled rule has actual > limit
    (STRICTLY greater than).  Only violating rows become red.

    Columns (RTL): التصنيف | منخفضة | متوسط | عالي | الحدود (ك/م/ع) | الحالة
    """
    enable_all  = bool(policy_snapshot.get('enable_low_severity_repetition_rule',       False))
    enable_med  = bool(policy_snapshot.get('enable_medium_severity_repetition_rule',    False))
    enable_high = bool(policy_snapshot.get('enable_high_severity_percentage_rule',      False))

    all_limit  = policy_snapshot.get('low_severity_limit',    0)
    med_limit  = policy_snapshot.get('medium_severity_limit', 0)
    high_limit = policy_snapshot.get('high_severity_limit',   0)

    if not any([enable_all, enable_med, enable_high]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True
        r = p.add_run("لا توجد قواعد تصنيف مفعّلة في السياسة الحالية.")
        r.font.size = Pt(11)
        r.font.name = 'Traditional Arabic'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        return

    if not classification_stats:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("لا توجد بيانات تصنيفية.")
        r.font.size = Pt(10)
        r.font.name = 'Traditional Arabic'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        return

    # Build limits display string (only enabled ones)
    parts = []
    if enable_all:  parts.append(str(all_limit))
    if enable_med:  parts.append(str(med_limit))
    if enable_high: parts.append(str(high_limit))
    limits_str = " / ".join(parts) if parts else "—"

    # Evaluate each classification
    rows_data = []
    for stat in classification_stats:
        total = stat.get('low_count', 0)
        med   = stat.get('medium_count', 0)
        high  = stat.get('high_count', 0)
        violating = (
            (enable_all  and total > all_limit)  or
            (enable_med  and med   > med_limit)  or
            (enable_high and high  > high_limit)
        )
        rows_data.append({
            'name_ar': stat.get('classification_name', ''),
            'total': total,
            'medium': med,
            'high': high,
            'violating': violating,
        })

    # Violating rows first, then alphabetical
    rows_data.sort(key=lambda x: (0 if x['violating'] else 1, x['name_ar']))

    # 6-column RTL table
    table = doc.add_table(rows=len(rows_data) + 1, cols=6)
    table.style = 'Table Grid'
    _tbl_rtl(table)

    for ci, txt in enumerate(['التصنيف', 'منخفضة', 'متوسط', 'عالي', 'الحدود (ك/م/ع)', 'الحالة']):
        _hdr_run(table.rows[0].cells[ci], txt, size=9)

    for ri, rd in enumerate(rows_data, start=1):
        row = table.rows[ri]
        violated = rd['violating']

        # Col 0 — Classification name
        row.cells[0].text = rd['name_ar']
        if row.cells[0].paragraphs[0].runs:
            r = row.cells[0].paragraphs[0].runs[0]
            r.font.size = Pt(9)
            r.font.name = 'Traditional Arabic'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            if violated:
                r.bold = True
                r.font.color.rgb = RGBColor(192, 0, 0)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        def _metric_cell(cell, value, rule_violated):
            cell.text = str(value)
            center_cell_content(cell)
            if cell.paragraphs[0].runs:
                r = cell.paragraphs[0].runs[0]
                r.font.size = Pt(9)
                if rule_violated:
                    r.bold = True
                    r.font.color.rgb = RGBColor(192, 0, 0)
                    set_cell_shading(cell, 'FFE6E6')

        _metric_cell(row.cells[1], rd['total'],  enable_all  and rd['total']  > all_limit)
        _metric_cell(row.cells[2], rd['medium'], enable_med  and rd['medium'] > med_limit)
        _metric_cell(row.cells[3], rd['high'],   enable_high and rd['high']   > high_limit)

        # Col 4 — Limits
        row.cells[4].text = limits_str
        center_cell_content(row.cells[4])
        if row.cells[4].paragraphs[0].runs:
            row.cells[4].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(row.cells[4], 'F2F2F2')

        # Col 5 — Status
        row.cells[5].text = "✗ مخالف" if violated else "✓ مطابق"
        center_cell_content(row.cells[5])
        if row.cells[5].paragraphs[0].runs:
            r = row.cells[5].paragraphs[0].runs[0]
            r.bold = True
            r.font.size = Pt(9)
            if violated:
                r.font.color.rgb = RGBColor(192, 0, 0)
                set_cell_shading(row.cells[5], 'FFC7CE')
            else:
                r.font.color.rgb = RGBColor(0, 128, 0)
                set_cell_shading(row.cells[5], 'C6EFCE')


def _render_domain_policy_table(
    doc: Document,
    header: Dict[str, Any],
    policy_snapshot: Dict[str, Any],
) -> None:
    """
    Administration / Department / Hospital policy: domain-level evaluation.

    Uses pre-computed violated_rules JSON from header when available;
    falls back to computing from header domain counts vs policy limits.

    Violation rule (STRICTLY greater than): DomainCases > DomainLimit.
    Only violating domain rows become red.

    Columns (RTL): المجال | الفعلي | الحد | الحالة
    """
    import json

    enable_domain = bool(policy_snapshot.get('enable_high_severity_percentage_by_domain_rule', False))
    if not enable_domain:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True
        r = p.add_run("لا توجد قواعد مجال مفعّلة في السياسة الحالية.")
        r.font.size = Pt(11)
        r.font.name = 'Traditional Arabic'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        return

    # Parse pre-computed violation details if available
    pre = {}
    raw = header.get('violated_rules')
    if raw:
        try:
            for v in json.loads(raw):
                pre[v.get('rule', '')] = v
        except Exception:
            pass

    clinical_count   = header.get('clinical_domain_count',    0)
    management_count = header.get('management_domain_count',  0)
    relational_count = header.get('relational_domain_count',  0)
    total_cases      = header.get('total_cases', 0)

    clinical_limit   = policy_snapshot.get('clinical_domain_limit',   0)
    management_limit = policy_snapshot.get('management_domain_limit', 0)
    relational_limit = policy_snapshot.get('relational_domain_limit', 0)

    def _build_domain_row(name_ar, count, limit, rule_key):
        if rule_key in pre:
            v         = pre[rule_key]
            actual    = v.get('actual', count)
            threshold = v.get('threshold', limit)
            unit      = v.get('threshold_unit', '')
            try:
                violated = float(actual) > float(threshold)
            except (TypeError, ValueError):
                violated = False
            return {
                'name_ar':  name_ar,
                'actual':   f"{actual}{unit}",
                'limit':    f"{threshold}{unit}",
                'violated': violated,
            }
        # Fallback: percentage comparison (STRICTLY greater than) to match violated_rules semantics
        total = header.get('total_cases', 0)
        if total > 0 and limit > 0:
            actual_pct = round((count / total) * 100, 1)
            violated   = actual_pct > limit
            actual_str = f"{actual_pct}%"
            limit_str  = f"{limit}%"
        else:
            violated   = limit > 0 and count > limit
            actual_str = str(count)
            limit_str  = str(limit)
        return {
            'name_ar':  name_ar,
            'actual':   actual_str,
            'limit':    limit_str,
            'violated': violated,
        }

    rows = []
    if clinical_limit > 0:
        rows.append(_build_domain_row('المجال السريري',  clinical_count,   clinical_limit,   'ClinicalDomainLimit'))
    if management_limit > 0:
        rows.append(_build_domain_row('المجال الإداري',  management_count, management_limit, 'ManagementDomainLimit'))
    if relational_limit > 0:
        rows.append(_build_domain_row('المجال العلائقي', relational_count, relational_limit, 'RelationalDomainLimit'))

    if not rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("لا توجد حدود مجال محددة في السياسة.")
        r.font.size = Pt(10)
        r.font.name = 'Traditional Arabic'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        return

    # 4-column RTL table
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = 'Table Grid'
    _tbl_rtl(table)

    for ci, txt in enumerate(['المجال', 'الفعلي', 'الحد', 'الحالة']):
        _hdr_run(table.rows[0].cells[ci], txt, size=10)

    for ri, rd in enumerate(rows, start=1):
        row = table.rows[ri]
        violated = rd['violated']

        # Col 0 — Domain name
        row.cells[0].text = rd['name_ar']
        if row.cells[0].paragraphs[0].runs:
            r = row.cells[0].paragraphs[0].runs[0]
            r.font.size = Pt(10)
            r.font.name = 'Traditional Arabic'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            if violated:
                r.bold = True
                r.font.color.rgb = RGBColor(192, 0, 0)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Col 1 — Actual
        row.cells[1].text = rd['actual']
        center_cell_content(row.cells[1])
        if row.cells[1].paragraphs[0].runs:
            r = row.cells[1].paragraphs[0].runs[0]
            r.font.size = Pt(10)
            if violated:
                r.bold = True
                r.font.color.rgb = RGBColor(192, 0, 0)
                set_cell_shading(row.cells[1], 'FFE6E6')

        # Col 2 — Limit
        row.cells[2].text = rd['limit']
        center_cell_content(row.cells[2])
        if row.cells[2].paragraphs[0].runs:
            row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)

        # Col 3 — Status
        row.cells[3].text = "✗ مخالف" if violated else "✓ مطابق"
        center_cell_content(row.cells[3])
        if row.cells[3].paragraphs[0].runs:
            r = row.cells[3].paragraphs[0].runs[0]
            r.bold = True
            r.font.size = Pt(10)
            if violated:
                r.font.color.rgb = RGBColor(192, 0, 0)
                set_cell_shading(row.cells[3], 'FFC7CE')
            else:
                r.font.color.rgb = RGBColor(0, 128, 0)
                set_cell_shading(row.cells[3], 'C6EFCE')


def _add_policy_compliance_section(
    doc,
    header: Dict[str, Any],
    policy_snapshot: Dict[str, Any],
    classification_stats: Optional[List[Dict[str, Any]]] = None,
):
    """
    Render policy compliance for a seasonal report.

    Section (orgunit_type==3): per-classification table — ONLY violating
    classification rows become red.

    Administration/Department/Hospital (orgunit_type 0/1/2): domain-level
    table — ONLY violating domain rows become red.
    """
    # ── Title ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_para.paragraph_format.right_to_left = True
    title_run = title_para.add_run("📊 تقييم الامتثال للسياسة (Policy Compliance Evaluation)")
    title_run.font.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Traditional Arabic'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    doc.add_paragraph()

    # ── No policy configured ────────────────────────────────────────────
    if not policy_snapshot:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True
        r = p.add_run("لا توجد سياسة محددة لهذه الوحدة التنظيمية.")
        r.font.size = Pt(11)
        r.font.name = 'Traditional Arabic'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        return

    orgunit_type = header.get('orgunit_type', 0)

    if orgunit_type == 3:
        # Section: per-classification violation table
        _render_section_policy_table(doc, policy_snapshot, classification_stats or [])
    else:
        # Administration / Department / Hospital: domain violation table
        _render_domain_policy_table(doc, header, policy_snapshot)

    # ── Overall compliance status (from pre-computed header flag) ───────
    doc.add_paragraph()
    is_compliant = bool(header.get('is_compliant', True))
    status_para = doc.add_paragraph()
    status_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    status_para.paragraph_format.right_to_left = True
    if is_compliant:
        sr = status_para.add_run("✓ الحالة العامة: مطابق للسياسة (COMPLIANT)")
        sr.font.color.rgb = RGBColor(0, 128, 0)
    else:
        sr = status_para.add_run("✗ الحالة العامة: غير مطابق (NON-COMPLIANT)")
        sr.font.color.rgb = RGBColor(192, 0, 0)
    sr.font.bold = True
    sr.font.size = Pt(12)
    sr.font.name = 'Traditional Arabic'
    sr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')


def _build_hierarchy(classification_stats: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Build hierarchical structure from flat classification statistics.
    
    Returns structure:
    {
        'Clinical': {
            'total': 6,
            'categories': {
                'Quality of Care': {
                    'total': 4,
                    'subcategories': {
                        'Examination & Monitoring': {
                            'total': 2,
                            'classifications': [...]
                        }
                    }
                }
            }
        }
    }
    """
    hierarchy = defaultdict(lambda: {
        'total': 0,
        'categories': defaultdict(lambda: {
            'total': 0,
            'subcategories': defaultdict(lambda: {
                'total': 0,
                'classifications': []
            })
        })
    })
    
    for stat in classification_stats:
        domain_name = stat.get('domain_name', 'Unknown Domain')
        category_name = stat.get('category_name', 'Unknown Category')
        subcategory_name = stat.get('subcategory_name', 'Unknown Sub-Category')
        
        # Add to hierarchy
        hierarchy[domain_name]['total'] += stat.get('total_count', 0)
        hierarchy[domain_name]['categories'][category_name]['total'] += stat.get('total_count', 0)
        hierarchy[domain_name]['categories'][category_name]['subcategories'][subcategory_name]['total'] += stat.get('total_count', 0)
        hierarchy[domain_name]['categories'][category_name]['subcategories'][subcategory_name]['classifications'].append(stat)
    
    return dict(hierarchy)


def _create_hierarchical_tables_by_domain_rtl(doc: Document, hierarchy: Dict[str, Any], language: str):
    """
    Create separate hierarchical tables for each domain with RTL layout.
    Each domain gets its own table with proper spacing.
    """
    
    if not hierarchy or len(hierarchy) == 0:
        doc.add_paragraph("لا توجد بيانات متاحة (No classification data available).")
        return
    
    # Process each domain separately
    for domain_idx, (domain_name, domain_data) in enumerate(sorted(hierarchy.items())):
        # Add domain title
        if domain_idx > 0:
            doc.add_paragraph()  # Spacing between domains
        
        domain_title = doc.add_paragraph()
        domain_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        domain_title.paragraph_format.right_to_left = True
        
        title_run = domain_title.add_run(f'📂 {domain_name} (n={domain_data["total"]})')
        title_run.font.bold = True
        title_run.font.size = Pt(13)
        title_run.font.name = 'Traditional Arabic'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        title_run.font.color.rgb = RGBColor(68, 114, 196)
        
        # Calculate rows for this domain
        domain_rows = 2  # Header rows
        for category_data in domain_data['categories'].values():
            for subcategory_data in category_data['subcategories'].values():
                domain_rows += len(subcategory_data['classifications'])
        
        # Create table for this domain (11 columns)
        table = doc.add_table(rows=domain_rows, cols=11)
        table.style = 'Table Grid'
        
        # Set RTL table direction
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)
        
        # Setup headers
        header_row1 = table.rows[0]
        header_row2 = table.rows[1]
        
        headers_main = [
            "Problem Category\nفئة المشكلة",
            "Sub-Category\nالفئة الفرعية",
            "التصنيف عربي\nClassification AR",
            "التصنيف إنجليزي\nClassification EN",
            "Total\nالمجموع",
            "الشدة Severity",
            "",
            "",
            "الإجراءات الوقائية\nPrevention",
            ""
        ]
        
        # Adjust: Removed Domain column, so now 10 columns effectively used
        # But table still has 11 cols - merge first col or skip it
        # Actually, let me keep it simpler - use all 11 cols but first is Category
        
        headers_main = [
            "Problem Category\nفئة المشكلة",
            "Sub-Category\nالفئة الفرعية",
            "التصنيف عربي\nClassification AR",
            "التصنيف إنجليزي\nClassification EN",
            "Total\nالمجموع",
            "الشدة Severity",
            "",
            "",
            "",
            "الإجراءات الوقائية\nPrevention",
            ""
        ]
        
        for idx, header_text in enumerate(headers_main):
            cell = header_row1.cells[idx]
            if header_text:
                cell.text = header_text
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                center_cell_content(cell)
            set_cell_shading(cell, '4472C4')
        
        subheaders = [
            "",
            "",
            "",
            "",
            "",
            "LOW\nمنخفض",
            "MED\nمتوسط",
            "HIGH\nعالي",
            "",
            "YES\nنعم",
            "NO\nلا"
        ]
        
        for idx, subheader_text in enumerate(subheaders):
            cell = header_row2.cells[idx]
            if subheader_text:
                cell.text = subheader_text
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                center_cell_content(cell)
            set_cell_shading(cell, '5B9BD5')
        
        # Merge headers
        try:
            for i in range(5):
                header_row1.cells[i].merge(header_row2.cells[i])
            header_row1.cells[5].merge(header_row1.cells[7])
            header_row1.cells[9].merge(header_row1.cells[10])
        except Exception as e:
            print(f"[FORMATTER] Warning: Could not merge headers: {e}")
        
        # Fill data rows for this domain
        current_row = 2
        for category_name, category_data in sorted(domain_data['categories'].items()):
            category_start_row = current_row
            category_total = category_data['total']
            
            for subcategory_name, subcategory_data in sorted(category_data['subcategories'].items()):
                subcategory_start_row = current_row
                subcategory_total = subcategory_data['total']
                
                for classification in subcategory_data['classifications']:
                    row = table.rows[current_row]
                    
                    row.cells[0].text = ""
                    row.cells[1].text = ""
                    
                    # Classification AR
                    row.cells[2].text = classification.get('classification_name', 'N/A')
                    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                    row.cells[2].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                    row.cells[2].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                    center_cell_content(row.cells[2])
                    
                    # Classification EN
                    row.cells[3].text = classification.get('classification_name_en', 'N/A')
                    row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                    center_cell_content(row.cells[3])
                    
                    # Total
                    row.cells[4].text = str(classification.get('total_count', 0))
                    row.cells[4].paragraphs[0].runs[0].bold = True
                    center_cell_content(row.cells[4])
                    
                    # Severity
                    low_val = classification.get('low_count', 0)
                    med_val = classification.get('medium_count', 0)
                    high_val = classification.get('high_count', 0)
                    
                    row.cells[5].text = str(low_val) if low_val > 0 else ""
                    center_cell_content(row.cells[5])
                    set_cell_shading(row.cells[5], 'C6EFCE' if low_val > 0 else 'F2F2F2')
                    
                    row.cells[6].text = str(med_val) if med_val > 0 else ""
                    center_cell_content(row.cells[6])
                    set_cell_shading(row.cells[6], 'FFEB9C' if med_val > 0 else 'F2F2F2')
                    
                    row.cells[7].text = str(high_val) if high_val > 0 else ""
                    center_cell_content(row.cells[7])
                    set_cell_shading(row.cells[7], 'FFC7CE' if high_val > 0 else 'F2F2F2')
                    
                    row.cells[8].text = ""  # Spacer column
                    
                    # Prevention
                    prev_yes = classification.get('preventive_yes_count', 0)
                    prev_no = classification.get('preventive_no_count', 0)
                    
                    row.cells[9].text = str(prev_yes)
                    center_cell_content(row.cells[9])
                    if prev_yes > 0:
                        set_cell_shading(row.cells[9], 'C6EFCE')
                    
                    row.cells[10].text = str(prev_no)
                    center_cell_content(row.cells[10])
                    if prev_no > 0:
                        set_cell_shading(row.cells[10], 'FFE6E6')
                    
                    current_row += 1
                
                # Merge subcategory
                if current_row - 1 >= subcategory_start_row:
                    table.rows[subcategory_start_row].cells[1].text = f"{subcategory_name}\n(n={subcategory_total})"
                    table.rows[subcategory_start_row].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                    table.rows[subcategory_start_row].cells[1].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                    table.rows[subcategory_start_row].cells[1].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                    table.rows[subcategory_start_row].cells[1].paragraphs[0].runs[0].bold = True
                    table.rows[subcategory_start_row].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_shading(table.rows[subcategory_start_row].cells[1], 'E7E6E6')
                    
                    if current_row - 1 > subcategory_start_row:
                        try:
                            for row_idx in range(subcategory_start_row + 1, current_row):
                                if row_idx < len(table.rows):
                                    table.rows[subcategory_start_row].cells[1].merge(table.rows[row_idx].cells[1])
                        except Exception as e:
                            print(f"[FORMATTER] Warning: subcategory merge: {e}")
            
            # Merge category
            if current_row - 1 >= category_start_row:
                table.rows[category_start_row].cells[0].text = f"{category_name}\n(n={category_total})"
                table.rows[category_start_row].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
                table.rows[category_start_row].cells[0].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                table.rows[category_start_row].cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                table.rows[category_start_row].cells[0].paragraphs[0].runs[0].bold = True
                table.rows[category_start_row].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(table.rows[category_start_row].cells[0], 'D9E2F3')
                
                if current_row - 1 > category_start_row:
                    try:
                        for row_idx in range(category_start_row + 1, current_row):
                            if row_idx < len(table.rows):
                                table.rows[category_start_row].cells[0].merge(table.rows[row_idx].cells[0])
                    except Exception as e:
                        print(f"[FORMATTER] Warning: category merge: {e}")


def _create_hierarchical_table_rtl(doc: Document, hierarchy: Dict[str, Any], language: str):
    """
    Create the main hierarchical classification table with RTL layout and 2-row header.
    
    Table structure (11 columns, RTL order):
    Row 1 (Main Headers):
    - Problem Domain | Problem Category | Sub-Category | Classification (عربي) | Classification (English) | Total | [Severity - 3 cols merged] | [Prevention - 2 cols merged]
    
    Row 2 (Sub-Headers):
    - [empty x6] | LOW | MEDIUM | HIGH | YES | NO
    
    RTL Order: Right-to-left for Arabic readers
    """
    
    # Calculate total rows needed (header=2 + data rows)
    total_rows = 2  # 2-row header
    for domain_name, domain_data in hierarchy.items():
        for category_name, category_data in domain_data['categories'].items():
            for subcategory_name, subcategory_data in category_data['subcategories'].items():
                total_rows += len(subcategory_data['classifications'])
    
    if total_rows == 2:
        doc.add_paragraph("لا توجد بيانات متاحة (No classification data available).")
        return
    
    # Create table with 11 columns (0-10)
    table = doc.add_table(rows=total_rows, cols=11)
    table.style = 'Table Grid'
    
    # Set RTL table direction
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Enable RTL
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)
    
    # ============================================================
    # HEADER SETUP - FILL BOTH ROWS FIRST, THEN MERGE
    # ============================================================
    header_row1 = table.rows[0]
    header_row2 = table.rows[1]
    
    # ============================================================
    # ROW 1: Set text for all main header cells (before any merging)
    # ============================================================
    headers_main = [
        "Problem Domain\nمجال المشكلة",
        "Problem Category\nفئة المشكلة",
        "Sub-Category\nالفئة الفرعية",
        "التصنيف عربي\nClassification AR",
        "التصنيف إنجليزي\nClassification EN",
        "Total\nالمجموع",
        "الشدة Severity",  # Col 6 (will span 6-8)
        "",  # Col 7 (will merge with 6)
        "",  # Col 8 (will merge with 6)
        "الإجراءات الوقائية\nPrevention Action",  # Col 9 (will span 9-10)
        ""  # Col 10 (will merge with 9)
    ]
    
    for idx, header_text in enumerate(headers_main):
        cell = header_row1.cells[idx]
        if header_text:  # Only set text if not empty
            cell.text = header_text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10 if idx == 6 else 9)
            cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')  # Dark blue for all
    
    # ============================================================
    # ROW 2: Set text for all sub-header cells (before any merging)
    # ============================================================
    subheaders = [
        "",  # Col 0 (will merge with row1)
        "",  # Col 1 (will merge with row1)
        "",  # Col 2 (will merge with row1)
        "",  # Col 3 (will merge with row1)
        "",  # Col 4 (will merge with row1)
        "",  # Col 5 (will merge with row1)
        "LOW\nمنخفضة",  # Col 6
        "MEDIUM\nمتوسطة",  # Col 7
        "HIGH\nعالية",  # Col 8
        "YES\nنعم",  # Col 9
        "NO\nلا"  # Col 10
    ]
    
    for idx, subheader_text in enumerate(subheaders):
        cell = header_row2.cells[idx]
        if subheader_text:  # Only set text if not empty
            cell.text = subheader_text
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '5B9BD5')  # Lighter blue
        else:
            set_cell_shading(cell, '4472C4')  # Dark blue for merged areas
    
    # ============================================================
    # NOW DO THE MERGING (after text is set)
    # ============================================================
    
    # Merge columns 0-5 vertically (row 1 with row 2)
    for col_idx in range(6):
        try:
            header_row1.cells[col_idx].merge(header_row2.cells[col_idx])
        except Exception as e:
            print(f"[FORMATTER] Warning: Could not merge col {col_idx}: {e}")
    
    # Merge Severity header horizontally in row 1 (cols 6, 7, 8)
    try:
        header_row1.cells[6].merge(header_row1.cells[7])
        header_row1.cells[6].merge(header_row1.cells[8])
    except Exception as e:
        print(f"[FORMATTER] Warning: Could not merge severity header: {e}")
    
    # Merge Prevention header horizontally in row 1 (cols 9, 10)
    try:
        header_row1.cells[9].merge(header_row1.cells[10])
    except Exception as e:
        print(f"[FORMATTER] Warning: Could not merge prevention header: {e}")
    
    # ============================================================
    # DATA ROWS WITH MERGED CELLS
    # ============================================================
    current_row = 2  # Start after 2-row header
    
    for domain_name, domain_data in sorted(hierarchy.items()):
        domain_start_row = current_row
        domain_total = domain_data['total']
        
        for category_name, category_data in sorted(domain_data['categories'].items()):
            category_start_row = current_row
            category_total = category_data['total']
            
            for subcategory_name, subcategory_data in sorted(category_data['subcategories'].items()):
                subcategory_start_row = current_row
                subcategory_total = subcategory_data['total']
                
                # Add all classifications under this subcategory
                for classification in subcategory_data['classifications']:
                    row = table.rows[current_row]
                    
                    # Column 0: Domain (will be merged later)
                    row.cells[0].text = ""
                    
                    # Column 1: Category (will be merged later)
                    row.cells[1].text = ""
                    
                    # Column 2: Sub-Category (will be merged later)
                    row.cells[2].text = ""
                    
                    # Column 3: Classification (Arabic)
                    classification_ar = classification.get('classification_name', 
                                                           classification.get('classification_name_ar', 'N/A'))
                    row.cells[3].text = classification_ar
                    row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                    row.cells[3].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                    row.cells[3].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Column 4: Classification (English)
                    classification_en = classification.get('classification_name_en', 
                                                          classification.get('classification_name', 'N/A'))
                    row.cells[4].text = classification_en
                    row.cells[4].paragraphs[0].runs[0].font.size = Pt(9)
                    row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Column 5: Total
                    row.cells[5].text = str(classification.get('total_count', 0))
                    row.cells[5].paragraphs[0].runs[0].bold = True
                    center_cell_content(row.cells[5])
                    
                    # Columns 6-8: Severity breakdown (empty cell if 0)
                    low_val = classification.get('low_count', 0)
                    med_val = classification.get('medium_count', 0)
                    high_val = classification.get('high_count', 0)
                    
                    row.cells[6].text = str(low_val) if low_val > 0 else ""
                    center_cell_content(row.cells[6])
                    if low_val > 0:
                        set_cell_shading(row.cells[6], 'C6EFCE')  # Light green
                    else:
                        set_cell_shading(row.cells[6], 'F2F2F2')  # Light gray for empty
                    
                    row.cells[7].text = str(med_val) if med_val > 0 else ""
                    center_cell_content(row.cells[7])
                    if med_val > 0:
                        set_cell_shading(row.cells[7], 'FFEB9C')  # Light yellow
                    else:
                        set_cell_shading(row.cells[7], 'F2F2F2')  # Light gray for empty
                    
                    row.cells[8].text = str(high_val) if high_val > 0 else ""
                    center_cell_content(row.cells[8])
                    if high_val > 0:
                        set_cell_shading(row.cells[8], 'FFC7CE')  # Light red
                    else:
                        set_cell_shading(row.cells[8], 'F2F2F2')  # Light gray for empty
                    
                    # Columns 9-10: Prevention Action (now properly separated)
                    prev_yes = classification.get('preventive_yes_count', 0)
                    prev_no = classification.get('preventive_no_count', 0)
                    
                    row.cells[9].text = str(prev_yes)
                    center_cell_content(row.cells[9])
                    if prev_yes > 0:
                        set_cell_shading(row.cells[9], 'C6EFCE')  # Light green
                    
                    row.cells[10].text = str(prev_no)
                    center_cell_content(row.cells[10])
                    if prev_no > 0:
                        set_cell_shading(row.cells[10], 'FFE6E6')  # Light red
                    
                    current_row += 1
                
                # Merge Sub-Category cells
                subcategory_end_row = current_row - 1
                if subcategory_end_row >= subcategory_start_row:
                    table.rows[subcategory_start_row].cells[2].text = f"{subcategory_name}\n(n={subcategory_total})"
                    table.rows[subcategory_start_row].cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                    table.rows[subcategory_start_row].cells[2].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                    table.rows[subcategory_start_row].cells[2].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                    table.rows[subcategory_start_row].cells[2].paragraphs[0].runs[0].bold = True
                    table.rows[subcategory_start_row].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_shading(table.rows[subcategory_start_row].cells[2], 'E7E6E6')
                    
                    if subcategory_end_row > subcategory_start_row:
                        try:
                            for row_idx in range(subcategory_start_row + 1, subcategory_end_row + 1):
                                if row_idx < len(table.rows):  # Safety check
                                    table.rows[subcategory_start_row].cells[2].merge(table.rows[row_idx].cells[2])
                        except Exception as e:
                            print(f"[FORMATTER] Warning: Could not merge subcategory cells: {e}")
            
            # Merge Category cells
            category_end_row = current_row - 1
            if category_end_row >= category_start_row:
                table.rows[category_start_row].cells[1].text = f"{category_name}\n(n={category_total})"
                table.rows[category_start_row].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
                table.rows[category_start_row].cells[1].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                table.rows[category_start_row].cells[1].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                table.rows[category_start_row].cells[1].paragraphs[0].runs[0].bold = True
                table.rows[category_start_row].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(table.rows[category_start_row].cells[1], 'D9E2F3')
                
                if category_end_row > category_start_row:
                    try:
                        for row_idx in range(category_start_row + 1, category_end_row + 1):
                            if row_idx < len(table.rows):  # Safety check
                                table.rows[category_start_row].cells[1].merge(table.rows[row_idx].cells[1])
                    except Exception as e:
                        print(f"[FORMATTER] Warning: Could not merge category cells: {e}")
        
        # Merge Domain cells
        domain_end_row = current_row - 1
        if domain_end_row >= domain_start_row:
            table.rows[domain_start_row].cells[0].text = f'"{domain_name}"\n(n={domain_total})'
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[domain_start_row].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(table.rows[domain_start_row].cells[0], 'BDD7EE')
            
            if domain_end_row > domain_start_row:
                try:
                    for row_idx in range(domain_start_row + 1, domain_end_row + 1):
                        if row_idx < len(table.rows):  # Safety check
                            table.rows[domain_start_row].cells[0].merge(table.rows[row_idx].cells[0])
                except Exception as e:
                    print(f"[FORMATTER] Warning: Could not merge domain cells: {e}")
    
    # Set column widths for landscape orientation (11 columns total)
    try:
        table.columns[0].width = Inches(1.1)   # Domain
        table.columns[1].width = Inches(1.2)   # Category
        table.columns[2].width = Inches(1.2)   # Sub-Category
        table.columns[3].width = Inches(1.3)   # Classification (AR)
        table.columns[4].width = Inches(1.3)   # Classification (EN)
        table.columns[5].width = Inches(0.6)   # Total
        table.columns[6].width = Inches(0.6)   # Low
        table.columns[7].width = Inches(0.7)   # Medium
        table.columns[8].width = Inches(0.6)   # High
        table.columns[9].width = Inches(0.6)   # Yes
        table.columns[10].width = Inches(0.6)  # No
    except:
        pass  # Column width adjustment may fail, continue anyway


def generate_seasonal_pdf_report(
    seasonal_data: Dict[str, Any],
    language: str = "en"
) -> bytes:
    """
    Generate a PDF document for a seasonal report.
    
    Currently returns Word format - PDF conversion can be added later.
    
    Args:
        seasonal_data: Seasonal report data from orchestrator
        language: Language for the report (en or ar)
    
    Returns:
        Bytes of the generated document
    """
    # For now, return Word format
    # TODO: Add PDF conversion using reportlab or weasyprint
    return generate_seasonal_word_report(seasonal_data, language)
    
    severity_para = doc.add_paragraph()
    severity_para.add_run(f"Low Severity: ").bold = True
    severity_para.add_run(f"{header.get('low_severity_count', 0)} cases\n")
    
    severity_para.add_run(f"Medium Severity: ").bold = True
    severity_para.add_run(f"{header.get('medium_severity_count', 0)} cases\n")
    
    severity_para.add_run(f"High Severity: ").bold = True
    severity_para.add_run(f"{header.get('high_severity_count', 0)} cases\n")
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def generate_seasonal_pdf_report(
    seasonal_data: Dict[str, Any],
    language: str = "en"
) -> bytes:
    """
    Generate a PDF for a seasonal report.
    Currently converts Word to PDF.
    
    Args:
        seasonal_data: Seasonal report data from orchestrator
        language: Language for the report (en or ar)
    
    Returns:
        Bytes of the generated PDF document
    """
    # For now, generate Word and note that PDF conversion would happen here
    # In production, use docx2pdf or similar
    word_bytes = generate_seasonal_word_report(seasonal_data, language)
    
    # TODO: Convert Word to PDF using docx2pdf or similar library
    # For now, return Word bytes with a note
    return word_bytes  # Placeholder - replace with actual PDF conversion


# ============================================================================
# VISUALIZATION FUNCTIONS FOR COMPARATIVE REPORTS
# ============================================================================

# ============================================================================
# ARABIC TEXT SHAPING
# matplotlib does not natively shape Arabic — characters appear disconnected.
# _ar() reshapes text and applies BiDi display order before passing to plt.
# ============================================================================
try:
    import arabic_reshaper
    from bidi.algorithm import get_display as _bidi_display
    _ARABIC_SHAPING = True
except ImportError:
    _ARABIC_SHAPING = False


def _ar(text: str) -> str:
    """Shape Arabic text for correct matplotlib rendering (joined letters, RTL)."""
    if not text or not _ARABIC_SHAPING:
        return text
    try:
        return _bidi_display(arabic_reshaper.reshape(text))
    except Exception:
        return text


_ARABIC_FONT_PATH = 'C:/Windows/Fonts/trado.ttf'

# Register the font once at import time so matplotlib can find it by name.
if os.path.exists(_ARABIC_FONT_PATH):
    try:
        fm.fontManager.addfont(_ARABIC_FONT_PATH)
        _ARABIC_FONT_FAMILY = fm.FontProperties(fname=_ARABIC_FONT_PATH).get_name()
    except Exception:
        _ARABIC_FONT_FAMILY = None
else:
    _ARABIC_FONT_FAMILY = None


def _configure_arabic_matplotlib():
    """Set up matplotlib for Arabic text — font + unicode minus fix."""
    if _ARABIC_FONT_FAMILY:
        plt.rcParams['font.family'] = _ARABIC_FONT_FAMILY
    plt.rcParams['axes.unicode_minus'] = False


def _split_bilingual_title(title: str):
    """
    Split a bilingual title at '|' into (en_part, ar_part).
    Handles both 'English | Arabic' and 'Arabic | English' ordering by
    detecting Arabic Unicode characters (U+0600–U+06FF) in each half.
    Returns (title, title) when no '|' separator is present.
    """
    if '|' not in title:
        return title, title
    first, second = title.split('|', 1)
    first, second = first.strip(), second.strip()
    first_has_arabic = any('؀' <= c <= 'ۿ' for c in first)
    if first_has_arabic:
        # Format is "Arabic | English"
        return second, first   # (en_part, ar_part)
    else:
        # Format is "English | Arabic"
        return first, second   # (en_part, ar_part)


# ============================================================================
# CHART GRID LAYOUT HELPERS
# Replaces stacked individual paragraphs with a structured 2-column table.
# ============================================================================

def _remove_table_borders(table) -> None:
    """Remove all visible borders from a docx table."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _add_chart_grid(doc: Document, chart_pairs: list, chart_width_inches: float = 4.5) -> None:
    """
    Insert a list of (image_buffer, caption_text) pairs into the document
    using a borderless 2-column table so charts sit side-by-side instead of
    being stacked as individual paragraphs.

    Pairs fill left→right, row by row.  Odd-count lists leave the last cell empty.
    """
    if not chart_pairs:
        return

    rows = [chart_pairs[i:i + 2] for i in range(0, len(chart_pairs), 2)]

    for row_items in rows:
        n_cols = len(row_items)
        table = doc.add_table(rows=1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        for col_idx, (img_buf, caption_text) in enumerate(row_items):
            cell = table.rows[0].cells[col_idx]

            # Image paragraph
            img_para = cell.paragraphs[0]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.space_before = Pt(0)
            img_para.paragraph_format.space_after = Pt(2)
            img_run = img_para.add_run()
            img_run.add_picture(img_buf, width=Inches(chart_width_inches))

            # Caption paragraph
            cap_para = cell.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para.paragraph_format.space_before = Pt(0)
            cap_para.paragraph_format.space_after = Pt(4)
            cap_run = cap_para.add_run(caption_text)
            cap_run.font.bold = True
            cap_run.font.size = Pt(9)
            cap_run.font.name = 'Traditional Arabic'
            cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

        doc.add_paragraph()  # single row spacer


def _generate_spider_chart(labels: List[str], prev_values: List[float],
                           curr_values: List[float], title: str,
                           prev_label: str, curr_label: str) -> io.BytesIO:
    """
    2-quarter spider chart.
    Converts the list-pair format to the unified dict format and delegates
    to _generate_nquarter_spider_chart.

    Args:
        labels:       Dimension labels (e.g. domain names).
        prev_values:  Previous-period counts.
        curr_values:  Current-period counts.
        title:        Chart title (used as both Arabic and English label).
        prev_label:   Legend label for previous period.
        curr_label:   Legend label for current period.

    Returns:
        BytesIO PNG image buffer.
    """
    if not labels:
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.text(0.5, 0.5, 'No Data Available', ha='center', va='center', fontsize=14)
        ax.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        return buf

    # Convert list-pair → dict format expected by unified engine
    data = {labels[i]: [int(prev_values[i]), int(curr_values[i])]
            for i in range(len(labels))}
    periods = [prev_label, curr_label]

    # Split mixed "English | Arabic" titles so the unified engine gets separate parts.
    title_en_part, title_ar_part = _split_bilingual_title(title)

    return _generate_nquarter_spider_chart(
        data=data,
        periods=periods,
        title_ar=title_ar_part,
        title_en=title_en_part,
        max_items=len(labels)  # 2Q callers already pre-filter to ≤10
    )


def _generate_diverging_bar_chart(labels: List[str], changes: List[float], 
                                  title: str) -> io.BytesIO:
    """
    Generate a diverging bar chart showing positive/negative changes.
    
    Args:
        labels: List of category labels
        changes: List of change values (can be positive or negative)
        title: Chart title
    
    Returns:
        BytesIO buffer containing PNG image
    """
    _configure_arabic_matplotlib()
    
    if len(labels) == 0:
        # Return empty chart
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.text(0.5, 0.5, 'No Data Available', ha='center', va='center', fontsize=14)
        ax.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        return buf
    
    # Create figure
    fig, ax = plt.subplots(figsize=(12, max(6, len(labels) * 0.4)))
    
    # Color code: green for negative (improvement), red for positive (worsening)
    colors = ['#70AD47' if c < 0 else '#C5504B' if c > 0 else '#BFBFBF' for c in changes]
    
    # Create horizontal bar chart
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, changes, color=colors, alpha=0.8)
    
    # Add value labels on bars
    for i, (bar, change) in enumerate(zip(bars, changes)):
        width = bar.get_width()
        label_x = width + (1 if width > 0 else -1)
        ax.text(label_x, i, f'{change:+.0f}', 
               ha='left' if width > 0 else 'right', 
               va='center', fontsize=10, weight='bold')
    
    # Customize — shape Arabic labels and title
    ax.set_yticks(y_pos)
    ax.set_yticklabels([_ar(l) for l in labels], fontsize=10)
    ax.set_xlabel('Change (← Decrease | Increase →)', fontsize=11)
    title_en_part, title_ar_part = _split_bilingual_title(title)
    ax.set_title(f"{_ar(title_ar_part)}\n{title_en_part}", fontsize=14, weight='bold', pad=15)
    ax.axvline(0, color='black', linewidth=0.8)
    ax.grid(axis='x', alpha=0.3)
    
    # Tight layout
    plt.tight_layout()
    
    # Save to buffer
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def _generate_heatmap(data: np.ndarray, row_labels: List[str], 
                      col_labels: List[str], title: str) -> io.BytesIO:
    """
    Generate a heatmap showing data intensity across two dimensions.
    
    Args:
        data: 2D numpy array of values
        row_labels: Labels for rows
        col_labels: Labels for columns
        title: Chart title
    
    Returns:
        BytesIO buffer containing PNG image
    """
    _configure_arabic_matplotlib()
    
    if data.size == 0:
        # Return empty chart
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.text(0.5, 0.5, 'No Data Available', ha='center', va='center', fontsize=14)
        ax.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        return buf
    
    # Create figure
    fig, ax = plt.subplots(figsize=(12, max(6, len(row_labels) * 0.5)))
    
    # Create heatmap
    im = ax.imshow(data, cmap='RdYlGn_r', aspect='auto')
    
    # Set ticks and labels
    ax.set_xticks(np.arange(len(col_labels)))
    ax.set_yticks(np.arange(len(row_labels)))
    ax.set_xticklabels(col_labels, fontsize=10)
    ax.set_yticklabels(row_labels, fontsize=10)
    
    # Rotate the tick labels for better readability
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right", rotation_mode="anchor")
    
    # Add text annotations
    for i in range(len(row_labels)):
        for j in range(len(col_labels)):
            text = ax.text(j, i, f'{data[i, j]:.0f}',
                         ha="center", va="center", color="black", fontsize=9, weight='bold')
    
    # Add colorbar
    cbar = plt.colorbar(im, ax=ax)
    cbar.set_label('Count', rotation=270, labelpad=15)
    
    # Set title
    ax.set_title(title, fontsize=14, weight='bold', pad=15)
    
    # Tight layout
    plt.tight_layout()
    
    # Save to buffer
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


# ============================================================================
# GRAPH USEFULNESS VALIDATORS
# ============================================================================

def _is_spider_2q_useful(labels: List[str], prev_values: List[float], curr_values: List[float]) -> Tuple[bool, str]:
    """
    Decide whether a 2-quarter spider/radar chart has analytical value.
    Returns (should_render, reason_if_skipped).

    A chart is skipped only when:
    - Fewer than 3 axes (radar charts need ≥3 to be meaningful)
    - BOTH series are entirely zero (nothing to show at all)

    A chart with one empty series is still rendered — it shows the shape of the
    non-empty period against a zero baseline, which IS analytically informative.
    """
    if len(labels) < 3:
        return False, f"only {len(labels)} dimension(s) - radar chart needs >=3"
    if not any(v > 0 for v in prev_values) and not any(v > 0 for v in curr_values):
        return False, "both periods have no data"
    return True, ""


def _is_spider_nq_useful(data: Dict[str, List[int]], max_items: int = 8) -> Tuple[bool, str]:
    """
    Decide whether a multi-quarter (3Q/4Q) spider chart has analytical value.
    data maps item names → list of per-quarter counts.
    Returns (should_render, reason_if_skipped).

    Skips when:
    - Fewer than 3 unique items (radar charts need ≥3 axes to be meaningful)
    - ALL values across ALL quarters are zero (nothing to plot)

    Sparse data (some quarters zero) still renders — zero axes are informative.
    """
    if not data:
        return False, "no data"
    renderable_count = min(len(data), max_items)
    if renderable_count < 3:
        return False, f"only {renderable_count} dimension(s) - radar chart needs >=3"
    if not any(sum(v) > 0 for v in data.values()):
        return False, "all values are zero across all quarters"
    return True, ""


def _is_bar_chart_useful(labels: List[str], changes: List[float]) -> Tuple[bool, str]:
    """
    Decide whether a diverging bar chart has analytical value.
    Returns (should_render, reason_if_skipped).
    """
    if not labels:
        return False, "no data"
    if not any(c != 0 for c in changes):
        return False, "all period-over-period changes are zero — nothing to compare"
    return True, ""


def _extract_nq_changes(
    data: Dict[str, List[int]],
    max_items: int = 10
) -> Tuple[List[str], List[float]]:
    """
    Derive labels and net changes (last_quarter - first_quarter) from N-quarter
    dict data.  Used to build diverging bar charts for 3Q/4Q comparisons.

    Args:
        data:      Dict mapping item names → list of N per-quarter counts.
        max_items: Keep only the top-N items by absolute change.

    Returns:
        (labels, changes) ready for _generate_diverging_bar_chart.
    """
    items = [
        (name, float(values[-1] - values[0]))
        for name, values in data.items()
        if values and len(values) >= 2
    ]
    items.sort(key=lambda x: abs(x[1]), reverse=True)
    items = items[:max_items]
    labels  = [it[0] for it in items]
    changes = [it[1] for it in items]
    return labels, changes


def _extract_domain_data(hierarchy: Dict) -> Tuple[List[str], List[int]]:
    """Extract domain names and counts from hierarchy"""
    domains = []
    counts = []
    
    for domain_name, domain_data in hierarchy.items():
        domains.append(domain_name)
        counts.append(domain_data.get('total', 0))
    
    return domains, counts


def _extract_category_data(hierarchy: Dict) -> Tuple[List[str], List[int]]:
    """Extract category names and counts from hierarchy"""
    categories = []
    counts = []
    
    for domain_name, domain_data in hierarchy.items():
        for cat_name, cat_data in domain_data.get('categories', {}).items():
            full_name = f"{domain_name} - {cat_name}"
            categories.append(full_name)
            counts.append(cat_data.get('total', 0))
    
    return categories, counts


def _extract_subcategory_data(hierarchy: Dict) -> Tuple[List[str], List[int]]:
    """Extract subcategory names and counts from hierarchy"""
    subcategories = []
    counts = []
    
    for domain_name, domain_data in hierarchy.items():
        for cat_name, cat_data in domain_data.get('categories', {}).items():
            for subcat_name, subcat_data in cat_data.get('subcategories', {}).items():
                full_name = f"{cat_name} - {subcat_name}"
                subcategories.append(full_name)
                counts.append(subcat_data.get('total', 0))
    
    return subcategories, counts


def _add_per_quarter_policy_sections(
    doc: Document,
    quarters: List[Tuple[str, Dict[str, Any], Any, List]]
) -> None:
    """
    Render a per-quarter policy compliance detail block.
    quarters: list of (period_label, header_dict, policy_snapshot_or_None, classification_stats)
    """
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading.paragraph_format.right_to_left = True
    h_run = heading.add_run("📋 تقييم الامتثال للسياسة لكل فصل | Per-Quarter Policy Compliance")
    h_run.font.bold = True
    h_run.font.size = Pt(13)
    h_run.font.name = 'Traditional Arabic'
    h_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    for period_label, header, policy_snapshot, classification_stats in quarters:
        doc.add_paragraph()
        lbl = doc.add_paragraph()
        lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lbl.paragraph_format.right_to_left = True
        lbl_run = lbl.add_run(f"◄ {period_label}")
        lbl_run.font.bold = True
        lbl_run.font.size = Pt(12)
        lbl_run.font.name = 'Traditional Arabic'
        lbl_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        lbl_run.font.color.rgb = RGBColor(68, 114, 196)
        _add_policy_compliance_section(doc, header, policy_snapshot, classification_stats)


def generate_comparative_seasonal_word_report(
    current_data: Dict[str, Any],
    previous_data: Dict[str, Any],
    language: str = "en"
) -> bytes:
    """
    Generate a comparative Word document showing current season vs previous season.
    
    The report mirrors the regular seasonal report structure but adds side-by-side
    comparison data in each table, showing both periods with delta indicators.
    
    Args:
        current_data: Current seasonal report data from orchestrator
        previous_data: Previous seasonal report data from orchestrator (may have zero data)
        language: Language for the report (en or ar)
    
    Returns:
        Bytes of the generated comparative Word document
    """
    # Utility functions
    def _safe(v):
        """Convert dimension values to int (python-docx requirement)"""
        return int(v)
    
    doc = Document()
    
    # ============================================================
    # DOCUMENT SETUP - A4 LANDSCAPE
    # ============================================================
    section = doc.sections[0]
    section.page_height = _safe(Mm(210))
    section.page_width = _safe(Mm(297))
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = _safe(Mm(15))
    section.right_margin = _safe(Mm(15))
    section.top_margin = _safe(Mm(15))
    section.bottom_margin = _safe(Mm(15))
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Traditional Arabic'
    font.size = Pt(11)
    
    # Extract data from both seasons
    current_header = current_data.get("header", {})
    previous_header = previous_data.get("header", {})
    
    current_period = current_header.get('period', 'N/A')
    previous_period = previous_header.get('period', 'N/A')
    
    # ============================================================
    # HEADER - LOGO (TOP RIGHT)
    # ============================================================
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            section.header_distance = Inches(0.1)
            header_section = section.header
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[FORMATTER] Could not add logo: {e}")
    
    # ============================================================
    # TITLE SECTION (ARABIC) - COMPARATIVE
    # ============================================================
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("تقرير المقارنة الموسمية | Seasonal Comparison Report")
    title_run.font.size = int(Pt(21))
    title_run.font.bold = True
    title_run.font.name = 'Traditional Arabic'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = int(Pt(3))
    
    # Subtitle with both periods
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(f"{current_period} مقابل {previous_period}")
    subtitle_run.font.size = int(Pt(16))
    subtitle_run.font.bold = True
    subtitle_run.font.name = 'Traditional Arabic'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.space_after = int(Pt(6))
    
    # Organization info
    orgunit_type = current_header.get('orgunit_type', 0)
    type_names = {0: "المستشفى", 1: "الإدارة", 2: "الدائرة", 3: "القسم"}
    type_name = type_names.get(orgunit_type, "الوحدة التنظيمية")
    
    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    org_run = org_para.add_run(f"الوحدة التنظيمية: {current_header.get('orgunit_name', 'N/A')} • النوع: {type_name}")
    org_run.font.size = int(Pt(14))
    org_run.font.bold = True
    org_run.font.name = 'Traditional Arabic'
    org_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    doc.add_paragraph()  # Spacer
    
    # ============================================================
    # SUMMARY COMPARISON TABLE
    # ============================================================
    
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary_heading.paragraph_format.right_to_left = True
    
    sh_run = summary_heading.add_run("📊 مقارنة الإحصائيات | Statistics Comparison")
    sh_run.font.bold = True
    sh_run.font.size = Pt(14)
    sh_run.font.name = 'Traditional Arabic'
    sh_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    doc.add_paragraph()
    
    # Create comparison table (3 columns: Category | Previous | Current) - NO DELTA
    comp_table = doc.add_table(rows=8, cols=3)
    comp_table.style = 'Table Grid'
    
    # Headers (RTL: Current, Previous, Category)
    headers_rtl = [current_period, previous_period, "الفئة\nCategory"]
    for idx, header_text in enumerate(headers_rtl):
        cell = comp_table.rows[0].cells[idx]
        cell.text = header_text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        center_cell_content(cell)
        set_cell_shading(cell, '4472C4')
    
    # Data rows
    current_total = current_header.get('total_cases', 0)
    previous_total = previous_header.get('total_cases', 0)
    
    metrics = [
        ("Total Cases | المجموع الكلي", previous_total, current_total),
        ("Clinical | السريرية", previous_header.get('clinical_domain_count', 0), current_header.get('clinical_domain_count', 0)),
        ("Management | الإدارية", previous_header.get('management_domain_count', 0), current_header.get('management_domain_count', 0)),
        ("Relational | العلائقية", previous_header.get('relational_domain_count', 0), current_header.get('relational_domain_count', 0)),
        ("Low Severity | منخفضة", previous_header.get('low_severity_count', 0), current_header.get('low_severity_count', 0)),
        ("Medium Severity | متوسطة", previous_header.get('medium_severity_count', 0), current_header.get('medium_severity_count', 0)),
        ("High Severity | عالية", previous_header.get('high_severity_count', 0), current_header.get('high_severity_count', 0))
    ]
    
    for idx, (label, prev_val, curr_val) in enumerate(metrics, start=1):
        row = comp_table.rows[idx]
        
        # Category (rightmost - col 2)
        row.cells[2].text = label
        row.cells[2].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        row.cells[2].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Previous value (col 1)
        row.cells[1].text = str(prev_val)
        center_cell_content(row.cells[1])
        
        # Current value (col 0 - leftmost)
        row.cells[0].text = str(curr_val)
        center_cell_content(row.cells[0])
    
    doc.add_paragraph()
    
    # ============================================================
    # HIERARCHICAL CLASSIFICATION COMPARISON BY DOMAIN
    # ============================================================
    
    # Build hierarchies for table generation (moved before graphs)
    current_hierarchy = _build_hierarchy(current_data.get("classification_stats", []))
    previous_hierarchy = _build_hierarchy(previous_data.get("classification_stats", []))
    
    # Merge all domains from both seasons
    all_domains = set(list(current_hierarchy.keys()) + list(previous_hierarchy.keys()))
    
    # Create comparative tables for each domain
    _create_comparative_hierarchical_tables_by_domain(
        doc, 
        current_hierarchy, 
        previous_hierarchy, 
        current_period,
        previous_period,
        all_domains,
        language
    )
    
    doc.add_paragraph()
    
    # ============================================================
    # POLICY COMPLIANCE COMPARISON
    # ============================================================
    
    policy_heading = doc.add_paragraph()
    policy_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    policy_heading.paragraph_format.right_to_left = True
    
    ph_run = policy_heading.add_run("📊 مقارنة الامتثال للسياسة | Policy Compliance Comparison")
    ph_run.font.bold = True
    ph_run.font.size = Pt(14)
    ph_run.font.name = 'Traditional Arabic'
    ph_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    doc.add_paragraph()
    
    current_compliant = current_header.get('is_compliant', False)
    previous_compliant = previous_header.get('is_compliant', False)
    
    compliance_para = doc.add_paragraph()
    compliance_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    compliance_para.paragraph_format.right_to_left = True
    
    cp_text = f"{previous_period}: {'✓ مطابق' if previous_compliant else '✗ غير مطابق'}  |  {current_period}: {'✓ مطابق' if current_compliant else '✗ غير مطابق'}"
    cp_run = compliance_para.add_run(cp_text)
    cp_run.font.size = Pt(12)
    cp_run.font.bold = True
    cp_run.font.name = 'Traditional Arabic'
    cp_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    # Per-quarter detailed policy compliance tables
    _add_per_quarter_policy_sections(doc, [
        (previous_period, previous_data['header'], previous_data.get('policy_snapshot'), previous_data.get('classification_stats', [])),
        (current_period,  current_data['header'],  current_data.get('policy_snapshot'),  current_data.get('classification_stats', [])),
    ])

    doc.add_paragraph()

    # ============================================================
    # PAGE BREAK BEFORE VISUALIZATION SECTION
    # ============================================================

    doc.add_page_break()
    
    # ============================================================
    # VISUALIZATION SECTION - GRAPHS AT THE END
    # ============================================================
    
    try:
        # Add visual analysis heading
        visual_heading = doc.add_paragraph()
        visual_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vh_run = visual_heading.add_run("📊 التحليل المرئي | Visual Analysis")
        vh_run.font.bold = True
        vh_run.font.size = Pt(16)
        vh_run.font.name = 'Traditional Arabic'
        vh_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        
        doc.add_paragraph()
        
        # =================== DOMAIN LEVEL CHARTS ===================
        
        # Extract domain data
        prev_domains, prev_domain_counts = _extract_domain_data(previous_hierarchy)
        curr_domains, curr_domain_counts = _extract_domain_data(current_hierarchy)
        
        # Merge domain lists
        all_domain_names = sorted(set(prev_domains + curr_domains))
        
        # Align data
        prev_domain_values = []
        curr_domain_values = []
        domain_changes = []
        
        for domain in all_domain_names:
            prev_val = prev_domain_counts[prev_domains.index(domain)] if domain in prev_domains else 0
            curr_val = curr_domain_counts[curr_domains.index(domain)] if domain in curr_domains else 0
            prev_domain_values.append(prev_val)
            curr_domain_values.append(curr_val)
            domain_changes.append(curr_val - prev_val)
        
        # Collect all chart buffers — inserted as a grid at the end
        chart_pairs_2q = []

        # Domain Spider
        _domain_spider_ok, _domain_spider_skip = _is_spider_2q_useful(all_domain_names, prev_domain_values, curr_domain_values)
        if not _domain_spider_ok:
            print(f"[FORMATTER] Domain spider skipped: {_domain_spider_skip}")
        if _domain_spider_ok:
            spider_buf = _generate_spider_chart(
                labels=all_domain_names,
                prev_values=prev_domain_values,
                curr_values=curr_domain_values,
                title="Domain Comparison | مقارنة المجالات",
                prev_label=previous_period,
                curr_label=current_period
            )
            chart_pairs_2q.append((spider_buf, "🕸️ مخطط العنكبوت - المجالات | Domain Spider Chart"))

        # Domain Bar
        _domain_bar_ok, _domain_bar_skip = _is_bar_chart_useful(all_domain_names, domain_changes)
        if not _domain_bar_ok:
            print(f"[FORMATTER] Domain bar chart skipped: {_domain_bar_skip}")
        if _domain_bar_ok:
            bar_buf = _generate_diverging_bar_chart(
                labels=all_domain_names,
                changes=domain_changes,
                title="Domain Change Analysis | تحليل تغيرات المجالات"
            )
            chart_pairs_2q.append((bar_buf, "📊 مخطط الأعمدة - الفروقات (المجالات) | Domain Bar Chart"))

        # =================== CATEGORY LEVEL CHARTS ===================

        prev_categories, prev_category_counts = _extract_category_data(previous_hierarchy)
        curr_categories, curr_category_counts = _extract_category_data(current_hierarchy)
        all_category_names = sorted(set(prev_categories + curr_categories))

        prev_category_values = []
        curr_category_values = []
        category_changes = []
        for category in all_category_names:
            prev_val = prev_category_counts[prev_categories.index(category)] if category in prev_categories else 0
            curr_val = curr_category_counts[curr_categories.index(category)] if category in curr_categories else 0
            prev_category_values.append(prev_val)
            curr_category_values.append(curr_val)
            category_changes.append(curr_val - prev_val)

        if len(all_category_names) > 10:
            sorted_indices = sorted(range(len(category_changes)),
                                    key=lambda i: abs(category_changes[i]), reverse=True)[:10]
            all_category_names    = [all_category_names[i]    for i in sorted_indices]
            prev_category_values  = [prev_category_values[i]  for i in sorted_indices]
            curr_category_values  = [curr_category_values[i]  for i in sorted_indices]
            category_changes      = [category_changes[i]      for i in sorted_indices]

        # Category Spider
        _cat_spider_ok, _cat_spider_skip = _is_spider_2q_useful(all_category_names, prev_category_values, curr_category_values)
        if not _cat_spider_ok:
            print(f"[FORMATTER] Category spider skipped: {_cat_spider_skip}")
        if _cat_spider_ok:
            cat_spider_buf = _generate_spider_chart(
                labels=all_category_names,
                prev_values=prev_category_values,
                curr_values=curr_category_values,
                title="Category Comparison (Top 10) | مقارنة الفئات (أعلى 10)",
                prev_label=previous_period,
                curr_label=current_period
            )
            chart_pairs_2q.append((cat_spider_buf, "🕸️ مخطط العنكبوت - الفئات | Category Spider Chart"))

        # Category Bar
        _cat_bar_ok, _cat_bar_skip = _is_bar_chart_useful(all_category_names, category_changes)
        if not _cat_bar_ok:
            print(f"[FORMATTER] Category bar chart skipped: {_cat_bar_skip}")
        if _cat_bar_ok:
            cat_bar_buf = _generate_diverging_bar_chart(
                labels=all_category_names,
                changes=category_changes,
                title="Category Change Analysis (Top 10) | تحليل تغيرات الفئات (أعلى 10)"
            )
            chart_pairs_2q.append((cat_bar_buf, "📊 مخطط الأعمدة - الفروقات (الفئات) | Category Bar Chart"))

        # =================== SUBCATEGORY LEVEL CHART ===================

        prev_subcategories, prev_subcategory_counts = _extract_subcategory_data(previous_hierarchy)
        curr_subcategories, curr_subcategory_counts = _extract_subcategory_data(current_hierarchy)
        all_subcategory_names = sorted(set(prev_subcategories + curr_subcategories))

        prev_subcategory_values = []
        curr_subcategory_values = []
        subcategory_changes = []
        for subcategory in all_subcategory_names:
            prev_val = prev_subcategory_counts[prev_subcategories.index(subcategory)] if subcategory in prev_subcategories else 0
            curr_val = curr_subcategory_counts[curr_subcategories.index(subcategory)] if subcategory in curr_subcategories else 0
            prev_subcategory_values.append(prev_val)
            curr_subcategory_values.append(curr_val)
            subcategory_changes.append(curr_val - prev_val)

        if len(all_subcategory_names) > 10:
            sorted_indices = sorted(range(len(subcategory_changes)),
                                    key=lambda i: abs(subcategory_changes[i]), reverse=True)[:10]
            all_subcategory_names    = [all_subcategory_names[i]    for i in sorted_indices]
            prev_subcategory_values  = [prev_subcategory_values[i]  for i in sorted_indices]
            curr_subcategory_values  = [curr_subcategory_values[i]  for i in sorted_indices]
            subcategory_changes      = [subcategory_changes[i]      for i in sorted_indices]

        # SubCategory Spider
        _subcat_spider_ok, _subcat_spider_skip = _is_spider_2q_useful(all_subcategory_names, prev_subcategory_values, curr_subcategory_values)
        if not _subcat_spider_ok:
            print(f"[FORMATTER] SubCategory spider skipped: {_subcat_spider_skip}")
        if _subcat_spider_ok:
            subcat_spider_buf = _generate_spider_chart(
                labels=all_subcategory_names,
                prev_values=prev_subcategory_values,
                curr_values=curr_subcategory_values,
                title="Subcategory Comparison (Top 10) | مقارنة الفئات الفرعية (أعلى 10)",
                prev_label=previous_period,
                curr_label=current_period
            )
            chart_pairs_2q.append((subcat_spider_buf, "🕸️ مخطط العنكبوت - الفئات الفرعية | SubCategory Spider Chart"))

        # Insert all collected charts in a structured 2-column grid
        _add_chart_grid(doc, chart_pairs_2q)

    except Exception as e:
        # If chart generation fails, add error message but continue with report
        error_para = doc.add_paragraph()
        error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        error_run = error_para.add_run(f"⚠️ Chart generation encountered an issue: {str(e)}")
        error_run.font.size = Pt(11)
        error_run.font.color.rgb = RGBColor(192, 0, 0)
        doc.add_paragraph()
    
    # ============================================================
    # SAVE AND RETURN
    # ============================================================

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _create_comparative_hierarchical_tables_by_domain(
    doc: Document,
    current_hierarchy: Dict[str, Any],
    previous_hierarchy: Dict[str, Any],
    current_period: str,
    previous_period: str,
    all_domains: set,
    language: str
):
    """
    Create comparative hierarchical tables for each domain showing current vs previous data.
    
    Table structure (RTL, 9 columns):
    - Domain
    - Classification (AR + EN)
    - Previous Period: Total + Severity (L/M/H)
    - Current Period: Total + Severity (L/M/H)
    """
    if not all_domains:
        doc.add_paragraph("لا توجد بيانات متاحة (No classification data available).")
        return
    
    for domain_idx, domain_name in enumerate(sorted(all_domains)):
        if domain_idx > 0:
            doc.add_paragraph()  # Spacing between domains
        
        current_domain_data = current_hierarchy.get(domain_name, {'total': 0, 'categories': {}})
        previous_domain_data = previous_hierarchy.get(domain_name, {'total': 0, 'categories': {}})
        
        # Domain title with comparison
        domain_title = doc.add_paragraph()
        domain_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        domain_title.paragraph_format.right_to_left = True
        
        title_text = f'📂 {domain_name} | {previous_period}: n={previous_domain_data["total"]} → {current_period}: n={current_domain_data["total"]}'
        title_run = domain_title.add_run(title_text)
        title_run.font.bold = True
        title_run.font.size = Pt(13)
        title_run.font.name = 'Traditional Arabic'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        title_run.font.color.rgb = RGBColor(68, 114, 196)
        
        # Collect all classifications from both periods
        classification_map = {}  # classification_id -> {current_data, previous_data}
        
        # Process current period
        for category_name, category_data in current_domain_data['categories'].items():
            for subcat_name, subcat_data in category_data['subcategories'].items():
                for classification in subcat_data['classifications']:
                    class_id = classification.get('classification_id')
                    if class_id not in classification_map:
                        classification_map[class_id] = {'current': None, 'previous': None}
                    classification_map[class_id]['current'] = classification
        
        # Process previous period
        for category_name, category_data in previous_domain_data['categories'].items():
            for subcat_name, subcat_data in category_data['subcategories'].items():
                for classification in subcat_data['classifications']:
                    class_id = classification.get('classification_id')
                    if class_id not in classification_map:
                        classification_map[class_id] = {'current': None, 'previous': None}
                    classification_map[class_id]['previous'] = classification
        
        if not classification_map:
            # No classifications in this domain
            no_data = doc.add_paragraph()
            no_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
            no_data_run = no_data.add_run("لا توجد بيانات لهذا المجال (No data for this domain)")
            no_data_run.font.size = Pt(10)
            no_data_run.font.italic = True
            continue
        
        # Create table: 11 columns (Domain | Class AR | Class EN | Prev[Total+L+M+H] | Curr[Total+L+M+H])
        num_rows = len(classification_map) + 2  # data rows + 2 header rows
        table = doc.add_table(rows=num_rows, cols=11)
        table.style = 'Table Grid'
        
        # Set RTL
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr.append(bidiVisual)
        
        header_row1 = table.rows[0]
        header_row2 = table.rows[1]
        
        # Row 1: Domain | Class AR | Class EN | Previous (4 cols merged) | Current (4 cols merged)
        headers_main = [
            f"{domain_name}",
            "التصنيف عربي",
            "EN",
            previous_period,
            "", "", "",
            current_period,
            "", "", ""
        ]
        
        for idx, header_text in enumerate(headers_main):
            cell = header_row1.cells[idx]
            if header_text:
                cell.text = header_text
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                center_cell_content(cell)
            set_cell_shading(cell, '4472C4')
        
        # Row 2: Empty | Empty | Empty | Total L M H | Total L M H
        subheaders = ["", "", "", "المجموع", "L", "M", "H", "المجموع", "L", "M", "H"]
        
        for idx, subheader_text in enumerate(subheaders):
            cell = header_row2.cells[idx]
            if subheader_text:
                cell.text = subheader_text
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
                cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                center_cell_content(cell)
            set_cell_shading(cell, '5B9BD5')
        
        # Merge headers
        try:
            # Merge domain, class AR, class EN vertically
            for i in range(3):
                header_row1.cells[i].merge(header_row2.cells[i])
            # Merge previous period header (cols 3-6)
            header_row1.cells[3].merge(header_row1.cells[6])
            # Merge current period header (cols 7-10)
            header_row1.cells[7].merge(header_row1.cells[10])
        except Exception as e:
            print(f"[FORMATTER] Warning: Could not merge headers: {e}")
        
        # Fill data rows
        current_row = 2
        domain_start_row = current_row
        
        for class_id, data in sorted(classification_map.items()):
            row = table.rows[current_row]
            
            current_class = data['current']
            previous_class = data['previous']
            
            # Get classification names (prefer current, fallback to previous)
            if current_class:
                class_name_ar = current_class.get('classification_name', 'N/A')
                class_name_en = current_class.get('classification_name_en', 'N/A')
            else:
                class_name_ar = previous_class.get('classification_name', 'N/A')
                class_name_en = previous_class.get('classification_name_en', 'N/A')
            
            # Column 0: Domain (will be merged later)
            row.cells[0].text = ""
            
            # Column 1: Classification AR
            row.cells[1].text = class_name_ar
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[1].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
            row.cells[1].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            center_cell_content(row.cells[1])
            
            # Column 2: Classification EN
            row.cells[2].text = class_name_en
            row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
            center_cell_content(row.cells[2])
            
            # Previous period data (cols 3-6)
            if previous_class:
                prev_total = previous_class.get('total_count', 0)
                prev_low = previous_class.get('low_count', 0)
                prev_med = previous_class.get('medium_count', 0)
                prev_high = previous_class.get('high_count', 0)
            else:
                prev_total = prev_low = prev_med = prev_high = 0
            
            row.cells[3].text = str(prev_total)
            row.cells[3].paragraphs[0].runs[0].bold = True
            center_cell_content(row.cells[3])
            
            row.cells[4].text = str(prev_low) if prev_low > 0 else ""
            center_cell_content(row.cells[4])
            set_cell_shading(row.cells[4], 'C6EFCE' if prev_low > 0 else 'F2F2F2')
            
            row.cells[5].text = str(prev_med) if prev_med > 0 else ""
            center_cell_content(row.cells[5])
            set_cell_shading(row.cells[5], 'FFEB9C' if prev_med > 0 else 'F2F2F2')
            
            row.cells[6].text = str(prev_high) if prev_high > 0 else ""
            center_cell_content(row.cells[6])
            set_cell_shading(row.cells[6], 'FFC7CE' if prev_high > 0 else 'F2F2F2')
            
            # Current period data (cols 7-10)
            if current_class:
                curr_total = current_class.get('total_count', 0)
                curr_low = current_class.get('low_count', 0)
                curr_med = current_class.get('medium_count', 0)
                curr_high = current_class.get('high_count', 0)
            else:
                curr_total = curr_low = curr_med = curr_high = 0
            
            row.cells[7].text = str(curr_total)
            row.cells[7].paragraphs[0].runs[0].bold = True
            center_cell_content(row.cells[7])
            
            row.cells[8].text = str(curr_low) if curr_low > 0 else ""
            center_cell_content(row.cells[8])
            set_cell_shading(row.cells[8], 'C6EFCE' if curr_low > 0 else 'F2F2F2')
            
            row.cells[9].text = str(curr_med) if curr_med > 0 else ""
            center_cell_content(row.cells[9])
            set_cell_shading(row.cells[9], 'FFEB9C' if curr_med > 0 else 'F2F2F2')
            
            row.cells[10].text = str(curr_high) if curr_high > 0 else ""
            center_cell_content(row.cells[10])
            set_cell_shading(row.cells[10], 'FFC7CE' if curr_high > 0 else 'F2F2F2')
            
            current_row += 1
        
        # Merge Domain column for all rows
        if current_row - 1 > domain_start_row:
            table.rows[domain_start_row].cells[0].text = domain_name
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            table.rows[domain_start_row].cells[0].paragraphs[0].runs[0].font.bold = True
            center_cell_content(table.rows[domain_start_row].cells[0])
            set_cell_shading(table.rows[domain_start_row].cells[0], 'BDD7EE')
            
            try:
                for row_idx in range(domain_start_row + 1, current_row):
                    if row_idx < len(table.rows):
                        table.rows[domain_start_row].cells[0].merge(table.rows[row_idx].cells[0])
            except Exception as e:
                print(f"[FORMATTER] Warning: Could not merge domain cells: {e}")
    
    doc.add_paragraph()


def _create_3quarter_hierarchical_tables_by_domain(
    doc: Document,
    hierarchies: List[Dict[str, Any]],
    periods: List[str],
):
    """
    Hospital-protocol hierarchical comparison table for 3 quarters.

    Table — 15 columns (A3 landscape):
      Domain | ClassAR | ClassEN | Q1(Total L M H) | Q2(Total L M H) | Q3(Total L M H)

    Args:
        hierarchies: [q1_hierarchy, q2_hierarchy, q3_hierarchy] from _build_hierarchy()
        periods:     ['Q4-2025', 'Q1-2026', 'Q2-2026']
    """
    all_domains: set = set()
    for hier in hierarchies:
        all_domains.update(hier.keys())

    if not all_domains:
        p = doc.add_paragraph("لا توجد بيانات تصنيفية (No classification data available).")
        return

    # Column widths (twips): Domain | ClassAR | ClassEN | Q×(Total L M H) ×3
    # Total ≈ 292mm — comfortable on A3 landscape (390mm usable)
    col_widths_cm = [
        2.0,                      # 0  Domain
        5.0,                      # 1  ClassAR
        4.0,                      # 2  ClassEN
        2.0, 1.3, 1.3, 1.3,      # 3-6  Q1
        2.0, 1.3, 1.3, 1.3,      # 7-10 Q2
        2.0, 1.3, 1.3, 1.3,      # 11-14 Q3
    ]

    def _w(cm_val):
        return int(Cm(cm_val))

    def _fill(cell, text, bold=False, size=9, arabic=True, color=None, bg=None, align_center=True):
        cell.text = str(text)
        if bg:
            set_cell_shading(cell, bg)
        if str(text) and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = bold
            run.font.size = Pt(size)
            if arabic:
                run.font.name = 'Traditional Arabic'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            if color:
                run.font.color.rgb = color
        if align_center:
            center_cell_content(cell)
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for domain_idx, domain_name in enumerate(sorted(all_domains)):
        if domain_idx > 0:
            doc.add_paragraph()

        # Domain totals line
        d_totals = [
            hier.get(domain_name, {'total': 0}).get('total', 0)
            for hier in hierarchies
        ]
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.paragraph_format.right_to_left = True
        h_run = heading.add_run(
            f"📂 {domain_name}  |  "
            f"{periods[0]}: n={d_totals[0]}  |  "
            f"{periods[1]}: n={d_totals[1]}  |  "
            f"{periods[2]}: n={d_totals[2]}"
        )
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.name = 'Traditional Arabic'
        h_run.font.color.rgb = RGBColor(68, 114, 196)

        # Build classification_map: class_id -> [q1_stat|None, q2_stat|None, q3_stat|None]
        classification_map: dict = {}
        for q_idx, hier in enumerate(hierarchies):
            dom = hier.get(domain_name, {'categories': {}})
            for cat_data in dom.get('categories', {}).values():
                for sub_data in cat_data.get('subcategories', {}).values():
                    for stat in sub_data.get('classifications', []):
                        cid = stat.get('classification_id')
                        if cid not in classification_map:
                            classification_map[cid] = [None, None, None]
                        classification_map[cid][q_idx] = stat

        if not classification_map:
            doc.add_paragraph("لا توجد بيانات لهذا المجال.")
            continue

        num_rows = len(classification_map) + 2   # 2 header rows + data rows
        table = doc.add_table(rows=num_rows, cols=15)
        table.style = 'Table Grid'

        # Set RTL
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblPr.append(OxmlElement('w:bidiVisual'))

        # Set column widths on every row
        for row in table.rows:
            for ci, w in enumerate(col_widths_cm):
                row.cells[ci].width = _w(w)

        # ── Header row 1 ──────────────────────────────────────────────────────
        hr1 = table.rows[0]
        h1_texts = [
            domain_name, 'التصنيف عربي', 'EN',
            periods[0], '', '', '',
            periods[1], '', '', '',
            periods[2], '', '', '',
        ]
        for ci, txt in enumerate(h1_texts):
            _fill(hr1.cells[ci], txt, bold=True, size=8,
                  color=RGBColor(255, 255, 255), bg='4472C4')

        # ── Header row 2 ──────────────────────────────────────────────────────
        hr2 = table.rows[1]
        h2_texts = ['', '', '',
                    'المجموع', 'L', 'M', 'H',
                    'المجموع', 'L', 'M', 'H',
                    'المجموع', 'L', 'M', 'H']
        for ci, txt in enumerate(h2_texts):
            _fill(hr2.cells[ci], txt, bold=True, size=7,
                  color=RGBColor(255, 255, 255), bg='5B9BD5')

        # Merge headers
        try:
            for ci in range(3):
                hr1.cells[ci].merge(hr2.cells[ci])
            hr1.cells[3].merge(hr1.cells[6])
            hr1.cells[7].merge(hr1.cells[10])
            hr1.cells[11].merge(hr1.cells[14])
        except Exception as e:
            print(f"[3Q FORMATTER] header merge warning: {e}")

        # ── Data rows ─────────────────────────────────────────────────────────
        domain_start = 2
        for ri, (cid, q_stats) in enumerate(sorted(classification_map.items()), start=2):
            row = table.rows[ri]
            q1, q2, q3 = q_stats
            ref = q1 or q2 or q3

            row.cells[0].text = ''

            # Arabic name
            _fill(row.cells[1], ref.get('classification_name', 'N/A'),
                  size=8, align_center=False)

            # English name
            _fill(row.cells[2], ref.get('classification_name_en', 'N/A'),
                  size=7, arabic=False)

            # Q1, Q2, Q3 severity blocks
            for q_idx, q_stat in enumerate([q1, q2, q3]):
                base = 3 + q_idx * 4
                tot = q_stat.get('total_count', 0) if q_stat else 0
                low = q_stat.get('low_count', 0)   if q_stat else 0
                med = q_stat.get('medium_count', 0) if q_stat else 0
                high = q_stat.get('high_count', 0)  if q_stat else 0

                _fill(row.cells[base],     str(tot),             bold=True, size=8, arabic=False)
                _fill(row.cells[base + 1], str(low)  if low  else '', size=7, arabic=False,
                      bg='C6EFCE' if low  else 'F2F2F2')
                _fill(row.cells[base + 2], str(med)  if med  else '', size=7, arabic=False,
                      bg='FFEB9C' if med  else 'F2F2F2')
                _fill(row.cells[base + 3], str(high) if high else '', size=7, arabic=False,
                      bg='FFC7CE' if high else 'F2F2F2')

        # Merge domain column vertically across all data rows
        n_data = len(classification_map)
        if n_data > 0:
            first = table.rows[domain_start].cells[0]
            _fill(first, domain_name, bold=True, size=9, bg='BDD7EE')
            try:
                for ri in range(domain_start + 1, domain_start + n_data):
                    first.merge(table.rows[ri].cells[0])
            except Exception as e:
                print(f"[3Q FORMATTER] domain merge warning: {e}")

    doc.add_paragraph()


def _add_policy_compliance_comparison(
    doc: Document,
    reports: List[Dict[str, Any]],
    periods: List[str],
) -> None:
    """
    Add a policy compliance comparison block.
    Works for any N quarters (2Q, 3Q, 4Q).

    Renders a compact 2-row table:
      Row 0 (header) : period labels  — blue
      Row 1 (data)   : ✓/✗ status     — green if compliant, red if not
    """
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading.paragraph_format.right_to_left = True
    h_run = heading.add_run("📊 مقارنة الامتثال للسياسة | Policy Compliance Comparison")
    h_run.font.bold = True
    h_run.font.size = Pt(13)
    h_run.font.name = 'Traditional Arabic'
    h_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    doc.add_paragraph()

    n = len(periods)
    table = doc.add_table(rows=2, cols=n)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, period in enumerate(periods):
        # Header row — period label
        hdr = table.rows[0].cells[i]
        hdr.text = period
        if hdr.paragraphs[0].runs:
            run = hdr.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Traditional Arabic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            run.font.color.rgb = RGBColor(255, 255, 255)
        center_cell_content(hdr)
        set_cell_shading(hdr, '4472C4')

        # Data row — compliance status
        is_compliant = reports[i]['header'].get('is_compliant', False)
        status_text = 'مطابق | Compliant' if is_compliant else 'غير مطابق | Non-Compliant'
        status_icon = '✓' if is_compliant else '✗'
        cell = table.rows[1].cells[i]
        cell.text = f"{status_icon}  {status_text}"
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Traditional Arabic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        center_cell_content(cell)
        set_cell_shading(cell, 'C6EFCE' if is_compliant else 'FFC7CE')

    doc.add_paragraph()


def generate_3_quarter_comparison_report(comparison_data: Dict[str, Any], language: str = 'ar') -> Document:
    """
    Generate a professional Word document comparing 3 seasonal quarters.
    
    Features:
    - Summary table with 3 quarters + trend indicators
    - Hierarchical tables showing domain/category/subcategory comparisons
    - Spider charts only (3 graphs: Domain, Category, SubCategory)
    - Trend indicators (↑↑, ↑, →, ↓, ↓↓)
    - A4 Landscape orientation with Arabic support
    
    Args:
        comparison_data: Dictionary containing comparison data from seasonal_comparison_service
        language: Language code ('ar' or 'en')
        
    Returns:
        Document object ready for saving
    """
    doc = Document()
    
    # Configure page layout (A3 Landscape — wider than A4 for 15-column protocol table)
    section = doc.sections[0]
    section.page_width = Mm(420)
    section.page_height = Mm(297)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    reports = comparison_data['reports']
    periods = comparison_data['periods']
    trends = comparison_data['trends']
    orgunit_name = comparison_data.get('orgunit_name', 'N/A')
    
    # =============================
    # 1. HEADER WITH LOGO
    # =============================
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            section.header_distance = Inches(0.1)
            header_section = section.header
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[FORMATTER] Could not add logo: {e}")
        pass
    
    # Title
    title_text = "تقرير المقارنة الموسمية - 3 أرباع" if language == 'ar' else "Seasonal Comparison Report - 3 Quarters"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Traditional Arabic'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    doc.add_paragraph()
    
    # =============================
    # 2. SUMMARY TABLE (3 QUARTERS + TRENDS)
    # =============================
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary_run = summary_heading.add_run("📊 ملخص المقارنة | Comparison Summary")
    summary_run.font.size = Pt(14)
    summary_run.font.bold = True
    summary_run.font.name = 'Traditional Arabic'
    summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    # Create summary table: 5 columns (Metric | Q1 | Q2 | Q3 | Trend)
    summary_table = doc.add_table(rows=11, cols=5)
    summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    summary_table.style = 'Table Grid'

    # Set RTL
    _tbl = summary_table._element
    _tblPr = _tbl.tblPr
    if _tblPr is None:
        _tblPr = OxmlElement('w:tblPr')
        _tbl.insert(0, _tblPr)
    _tblPr.append(OxmlElement('w:bidiVisual'))

    # Header row
    headers = ['الاتجاه | Trend', periods[2], periods[1], periods[0], 'المؤشر | Metric']
    for i, header_text in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        center_cell_content(cell)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    metrics = [
        ('إجمالي الحالات | Total Cases', 'total_cases', 'total_cases'),
        ('السريرية | Clinical', 'clinical_domain_count', 'clinical'),
        ('الإدارية | Management', 'management_domain_count', 'management'),
        ('العلاقاتية | Relational', 'relational_domain_count', 'relational'),
        ('منخفضة الخطورة | Low Severity', 'low_severity_count', 'low_severity'),
        ('متوسطة الخطورة | Medium Severity', 'medium_severity_count', 'medium_severity'),
        ('عالية الخطورة | High Severity', 'high_severity_count', 'high_severity'),
        ('إجراءات وقائية | Prevention Actions', 'prevention_action_count', None),
        ('تفسيرات مقدمة | Explanations Submitted', 'explanation_count', None),
        ('حالات مفتوحة | Open Cases', 'open_cases_count', None)
    ]
    
    for row_idx, (metric_label, metric_key, trend_key) in enumerate(metrics, start=1):
        row = summary_table.rows[row_idx]
        
        # Metric name
        row.cells[4].text = metric_label
        row.cells[4].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[4].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        row.cells[4].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_shading(row.cells[4], 'D9E2F3')
        
        # Q1, Q2, Q3 values
        for i in range(3):
            value = reports[i]['header'].get(metric_key, 0)
            row.cells[3-i].text = str(value)
            row.cells[3-i].paragraphs[0].runs[0].font.size = Pt(10)
            center_cell_content(row.cells[3-i])
        
        # Trend indicator
        if trend_key and trend_key in trends:
            row.cells[0].text = trends[trend_key]
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(14)
            center_cell_content(row.cells[0])
            
            # Color code trends
            trend_value = trends[trend_key]
            if trend_value in ['↑', '↑↑']:
                set_cell_shading(row.cells[0], 'C6EFCE')  # Green
            elif trend_value in ['↓', '↓↓']:
                set_cell_shading(row.cells[0], 'FFC7CE')  # Red
            else:
                set_cell_shading(row.cells[0], 'FFEB9C')  # Yellow
        else:
            row.cells[0].text = '—'
            center_cell_content(row.cells[0])
    
    doc.add_paragraph()
    
    # =============================
    # 3. HIERARCHICAL COMPARISON TABLES (Hospital Protocol)
    # =============================

    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ph_run = ph.add_run("📋 التفاصيل التصنيفية | Classification Details")
    ph_run.font.size = Pt(13)
    ph_run.font.bold = True
    ph_run.font.name = 'Traditional Arabic'
    ph_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    doc.add_paragraph()

    hierarchies = [_build_hierarchy(r.get('classification_stats', [])) for r in reports]
    domain_data = comparison_data['domain_comparison']
    category_data = comparison_data['category_comparison']
    subcategory_data = comparison_data['subcategory_comparison']
    _create_3quarter_hierarchical_tables_by_domain(doc, hierarchies, periods)

    # =============================
    # 4. POLICY COMPLIANCE COMPARISON
    # =============================
    _add_policy_compliance_comparison(doc, reports, periods)

    # Per-quarter detailed policy compliance tables
    _add_per_quarter_policy_sections(doc, [
        (periods[i], reports[i]['header'], reports[i].get('policy_snapshot'), reports[i].get('classification_stats', []))
        for i in range(len(reports))
    ])

    # =============================
    # 5. PAGE BREAK BEFORE GRAPHS
    # =============================
    doc.add_page_break()

    # =============================
    # 5. VISUAL ANALYSIS - SPIDER CHARTS ONLY
    # =============================
    visual_heading = doc.add_paragraph()
    visual_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    visual_run = visual_heading.add_run("📊 التحليل البصري | Visual Analysis")
    visual_run.font.size = Pt(16)
    visual_run.font.bold = True
    visual_run.font.name = 'Traditional Arabic'
    visual_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    doc.add_paragraph()

    # Build chart list in interleaved order so the 2-column grid renders as:
    #   Row 1: [Domain Spider  | Domain Bar   ]
    #   Row 2: [Category Spider| Category Bar ]
    #   Row 3: [SubCat Spider  | (empty)      ]
    chart_pairs_3q = []
    _bar_title_suffix = f"{periods[0]} → {periods[-1]}"

    # ── Domain level ─────────────────────────────────────────────────────────
    _3q_domain_ok, _3q_domain_skip = _is_spider_nq_useful(domain_data, max_items=8)
    if not _3q_domain_ok:
        print(f"[FORMATTER] 3Q domain spider skipped: {_3q_domain_skip}")
    if _3q_domain_ok:
        chart_pairs_3q.append((
            _generate_3quarter_spider_chart(domain_data, periods,
                                            title_ar="مخطط العنكبوت - المجالات",
                                            title_en="Domain Spider Chart"),
            "🕸️ مخطط العنكبوت - المجالات | Domain Spider Chart"
        ))

    _dom_bar_labels, _dom_bar_changes = _extract_nq_changes(domain_data, max_items=8)
    _3q_dom_bar_ok, _3q_dom_bar_skip = _is_bar_chart_useful(_dom_bar_labels, _dom_bar_changes)
    if not _3q_dom_bar_ok:
        print(f"[FORMATTER] 3Q domain bar skipped: {_3q_dom_bar_skip}")
    if _3q_dom_bar_ok:
        chart_pairs_3q.append((
            _generate_diverging_bar_chart(
                labels=_dom_bar_labels,
                changes=_dom_bar_changes,
                title=f"تغيير المجالات: {_bar_title_suffix} | Domain Net Change"
            ),
            f"📊 تغيير المجالات | Domain Net Change ({_bar_title_suffix})"
        ))

    # ── Category level ────────────────────────────────────────────────────────
    _3q_cat_ok, _3q_cat_skip = _is_spider_nq_useful(category_data, max_items=10)
    if not _3q_cat_ok:
        print(f"[FORMATTER] 3Q category spider skipped: {_3q_cat_skip}")
    if _3q_cat_ok:
        chart_pairs_3q.append((
            _generate_3quarter_spider_chart(category_data, periods,
                                            title_ar="مخطط العنكبوت - الفئات",
                                            title_en="Category Spider Chart",
                                            max_items=10),
            "🕸️ مخطط العنكبوت - الفئات | Category Spider Chart"
        ))

    _cat_bar_labels, _cat_bar_changes = _extract_nq_changes(category_data, max_items=10)
    _3q_cat_bar_ok, _3q_cat_bar_skip = _is_bar_chart_useful(_cat_bar_labels, _cat_bar_changes)
    if not _3q_cat_bar_ok:
        print(f"[FORMATTER] 3Q category bar skipped: {_3q_cat_bar_skip}")
    if _3q_cat_bar_ok:
        chart_pairs_3q.append((
            _generate_diverging_bar_chart(
                labels=_cat_bar_labels,
                changes=_cat_bar_changes,
                title=f"تغيير الفئات: {_bar_title_suffix} | Category Net Change"
            ),
            f"📊 تغيير الفئات | Category Net Change ({_bar_title_suffix})"
        ))

    # ── SubCategory level (spider only) ──────────────────────────────────────
    _3q_subcat_ok, _3q_subcat_skip = _is_spider_nq_useful(subcategory_data, max_items=12)
    if not _3q_subcat_ok:
        print(f"[FORMATTER] 3Q subcategory spider skipped: {_3q_subcat_skip}")
    if _3q_subcat_ok:
        chart_pairs_3q.append((
            _generate_3quarter_spider_chart(subcategory_data, periods,
                                            title_ar="مخطط العنكبوت - الفئات الفرعية",
                                            title_en="SubCategory Spider Chart",
                                            max_items=12),
            "🕸️ مخطط العنكبوت - الفئات الفرعية | SubCategory Spider Chart"
        ))

    _add_chart_grid(doc, chart_pairs_3q)
    return doc


def _create_3quarter_hierarchical_table(doc: Document, data: Dict[str, List[int]], periods: List[str], level: str):
    """
    Create a hierarchical table comparing 3 quarters for a specific level (domain/category/subcategory).
    
    Args:
        doc: Document object
        data: Dictionary mapping item names to list of 3 values
        periods: List of 3 period labels
        level: 'domain', 'category', or 'subcategory'
    """
    # Create table: 5 columns (Item | Q1 | Q2 | Q3 | Trend)
    num_rows = len(data) + 2  # +1 for header, +1 for totals
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = 'Table Grid'

    # Set RTL
    _tbl = table._element
    _tblPr = _tbl.tblPr
    if _tblPr is None:
        _tblPr = OxmlElement('w:tblPr')
        _tbl.insert(0, _tblPr)
    _tblPr.append(OxmlElement('w:bidiVisual'))

    # Header row
    headers = ['الاتجاه | Trend', periods[2], periods[1], periods[0], f'{"المجال" if level == "domain" else "الفئة" if level == "category" else "الفئة الفرعية"} | {level.title()}']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        center_cell_content(cell)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    totals = [0, 0, 0]
    for row_idx, (item_name, values) in enumerate(sorted(data.items()), start=1):
        row = table.rows[row_idx]
        
        # Item name
        row.cells[4].text = item_name
        row.cells[4].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[4].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        row.cells[4].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Quarter values
        for i in range(3):
            row.cells[3-i].text = str(values[i])
            row.cells[3-i].paragraphs[0].runs[0].font.size = Pt(10)
            center_cell_content(row.cells[3-i])
            totals[i] += values[i]
        
        # Trend indicator
        trend = _calculate_trend_indicator(values[0], values[-1])
        row.cells[0].text = trend
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(14)
        center_cell_content(row.cells[0])
        
        # Color code trends
        if trend in ['↑', '↑↑']:
            set_cell_shading(row.cells[0], 'C6EFCE')  # Green
        elif trend in ['↓', '↓↓']:
            set_cell_shading(row.cells[0], 'FFC7CE')  # Red
        else:
            set_cell_shading(row.cells[0], 'FFEB9C')  # Yellow
    
    # Totals row
    totals_row = table.rows[-1]
    totals_row.cells[4].text = 'الإجمالي | Total'
    totals_row.cells[4].paragraphs[0].runs[0].font.bold = True
    totals_row.cells[4].paragraphs[0].runs[0].font.size = Pt(11)
    totals_row.cells[4].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
    totals_row.cells[4].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    totals_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(totals_row.cells[4], 'D9E2F3')
    
    for i in range(3):
        totals_row.cells[3-i].text = str(totals[i])
        totals_row.cells[3-i].paragraphs[0].runs[0].font.bold = True
        totals_row.cells[3-i].paragraphs[0].runs[0].font.size = Pt(11)
        center_cell_content(totals_row.cells[3-i])
        set_cell_shading(totals_row.cells[3-i], 'D9E2F3')
    
    # Overall trend
    overall_trend = _calculate_trend_indicator(totals[0], totals[-1])
    totals_row.cells[0].text = overall_trend
    totals_row.cells[0].paragraphs[0].runs[0].font.size = Pt(14)
    totals_row.cells[0].paragraphs[0].runs[0].font.bold = True
    center_cell_content(totals_row.cells[0])
    set_cell_shading(totals_row.cells[0], 'BDD7EE')


def _calculate_trend_indicator(first_value: int, last_value: int) -> str:
    """
    Calculate trend indicator for 3-quarter comparison.
    
    Returns:
        Trend string: '↑↑', '↑', '→', '↓', '↓↓'
    """
    if first_value == 0:
        if last_value == 0:
            return '→'
        else:
            return '↑↑'
    
    change_percent = ((last_value - first_value) / first_value) * 100
    
    if change_percent > 20:
        return '↑↑'
    elif change_percent > 5:
        return '↑'
    elif change_percent < -20:
        return '↓↓'
    elif change_percent < -5:
        return '↓'
    else:
        return '→'


# ============================================================================
# UNIFIED SPIDER CHART ENGINE
# All comparison modes (2Q / 3Q / 4Q) funnel through this single renderer.
# ============================================================================

_SPIDER_COLORS = ['#4472C4', '#ED7D31', '#A5A5A5', '#FFC000', '#70AD47', '#FF7F7F']


def _generate_nquarter_spider_chart(
    data: Dict[str, List[int]],
    periods: List[str],
    title_ar: str,
    title_en: str,
    max_items: int = 8
) -> io.BytesIO:
    """
    Unified spider/radar chart renderer for any N-quarter comparison (N ≥ 2).

    Args:
        data:      Dict mapping item names → list of N per-quarter counts.
        periods:   N period label strings (one per quarter series).
        title_ar:  Arabic chart title.
        title_en:  English chart title.
        max_items: Maximum radar axes to display (top by total count).

    Returns:
        BytesIO PNG image buffer.
    """
    n_quarters = len(periods)

    _configure_arabic_matplotlib()

    # Select top items by total count
    sorted_items = sorted(data.items(), key=lambda x: sum(x[1]), reverse=True)[:max_items]

    if not sorted_items:
        fig, ax = plt.subplots(figsize=(10, 8), subplot_kw=dict(projection='polar'))
        ax.text(0.5, 0.5, _ar('لا توجد بيانات') + '\nNo Data', ha='center', va='center',
                transform=ax.transAxes, fontsize=16)
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    categories = [item[0] for item in sorted_items]
    num_vars = len(categories)

    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(12, 10), subplot_kw=dict(projection='polar'))

    colors = (_SPIDER_COLORS * ((n_quarters // len(_SPIDER_COLORS)) + 1))[:n_quarters]

    for i in range(n_quarters):
        values = [item[1][i] if i < len(item[1]) else 0 for item in sorted_items]
        values += values[:1]
        ax.plot(angles, values, 'o-', linewidth=2, label=_ar(periods[i]), color=colors[i])
        ax.fill(angles, values, alpha=0.15, color=colors[i])

    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([_ar(c) for c in categories], fontsize=10)

    plt.title(f"{_ar(title_ar)}\n{title_en}", size=14, weight='bold', pad=20)
    plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=10)
    ax.grid(True, linestyle='--', alpha=0.7)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _generate_3quarter_spider_chart(
    data: Dict[str, List[int]],
    periods: List[str],
    title_ar: str,
    title_en: str,
    max_items: int = 8
) -> io.BytesIO:
    """3-quarter spider chart — delegates to unified engine."""
    return _generate_nquarter_spider_chart(data, periods, title_ar, title_en, max_items)


def _create_4quarter_hierarchical_tables_by_domain(
    doc: Document,
    hierarchies: List[Dict[str, Any]],
    periods: List[str],
):
    """
    Hospital-protocol hierarchical comparison table for 4 quarters.

    Table — 19 columns (A3 landscape):
      Domain | ClassAR | ClassEN | Q1(T L M H) | Q2(T L M H) | Q3(T L M H) | Q4(T L M H)
    """
    all_domains: set = set()
    for hier in hierarchies:
        all_domains.update(hier.keys())

    if not all_domains:
        p = doc.add_paragraph("لا توجد بيانات تصنيفية (No classification data available).")
        return

    col_widths_cm = [
        1.8,                              # 0  Domain
        4.0,                              # 1  ClassAR
        3.2,                              # 2  ClassEN
        1.8, 1.1, 1.1, 1.1,              # 3-6   Q1
        1.8, 1.1, 1.1, 1.1,              # 7-10  Q2
        1.8, 1.1, 1.1, 1.1,              # 11-14 Q3
        1.8, 1.1, 1.1, 1.1,              # 15-18 Q4
    ]

    def _w(cm_val):
        return int(Cm(cm_val))

    def _fill(cell, text, bold=False, size=9, arabic=True, color=None, bg=None, align_center=True):
        cell.text = str(text)
        if bg:
            set_cell_shading(cell, bg)
        if str(text) and cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = bold
            run.font.size = Pt(size)
            if arabic:
                run.font.name = 'Traditional Arabic'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
            if color:
                run.font.color.rgb = color
        if align_center:
            center_cell_content(cell)
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for domain_idx, domain_name in enumerate(sorted(all_domains)):
        if domain_idx > 0:
            doc.add_paragraph()

        d_totals = [
            hier.get(domain_name, {'total': 0}).get('total', 0)
            for hier in hierarchies
        ]
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.paragraph_format.right_to_left = True
        h_run = heading.add_run(
            f"📂 {domain_name}  |  "
            + "  |  ".join(f"{p}: n={t}" for p, t in zip(periods, d_totals))
        )
        h_run.font.bold = True
        h_run.font.size = Pt(11)
        h_run.font.name = 'Traditional Arabic'
        h_run.font.color.rgb = RGBColor(68, 114, 196)

        # Build classification_map: class_id -> [q1|None, q2|None, q3|None, q4|None]
        classification_map: dict = {}
        for q_idx, hier in enumerate(hierarchies):
            dom = hier.get(domain_name, {'categories': {}})
            for cat_data in dom.get('categories', {}).values():
                for sub_data in cat_data.get('subcategories', {}).values():
                    for stat in sub_data.get('classifications', []):
                        cid = stat.get('classification_id')
                        if cid not in classification_map:
                            classification_map[cid] = [None, None, None, None]
                        classification_map[cid][q_idx] = stat

        if not classification_map:
            doc.add_paragraph("لا توجد بيانات لهذا المجال.")
            continue

        num_rows = len(classification_map) + 2
        table = doc.add_table(rows=num_rows, cols=19)
        table.style = 'Table Grid'

        # Set RTL
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblPr.append(OxmlElement('w:bidiVisual'))

        for row in table.rows:
            for ci, w in enumerate(col_widths_cm):
                row.cells[ci].width = _w(w)

        # Header row 1
        hr1 = table.rows[0]
        h1_texts = [
            domain_name, 'التصنيف عربي', 'EN',
            periods[0], '', '', '',
            periods[1], '', '', '',
            periods[2], '', '', '',
            periods[3], '', '', '',
        ]
        for ci, txt in enumerate(h1_texts):
            _fill(hr1.cells[ci], txt, bold=True, size=8,
                  color=RGBColor(255, 255, 255), bg='4472C4')

        # Header row 2
        hr2 = table.rows[1]
        h2_texts = ['', '', '',
                    'المجموع', 'L', 'M', 'H',
                    'المجموع', 'L', 'M', 'H',
                    'المجموع', 'L', 'M', 'H',
                    'المجموع', 'L', 'M', 'H']
        for ci, txt in enumerate(h2_texts):
            _fill(hr2.cells[ci], txt, bold=True, size=7,
                  color=RGBColor(255, 255, 255), bg='5B9BD5')

        try:
            for ci in range(3):
                hr1.cells[ci].merge(hr2.cells[ci])
            hr1.cells[3].merge(hr1.cells[6])
            hr1.cells[7].merge(hr1.cells[10])
            hr1.cells[11].merge(hr1.cells[14])
            hr1.cells[15].merge(hr1.cells[18])
        except Exception as e:
            print(f"[4Q FORMATTER] header merge warning: {e}")

        domain_start = 2
        for ri, (cid, q_stats) in enumerate(sorted(classification_map.items()), start=2):
            row = table.rows[ri]
            ref = next((s for s in q_stats if s), None)

            row.cells[0].text = ''
            _fill(row.cells[1], ref.get('classification_name', 'N/A') if ref else '',
                  size=8, align_center=False)
            _fill(row.cells[2], ref.get('classification_name_en', 'N/A') if ref else '',
                  size=7, arabic=False)

            for q_idx, q_stat in enumerate(q_stats):
                base = 3 + q_idx * 4
                tot  = q_stat.get('total_count', 0)  if q_stat else 0
                low  = q_stat.get('low_count', 0)    if q_stat else 0
                med  = q_stat.get('medium_count', 0) if q_stat else 0
                high = q_stat.get('high_count', 0)   if q_stat else 0

                _fill(row.cells[base],     str(tot),            bold=True, size=8, arabic=False)
                _fill(row.cells[base + 1], str(low)  if low  else '', size=7, arabic=False,
                      bg='C6EFCE' if low  else 'F2F2F2')
                _fill(row.cells[base + 2], str(med)  if med  else '', size=7, arabic=False,
                      bg='FFEB9C' if med  else 'F2F2F2')
                _fill(row.cells[base + 3], str(high) if high else '', size=7, arabic=False,
                      bg='FFC7CE' if high else 'F2F2F2')

        n_data = len(classification_map)
        if n_data > 0:
            first = table.rows[domain_start].cells[0]
            _fill(first, domain_name, bold=True, size=9, bg='BDD7EE')
            try:
                for ri in range(domain_start + 1, domain_start + n_data):
                    first.merge(table.rows[ri].cells[0])
            except Exception as e:
                print(f"[4Q FORMATTER] domain merge warning: {e}")

    doc.add_paragraph()


def generate_4_quarter_comparison_report(comparison_data: Dict[str, Any], language: str = 'ar') -> Document:
    """
    Generate a professional Word document comparing 4 seasonal quarters (full year).
    
    Features:
    - Summary table with 4 quarters + yearly total + trend indicators
    - Hierarchical tables showing domain/category/subcategory comparisons
    - Spider charts only (3 graphs: Domain, Category, SubCategory)
    - Trend indicators (↑↑, ↑, →, ↓, ↓↓)
    - A4 Landscape orientation with Arabic support
    
    Args:
        comparison_data: Dictionary containing comparison data from seasonal_comparison_service
        language: Language code ('ar' or 'en')
        
    Returns:
        Document object ready for saving
    """
    doc = Document()
    
    # Configure page layout (A3 Landscape — wider than A4 for 19-column protocol table)
    section = doc.sections[0]
    section.page_width = Mm(420)
    section.page_height = Mm(297)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    reports = comparison_data['reports']
    periods = comparison_data['periods']
    trends = comparison_data['trends']
    yearly_totals = comparison_data['yearly_totals']
    orgunit_name = comparison_data.get('orgunit_name', 'N/A')
    
    # =============================
    # 1. HEADER WITH LOGO
    # =============================
    try:
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            section.header_distance = Inches(0.1)
            header_section = section.header
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[FORMATTER] Could not add logo: {e}")
        pass
    
    # Title
    title_text = "تقرير المقارنة الموسمية - التقرير السنوي (4 أرباع)" if language == 'ar' else "Seasonal Comparison Report - Annual (4 Quarters)"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Traditional Arabic'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    doc.add_paragraph()
    
    # =============================
    # 2. SUMMARY TABLE (4 QUARTERS + YEARLY TOTAL + TRENDS)
    # =============================
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary_run = summary_heading.add_run("📊 ملخص المقارنة السنوية | Annual Comparison Summary")
    summary_run.font.size = Pt(14)
    summary_run.font.bold = True
    summary_run.font.name = 'Traditional Arabic'
    summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    # Create summary table: 7 columns (Metric | Q1 | Q2 | Q3 | Q4 | Yearly Total | Trend)
    summary_table = doc.add_table(rows=11, cols=7)
    summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    summary_table.style = 'Table Grid'

    # Set RTL
    _tbl = summary_table._element
    _tblPr = _tbl.tblPr
    if _tblPr is None:
        _tblPr = OxmlElement('w:tblPr')
        _tbl.insert(0, _tblPr)
    _tblPr.append(OxmlElement('w:bidiVisual'))

    # Header row
    headers = ['الاتجاه | Trend', 'الإجمالي السنوي | Yearly Total', periods[3], periods[2], periods[1], periods[0], 'المؤشر | Metric']
    for i, header_text in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        center_cell_content(cell)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    metrics = [
        ('إجمالي الحالات | Total Cases', 'total_cases', 'total_cases'),
        ('السريرية | Clinical', 'clinical_domain_count', 'clinical'),
        ('الإدارية | Management', 'management_domain_count', 'management'),
        ('العلاقاتية | Relational', 'relational_domain_count', 'relational'),
        ('منخفضة الخطورة | Low Severity', 'low_severity_count', 'low_severity'),
        ('متوسطة الخطورة | Medium Severity', 'medium_severity_count', 'medium_severity'),
        ('عالية الخطورة | High Severity', 'high_severity_count', 'high_severity'),
        ('إجراءات وقائية | Prevention Actions', 'prevention_action_count', None),
        ('تفسيرات مقدمة | Explanations Submitted', 'explanation_count', None),
        ('حالات مفتوحة | Open Cases', 'open_cases_count', None)
    ]
    
    for row_idx, (metric_label, metric_key, trend_key) in enumerate(metrics, start=1):
        row = summary_table.rows[row_idx]
        
        # Metric name
        row.cells[6].text = metric_label
        row.cells[6].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[6].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        row.cells[6].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_shading(row.cells[6], 'D9E2F3')
        
        # Q1, Q2, Q3, Q4 values
        for i in range(4):
            value = reports[i]['header'].get(metric_key, 0)
            row.cells[5-i].text = str(value)
            row.cells[5-i].paragraphs[0].runs[0].font.size = Pt(9)
            center_cell_content(row.cells[5-i])
        
        # Yearly total
        if trend_key and trend_key in yearly_totals:
            yearly_value = yearly_totals[trend_key]
        else:
            yearly_value = sum(reports[i]['header'].get(metric_key, 0) for i in range(4))
        
        row.cells[1].text = str(yearly_value)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        center_cell_content(row.cells[1])
        set_cell_shading(row.cells[1], 'E7E6E6')
        
        # Trend indicator
        if trend_key and trend_key in trends:
            row.cells[0].text = trends[trend_key]
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(14)
            center_cell_content(row.cells[0])
            
            # Color code trends
            trend_value = trends[trend_key]
            if trend_value in ['↑', '↑↑']:
                set_cell_shading(row.cells[0], 'C6EFCE')  # Green
            elif trend_value in ['↓', '↓↓']:
                set_cell_shading(row.cells[0], 'FFC7CE')  # Red
            else:
                set_cell_shading(row.cells[0], 'FFEB9C')  # Yellow
        else:
            row.cells[0].text = '—'
            center_cell_content(row.cells[0])
    
    doc.add_paragraph()
    
    # =============================
    # 3. HIERARCHICAL COMPARISON TABLES (Hospital Protocol)
    # =============================

    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ph_run = ph.add_run("📋 التفاصيل التصنيفية | Classification Details")
    ph_run.font.size = Pt(13)
    ph_run.font.bold = True
    ph_run.font.name = 'Traditional Arabic'
    ph_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    doc.add_paragraph()

    hierarchies = [_build_hierarchy(r.get('classification_stats', [])) for r in reports]
    domain_data = comparison_data['domain_comparison']
    category_data = comparison_data['category_comparison']
    subcategory_data = comparison_data['subcategory_comparison']
    _create_4quarter_hierarchical_tables_by_domain(doc, hierarchies, periods)

    # =============================
    # 4. POLICY COMPLIANCE COMPARISON
    # =============================
    _add_policy_compliance_comparison(doc, reports, periods)

    # Per-quarter detailed policy compliance tables
    _add_per_quarter_policy_sections(doc, [
        (periods[i], reports[i]['header'], reports[i].get('policy_snapshot'), reports[i].get('classification_stats', []))
        for i in range(len(reports))
    ])

    # =============================
    # 5. PAGE BREAK BEFORE GRAPHS
    # =============================
    doc.add_page_break()

    # =============================
    # 5. VISUAL ANALYSIS - SPIDER CHARTS ONLY
    # =============================
    visual_heading = doc.add_paragraph()
    visual_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    visual_run = visual_heading.add_run("📊 التحليل البصري | Visual Analysis")
    visual_run.font.size = Pt(16)
    visual_run.font.bold = True
    visual_run.font.name = 'Traditional Arabic'
    visual_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')

    doc.add_paragraph()

    # Build chart list in interleaved order:
    #   Row 1: [Domain Spider  | Domain Bar   ]
    #   Row 2: [Category Spider| Category Bar ]
    #   Row 3: [SubCat Spider  | (empty)      ]
    chart_pairs_4q = []
    _bar_title_suffix = f"{periods[0]} → {periods[-1]}"

    # ── Domain level ─────────────────────────────────────────────────────────
    _4q_domain_ok, _4q_domain_skip = _is_spider_nq_useful(domain_data, max_items=8)
    if not _4q_domain_ok:
        print(f"[FORMATTER] 4Q domain spider skipped: {_4q_domain_skip}")
    if _4q_domain_ok:
        chart_pairs_4q.append((
            _generate_4quarter_spider_chart(domain_data, periods,
                                            title_ar="مخطط العنكبوت - المجالات",
                                            title_en="Domain Spider Chart"),
            "🕸️ مخطط العنكبوت - المجالات | Domain Spider Chart"
        ))

    _dom_bar_labels, _dom_bar_changes = _extract_nq_changes(domain_data, max_items=8)
    _4q_dom_bar_ok, _4q_dom_bar_skip = _is_bar_chart_useful(_dom_bar_labels, _dom_bar_changes)
    if not _4q_dom_bar_ok:
        print(f"[FORMATTER] 4Q domain bar skipped: {_4q_dom_bar_skip}")
    if _4q_dom_bar_ok:
        chart_pairs_4q.append((
            _generate_diverging_bar_chart(
                labels=_dom_bar_labels,
                changes=_dom_bar_changes,
                title=f"تغيير المجالات: {_bar_title_suffix} | Domain Net Change"
            ),
            f"📊 تغيير المجالات | Domain Net Change ({_bar_title_suffix})"
        ))

    # ── Category level ────────────────────────────────────────────────────────
    _4q_cat_ok, _4q_cat_skip = _is_spider_nq_useful(category_data, max_items=10)
    if not _4q_cat_ok:
        print(f"[FORMATTER] 4Q category spider skipped: {_4q_cat_skip}")
    if _4q_cat_ok:
        chart_pairs_4q.append((
            _generate_4quarter_spider_chart(category_data, periods,
                                            title_ar="مخطط العنكبوت - الفئات",
                                            title_en="Category Spider Chart",
                                            max_items=10),
            "🕸️ مخطط العنكبوت - الفئات | Category Spider Chart"
        ))

    _cat_bar_labels, _cat_bar_changes = _extract_nq_changes(category_data, max_items=10)
    _4q_cat_bar_ok, _4q_cat_bar_skip = _is_bar_chart_useful(_cat_bar_labels, _cat_bar_changes)
    if not _4q_cat_bar_ok:
        print(f"[FORMATTER] 4Q category bar skipped: {_4q_cat_bar_skip}")
    if _4q_cat_bar_ok:
        chart_pairs_4q.append((
            _generate_diverging_bar_chart(
                labels=_cat_bar_labels,
                changes=_cat_bar_changes,
                title=f"تغيير الفئات: {_bar_title_suffix} | Category Net Change"
            ),
            f"📊 تغيير الفئات | Category Net Change ({_bar_title_suffix})"
        ))

    # ── SubCategory level (spider only) ──────────────────────────────────────
    _4q_subcat_ok, _4q_subcat_skip = _is_spider_nq_useful(subcategory_data, max_items=12)
    if not _4q_subcat_ok:
        print(f"[FORMATTER] 4Q subcategory spider skipped: {_4q_subcat_skip}")
    if _4q_subcat_ok:
        chart_pairs_4q.append((
            _generate_4quarter_spider_chart(subcategory_data, periods,
                                            title_ar="مخطط العنكبوت - الفئات الفرعية",
                                            title_en="SubCategory Spider Chart",
                                            max_items=12),
            "🕸️ مخطط العنكبوت - الفئات الفرعية | SubCategory Spider Chart"
        ))

    _add_chart_grid(doc, chart_pairs_4q)
    return doc


def _create_4quarter_hierarchical_table(doc: Document, data: Dict[str, List[int]], periods: List[str], yearly_totals: Optional[Dict] = None, level: str = 'domain'):
    """
    Create a hierarchical table comparing 4 quarters for a specific level (domain/category/subcategory).
    
    Args:
        doc: Document object
        data: Dictionary mapping item names to list of 4 values
        periods: List of 4 period labels
        yearly_totals: Optional dictionary of yearly totals
        level: 'domain', 'category', or 'subcategory'
    """
    # Create table: 7 columns (Item | Q1 | Q2 | Q3 | Q4 | Yearly Total | Trend)
    num_rows = len(data) + 2  # +1 for header, +1 for totals
    table = doc.add_table(rows=num_rows, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = 'Table Grid'

    # Set RTL
    _tbl = table._element
    _tblPr = _tbl.tblPr
    if _tblPr is None:
        _tblPr = OxmlElement('w:tblPr')
        _tbl.insert(0, _tblPr)
    _tblPr.append(OxmlElement('w:bidiVisual'))

    # Header row
    level_label = "المجال" if level == "domain" else "الفئة" if level == "category" else "الفئة الفرعية"
    headers = ['الاتجاه | Trend', 'الإجمالي السنوي | Yearly', periods[3], periods[2], periods[1], periods[0], f'{level_label} | {level.title()}']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        center_cell_content(cell)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    totals = [0, 0, 0, 0]
    for row_idx, (item_name, values) in enumerate(sorted(data.items()), start=1):
        row = table.rows[row_idx]
        
        # Item name
        row.cells[6].text = item_name
        row.cells[6].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[6].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
        row.cells[6].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Quarter values
        for i in range(4):
            row.cells[5-i].text = str(values[i])
            row.cells[5-i].paragraphs[0].runs[0].font.size = Pt(9)
            center_cell_content(row.cells[5-i])
            totals[i] += values[i]
        
        # Yearly total for this item
        yearly = sum(values)
        row.cells[1].text = str(yearly)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        center_cell_content(row.cells[1])
        set_cell_shading(row.cells[1], 'F2F2F2')
        
        # Trend indicator
        trend = _calculate_trend_indicator(values[0], values[-1])
        row.cells[0].text = trend
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(14)
        center_cell_content(row.cells[0])
        
        # Color code trends
        if trend in ['↑', '↑↑']:
            set_cell_shading(row.cells[0], 'C6EFCE')  # Green
        elif trend in ['↓', '↓↓']:
            set_cell_shading(row.cells[0], 'FFC7CE')  # Red
        else:
            set_cell_shading(row.cells[0], 'FFEB9C')  # Yellow
    
    # Totals row
    totals_row = table.rows[-1]
    totals_row.cells[6].text = 'الإجمالي | Total'
    totals_row.cells[6].paragraphs[0].runs[0].font.bold = True
    totals_row.cells[6].paragraphs[0].runs[0].font.size = Pt(10)
    totals_row.cells[6].paragraphs[0].runs[0].font.name = 'Traditional Arabic'
    totals_row.cells[6].paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    totals_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_cell_shading(totals_row.cells[6], 'D9E2F3')
    
    for i in range(4):
        totals_row.cells[5-i].text = str(totals[i])
        totals_row.cells[5-i].paragraphs[0].runs[0].font.bold = True
        totals_row.cells[5-i].paragraphs[0].runs[0].font.size = Pt(10)
        center_cell_content(totals_row.cells[5-i])
        set_cell_shading(totals_row.cells[5-i], 'D9E2F3')
    
    # Yearly grand total
    grand_yearly = sum(totals)
    totals_row.cells[1].text = str(grand_yearly)
    totals_row.cells[1].paragraphs[0].runs[0].font.bold = True
    totals_row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    center_cell_content(totals_row.cells[1])
    set_cell_shading(totals_row.cells[1], 'BDD7EE')
    
    # Overall trend
    overall_trend = _calculate_trend_indicator(totals[0], totals[-1])
    totals_row.cells[0].text = overall_trend
    totals_row.cells[0].paragraphs[0].runs[0].font.size = Pt(14)
    totals_row.cells[0].paragraphs[0].runs[0].font.bold = True
    center_cell_content(totals_row.cells[0])
    set_cell_shading(totals_row.cells[0], 'BDD7EE')


def _generate_4quarter_spider_chart(
    data: Dict[str, List[int]],
    periods: List[str],
    title_ar: str,
    title_en: str,
    max_items: int = 8
) -> io.BytesIO:
    """4-quarter spider chart — delegates to unified engine."""
    return _generate_nquarter_spider_chart(data, periods, title_ar, title_en, max_items)



