"""
D-B8: Seasonal Word Adapter

Formats doctor and worker seasonal report data into professional Word documents.
Uses python-docx to generate structured reports with tables, charts, and RTL support.

Architecture:
- Takes data payload from D-B6/D-B7 seasonal builders
- Generates Word document bytes (no file I/O)
- Reuses existing Word generation patterns from seasonal_report_formatter.py
"""

from typing import Dict, Any
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import io


class SeasonalWordAdapter:
    """
    Adapter for converting seasonal report data into Word documents.
    
    Supports both doctor and worker reports with consistent formatting.
    """
    
    @staticmethod
    def generate_doctor_seasonal_word(data: Dict[str, Any]) -> bytes:
        """
        Generate Word document for doctor seasonal report.
        
        Args:
            data: Doctor seasonal report data from D-B6
                  Structure: {
                      'doctor_identity': {...},
                      'period': {...},
                      'metrics': {...},
                      'performance': {...},
                      'category_breakdown': [...],
                      'monthly_trend': [...],
                      'incidents_summary': {...}
                  }
        
        Returns:
            bytes: Word document content
        """
        doc = Document()
        
        # Document setup - A4 Portrait
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # ============================================================
        # TITLE (Changed to All Time Doctor Report)
        # ============================================================
        title = doc.add_heading('All Time Doctor Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ============================================================
        # DOCTOR IDENTITY
        # ============================================================
        doc.add_heading('Doctor Information', level=2)
        identity = data.get('doctor_identity', {})
        
        identity_table = doc.add_table(rows=3, cols=2)
        identity_table.style = 'Light Grid Accent 1'
        
        identity_table.cell(0, 0).text = 'Doctor ID'
        identity_table.cell(0, 1).text = str(identity.get('id', 'N/A'))
        
        identity_table.cell(1, 0).text = 'Name (Arabic)'
        identity_table.cell(1, 1).text = identity.get('name_ar', 'N/A')
        
        identity_table.cell(2, 0).text = 'Name (English)'
        identity_table.cell(2, 1).text = identity.get('name_en', 'N/A')
        
        doc.add_paragraph()  # Spacing
        
        # ============================================================
        # REPORTING PERIOD
        # ============================================================
        period = data.get('period', {})
        doc.add_heading('Reporting Period', level=2)
        
        period_table = doc.add_table(rows=3, cols=2)
        period_table.style = 'Light Grid Accent 1'
        
        period_table.cell(0, 0).text = 'Season'
        period_table.cell(0, 1).text = period.get('season_name', 'N/A')
        
        period_table.cell(1, 0).text = 'Start Date'
        period_table.cell(1, 1).text = period.get('season_start', 'N/A')
        
        period_table.cell(2, 0).text = 'End Date'
        period_table.cell(2, 1).text = period.get('season_end', 'N/A')
        
        doc.add_paragraph()
        
        # ============================================================
        # (REMOVED) PERFORMANCE SCORE SECTION
        # ============================================================
        # This section intentionally left blank as per requirements.
        doc.add_paragraph()
        
        # ============================================================
        # METRICS SUMMARY
        # ============================================================
        metrics = data.get('metrics', {})
        doc.add_heading('Performance Metrics', level=2)
        
        metrics_table = doc.add_table(rows=13, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        
        # Incident metrics
        metrics_table.cell(0, 0).text = 'Total Incidents'
        metrics_table.cell(0, 1).text = str(metrics.get('total_incidents', 0))
        
        metrics_table.cell(1, 0).text = 'High Severity'
        metrics_table.cell(1, 1).text = str(metrics.get('high_severity', 0))
        
        metrics_table.cell(2, 0).text = 'Medium Severity'
        metrics_table.cell(2, 1).text = str(metrics.get('medium_severity', 0))
        
        metrics_table.cell(3, 0).text = 'Low Severity'
        metrics_table.cell(3, 1).text = str(metrics.get('low_severity', 0))
        
        # Intent classification
        metrics_table.cell(4, 0).text = 'Good Feedback (Notice/تنويه)'
        metrics_table.cell(4, 1).text = str(metrics.get('good_feedback_count', 0))
        
        metrics_table.cell(5, 0).text = 'Bad Feedback (Critique/نقد)'
        metrics_table.cell(5, 1).text = str(metrics.get('bad_feedback_count', 0))
        
        metrics_table.cell(6, 0).text = 'Neutral Feedback'
        metrics_table.cell(6, 1).text = str(metrics.get('neutral_feedback_count', 0))
        
        # Action metrics
        metrics_table.cell(7, 0).text = 'Total Actions'
        metrics_table.cell(7, 1).text = str(metrics.get('total_actions', 0))
        
        metrics_table.cell(8, 0).text = 'Overdue Actions'
        metrics_table.cell(8, 1).text = str(metrics.get('overdue_actions', 0))
        
        metrics_table.cell(9, 0).text = 'Completed Actions'
        metrics_table.cell(9, 1).text = str(metrics.get('completed_actions', 0))
        
        # Explanation metrics
        metrics_table.cell(10, 0).text = 'Accepted Explanations'
        metrics_table.cell(10, 1).text = str(metrics.get('accepted_explanations', 0))
        
        metrics_table.cell(11, 0).text = 'Rejected Explanations'
        metrics_table.cell(11, 1).text = str(metrics.get('rejected_explanations', 0))
        
        # Calculated metrics
        total_explanations = metrics.get('accepted_explanations', 0) + metrics.get('rejected_explanations', 0)
        acceptance_rate = (metrics.get('accepted_explanations', 0) / total_explanations * 100) if total_explanations > 0 else 0
        
        metrics_table.cell(12, 0).text = 'Explanation Acceptance Rate'
        metrics_table.cell(12, 1).text = f"{acceptance_rate:.1f}%"
        
        doc.add_paragraph()
        
        # ============================================================
        # DETAILED INCIDENTS TABLE (Always show if incidents exist, now with Complaint Text)
        # ============================================================
        doc.add_heading('Incident Details', level=2)
        incidents = data.get('incidents', [])
        if incidents:
            detail_table = doc.add_table(rows=len(incidents) + 1, cols=7)
            detail_table.style = 'Light List Accent 1'
            # Header row
            header_cells = detail_table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Patient'
            header_cells[2].text = 'Category'
            header_cells[3].text = 'Severity'
            header_cells[4].text = 'Status'
            header_cells[5].text = 'Case #'
            header_cells[6].text = 'Complaint Text'
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            for idx, incident in enumerate(incidents, start=1):
                row_cells = detail_table.rows[idx].cells
                row_cells[0].text = str(incident.get('Date', 'N/A'))
                row_cells[1].text = str(incident.get('PatientName', 'N/A'))[:30]
                row_cells[2].text = str(incident.get('Category', 'N/A'))
                row_cells[3].text = str(incident.get('Severity', 'N/A'))
                row_cells[4].text = str(incident.get('Status', 'N/A'))
                row_cells[5].text = str(incident.get('RecordID', 'N/A'))
                row_cells[6].text = str(incident.get('Description', ''))[:200]
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
        else:
            # Only show the zero-complaint message if no incidents
            zero_para = doc.add_paragraph()
            zero_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            zero_run = zero_para.add_run('No incidents recorded in this period')
            zero_run.font.size = Pt(12)
            zero_run.font.color.rgb = RGBColor(0, 128, 0)
            zero_run.bold = True
            praise_para = doc.add_paragraph()
            praise_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            praise_run = praise_para.add_run('Clean record — Excellent performance')
            praise_run.font.size = Pt(11)
            praise_run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()
        
        # ============================================================
        # FOOTER NOTE
        # ============================================================
        doc.add_paragraph()
        footer_note = doc.add_paragraph('Generated by Patient Feedback System - Seasonal Report Module')
        footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_note.runs[0].font.size = Pt(9)
        footer_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    @staticmethod
    def generate_worker_seasonal_word(data: Dict[str, Any]) -> bytes:
        """
        Generate Word document for worker seasonal report.
        
        Args:
            data: Worker seasonal report data from D-B7
                  Structure: {
                      'worker_identity': {...},
                      'period': {...},
                      'metrics': {...},
                      'performance': {...}
                  }
        
        Returns:
            bytes: Word document content
        """
        doc = Document()
        
        # Document setup - A4 Portrait
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # ============================================================
        # TITLE
        # ============================================================
        title = doc.add_heading('Worker Seasonal Performance Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ============================================================
        # WORKER IDENTITY
        # ============================================================
        doc.add_heading('Worker Information', level=2)
        identity = data.get('worker_identity', {})
        
        identity_table = doc.add_table(rows=4, cols=2)
        identity_table.style = 'Light Grid Accent 1'
        
        identity_table.cell(0, 0).text = 'Employee ID'
        identity_table.cell(0, 1).text = str(identity.get('employee_id', 'N/A'))
        
        identity_table.cell(1, 0).text = 'Full Name'
        identity_table.cell(1, 1).text = identity.get('full_name', 'N/A')
        
        identity_table.cell(2, 0).text = 'Department'
        identity_table.cell(2, 1).text = identity.get('department', 'N/A')
        
        identity_table.cell(3, 0).text = 'Section'
        identity_table.cell(3, 1).text = identity.get('section', 'N/A')
        
        doc.add_paragraph()  # Spacing
        
        # ============================================================
        # REPORTING PERIOD
        # ============================================================
        period = data.get('period', {})
        doc.add_heading('Reporting Period', level=2)
        
        period_table = doc.add_table(rows=3, cols=2)
        period_table.style = 'Light Grid Accent 1'
        
        period_table.cell(0, 0).text = 'Season'
        period_table.cell(0, 1).text = period.get('season_name', 'N/A')
        
        period_table.cell(1, 0).text = 'Start Date'
        period_table.cell(1, 1).text = period.get('season_start', 'N/A')
        
        period_table.cell(2, 0).text = 'End Date'
        period_table.cell(2, 1).text = period.get('season_end', 'N/A')
        
        doc.add_paragraph()
        
        # ============================================================
        # PERFORMANCE SCORE
        # ============================================================
        performance = data.get('performance', {})
        doc.add_heading('Performance Score', level=2)
        
        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"Score: {performance.get('score', 0)}/100")
        score_run.bold = True
        score_run.font.size = Pt(14)
        
        # Color code based on praise level
        praise_level = performance.get('praise_level', 'average')
        if praise_level == 'excellent':
            score_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif praise_level == 'good':
            score_run.font.color.rgb = RGBColor(0, 100, 200)  # Blue
        elif praise_level == 'average':
            score_run.font.color.rgb = RGBColor(200, 100, 0)  # Orange
        else:
            score_run.font.color.rgb = RGBColor(200, 0, 0)  # Red
        
        doc.add_paragraph(f"Praise Level: {praise_level.upper()}")
        doc.add_paragraph(f"Risk Level: {performance.get('risk_level', 'N/A').upper()}")
        
        # Flags
        flags = performance.get('flags', [])
        if flags:
            doc.add_paragraph('Flags:')
            for flag in flags:
                doc.add_paragraph(f'  • {flag}', style='List Bullet')
        
        doc.add_paragraph()
        
        # ============================================================
        # METRICS SUMMARY
        # ============================================================
        metrics = data.get('metrics', {})
        doc.add_heading('Performance Metrics', level=2)
        
        metrics_table = doc.add_table(rows=13, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        
        # Incident metrics
        metrics_table.cell(0, 0).text = 'Total Incidents'
        metrics_table.cell(0, 1).text = str(metrics.get('total_incidents', 0))
        
        # Severity breakdown
        metrics_table.cell(1, 0).text = 'High Severity'
        metrics_table.cell(1, 1).text = str(metrics.get('high_severity', 0))
        
        metrics_table.cell(2, 0).text = 'Medium Severity'
        metrics_table.cell(2, 1).text = str(metrics.get('medium_severity', 0))
        
        metrics_table.cell(3, 0).text = 'Low Severity'
        metrics_table.cell(3, 1).text = str(metrics.get('low_severity', 0))
        
        # Intent classification
        metrics_table.cell(4, 0).text = 'Good Feedback (Notice/تنويه)'
        metrics_table.cell(4, 1).text = str(metrics.get('good_feedback_count', 0))
        
        metrics_table.cell(5, 0).text = 'Bad Feedback (Critique/نقد)'
        metrics_table.cell(5, 1).text = str(metrics.get('bad_feedback_count', 0))
        
        metrics_table.cell(6, 0).text = 'Neutral Feedback'
        metrics_table.cell(6, 1).text = str(metrics.get('neutral_feedback_count', 0))
        
        # Action metrics
        metrics_table.cell(7, 0).text = 'Total Action Items'
        metrics_table.cell(7, 1).text = str(metrics.get('total_action_items', 0))
        
        metrics_table.cell(8, 0).text = 'Overdue Actions'
        metrics_table.cell(8, 1).text = str(metrics.get('overdue_action_items', 0))
        
        metrics_table.cell(9, 0).text = 'Completed Actions'
        metrics_table.cell(9, 1).text = str(metrics.get('completed_action_items', 0))
        
        # Explanation metrics
        metrics_table.cell(10, 0).text = 'Accepted Explanations'
        metrics_table.cell(10, 1).text = str(metrics.get('explanation_accepted_count', 0))
        
        metrics_table.cell(11, 0).text = 'Rejected Explanations'
        metrics_table.cell(11, 1).text = str(metrics.get('explanation_rejected_count', 0))
        
        # Calculated metrics
        total_actions = metrics.get('total_action_items', 0)
        completion_rate = (metrics.get('completed_action_items', 0) / total_actions * 100) if total_actions > 0 else 0
        
        metrics_table.cell(12, 0).text = 'Action Completion Rate'
        metrics_table.cell(12, 1).text = f"{completion_rate:.1f}%"
        
        doc.add_paragraph()
        
        # ============================================================
        # DETAILED INCIDENTS TABLE
        # ============================================================
        doc.add_heading('Incident Details', level=2)
        
        incidents = data.get('incidents_details', [])
        if incidents:
            incidents_table = doc.add_table(rows=len(incidents) + 1, cols=8)
            incidents_table.style = 'Light List Accent 1'
            # Header row
            header_cells = incidents_table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Patient'
            header_cells[2].text = 'Category'
            header_cells[3].text = 'Severity'
            header_cells[4].text = 'Classification'
            header_cells[5].text = 'Status'
            header_cells[6].text = 'Case #'
            header_cells[7].text = 'Complaint Text'
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            classification_labels = {
                'good': 'Good ✓',
                'bad': 'Bad ✗',
                'neutral': 'Neutral ―'
            }
            # Data rows
            for idx, incident in enumerate(incidents, start=1):
                row_cells = incidents_table.rows[idx].cells
                row_cells[0].text = str(incident.get('date', 'N/A'))
                row_cells[1].text = str(incident.get('patient_name', 'N/A'))[:30]
                row_cells[2].text = str(incident.get('category', 'N/A'))
                row_cells[3].text = str(incident.get('severity', 'N/A'))
                classification = incident.get('classification', 'neutral')
                row_cells[4].text = classification_labels.get(classification, 'Neutral ―')
                row_cells[5].text = str(incident.get('status', 'N/A'))
                row_cells[6].text = str(incident.get('id', 'N/A'))
                row_cells[7].text = str(incident.get('Description', ''))[:200]
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
        else:
            # Elegant zero-complaint message
            zero_para = doc.add_paragraph()
            zero_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            zero_run = zero_para.add_run('No incidents recorded in this period')
            zero_run.font.size = Pt(12)
            zero_run.font.color.rgb = RGBColor(0, 128, 0)
            zero_run.bold = True
            praise_para = doc.add_paragraph()
            praise_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            praise_run = praise_para.add_run('Clean record — Excellent performance')
            praise_run.font.size = Pt(11)
            praise_run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()
        
        # ============================================================
        # FOOTER NOTE
        # ============================================================
        doc.add_paragraph()
        footer_note = doc.add_paragraph('Generated by Patient Feedback System - Seasonal Report Module')
        footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_note.runs[0].font.size = Pt(9)
        footer_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    @staticmethod
    def generate_patient_word_report(data: Dict[str, Any]) -> bytes:
        """
        Generate Word document for patient history report.
        
        Args:
            data: Patient export data structure: {
                'patient': {patient profile dict},
                'incidents': [list of incident dicts],
                'export_date': ISO timestamp
            }
        
        Returns:
            bytes: Word document content in Arabic with hospital branding
        """
        doc = Document()

        # Document setup - A4 Portrait
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.orientation = WD_ORIENT.PORTRAIT
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # ============================================================
        # LOGO HEADER
        # ============================================================
        try:
            logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
            if os.path.exists(logo_path):
                section.header_distance = Inches(0.1)
                header_section = section.header
                header_para = header_section.paragraphs[0]
                header_para.clear()
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = header_para.add_run()
                run.add_picture(logo_path, width=Inches(0.9))
        except Exception as e:
            print(f"[PATIENT_HISTORY] Could not add logo: {e}")

        # ============================================================
        # TITLE (Arabic)
        # ============================================================
        title = doc.add_heading('تقرير تاريخ المريض', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.name = 'Calibri'
        title_run.bold = True
        title_run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

        # Subtitle
        subtitle = doc.add_paragraph('نظام ملاحظات المرضى')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.name = 'Calibri'
        subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        subtitle.runs[0]._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
        
        doc.add_paragraph()  # Spacing
        
        # ============================================================
        # PATIENT INFORMATION (Arabic)
        # ============================================================
        patient = data.get('patient', {})
        
        info_heading = doc.add_heading('معلومات المريض', level=2)
        for r in info_heading.runs:
            r.font.name = 'Calibri'
            r._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

        patient_table = doc.add_table(rows=4, cols=2)
        patient_table.style = 'Light Grid Accent 1'
        
        patient_table.cell(0, 0).text = 'رقم المريض'
        patient_table.cell(0, 1).text = str(patient.get('patient_id', 'N/A'))
        
        patient_table.cell(1, 0).text = 'رقم الملف الطبي'
        patient_table.cell(1, 1).text = str(patient.get('mrn', 'غير متوفر'))
        
        patient_table.cell(2, 0).text = 'الاسم الكامل'
        patient_table.cell(2, 1).text = str(patient.get('full_name', 'غير متوفر'))
        
        patient_table.cell(3, 0).text = 'عدد الشكاوى'
        patient_table.cell(3, 1).text = str(len(data.get('incidents', [])))
        
        doc.add_paragraph()  # Spacing
        
        # ============================================================
        # COMPLAINTS TABLE (Arabic)
        # ============================================================
        incidents = data.get('incidents', [])
        
        complaints_heading = doc.add_heading('سجل الشكاوى', level=2)
        for r in complaints_heading.runs:
            r.font.name = 'Calibri'
            r._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
        
        if incidents:
            # Create table with header
            complaints_table = doc.add_table(rows=len(incidents) + 1, cols=6)
            complaints_table.style = 'Light List Accent 1'
            
            # Header row (Arabic)
            header_cells = complaints_table.rows[0].cells
            header_cells[0].text = 'التاريخ'
            header_cells[1].text = 'القسم'
            header_cells[2].text = 'التصنيف'
            header_cells[3].text = 'الخطورة'
            header_cells[4].text = 'الحالة'
            header_cells[5].text = 'نص الشكوى'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                        run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

            # Data rows
            for idx, incident in enumerate(incidents, start=1):
                row_cells = complaints_table.rows[idx].cells

                row_cells[0].text = str(incident.get('Date', 'N/A'))
                row_cells[1].text = str(incident.get('Department', 'غير محدد'))
                row_cells[2].text = str(incident.get('Category', 'غير محدد'))
                row_cells[3].text = str(incident.get('Severity', 'غير محدد'))
                row_cells[4].text = str(incident.get('Status', 'مفتوح'))

                complaint_text = str(incident.get('ComplaintText', 'لا يوجد نص'))
                if len(complaint_text) > 100:
                    complaint_text = complaint_text[:97] + '...'
                row_cells[5].text = complaint_text

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Calibri'
                            run._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
        else:
            no_complaints = doc.add_paragraph('لا توجد شكاوى مسجلة لهذا المريض.')
            no_complaints.runs[0].font.name = 'Calibri'
            no_complaints.runs[0]._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

        doc.add_paragraph()

        # ============================================================
        # EXPORT METADATA (Arabic)
        # ============================================================
        export_date = data.get('export_date', 'N/A')

        metadata_para = doc.add_paragraph(f'{export_date} :تاريخ التصدير')
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        metadata_para.runs[0].font.size = Pt(9)
        metadata_para.runs[0].font.name = 'Calibri'
        metadata_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        metadata_para.runs[0]._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')

        # ============================================================
        # FOOTER NOTE
        # ============================================================
        doc.add_paragraph()
        footer_note = doc.add_paragraph('تم إنشاؤه بواسطة نظام ملاحظات المرضى')
        footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_note.runs[0].font.size = Pt(9)
        footer_note.runs[0].font.name = 'Calibri'
        footer_note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        footer_note.runs[0]._element.rPr.rFonts.set(qn('w:cs'), 'Calibri')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
