"""
Smoke test for the rewritten Stylish monthly Word formatter
(backend/api/services/monthly_stylish_word_formatter.py).

Pure unit-level check — no DB, no server. Builds report_data fixtures by
hand (matching the real get_detailed_monthly_report() shape) and verifies
the generator produces a well-formed, correctly-geometried DOCX. Kept in
the repo as a lightweight regression guard since this formatter has no
other test coverage.

Run from repo root:
    python backend/test_monthly_stylish_formatter_smoke.py
"""

import os
import sys
from io import BytesIO

_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.shared import Mm

from backend.api.services.monthly_stylish_word_formatter import (
    generate_monthly_stylish_docx,
    MAX_TABLE_WIDTH_MM,
)

FAILURES = []


def check(label, condition):
    status = "PASS" if condition else "FAIL"
    print(f"[{status}] {label}")
    if not condition:
        FAILURES.append(label)


def _complaint_fixture(**overrides):
    base = {
        "id": 1001,
        "incident_id": 1001,
        "received_date": "2026-07-05",
        "incident_date": "2026-07-04",
        "publication_date": "2026-07-06",
        "domain_name": "CLINICAL",
        "category_name": "Quality of Care",
        "subcategory_name": "Examination/Monitoring Problems",
        "classification_name": "خلل في متابعة حالة المريض",
        "classification_name_en": "Examination/Monitoring Problems",
        "severity_name": "LOW",
        "stage_name": "Care on The Ward",
        "harm_level": "Moderate",
        "status_name": "Closed",
        "clinical_risk_type_name": "Ordinary",
        "section_name": "Post CSU",
        "source_name": "صندوق",
        "patient_name": "علي محمود نحلة",
        "complaint_text": "ذكرت المرافقة المريض أن الممرض سجل نسبة أوكسجين غير صحيحة، تم تصحيحها لاحقاً.",
        "immediate_action": "تم إرسال رسالة إلى رئيس اللجنة.",
        "taken_action": "تم الرد بأن الموضوع غير قابل للتطبيق حالياً.",
        "feedback_intent_type_name_ar": "شكوى",
        "target_departments": [
            {"section_id": 1, "section_name": "قسم الطوارئ",
             "department_id": 2, "department_name": "دائرة الطوارئ",
             "administration_id": 3, "administration_name": "الإدارة الطبية",
             "is_primary": True}
        ],
    }
    base.update(overrides)
    return base


def _notice_fixture(**overrides):
    base = {
        "id": 2001,
        "incident_id": 2001,
        "received_date": "2026-07-05",
        "section_name": "Post CSU",
        "source_name": "صندوق",
        "patient_name": "علي محمود نحلة",
        "notice_text": "أنوّه بالخدمة الرائعة والاداء المميز مع شكري للجميع.",
        "feedback_intent_type_name_ar": "تنويه",
        "status_name": "Closed",
        "target_departments": [
            {"section_id": 1, "section_name": "قسم الطوارئ",
             "department_id": 2, "department_name": "دائرة الطوارئ",
             "administration_id": 3, "administration_name": "الإدارة الطبية",
             "is_primary": True}
        ],
    }
    base.update(overrides)
    return base


def _period_fixture():
    return {"label": "July 2026", "label_ar": "يوليو 2026",
            "start_date": "2026-07-01", "end_date": "2026-07-31"}


def _intent_counts_fixture():
    return {
        "sections": [{"unit_name": "قسم الطوارئ", "complaint_count": 3, "notice_count": 1, "total_count": 4}],
        "departments": [{"unit_name": "دائرة الطوارئ", "complaint_count": 3, "notice_count": 1, "total_count": 4}],
        "administrations": [{"unit_name": "الإدارة الطبية", "complaint_count": 3, "notice_count": 1, "total_count": 4}],
    }


def _full_text(doc) -> str:
    """All paragraph text in the document, including inside table cells
    (doc.paragraphs alone only covers top-level body paragraphs)."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def _open_and_basic_checks(label, content: bytes):
    check(f"{label}: non-empty bytes", len(content) > 0)
    check(f"{label}: ZIP signature", content[:4] == b"PK\x03\x04")

    try:
        doc = Document(BytesIO(content))
        check(f"{label}: python-docx opens file", True)
    except Exception as e:
        check(f"{label}: python-docx opens file ({e})", False)
        return None
    return doc



# python-docx round-trips lengths through twips (1/20 pt) internally, so a
# value set as e.g. int(Mm(12)) can read back a few dozen EMU off from the
# nominal value — not a real geometry bug, just Mm<->twips<->EMU rounding.
# EMU_TOLERANCE covers that noise (1500 EMU ~= 0.042mm, imperceptible in
# print — and the 270mm design ceiling already has ~3mm slack versus the
# 273mm hard usable-width limit, so this margin costs nothing real).
EMU_TOLERANCE = 1500


def _check_geometry(label, doc):
    max_emu = Mm(MAX_TABLE_WIDTH_MM)
    for si, sec in enumerate(doc.sections):
        check(f"{label}: section {si} landscape orientation",
              sec.orientation == WD_ORIENT.LANDSCAPE)
        check(f"{label}: section {si} page wider than tall",
              sec.page_width > sec.page_height)
        check(f"{label}: section {si} left margin == 12mm",
              abs(sec.left_margin - Mm(12)) < EMU_TOLERANCE)
        check(f"{label}: section {si} top margin == 13mm",
              abs(sec.top_margin - Mm(13)) < EMU_TOLERANCE)
        check(f"{label}: section {si} bottom margin == 3mm",
              abs(sec.bottom_margin - Mm(3)) < EMU_TOLERANCE)

    for ti, table in enumerate(doc.tables):
        col_widths = [col.width for col in table.columns if col.width is not None]
        total_w = sum(col_widths)
        if total_w == 0:
            continue
        check(f"{label}: table {ti} width <= {MAX_TABLE_WIDTH_MM}mm target",
              total_w <= max_emu + EMU_TOLERANCE)

        # tblGrid gridCol count must match the table's real column count —
        # the core gotcha this file's _set_row_col_width wrapper exists to prevent
        # (Word silently ignores tcW-only width changes under tblLayout=fixed).
        tblGrid = table._tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            grid_cols = tblGrid.findall(qn('w:gridCol'))
            check(f"{label}: table {ti} tblGrid column count matches table columns",
                  len(grid_cols) == len(table.columns))


def run():
    print("=" * 70)
    print("STYLISH MONTHLY FORMATTER - SMOKE TEST")
    print("=" * 70)

    # (a) Realistic fixture: complaints + notices + intent_counts
    report_data_full = {
        "complaints": [_complaint_fixture(), _complaint_fixture(id=1002, incident_id=1002)],
        "notices": [_notice_fixture(), _notice_fixture(id=2002, incident_id=2002)],
        "period": _period_fixture(),
        "intent_counts": _intent_counts_fixture(),
    }
    content_a = generate_monthly_stylish_docx(
        report_data_full, filename="test_full.docx", language="ar",
        report_entity_name="قسم الطوارئ", report_entity_type="section",
    )
    doc_a = _open_and_basic_checks("Full fixture", content_a)
    if doc_a is not None:
        _check_geometry("Full fixture", doc_a)
        # Appendix page was removed entirely — only complaints + notices
        # sections remain (each a distinct page-geometry section).
        check("Full fixture: has >= 2 sections (complaints+notices)",
              len(doc_a.sections) >= 2)
        full_text_a = _full_text(doc_a)
        check("Full fixture: no leftover appendix title text",
              "ملحق: توزيع السجلات" not in full_text_a)
        # Both fixture complaints (and both fixture notices) share the same
        # primary target unit, so grouping collapses each into exactly one
        # batch -> exactly one signature grid each (2 total), not one per
        # record (which would have given 3+ before this change).
        check("Full fixture: exactly one signature grid per batch (2 total)",
              full_text_a.count("خاص خدمات المرضى") == 2)
        check("Full fixture: batch signature page caption present",
              "جدول التوقيع" in full_text_a)
        # Round 4 checks: yellow RCA box removed, scope strip is Arabic-only,
        # date columns dropped from the classification table.
        check("Full fixture: no leftover RCA instruction box text",
              "يلزم ملء استمارة تحليل السبب الجذري" not in full_text_a)
        check("Full fixture: scope strip has no English label text",
              "Administration" not in full_text_a and "Circle" not in full_text_a
              and "Section" not in full_text_a and "Month" not in full_text_a)
        check("Full fixture: no date-column headers in classification table",
              "الاستلام" not in full_text_a and "الحادثة" not in full_text_a
              and "النشر" not in full_text_a)
        check("Full fixture: relabeled classification headers are English-only",
              "Problem Domain" in full_text_a and "المجال" not in full_text_a
              and "Complaint Field Type" in full_text_a and "نوع السجل" not in full_text_a)

    # (b) Empty complaints + empty notices
    report_data_empty = {
        "complaints": [],
        "notices": [],
        "period": _period_fixture(),
        "intent_counts": {},
    }
    content_b = generate_monthly_stylish_docx(report_data_empty, filename="test_empty.docx")
    doc_b = _open_and_basic_checks("Empty fixture", content_b)
    if doc_b is not None:
        _check_geometry("Empty fixture", doc_b)
        full_text = _full_text(doc_b)
        check("Empty fixture: shows no-records message",
              "لا توجد سجلات" in full_text)

    # (c) Pathological long complaint (> 6000 chars) — must truncate, not crash
    long_text = "هذا نص طويل جداً لاختبار آلية القطع عند الحد الأقصى. " * 200
    check("Fixture (c): long_text actually exceeds truncation ceiling",
          len(long_text) > 6000)
    report_data_long = {
        "complaints": [_complaint_fixture(complaint_text=long_text)],
        "notices": [],
        "period": _period_fixture(),
        "intent_counts": {},
    }
    content_c = generate_monthly_stylish_docx(report_data_long, filename="test_long.docx")
    doc_c = _open_and_basic_checks("Long-complaint fixture", content_c)
    if doc_c is not None:
        _check_geometry("Long-complaint fixture", doc_c)
        full_text = _full_text(doc_c)
        check("Long-complaint fixture: truncation marker present",
              "النص الكامل في النظام" in full_text)

    # (d) Multiple distinct units within one FILTERED export (e.g. a
    # department-level export whose complaints span >1 section beneath it):
    # 2 complaints in unit A, 1 in unit B, interleaved (not pre-sorted) —
    # must still yield exactly 2 signature pages (one per unit, batch
    # contiguous regardless of input order), not 3 (one per complaint), and
    # not 0 (department is a real filtered scope, not whole-hospital — see
    # fixture (e) below for the whole-hospital/no-signature case).
    unit_a = {"section_id": 10, "section_name": "قسم أ",
              "department_id": 20, "department_name": "دائرة أ",
              "administration_id": 30, "administration_name": "إدارة أ",
              "is_primary": True}
    unit_b = {"section_id": 11, "section_name": "قسم ب",
              "department_id": 21, "department_name": "دائرة ب",
              "administration_id": 31, "administration_name": "إدارة ب",
              "is_primary": True}
    report_data_multi_unit = {
        "complaints": [
            _complaint_fixture(id=3001, incident_id=3001, target_departments=[unit_a]),
            _complaint_fixture(id=3002, incident_id=3002, target_departments=[unit_b]),
            _complaint_fixture(id=3003, incident_id=3003, target_departments=[unit_a]),
        ],
        "notices": [],
        "period": _period_fixture(),
        "intent_counts": {},
    }
    content_d = generate_monthly_stylish_docx(
        report_data_multi_unit, filename="test_multi_unit.docx",
        report_entity_name="دائرة أ", report_entity_type="department",
    )
    doc_d = _open_and_basic_checks("Multi-unit fixture", content_d)
    if doc_d is not None:
        _check_geometry("Multi-unit fixture", doc_d)
        full_text_d = _full_text(doc_d)
        check("Multi-unit fixture: exactly 2 signature grids (one per unit, not per complaint)",
              full_text_d.count("خاص خدمات المرضى") == 2)
        check("Multi-unit fixture: both unit names appear in batch captions",
              "قسم أ" in full_text_d and "قسم ب" in full_text_d)

    # (e) Whole-hospital / unfiltered export (report_entity_type is None,
    # or explicitly "hospital") spanning multiple sections — must have ZERO
    # signature grids: this file isn't routed to one physical unit to sign.
    # Complaint/notice content itself still renders normally.
    report_data_hospital_wide = {
        "complaints": [
            _complaint_fixture(id=4001, incident_id=4001, target_departments=[unit_a]),
            _complaint_fixture(id=4002, incident_id=4002, target_departments=[unit_b]),
        ],
        "notices": [_notice_fixture(id=4003, incident_id=4003)],
        "period": _period_fixture(),
        "intent_counts": {},
    }
    for _type_label, _entity_type in (("None", None), ("'hospital'", "hospital")):
        content_e = generate_monthly_stylish_docx(
            report_data_hospital_wide, filename="test_hospital_wide.docx",
            report_entity_name="مستوى المستشفى", report_entity_type=_entity_type,
        )
        doc_e = _open_and_basic_checks(f"Hospital-wide fixture (type={_type_label})", content_e)
        if doc_e is not None:
            _check_geometry(f"Hospital-wide fixture (type={_type_label})", doc_e)
            full_text_e = _full_text(doc_e)
            check(f"Hospital-wide fixture (type={_type_label}): zero signature grids",
                  full_text_e.count("خاص خدمات المرضى") == 0)
            check(f"Hospital-wide fixture (type={_type_label}): complaint content still present",
                  "قسم أ" in full_text_e and "قسم ب" in full_text_e)

    # (f) Mixed Arabic/English bidi isolation — the exact reported bug: Word's
    # own bidi resolver visually scrambles embedded English words/numbers
    # relative to surrounding Arabic when they all share one unmarked run.
    # Verifies the fix at the OOXML level: each LTR "island" (suction, tubes,
    # 2.) must land in its OWN run carrying an explicit w:rtl val="0"
    # override, and reconstructing the paragraph's runs in order must
    # reproduce the source text exactly (segmentation must be lossless).
    mixed_text = "2. ذكر أنه يريد إجراء suction للمريض، فطلب من التمريض إحضار tubes لسحب البلغم."
    report_data_bidi = {
        "complaints": [_complaint_fixture(id=5001, incident_id=5001, complaint_text=mixed_text)],
        "notices": [],
        "period": _period_fixture(),
        "intent_counts": {},
    }
    content_f = generate_monthly_stylish_docx(
        report_data_bidi, filename="test_bidi.docx",
        report_entity_name="قسم الطوارئ", report_entity_type="section",
    )
    doc_f = _open_and_basic_checks("Bidi-isolation fixture", content_f)
    if doc_f is not None:
        _check_geometry("Bidi-isolation fixture", doc_f)

        RTL_TAG = qn("w:rtl")

        def rtl_val(run):
            rPr = run._element.rPr
            el = rPr.find(RTL_TAG) if rPr is not None else None
            return el.get(qn("w:val")) if el is not None else None

        found_paragraph = False
        found_isolated_island = False
        found_marked_rtl_segment = False
        for table in doc_f.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        joined = "".join(r.text for r in p.runs)
                        if joined != mixed_text:
                            continue
                        found_paragraph = True
                        for r in p.runs:
                            if r.text in ("suction", "tubes", "2."):
                                is_isolated = rtl_val(r) == "0"
                                if r.text in ("suction", "tubes") and is_isolated:
                                    found_isolated_island = True
                                check(f"Bidi-isolation fixture: run {r.text!r} carries w:rtl=0",
                                      is_isolated)
                            elif r.text.strip():
                                # Every non-empty RTL/neutral segment must be
                                # EXPLICITLY marked rtl=1, not left to inherit
                                # from the paragraph -- an unmarked segment
                                # sandwiched between two w:rtl=0 overrides is
                                # exactly what Word's bidi algorithm reverses
                                # (confirmed against a real corrupted
                                # example). Leaving this implicit is the bug,
                                # not a style choice.
                                if rtl_val(r) == "1":
                                    found_marked_rtl_segment = True
                                check(f"Bidi-isolation fixture: RTL segment {r.text!r} carries explicit w:rtl=1",
                                      rtl_val(r) == "1")

        check("Bidi-isolation fixture: paragraph found with lossless run reconstruction",
              found_paragraph)
        check("Bidi-isolation fixture: at least one English island actually isolated",
              found_isolated_island)
        check("Bidi-isolation fixture: at least one RTL segment explicitly marked",
              found_marked_rtl_segment)

    print("=" * 70)
    if FAILURES:
        print(f"RESULT: {len(FAILURES)} check(s) FAILED")
        for f in FAILURES:
            print(f"  - {f}")
        sys.exit(1)
    else:
        print("RESULT: all checks passed")
        sys.exit(0)


if __name__ == "__main__":
    run()
