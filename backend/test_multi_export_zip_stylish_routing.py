"""
Regression test: the ZIP multi-export path (multiple sections/departments/
administrations selected at once) must respect the monthly_report_format
setting (Classical vs Stylish) the same way the single-file export path
does, and the download endpoint must not crash on non-ASCII (Arabic)
filenames.

Two real bugs found and fixed in this session, both only reproducible via
this exact path (a ZIP export, not a single-unit export):

  1. multi_report_export_service.py's docx branch hardcoded a call to the
     Classical formatter (reports_service.generate_docx_export) and never
     checked monthly_report_format at all — so switching the setting to
     "stylish" had no effect on ZIP downloads, only single-file ones.
     Fixed by mirroring report_export_service.py's routing logic.

  2. GET /api/reports/download/{export_id} built the Content-Disposition
     header directly from the export filename, which for multi-exports
     embeds the org unit's (often Arabic) name — HTTP headers must be
     latin-1 encodable, so any such download crashed with
     UnicodeEncodeError before the fix. Fixed via RFC 5987
     filename*=UTF-8''... encoding with an ASCII-safe fallback.

Temporarily flips monthly_report_format to "stylish" and ALWAYS restores
the original value in a finally block — this is a shared, persisted
setting used by the real app.

Run: python backend/test_multi_export_zip_stylish_routing.py
"""

import io
import os
import sys
import zipfile
import requests
from docx import Document

_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
_BACKEND_DIR = os.path.dirname(__file__)
for _p in (_REPO_ROOT, _BACKEND_DIR):
    if _p not in sys.path:
        sys.path.insert(0, _p)

BASE_URL = "http://127.0.0.1:8000"
USERNAME = "complaint_supervisor"
PASSWORD = "5bb5a339"

FAILURES = []


def check(label, condition):
    status = "PASS" if condition else "FAIL"
    print(f"[{status}] {label}")
    if not condition:
        FAILURES.append(label)


def find_two_nonempty_sections(s: requests.Session):
    depts = s.get(f"{BASE_URL}/api/org-units/departments", timeout=30).json()["departments"]
    sections = s.get(f"{BASE_URL}/api/org-units/sections", timeout=30).json()["sections"]
    sec_by_dept = {}
    for sec in sections:
        sec_by_dept.setdefault(sec["department_id"], []).append(sec)

    picked = []
    for dept in depts:
        for sec in sec_by_dept.get(dept["id"], []):
            r = s.post(f"{BASE_URL}/api/reports/monthly/view",
                       json={"year": 2026, "month": 1, "mode": "detailed", "section_ids": str(sec["id"])},
                       timeout=60)
            if r.status_code == 200 and r.json().get("pagination", {}).get("total_records", 0) > 0:
                picked.append(sec["id"])
                if len(picked) == 2:
                    return picked
    return picked


def run():
    print("=" * 70)
    print("MULTI-EXPORT ZIP: STYLISH ROUTING + FILENAME-ENCODING REGRESSION TEST")
    print("=" * 70)

    from api.db_layer.report_config_db import get_report_config, set_report_config

    s = requests.Session()
    r = s.post(f"{BASE_URL}/api/auth/login", json={"username": USERNAME, "password": PASSWORD}, timeout=30)
    check("login succeeds", r.status_code == 200 and r.json().get("success") is True)

    picked = find_two_nonempty_sections(s)
    check("found two sections with real data", len(picked) == 2)
    if len(picked) < 2:
        sys.exit(1)
    print(f"  sections under test: {picked}")

    original_format = get_report_config().get("monthly_report_format", "classical")
    print(f"  original monthly_report_format = '{original_format}' (will be restored)")

    try:
        set_report_config({"monthly_report_format": "stylish"})

        r = s.post(
            f"{BASE_URL}/api/reports/monthly/export",
            params={
                "year": 2026, "month": 1, "format": "docx", "display_mode": "detailed",
                "scope": "section",
                "section_ids": ",".join(str(x) for x in picked),
                "language": "ar",
            },
            timeout=180,
        )
        check("export request succeeds", r.status_code == 200)
        meta = r.json() if r.status_code == 200 else {}
        check("multi-export (ZIP) path was actually used", meta.get("is_multi_export") is True)

        download_url = meta.get("download_url")
        check("export descriptor includes a download_url", bool(download_url))
        if not download_url:
            sys.exit(1)

        dl = s.get(f"{BASE_URL}{download_url}", timeout=60)
        check("download does not crash (no UnicodeEncodeError on Arabic filename)", dl.status_code == 200)
        if dl.status_code != 200:
            sys.exit(1)

        content = dl.content
        check("downloaded content is a valid ZIP", content[:4] == b"PK\x03\x04")

        zf = zipfile.ZipFile(io.BytesIO(content))
        docx_names = [n for n in zf.namelist() if n.endswith(".docx") and not n.startswith("_SUMMARY")]
        check("ZIP contains one docx per selected section", len(docx_names) == len(picked))

        for name in docx_names:
            doc = Document(io.BytesIO(zf.read(name)))
            has_classification_table = any(len(t.columns) == 14 for t in doc.tables)
            has_scope_strip = any(len(t.columns) == 4 for t in doc.tables)
            check(f"[{name}] uses the Stylish formatter (14-col classification table present)",
                  has_classification_table)
            check(f"[{name}] uses the Stylish formatter (4-col scope strip present)",
                  has_scope_strip)
    finally:
        set_report_config({"monthly_report_format": original_format})
        restored = get_report_config().get("monthly_report_format")
        check(f"monthly_report_format restored to original ('{original_format}')", restored == original_format)

    print("=" * 70)
    if FAILURES:
        print(f"RESULT: {len(FAILURES)} check(s) FAILED")
        for f in FAILURES:
            print(f"  - {f}")
        sys.exit(1)
    else:
        print("RESULT: all checks passed")
        sys.exit(0)


if __name__ == "__main__":
    run()
