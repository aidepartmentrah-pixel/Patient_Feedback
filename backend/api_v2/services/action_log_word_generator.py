"""
Action Log Word Generator (API V2)
Generates Arabic Word documents for Action Item Accountability Log reports.

PHASE F — ACTION LOG REPORT
Cloned pattern from seasonal_report_formatter.py (DO NOT MODIFY ORIGINAL).

Features:
- A4 Landscape orientation with Arabic header and logo
- RTL table layout for Arabic readers
- Two tables: Completed and Not Completed action items
- Overdue items highlighted in red
- Hospital logo header
- Arabic formatting throughout
"""

from typing import Dict, Any
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
import io
from datetime import date


# ============================================================================
# HELPER FUNCTIONS (CLONED FROM SEASONAL GENERATOR)
# ============================================================================

def set_cell_shading(cell: _Cell, color: str):
    """
    Set background color for a table cell.
    
    Args:
        cell: The cell to shade
        color: Hex color code (e.g., 'D9E2F3' for light blue, 'FFCCCC' for red)
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def center_cell_content(cell: _Cell):
    """
    Center cell content both horizontally and vertically.
    
    Args:
        cell: The cell to center
    """
    # Horizontal centering
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Vertical centering
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_date_arabic(d: date) -> str:
    """
    Format date for Arabic display.
    
    Args:
        d: Date object
        
    Returns:
        Formatted date string (YYYY-MM-DD for clarity)
    """
    if d is None:
        return "—"
    return d.strftime("%Y-%m-%d")


# ============================================================================
# MAIN GENERATOR
# ============================================================================

def generate_action_log_word(report_data: Dict[str, Any]) -> bytes:
    """
    Generate Arabic Word document for Action Log report.
    
    This generator is CLONED from seasonal_report_formatter.py pattern.
    DO NOT modify the original seasonal generator.
    
    Input Structure:
        {
            "meta": {
                "season_id": int,
                "season_name": str | None,
                "start_date": date,
                "end_date": date,
                "generated_at": date,
                "generated_by": str | None
            },
            "completed_items": list[dict],
            "not_completed_items": list[dict],
            "totals": {
                "completed_count": int,
                "not_completed_count": int,
                "overdue_count": int
            }
        }
    
    Args:
        report_data: Action log report dataset from build_action_log_report
        
    Returns:
        Bytes of the generated Word document
    """
    
    # Utility for safe dimension conversion
    def _safe(v):
        """Convert dimension values to int (python-docx requirement)"""
        return int(v)
    
    doc = Document()
    
    # ========================================================================
    # DOCUMENT SETUP - A4 LANDSCAPE
    # ========================================================================
    
    section = doc.sections[0]
    section.page_height = _safe(Mm(210))  # A4 width becomes height in landscape
    section.page_width = _safe(Mm(297))   # A4 height becomes width in landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = _safe(Mm(15))
    section.right_margin = _safe(Mm(15))
    section.top_margin = _safe(Mm(15))
    section.bottom_margin = _safe(Mm(15))
    
    # Set default font to Traditional Arabic
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Traditional Arabic'
    font.size = Pt(11)
    
    # Extract data
    meta = report_data.get("meta", {})
    completed_items = report_data.get("completed_items", [])
    not_completed_items = report_data.get("not_completed_items", [])
    totals = report_data.get("totals", {})
    
    # ========================================================================
    # HEADER - LOGO (TOP RIGHT)
    # ========================================================================
    
    try:
        # Logo path relative to this file: backend/api_v2/services/
        # Assets at: backend/assets/logo.png
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            section.header_distance = Inches(0.1)
            header_section = section.header
            
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[ACTION_LOG_WORD] Could not add logo: {e}")
        pass
    
    # ========================================================================
    # TITLE SECTION (ARABIC)
    # ========================================================================
    
    # Main Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("تقرير سجل الإجراءات")
    title_run.font.size = int(Pt(21))
    title_run.font.bold = True
    title_run.font.name = 'Traditional Arabic'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = int(Pt(3))
    
    # Subtitle - Period
    start_date = meta.get("start_date")
    end_date = meta.get("end_date")
    period_str = f"الفترة: {format_date_arabic(start_date)} — {format_date_arabic(end_date)}"
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(period_str)
    subtitle_run.font.size = int(Pt(14))
    subtitle_run.font.name = 'Traditional Arabic'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.space_after = int(Pt(12))
    
    # ========================================================================
    # SUMMARY LINE
    # ========================================================================
    
    completed_count = totals.get("completed_count", 0)
    not_completed_count = totals.get("not_completed_count", 0)
    overdue_count = totals.get("overdue_count", 0)
    
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(
        f"إجمالي المنجز: {completed_count} | غير المنجز: {not_completed_count} (متأخر: {overdue_count})"
    )
    summary_run.font.size = int(Pt(12))
    summary_run.font.bold = True
    summary_run.font.name = 'Traditional Arabic'
    summary_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_para.space_after = int(Pt(12))
    
    # ========================================================================
    # TABLE 1 — COMPLETED ITEMS
    # ========================================================================
    
    # Section header
    completed_header_para = doc.add_paragraph()
    completed_header_run = completed_header_para.add_run("الإجراءات المنجزة")
    completed_header_run.font.size = int(Pt(16))
    completed_header_run.font.bold = True
    completed_header_run.font.name = 'Traditional Arabic'
    completed_header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    completed_header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    completed_header_para.space_after = int(Pt(6))
    
    # Create table: header + data rows
    num_completed = len(completed_items)
    completed_table = doc.add_table(rows=1 + num_completed, cols=6)
    completed_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    completed_table.style = 'Table Grid'
    
    # Header row
    completed_header_cells = completed_table.rows[0].cells
    completed_column_names = [
        "تاريخ الإنجاز",
        "تاريخ الاستحقاق",
        "القسم",
        "الموظف المسؤول",
        "العنوان",
        "رقم الإجراء"
    ]
    
    for idx, col_name in enumerate(completed_column_names):
        cell = completed_header_cells[idx]
        cell.text = col_name
        set_cell_shading(cell, 'B4C7E7')  # Light blue
        center_cell_content(cell)
        
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.font.bold = True
        run.font.name = 'Traditional Arabic'
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    # Data rows
    for row_idx, item in enumerate(completed_items, start=1):
        cells = completed_table.rows[row_idx].cells
        
        # Column values (RTL order - rightmost first)
        values = [
            format_date_arabic(item.get("completed_at").date() if item.get("completed_at") else None),
            format_date_arabic(item.get("due_date")),
            item.get("org_unit_name", "—"),
            item.get("assigned_to_display_name", "—"),
            item.get("title", "—"),
            str(item.get("action_item_id", ""))
        ]
        
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            cell.text = str(value)
            center_cell_content(cell)
            
            para = cell.paragraphs[0]
            if para.runs:
                run = para.runs[0]
                run.font.name = 'Traditional Arabic'
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    
    # Space after table
    doc.add_paragraph()
    
    # ========================================================================
    # TABLE 2 — NOT COMPLETED ITEMS
    # ========================================================================
    
    # Section header
    not_completed_header_para = doc.add_paragraph()
    not_completed_header_run = not_completed_header_para.add_run("الإجراءات غير المنجزة")
    not_completed_header_run.font.size = int(Pt(16))
    not_completed_header_run.font.bold = True
    not_completed_header_run.font.name = 'Traditional Arabic'
    not_completed_header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    not_completed_header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    not_completed_header_para.space_after = int(Pt(6))
    
    if num_not_completed := len(not_completed_items):
        # Create table
        not_completed_table = doc.add_table(rows=1 + num_not_completed, cols=6)
        not_completed_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        not_completed_table.style = 'Table Grid'
        
        # Header row
        not_completed_header_cells = not_completed_table.rows[0].cells
        not_completed_column_names = [
            "متأخر؟",
            "تاريخ الاستحقاق",
            "القسم",
            "الموظف المسؤول",
            "العنوان",
            "رقم الإجراء"
        ]
        
        for idx, col_name in enumerate(not_completed_column_names):
            cell = not_completed_header_cells[idx]
            cell.text = col_name
            set_cell_shading(cell, 'FFD966')  # Light orange
            center_cell_content(cell)
            
            para = cell.paragraphs[0]
            run = para.runs[0]
            run.font.bold = True
            run.font.name = 'Traditional Arabic'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        
        # Data rows
        for row_idx, item in enumerate(not_completed_items, start=1):
            cells = not_completed_table.rows[row_idx].cells
            
            is_overdue = item.get("is_overdue", False)
            days_overdue = item.get("days_overdue")
            
            # Overdue status text
            if is_overdue:
                overdue_text = f"نعم ({days_overdue} يوم)"
            else:
                overdue_text = "لا"
            
            # Column values (RTL order)
            values = [
                overdue_text,
                format_date_arabic(item.get("due_date")),
                item.get("org_unit_name", "—"),
                item.get("assigned_to_display_name", "—"),
                item.get("title", "—"),
                str(item.get("action_item_id", ""))
            ]
            
            for col_idx, value in enumerate(values):
                cell = cells[col_idx]
                cell.text = str(value)
                center_cell_content(cell)
                
                # Apply red shading if overdue
                if is_overdue:
                    set_cell_shading(cell, 'FFCCCC')  # Light red
                
                para = cell.paragraphs[0]
                if para.runs:
                    run = para.runs[0]
                    run.font.name = 'Traditional Arabic'
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    else:
        # No items message
        no_items_para = doc.add_paragraph()
        no_items_run = no_items_para.add_run("لا توجد إجراءات غير منجزة")
        no_items_run.font.name = 'Traditional Arabic'
        no_items_run.font.size = Pt(12)
        no_items_run.font.italic = True
        no_items_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
        no_items_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================================================
    # FOOTER - GENERATION INFO
    # ========================================================================
    
    doc.add_paragraph()
    
    footer_para = doc.add_paragraph()
    generated_at = format_date_arabic(meta.get("generated_at"))
    generated_by = meta.get("generated_by", "النظام")
    footer_text = f"تم إنشاء التقرير في {generated_at} بواسطة: {generated_by}"
    
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = 'Traditional Arabic'
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Traditional Arabic')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================================================
    # SAVE TO BYTES
    # ========================================================================
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()
