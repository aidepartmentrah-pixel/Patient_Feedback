"""
Drawer Notes Word Export Service (API V2 - Phase G-B9)

Generates a Word document export containing all non-deleted drawer notes with their labels.
Uses existing Word export patterns from action_log_word_generator.py.

PHASE G — G-B9 — WORD EXPORT BUILDER
"""

from typing import Dict, Any, List
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
import io
from datetime import datetime, timezone

from api_v2.db_layer import drawer_note_db


# ============================================================================
# MAIN EXPORT BUILDER
# ============================================================================

def build_drawer_notes_word_export() -> bytes:
    """
    Generate Word document export of all non-deleted drawer notes.
    
    Document structure:
    - Header: System name, title "Drawer Notes Registry", timestamp
    - For each note (DESC by created_at):
      - Note ID
      - Created At
      - Author
      - Labels (comma-separated or "-")
      - Note Text
      - Separator line
    
    Returns:
        Bytes of the generated Word document
    """
    
    # Utility for safe dimension conversion
    def _safe(v):
        """Convert dimension values to int (python-docx requirement)"""
        return int(v)
    
    # Load all notes with labels from DB
    notes = drawer_note_db.get_all_notes_with_labels()
    
    doc = Document()
    
    # ========================================================================
    # DOCUMENT SETUP - A4 PORTRAIT
    # ========================================================================
    
    section = doc.sections[0]
    section.page_height = _safe(Mm(297))  # A4 height
    section.page_width = _safe(Mm(210))   # A4 width
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = _safe(Mm(20))
    section.right_margin = _safe(Mm(20))
    section.top_margin = _safe(Mm(20))
    section.bottom_margin = _safe(Mm(20))
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========================================================================
    # HEADER - LOGO (TOP RIGHT)
    # ========================================================================
    
    try:
        # Logo path relative to this file: backend/api_v2/services/
        # Assets at: backend/assets/logo.png
        logo_path = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'logo.png')
        if os.path.exists(logo_path):
            section.header_distance = Inches(0.1)
            header_section = section.header
            
            header_para = header_section.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[DRAWER_NOTES_EXPORT] Could not add logo: {e}")
        pass
    
    # ========================================================================
    # TITLE SECTION
    # ========================================================================
    
    # System/Hospital Name
    system_para = doc.add_paragraph()
    system_run = system_para.add_run("Al-Rasoul Al-Adham Hospital")
    system_run.font.size = int(Pt(12))
    system_run.font.bold = True
    system_run.font.name = 'Calibri'
    system_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    system_para.space_after = int(Pt(3))
    
    # Main Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Drawer Notes Registry")
    title_run.font.size = int(Pt(18))
    title_run.font.bold = True
    title_run.font.name = 'Calibri'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = int(Pt(3))
    
    # Generated at timestamp
    generated_at = datetime.now(timezone.utc)
    timestamp_str = f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M:%S')} UTC"
    
    timestamp_para = doc.add_paragraph()
    timestamp_run = timestamp_para.add_run(timestamp_str)
    timestamp_run.font.size = int(Pt(10))
    timestamp_run.font.name = 'Calibri'
    timestamp_run.font.italic = True
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_para.space_after = int(Pt(16))
    
    # ========================================================================
    # NOTES CONTENT
    # ========================================================================
    
    if not notes:
        # No notes case
        no_notes_para = doc.add_paragraph()
        no_notes_run = no_notes_para.add_run("No notes available.")
        no_notes_run.font.size = int(Pt(11))
        no_notes_run.font.italic = True
        no_notes_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for idx, note in enumerate(notes):
            # Note header section
            header_para = doc.add_paragraph()
            header_para.space_before = int(Pt(12)) if idx > 0 else int(Pt(6))
            header_para.space_after = int(Pt(6))
            
            # Note ID
            id_run = header_para.add_run(f"Note ID: {note['note_id']}")
            id_run.font.bold = True
            id_run.font.size = int(Pt(11))
            
            # Created At
            created_at = note.get('created_at')
            if created_at:
                created_str = created_at.strftime('%Y-%m-%d %H:%M:%S') if isinstance(created_at, datetime) else str(created_at)
            else:
                created_str = "N/A"
            
            created_para = doc.add_paragraph()
            created_run = created_para.add_run(f"Created At: {created_str}")
            created_run.font.size = int(Pt(10))
            created_para.space_after = int(Pt(3))
            
            # Author
            author_name = note.get('created_by_name', 'Unknown')
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"Author: {author_name}")
            author_run.font.size = int(Pt(10))
            author_para.space_after = int(Pt(3))
            
            # Labels
            label_names = note.get('label_names', [])
            if label_names:
                labels_str = ", ".join(label_names)
            else:
                labels_str = "-"
            
            labels_para = doc.add_paragraph()
            labels_run = labels_para.add_run(f"Labels: {labels_str}")
            labels_run.font.size = int(Pt(10))
            labels_para.space_after = int(Pt(6))
            
            # Note text heading
            text_heading_para = doc.add_paragraph()
            text_heading_run = text_heading_para.add_run("Text:")
            text_heading_run.font.bold = True
            text_heading_run.font.size = int(Pt(10))
            text_heading_para.space_after = int(Pt(3))
            
            # Note text content
            note_text = note.get('note_text', '')
            text_para = doc.add_paragraph()
            text_run = text_para.add_run(note_text)
            text_run.font.size = int(Pt(11))
            text_para.space_after = int(Pt(8))
            
            # Separator line (except for last note)
            if idx < len(notes) - 1:
                separator_para = doc.add_paragraph()
                separator_run = separator_para.add_run("─" * 80)
                separator_run.font.size = int(Pt(8))
                separator_run.font.color.rgb = RGBColor(200, 200, 200)
                separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator_para.space_after = int(Pt(8))
    
    # ========================================================================
    # FOOTER
    # ========================================================================
    
    total_notes = len(notes)
    footer_text = f"Total Notes: {total_notes}"
    
    footer_para = doc.add_paragraph()
    footer_para.space_before = int(Pt(16))
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========================================================================
    # SAVE TO BYTES
    # ========================================================================
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()
