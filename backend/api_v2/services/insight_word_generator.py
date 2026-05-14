"""
Insight Word Generator
Generates Arabic RTL Word document representing the Insight page operational state.

This is NOT a historical analytics report.
It represents current pending workload, waiting states, and follow-up responsibilities.

Document sections mirror the three Insight page tabs:
  1. Active Cases (grouped by Sections / Departments / Administrations)
  2. Force Closed Drafts (data incomplete)
  3. Force Closed Completed
"""

import os
import io
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell


# ============================================================================
# STYLE HELPERS
# ============================================================================

def _set_cell_shading(cell: _Cell, color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_para_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_table_rtl(table):
    tbl_elem = table._tbl
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)


def _apply_font(run, size: int, bold: bool = False, color: RGBColor = None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = False
    if color:
        run.font.color.rgb = color
    # Set complex-script (Arabic) font explicitly via XML
    try:
        rPr = run._element.rPr
        if rPr is not None and rPr.rFonts is not None:
            rPr.rFonts.set(qn('w:cs'), 'Calibri')
    except Exception:
        pass


def _write_cell(cell: _Cell, text: str, bold: bool = False, size: int = 9,
                center: bool = True, bg_color: str = None):
    cell.text = str(text)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    if not center:
        _set_para_rtl(para)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if para.runs:
        _apply_font(para.runs[0], size=size, bold=bold)
    if bg_color:
        _set_cell_shading(cell, bg_color)


def _severity_arabic(severity: str) -> str:
    return {
        'HIGH': 'عالية',
        'MEDIUM': 'متوسطة',
        'LOW': 'منخفضة',
        'NEUTRAL': 'محايدة',
    }.get((severity or '').upper(), severity or '—')


def _status_arabic(status: str) -> str:
    return {
        'SUBMITTED_TO_SECTION': 'مقدم للقسم',
        'SECTION_ACCEPTED': 'مقبول من القسم',
        'SECTION_ACCEPTED_PENDING_DEPT': 'بانتظار الدائرة',
        'DEPT_ACCEPTED_PENDING_ADMIN': 'بانتظار الإدارة',
        'ADMIN_APPROVED': 'معتمد',
        'SECTION_DENIED': 'مرفوض',
        'FORCE_CLOSED': 'مغلق بالإجبار',
        'FORCE_CLOSED_DRAFT': 'مغلق - مسودة',
        'FORCE_CLOSED_COMPLETE': 'مغلق - مكتمل',
    }.get(status or '', status or '—')


def _format_date(value) -> str:
    if value is None:
        return '—'
    try:
        if hasattr(value, 'strftime'):
            return value.strftime('%Y-%m-%d')
        return str(value)[:10]
    except Exception:
        return str(value)


# ============================================================================
# SECTION RENDERERS
# ============================================================================

def _render_section_heading(doc: Document, text: str, font_color: RGBColor,
                            fill_color: str = None):
    para = doc.add_paragraph()
    _set_para_rtl(para)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    _apply_font(run, size=15, bold=True, color=font_color)
    if fill_color:
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill_color)
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)


def _render_active_cases(doc: Document, grouped_inbox: List[Dict[str, Any]]):
    _render_section_heading(
        doc, "الحالات النشطة — حالات بانتظار إجراء",
        font_color=RGBColor(0x1F, 0x38, 0x64),
        fill_color='DAEEF3'
    )

    if not grouped_inbox:
        para = doc.add_paragraph()
        _set_para_rtl(para)
        run = para.add_run("لا توجد حالات نشطة معلقة حالياً ✓")
        _apply_font(run, size=11, color=RGBColor(0x00, 0x80, 0x00))
        return

    ORG_CONFIG = {
        'SECTION': {
            'label': 'الأقسام',
            'header_fill': 'CDFBF0',
            'row_fill': '00B894',
            'font_color': RGBColor(0x00, 0x70, 0x60),
        },
        'DEPARTMENT': {
            'label': 'الدوائر',
            'header_fill': 'E8E4FF',
            'row_fill': '6C5CE7',
            'font_color': RGBColor(0x40, 0x30, 0xA0),
        },
        'ADMINISTRATION': {
            'label': 'الإدارات',
            'header_fill': 'FFE5DF',
            'row_fill': 'E17055',
            'font_color': RGBColor(0xA0, 0x40, 0x20),
        },
    }

    # Group units by org type
    org_groups: Dict[str, List] = {k: [] for k in ORG_CONFIG}
    for unit in grouped_inbox:
        key = (unit.get('org_type') or 'SECTION').upper()
        if key not in org_groups:
            key = 'SECTION'
        org_groups[key].append(unit)

    for org_key in ['SECTION', 'DEPARTMENT', 'ADMINISTRATION']:
        units = org_groups[org_key]
        if not units:
            continue

        cfg = ORG_CONFIG[org_key]

        # Org type group header
        grp_para = doc.add_paragraph()
        _set_para_rtl(grp_para)
        grp_para.paragraph_format.space_before = Pt(8)
        grp_para.paragraph_format.space_after = Pt(4)
        total_pending = sum(u.get('pending_count', 0) for u in units)
        grp_run = grp_para.add_run(
            f"  {cfg['label']}  ({len(units)} وحدة — {total_pending} حالة معلقة)  "
        )
        _apply_font(grp_run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        pPr = grp_para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), cfg['row_fill'])
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)

        for unit in units:
            unit_name = unit.get('section_name', '—')
            supervisor = unit.get('supervisor_name', 'غير محدد')
            pending_count = unit.get('pending_count', 0)
            subcases = unit.get('subcases', [])

            # Unit sub-heading
            unit_para = doc.add_paragraph()
            _set_para_rtl(unit_para)
            unit_para.paragraph_format.space_before = Pt(6)
            unit_para.paragraph_format.space_after = Pt(2)
            unit_run = unit_para.add_run(
                f"{unit_name}  |  المشرف: {supervisor}  |  الحالات المعلقة: {pending_count}"
            )
            _apply_font(unit_run, size=11, bold=True, color=cfg['font_color'])

            if not subcases:
                continue

            cols = [
                "رقم الحادثة",
                "اسم المريض",
                "التصنيف",
                "الخطورة",
                "الحالة الحالية",
                "أيام الانتظار",
            ]
            tbl = doc.add_table(rows=1 + len(subcases), cols=len(cols))
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
            _set_table_rtl(tbl)

            # Header row
            for ci, col_name in enumerate(cols):
                _write_cell(tbl.rows[0].cells[ci], col_name,
                            bold=True, size=10, bg_color=cfg['header_fill'])

            # Data rows
            for ri, subcase in enumerate(subcases, start=1):
                row_cells = tbl.rows[ri].cells
                incident_num = (subcase.get('incident_number') or
                                subcase.get('incident_id') or '—')
                values = [
                    incident_num,
                    subcase.get('patient_name', '—'),
                    subcase.get('category', '—'),
                    _severity_arabic(subcase.get('severity', '')),
                    _status_arabic(subcase.get('status', '')),
                    str(subcase.get('waiting_days', 0)),
                ]
                alt_bg = 'F7F7F7' if ri % 2 == 0 else None
                for ci, val in enumerate(values):
                    _write_cell(row_cells[ci], val, size=9, bg_color=alt_bg)

            doc.add_paragraph()


def _render_fc_section(doc: Document, cases: List[Dict[str, Any]],
                       heading: str, font_color: RGBColor, fill_color: str,
                       header_bg: str):
    _render_section_heading(doc, heading, font_color=font_color, fill_color=fill_color)

    if not cases:
        para = doc.add_paragraph()
        _set_para_rtl(para)
        run = para.add_run("لا توجد حالات في هذه الفئة")
        _apply_font(run, size=10, color=RGBColor(0x80, 0x80, 0x80))
        return

    cols = [
        "رقم الحادثة",
        "اسم المريض",
        "الوحدة المستهدفة",
        "سبب الإغلاق",
        "تاريخ الإغلاق",
        "أيام الانتظار",
    ]
    tbl = doc.add_table(rows=1 + len(cases), cols=len(cols))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_rtl(tbl)

    for ci, col_name in enumerate(cols):
        _write_cell(tbl.rows[0].cells[ci], col_name,
                    bold=True, size=10, bg_color=header_bg)

    for ri, case in enumerate(cases, start=1):
        row_cells = tbl.rows[ri].cells
        incident_num = (case.get('incident_number') or
                        case.get('incident_request_case_id') or '—')
        fc_date = _format_date(case.get('force_closed_at') or case.get('created_at'))
        reason = case.get('force_close_reason') or '—'
        alt_bg = 'F9F9F9' if ri % 2 == 0 else None
        values = [
            incident_num,
            case.get('patient_name', '—'),
            case.get('org_unit_name', '—'),
            reason,
            fc_date,
            str(case.get('waiting_days', 0)),
        ]
        for ci, val in enumerate(values):
            _write_cell(row_cells[ci], val, size=9, bg_color=alt_bg)


# ============================================================================
# MAIN GENERATOR
# ============================================================================

def generate_insight_word_report(
    grouped_inbox: List[Dict[str, Any]],
    fc_draft_cases: List[Dict[str, Any]],
    fc_complete_cases: List[Dict[str, Any]],
    generated_by: str = "النظام",
) -> bytes:
    """
    Generate Arabic RTL Word document for Insight page operational state.

    Args:
        grouped_inbox: Active cases grouped by org unit
        fc_draft_cases: Force-closed draft subcases
        fc_complete_cases: Force-closed completed subcases
        generated_by: Display name of the generating user

    Returns:
        bytes of the .docx document
    """

    def _safe(v):
        return int(v)

    doc = Document()

    # ── Document setup: A4 Portrait ──────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_height = _safe(Mm(297))
    sec.page_width = _safe(Mm(210))
    sec.left_margin = _safe(Mm(20))
    sec.right_margin = _safe(Mm(20))
    sec.top_margin = _safe(Mm(20))
    sec.bottom_margin = _safe(Mm(20))

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.italic = False

    # ── Logo (top right in header) ────────────────────────────────────────────
    try:
        logo_path = os.path.join(
            os.path.dirname(__file__), '..', '..', 'assets', 'logo.png'
        )
        if os.path.exists(logo_path):
            header_para = sec.header.paragraphs[0]
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header_para.add_run().add_picture(logo_path, width=Inches(0.9))
    except Exception as e:
        print(f"[INSIGHT_WORD] Logo error: {e}")

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_rtl(title_para)
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run("تقرير الأعمال المعلقة الحالية")
    _apply_font(title_run, size=20, bold=True, color=RGBColor(0x1F, 0x38, 0x64))

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(2)
    sub_run = sub_para.add_run(
        "نظام ملاحظات المرضى — عرض أعباء العمل الحالية والحالات المعلقة"
    )
    _apply_font(sub_run, size=11, color=RGBColor(0x70, 0x70, 0x70))

    now_str = datetime.now().strftime("%Y-%m-%d  %H:%M")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(8)
    date_run = date_para.add_run(
        f"تاريخ الإنشاء: {now_str}  |  أُنشئ بواسطة: {generated_by}"
    )
    _apply_font(date_run, size=10, color=RGBColor(0x50, 0x50, 0x50))

    # ── Summary cards row ─────────────────────────────────────────────────────
    total_active = sum(u.get('pending_count', 0) for u in grouped_inbox)
    total_units = len(grouped_inbox)

    summary_tbl = doc.add_table(rows=1, cols=4)
    summary_tbl.style = 'Table Grid'
    summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    cards = [
        ("الحالات النشطة", str(total_active), 'DAEEF3'),
        ("الوحدات المعلقة", str(total_units), 'E2EFDA'),
        ("مغلقة إجباراً — مسودة", str(len(fc_draft_cases)), 'FFF2CC'),
        ("مغلقة إجباراً — مكتملة", str(len(fc_complete_cases)), 'F2F2F2'),
    ]

    for ci, (label, value, color) in enumerate(cards):
        cell = summary_tbl.rows[0].cells[ci]
        cell.text = ''
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, color)

        p = cell.paragraphs[0]
        lbl_run = p.add_run(label + "\n")
        _apply_font(lbl_run, size=9, color=RGBColor(0x40, 0x40, 0x40))

        val_run = p.add_run(value)
        _apply_font(val_run, size=18, bold=True)

    doc.add_paragraph()

    # ── Section 1: Active Cases ───────────────────────────────────────────────
    _render_active_cases(doc, grouped_inbox)

    # ── Page break before FC sections ────────────────────────────────────────
    if fc_draft_cases or fc_complete_cases:
        doc.add_page_break()

    # ── Section 2: Force Closed Drafts ────────────────────────────────────────
    _render_fc_section(
        doc, fc_draft_cases,
        heading="الحالات المغلقة بالإجبار — مسودة (البيانات غير مكتملة)",
        font_color=RGBColor(0xB7, 0x6E, 0x00),
        fill_color='FFF2CC',
        header_bg='FFE599',
    )

    doc.add_paragraph()

    # ── Section 3: Force Closed Completed ────────────────────────────────────
    _render_fc_section(
        doc, fc_complete_cases,
        heading="الحالات المغلقة بالإجبار — مكتملة",
        font_color=RGBColor(0x40, 0x40, 0x40),
        fill_color='F2F2F2',
        header_bg='D9D9D9',
    )

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"تم إنشاء هذا التقرير في {now_str}  |  نظام ملاحظات المرضى"
    )
    _apply_font(footer_run, size=9, color=RGBColor(0xA0, 0xA0, 0xA0))

    # ── Save ──────────────────────────────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
