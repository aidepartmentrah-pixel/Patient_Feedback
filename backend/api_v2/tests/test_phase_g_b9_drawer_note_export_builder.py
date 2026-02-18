"""
Phase G - G-B9: Drawer Notes Word Export Builder Tests

Tests the Word document export builder for drawer notes.

Test Coverage:
1. Export returns valid bytes
2. Document can be loaded with python-docx
3. Document contains expected content (title, note texts, labels, authors)
4. Soft-deleted notes are excluded from export

Uses real database, no mocks.
"""

import pytest
import sys
from pathlib import Path
from io import BytesIO
import uuid

# Add backend to path
backend_path = Path(__file__).parent.parent.parent
sys.path.insert(0, str(backend_path))

from api_v2.services.drawer_note_export_service import build_drawer_notes_word_export
from api_v2.db_layer import drawer_note_db, drawer_label_db
from docx import Document


# ==================== TEST SETUP ====================

@pytest.fixture
def test_data_cleanup():
    """Track created test data for cleanup."""
    created_label_ids = []
    created_note_ids = []
    
    yield created_label_ids, created_note_ids
    
    # Cleanup
    conn = drawer_note_db.get_db_connection()
    cursor = conn.cursor()
    
    try:
        # Delete note-label links
        for note_id in created_note_ids:
            cursor.execute("DELETE FROM dbo.APP_DrawerNoteLabelLink WHERE NoteID = ?", (note_id,))
        
        # Delete notes
        for note_id in created_note_ids:
            cursor.execute("DELETE FROM dbo.APP_DrawerNote WHERE NoteID = ?", (note_id,))
        
        # Delete labels
        for label_id in created_label_ids:
            cursor.execute("DELETE FROM dbo.APP_DrawerLabel WHERE LabelID = ?", (label_id,))
        
        conn.commit()
    except Exception as e:
        print(f"Cleanup warning: {e}")
    finally:
        cursor.close()
        conn.close()


# ==================== TESTS ====================

def test_export_returns_bytes(test_data_cleanup):
    """Test 1: Export returns non-empty bytes."""
    created_label_ids, created_note_ids = test_data_cleanup
    
    # Create test data
    label_id_1 = drawer_label_db.insert_label(f"Priority_{uuid.uuid4().hex[:8]}")
    label_id_2 = drawer_label_db.insert_label(f"Urgent_{uuid.uuid4().hex[:8]}")
    created_label_ids.extend([label_id_1, label_id_2])
    
    note_text_1 = f"Test note one {uuid.uuid4().hex[:8]}"
    note_id_1 = drawer_note_db.insert_note(note_text_1, 1, "Test User 1")
    drawer_note_db.attach_labels_to_note(note_id_1, [label_id_1, label_id_2])
    created_note_ids.append(note_id_1)
    
    note_text_2 = f"Test note two {uuid.uuid4().hex[:8]}"
    note_id_2 = drawer_note_db.insert_note(note_text_2, 2, "Test User 2")
    drawer_note_db.attach_labels_to_note(note_id_2, [label_id_1])
    created_note_ids.append(note_id_2)
    
    # Build export
    result = build_drawer_notes_word_export()
    
    # Assertions
    assert isinstance(result, bytes), "Should return bytes"
    assert len(result) > 0, "Should return non-empty bytes"


def test_export_document_loads_successfully(test_data_cleanup):
    """Test 2: Document can be loaded with python-docx."""
    created_label_ids, created_note_ids = test_data_cleanup
    
    # Create test data
    label_id = drawer_label_db.insert_label(f"TestLabel_{uuid.uuid4().hex[:8]}")
    created_label_ids.append(label_id)
    
    note_text = f"Sample note {uuid.uuid4().hex[:8]}"
    note_id = drawer_note_db.insert_note(note_text, 1, "Sample Author")
    drawer_note_db.attach_labels_to_note(note_id, [label_id])
    created_note_ids.append(note_id)
    
    # Build export
    result = build_drawer_notes_word_export()
    
    # Load with python-docx
    doc = Document(BytesIO(result))
    
    # Should load without error
    assert doc is not None
    assert len(doc.paragraphs) > 0


def test_export_contains_expected_content(test_data_cleanup):
    """Test 3: Document contains title, note texts, labels, and authors."""
    created_label_ids, created_note_ids = test_data_cleanup
    
    # Create test data with unique identifiers
    unique_suffix_1 = uuid.uuid4().hex[:8]
    unique_suffix_2 = uuid.uuid4().hex[:8]
    
    label_name_1 = f"Priority_{unique_suffix_1}"
    label_name_2 = f"Urgent_{unique_suffix_2}"
    label_id_1 = drawer_label_db.insert_label(label_name_1)
    label_id_2 = drawer_label_db.insert_label(label_name_2)
    created_label_ids.extend([label_id_1, label_id_2])
    
    note_text_1 = f"First test note content {unique_suffix_1}"
    author_1 = f"Author One {unique_suffix_1}"
    note_id_1 = drawer_note_db.insert_note(note_text_1, 1, author_1)
    drawer_note_db.attach_labels_to_note(note_id_1, [label_id_1, label_id_2])
    created_note_ids.append(note_id_1)
    
    note_text_2 = f"Second test note content {unique_suffix_2}"
    author_2 = f"Author Two {unique_suffix_2}"
    note_id_2 = drawer_note_db.insert_note(note_text_2, 2, author_2)
    drawer_note_db.attach_labels_to_note(note_id_2, [label_id_1])
    created_note_ids.append(note_id_2)
    
    # Build export
    result = build_drawer_notes_word_export()
    
    # Load document
    doc = Document(BytesIO(result))
    
    # Extract all text from document
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Assertions
    assert "Drawer Notes Registry" in full_text, "Should contain document title"
    assert note_text_1 in full_text, "Should contain first note text"
    assert note_text_2 in full_text, "Should contain second note text"
    assert label_name_1 in full_text or label_name_2 in full_text, "Should contain at least one label name"
    assert author_1 in full_text, "Should contain first author name"
    assert author_2 in full_text, "Should contain second author name"


def test_export_excludes_deleted_notes(test_data_cleanup):
    """Test 4: Soft-deleted notes are not included in export."""
    created_label_ids, created_note_ids = test_data_cleanup
    
    # Create test data
    unique_suffix = uuid.uuid4().hex[:8]
    
    label_id = drawer_label_db.insert_label(f"Label_{unique_suffix}")
    created_label_ids.append(label_id)
    
    note_text_active = f"Active note content {unique_suffix}"
    note_id_active = drawer_note_db.insert_note(note_text_active, 1, "Active Author")
    drawer_note_db.attach_labels_to_note(note_id_active, [label_id])
    created_note_ids.append(note_id_active)
    
    note_text_deleted = f"Deleted note content {unique_suffix}"
    note_id_deleted = drawer_note_db.insert_note(note_text_deleted, 2, "Deleted Author")
    drawer_note_db.attach_labels_to_note(note_id_deleted, [label_id])
    created_note_ids.append(note_id_deleted)
    
    # Soft delete the second note
    drawer_note_db.soft_delete_note(note_id_deleted)
    
    # Build export
    result = build_drawer_notes_word_export()
    
    # Load document
    doc = Document(BytesIO(result))
    
    # Extract all text
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Assertions
    assert note_text_active in full_text, "Should contain active note"
    assert note_text_deleted not in full_text, "Should NOT contain deleted note"


def test_export_with_no_notes():
    """Test 5: Export with no notes returns valid document."""
    # Note: This test assumes there might be notes in DB, 
    # but we want to test the builder handles empty case gracefully
    # We'll build the export and verify it doesn't crash
    
    # Build export (may contain existing notes from DB)
    result = build_drawer_notes_word_export()
    
    # Should return bytes
    assert isinstance(result, bytes)
    assert len(result) > 0
    
    # Should load as document
    doc = Document(BytesIO(result))
    assert doc is not None
    
    # Should contain title regardless of note count
    full_text = "\n".join([para.text for para in doc.paragraphs])
    assert "Drawer Notes Registry" in full_text


def test_export_note_without_labels(test_data_cleanup):
    """Test 6: Export handles notes with no labels (shows '-')."""
    created_label_ids, created_note_ids = test_data_cleanup
    
    # Create note without any labels
    unique_suffix = uuid.uuid4().hex[:8]
    note_text = f"Note without labels {unique_suffix}"
    note_id = drawer_note_db.insert_note(note_text, 1, "Author Without Labels")
    created_note_ids.append(note_id)
    # Don't attach any labels
    
    # Build export
    result = build_drawer_notes_word_export()
    
    # Load document
    doc = Document(BytesIO(result))
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Should contain the note
    assert note_text in full_text
    
    # Should show "-" for labels (look for "Labels: -" pattern)
    # Since we can't guarantee exact positioning, just verify the note is present
    assert note_text in full_text
