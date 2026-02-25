"""
Phase 6 Test: Doctor Good/Bad Classification

Tests:
1. doctors_db.get_doctor_statistics() returns good/bad/neutral feedback counts
2. doctors_db.get_doctor_incidents() returns classification field per incident
3. doctor_seasonal_reporting_service builds correct metrics with intent classification
4. seasonal_word_adapter.generate_doctor_seasonal_word() includes classification in metrics + incidents  
5. Fixed bug: DB keys (total/high/medium/low) correctly mapped in seasonal service
6. Zero-incident doctor gets elegant message
"""
import sys
import os
sys.path.insert(0, '.')
sys.path.insert(0, os.path.join('.', 'backend'))

from backend.core.database import get_connection
from backend.api.db_layer.doctors_db import get_doctor_statistics, get_doctor_incidents
from backend.api.services.doctor_seasonal_reporting_service import DoctorSeasonalReportingService
from backend.api.services.seasonal_word_adapter import SeasonalWordAdapter
from backend.api.services.person_report_word_adapter import generate_person_seasonal_word_report
from docx import Document
import io
from datetime import date

PASS = 0
FAIL = 0

def test(name, condition, detail=""):
    global PASS, FAIL
    if condition:
        PASS += 1
        icon = "✅"
    else:
        FAIL += 1
        icon = "❌"
    print(f"  {icon} {name}" + (f" — {detail}" if detail else ""))


def run_all_tests():
    global PASS, FAIL

    print("=" * 70)
    print("PHASE 6 TEST SUITE: Doctor Good/Bad Classification")
    print("=" * 70)

    # =========================================================================
    # SETUP: Find a doctor with incidents
    # =========================================================================
    conn = get_connection()
    cursor = conn.cursor()

    # Find doctor linked to incidents via APP_IncidentCaseDoctor
    cursor.execute("""
        SELECT TOP 1 icd.DoctorID, COUNT(*) as cnt
        FROM dbo.APP_IncidentCaseDoctor icd
        GROUP BY icd.DoctorID
        ORDER BY cnt DESC
    """)
    row = cursor.fetchone()
    if not row:
        print("❌ FATAL: No doctors linked in APP_IncidentCaseDoctor")
        return
    doc_id = row.DoctorID
    doc_count = row.cnt
    print(f"\n--- Doctor with incidents: DoctorID={doc_id} (count={doc_count}) ---")

    # Get raw SQL expected values for validation
    cursor.execute("""
        SELECT
            COUNT(*) as total,
            SUM(CASE WHEN ic.FeedbackIntentTypeID = 2 THEN 1 ELSE 0 END) as good,
            SUM(CASE WHEN ic.FeedbackIntentTypeID = 3 THEN 1 ELSE 0 END) as bad,
            SUM(CASE WHEN ic.FeedbackIntentTypeID IN (1,4) OR ic.FeedbackIntentTypeID IS NULL THEN 1 ELSE 0 END) as neutral
        FROM dbo.APP_IncidentCaseDoctor icd
        INNER JOIN dbo.APP_IncidentCase ic ON icd.IncidentRequestCaseID = ic.IncidentRequestCaseID
        WHERE icd.DoctorID = ?
    """, (doc_id,))
    raw = cursor.fetchone()
    expected_good = raw.good or 0
    expected_bad = raw.bad or 0
    expected_neutral = raw.neutral or 0
    print(f"   Expected intent: Good={expected_good}, Bad={expected_bad}, Neutral={expected_neutral}")
    conn.close()

    # =========================================================================
    # TEST GROUP 1: DB layer — get_doctor_statistics() with intent fields
    # =========================================================================
    print("\n--- TEST GROUP 1: doctors_db.get_doctor_statistics() ---")

    stats = get_doctor_statistics(
        doctor_id=doc_id,
        from_date='2020-01-01',
        to_date='2030-12-31'
    )

    test("stats has 'total' key", 'total' in stats, f"keys={list(stats.keys())}")
    test("stats has 'high' key", 'high' in stats)
    test("stats has 'good_feedback' key", 'good_feedback' in stats)
    test("stats has 'bad_feedback' key", 'bad_feedback' in stats)
    test("stats has 'neutral_feedback' key", 'neutral_feedback' in stats)

    test("good_feedback matches raw SQL",
         stats.get('good_feedback', -1) == expected_good,
         f"got={stats.get('good_feedback')}, expected={expected_good}")
    test("bad_feedback matches raw SQL",
         stats.get('bad_feedback', -1) == expected_bad,
         f"got={stats.get('bad_feedback')}, expected={expected_bad}")
    test("neutral_feedback matches raw SQL",
         stats.get('neutral_feedback', -1) == expected_neutral,
         f"got={stats.get('neutral_feedback')}, expected={expected_neutral}")

    # =========================================================================
    # TEST GROUP 2: DB layer — get_doctor_incidents() has classification
    # =========================================================================
    print("\n--- TEST GROUP 2: doctors_db.get_doctor_incidents() classification ---")

    incidents_result = get_doctor_incidents(
        doctor_id=doc_id,
        from_date='2020-01-01',
        to_date='2030-12-31',
        limit=100,
        offset=0
    )
    incidents = incidents_result.get('incidents', [])

    test("incidents returned", len(incidents) > 0, f"count={len(incidents)}")

    if incidents:
        first = incidents[0]
        test("incident has 'classification' field", 'classification' in first,
             f"keys={list(first.keys())}")
        test("incident has 'intent_type_ar' field", 'intent_type_ar' in first)
        test("incident has 'intent_type_en' field", 'intent_type_en' in first)
        test("incident has 'patient_name' field", 'patient_name' in first)

        valid_classes = {'good', 'bad', 'neutral'}
        test("classification is valid",
             first.get('classification') in valid_classes,
             f"got={first.get('classification')}")

        # Count classifications from incidents list
        good_from_list = sum(1 for i in incidents if i.get('classification') == 'good')
        bad_from_list = sum(1 for i in incidents if i.get('classification') == 'bad')
        neutral_from_list = sum(1 for i in incidents if i.get('classification') == 'neutral')

        test("classification counts from list match stats good",
             good_from_list == stats.get('good_feedback', -1),
             f"list={good_from_list}, stats={stats.get('good_feedback')}")
        test("classification counts from list match stats bad",
             bad_from_list == stats.get('bad_feedback', -1),
             f"list={bad_from_list}, stats={stats.get('bad_feedback')}")
    else:
        for _ in range(7):
            test("(skipped — no incidents)", True)

    # =========================================================================
    # TEST GROUP 3: Seasonal service — correct key mapping + intent
    # =========================================================================
    print("\n--- TEST GROUP 3: Doctor seasonal service ---")

    try:
        seasonal_data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=doc_id,
            season_start=date(2020, 1, 1),
            season_end=date(2030, 12, 31)
        )
        test("seasonal returns dict", isinstance(seasonal_data, dict))

        m = seasonal_data.get('metrics', {})
        # Fix verification: total_incidents should now be > 0 (was 0 before bugfix)
        test("total_incidents > 0 (bugfix verified)",
             m.get('total_incidents', 0) > 0,
             f"total_incidents={m.get('total_incidents')}")

        test("high_severity present and >= 0",
             m.get('high_severity', -1) >= 0,
             f"high_severity={m.get('high_severity')}")

        test("good_feedback_count in seasonal metrics",
             'good_feedback_count' in m,
             f"keys={list(m.keys())}")
        test("bad_feedback_count in seasonal metrics",
             'bad_feedback_count' in m)
        test("neutral_feedback_count in seasonal metrics",
             'neutral_feedback_count' in m)

        test("good_feedback_count matches stats",
             m.get('good_feedback_count') == stats.get('good_feedback'),
             f"seasonal={m.get('good_feedback_count')}, db={stats.get('good_feedback')}")

        test("incidents_details in payload",
             'incidents_details' in seasonal_data)

        # Check incidents have classification
        inc_details = seasonal_data.get('incidents_details', [])
        if inc_details:
            test("first incident has classification",
                 'classification' in inc_details[0],
                 f"keys={list(inc_details[0].keys())}")
        else:
            test("(no incidents_details to check)", True)

    except Exception as e:
        test("seasonal returns dict", False, f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        for _ in range(8):
            test("(skipped)", False)

    # =========================================================================
    # TEST GROUP 4: Doctor Word generation with intent classification
    # =========================================================================
    print("\n--- TEST GROUP 4: Doctor Word generation ---")

    try:
        doc_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(seasonal_data)
        test("doctor seasonal Word generates bytes",
             isinstance(doc_bytes, bytes) and len(doc_bytes) > 100,
             f"size={len(doc_bytes)}")

        doc_obj = Document(io.BytesIO(doc_bytes))

        # Find 13-row metrics table (with intent classification rows)
        metrics_table_found = False
        for t in doc_obj.tables:
            if len(t.columns) == 2 and len(t.rows) >= 13:
                first_cell = t.cell(0, 0).text
                if 'Total Incidents' in first_cell:
                    metrics_table_found = True
                    test("doctor metrics: Good Feedback row",
                         'Good Feedback' in t.cell(4, 0).text)
                    test("doctor metrics: Bad Feedback row",
                         'Bad Feedback' in t.cell(5, 0).text)
                    test("doctor metrics: Neutral Feedback row",
                         'Neutral Feedback' in t.cell(6, 0).text)
                    break
        test("doctor 13-row metrics table found", metrics_table_found)

        # Find 7-col incidents detail table
        incidents_table_found = False
        for t in doc_obj.tables:
            if len(t.columns) == 7 and len(t.rows) > 1:
                header = t.cell(0, 0).text
                if 'Date' in header:
                    incidents_table_found = True
                    test("doctor incidents table: Classification column",
                         'Classification' in t.cell(0, 4).text)
                    break
        test("doctor 7-col incidents table found", incidents_table_found)

    except Exception as e:
        test("doctor seasonal Word generates bytes", False, f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        for _ in range(5):
            test("(skipped)", False)

    # =========================================================================
    # TEST GROUP 5: person_report_word_adapter doctor report
    # =========================================================================
    print("\n--- TEST GROUP 5: person_report_word_adapter doctor ---")

    try:
        person_bytes = generate_person_seasonal_word_report(
            person_type="doctor",
            payload=seasonal_data
        )
        test("person_report doctor generates bytes",
             isinstance(person_bytes, bytes) and len(person_bytes) > 100,
             f"size={len(person_bytes)}")

        person_doc = Document(io.BytesIO(person_bytes))

        # Should have 7-row metrics table with intent
        found_intent_metrics = False
        for t in person_doc.tables:
            if len(t.columns) == 2 and len(t.rows) == 7:
                found_intent_metrics = True
                break
        test("person_report doctor has 7-row metrics", found_intent_metrics)

    except Exception as e:
        test("person_report doctor generates bytes", False, f"ERROR: {e}")
        test("(skipped)", False)

    # =========================================================================
    # TEST GROUP 6: Zero-incident doctor
    # =========================================================================
    print("\n--- TEST GROUP 6: Doctor zero-incident Word ---")

    zero_doctor_payload = {
        'doctor_identity': {
            'id': 99999,
            'name': 'Clean Doctor',
            'name_ar': 'طبيب نظيف',
            'specialty': 'Cardiology'
        },
        'period': {'start': date(2025, 1, 1), 'end': date(2025, 6, 30)},
        'metrics': {
            'total_incidents': 0,
            'high_severity': 0, 'medium_severity': 0, 'low_severity': 0,
            'good_feedback_count': 0, 'bad_feedback_count': 0, 'neutral_feedback_count': 0,
            'red_flags': 0
        },
        'incidents_summary': {'count': 0, 'top_categories': []},
        'incidents_details': [],
        'performance': {'score': 100, 'praise_level': 'excellent', 'risk_level': 'none', 'flags': []}
    }

    try:
        zero_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(zero_doctor_payload)
        test("zero-incident doctor Word generates bytes",
             isinstance(zero_bytes, bytes) and len(zero_bytes) > 100)

        zero_doc = Document(io.BytesIO(zero_bytes))
        all_text = '\n'.join([p.text for p in zero_doc.paragraphs])
        test("zero-incident: elegant message present",
             'No incidents recorded' in all_text or 'Clean record' in all_text)

        # No 7-col incidents table
        has_inc_table = any(
            len(t.columns) == 7 and len(t.rows) > 1 and 'Date' in t.cell(0, 0).text
            for t in zero_doc.tables
        )
        test("zero-incident: NO incidents table", not has_inc_table)

    except Exception as e:
        test("zero-incident doctor Word generates bytes", False, f"ERROR: {e}")
        for _ in range(2):
            test("(skipped)", False)

    # =========================================================================
    # FINAL RESULTS
    # =========================================================================
    print(f"\n{'='*70}")
    print(f"PHASE 6 RESULTS: {PASS} passed, {FAIL} failed, {PASS+FAIL} total")
    print(f"{'='*70}")

    if FAIL == 0:
        print("🎉 ALL TESTS PASSED! Phase 6 complete — Doctor good/bad classification done.")
    else:
        print(f"⚠ {FAIL} test(s) failed. Review output above.")


if __name__ == "__main__":
    run_all_tests()
