"""
Phase 4 Test: Word Template Upgrade

Tests:
1. person_report_word_adapter generates valid Word bytes with new fields
2. Metrics table has 7 rows (total + severity + intent classification)
3. Incidents table has 7 columns (including classification)
4. Zero-incident worker gets elegant message (not 404)
5. seasonal_word_adapter generate_worker_seasonal_word includes new metrics + incidents table
6. Word documents are valid (can be parsed back)
"""
import sys
import os
sys.path.insert(0, '.')
sys.path.insert(0, os.path.join('.', 'backend'))

from backend.api.services.person_report_word_adapter import generate_person_seasonal_word_report
from backend.api.services.seasonal_word_adapter import SeasonalWordAdapter
from docx import Document
import io
from datetime import date

PASS = 0
FAIL = 0

def test(name, condition, detail=""):
    global PASS, FAIL
    if condition:
        PASS += 1
        icon = "✅"
    else:
        FAIL += 1
        icon = "❌"
    print(f"  {icon} {name}" + (f" — {detail}" if detail else ""))


def build_worker_payload_with_incidents():
    """Build a realistic worker payload with incidents for testing."""
    return {
        'worker_identity': {
            'employee_id': 1,
            'full_name': 'Test Worker',
            'job_title': 'Quality Specialist',
            'department_id': 42,
            'section_id': 8,
            'administration_id': 3
        },
        'period': {
            'start': date(2025, 1, 1),
            'end': date(2025, 6, 30)
        },
        'metrics': {
            'total_incidents': 3,
            'high_severity': 1,
            'medium_severity': 1,
            'low_severity': 1,
            'good_feedback_count': 1,
            'bad_feedback_count': 1,
            'neutral_feedback_count': 1,
            'total_action_items': 5,
            'completed_action_items': 3,
            'overdue_action_items': 1,
            'explanation_accepted_count': 2,
            'explanation_rejected_count': 1
        },
        'incidents_details': [
            {
                'id': 490,
                'date': '2025-01-15',
                'patient_name': 'Ahmed Patient',
                'complaint_text': 'Test complaint 1',
                'category': 'Quality',
                'severity': 'High',
                'status': 'Closed',
                'intent_type_ar': 'نقد/اقتراح',
                'intent_type_en': 'Critique/Suggestion',
                'classification': 'bad',
                'is_primary': True
            },
            {
                'id': 491,
                'date': '2025-02-20',
                'patient_name': 'Sara Patient',
                'complaint_text': 'Test complaint 2',
                'category': 'Service',
                'severity': 'Medium',
                'status': 'Open',
                'intent_type_ar': 'تنويه',
                'intent_type_en': 'Notice',
                'classification': 'good',
                'is_primary': False
            },
            {
                'id': 492,
                'date': '2025-03-10',
                'patient_name': 'Mohammed Patient',
                'complaint_text': 'Test complaint 3',
                'category': 'Facilities',
                'severity': 'Low',
                'status': 'In Progress',
                'intent_type_ar': 'فرصة تحسين',
                'intent_type_en': 'Improvement Opportunity',
                'classification': 'neutral',
                'is_primary': False
            }
        ],
        'performance': {
            'score': 75,
            'praise_level': 'good',
            'risk_level': 'low',
            'flags': []
        }
    }


def build_worker_payload_zero_incidents():
    """Build a worker payload with zero incidents."""
    return {
        'worker_identity': {
            'employee_id': 999,
            'full_name': 'Clean Worker',
            'job_title': 'Perfect Employee',
            'department_id': 1,
            'section_id': 1,
            'administration_id': 1
        },
        'period': {
            'start': date(2025, 1, 1),
            'end': date(2025, 6, 30)
        },
        'metrics': {
            'total_incidents': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'good_feedback_count': 0,
            'bad_feedback_count': 0,
            'neutral_feedback_count': 0,
            'total_action_items': 0,
            'completed_action_items': 0,
            'overdue_action_items': 0,
            'explanation_accepted_count': 0,
            'explanation_rejected_count': 0
        },
        'incidents_details': [],
        'performance': {
            'score': 100,
            'praise_level': 'excellent',
            'risk_level': 'none',
            'flags': []
        }
    }


def run_all_tests():
    global PASS, FAIL

    print("=" * 70)
    print("PHASE 4 TEST SUITE: Word Template Upgrade")
    print("=" * 70)

    payload_with = build_worker_payload_with_incidents()
    payload_zero = build_worker_payload_zero_incidents()

    # =========================================================================
    # TEST GROUP 1: person_report_word_adapter — Worker with incidents
    # =========================================================================
    print("\n--- TEST GROUP 1: person_report_word_adapter with incidents ---")

    try:
        doc_bytes = generate_person_seasonal_word_report(
            person_type="worker",
            payload=payload_with
        )
        test("generates bytes successfully", isinstance(doc_bytes, bytes) and len(doc_bytes) > 0,
             f"size={len(doc_bytes)} bytes")

        # Parse back to verify structure
        doc = Document(io.BytesIO(doc_bytes))
        
        # Find all tables
        tables = doc.tables
        test("document has tables", len(tables) >= 2, f"found {len(tables)} tables")

        # Find the metrics table (should have 7 rows for: total, high, med, low, good, bad, neutral)
        metrics_found = False
        for t in tables:
            if len(t.rows) == 7 and len(t.columns) == 2:
                # Check if first cell contains total incidents label
                first_cell_text = t.cell(0, 0).text
                if 'إجمالي' in first_cell_text or 'الحالات' in first_cell_text:
                    metrics_found = True
                    # Check severity rows
                    test("metrics row 1: high severity label",
                         'عالية' in t.cell(1, 0).text)
                    test("metrics row 4: good feedback label",
                         'إيجابية' in t.cell(4, 0).text or 'تنويه' in t.cell(4, 0).text)
                    test("metrics row 5: bad feedback label",
                         'سلبية' in t.cell(5, 0).text or 'نقد' in t.cell(5, 0).text)
                    test("metrics row 6: neutral feedback label",
                         'محايد' in t.cell(6, 0).text)
                    break
        
        test("7-row metrics table found with intent classification", metrics_found)

        # Find the incidents detail table (should have 7 columns)
        incidents_table_found = False
        for t in tables:
            if len(t.columns) == 7 and len(t.rows) > 1:
                header_texts = [t.cell(0, i).text for i in range(7)]
                if 'التاريخ' in header_texts[0]:
                    incidents_table_found = True
                    test("incidents table has 7 columns",
                         len(t.columns) == 7, f"cols={len(t.columns)}")
                    test("incidents table has classification header",
                         'نوع الملاحظة' in header_texts[4] or 'الملاحظة' in header_texts[4],
                         f"header[4]='{header_texts[4]}'")
                    test("incidents table data rows match",
                         len(t.rows) - 1 == 3, f"data_rows={len(t.rows) - 1}")
                    
                    # Check classification values in data
                    class_col_values = [t.cell(r, 4).text for r in range(1, len(t.rows))]
                    test("classification values present in data",
                         all(v != '' for v in class_col_values),
                         f"values={class_col_values}")
                    break

        test("7-column incidents detail table found", incidents_table_found)

    except Exception as e:
        test("generates bytes successfully", False, f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        for _ in range(9):
            test("(skipped — generation error)", False)

    # =========================================================================
    # TEST GROUP 2: person_report_word_adapter — Zero incidents (elegant msg)
    # =========================================================================
    print("\n--- TEST GROUP 2: person_report_word_adapter zero incidents ---")

    try:
        doc_bytes_zero = generate_person_seasonal_word_report(
            person_type="worker",
            payload=payload_zero
        )
        test("zero-incident doc generates bytes", isinstance(doc_bytes_zero, bytes) and len(doc_bytes_zero) > 0)

        doc_zero = Document(io.BytesIO(doc_bytes_zero))

        # Check for elegant message instead of plain text
        all_text = '\n'.join([p.text for p in doc_zero.paragraphs])
        test("contains elegant zero-complaint message",
             'لا توجد حالات مسجلة' in all_text,
             f"found in paragraphs")
        test("contains praise message for clean record",
             'سجل نظيف' in all_text or 'أداء ممتاز' in all_text,
             f"found praise")

        # Should NOT have an incidents detail table
        has_incidents_table = False
        for t in doc_zero.tables:
            if len(t.columns) == 7 and len(t.rows) > 1:
                header_text = t.cell(0, 0).text
                if 'التاريخ' in header_text:
                    has_incidents_table = True
                    break
        test("zero incidents does NOT create incidents table", not has_incidents_table)

    except Exception as e:
        test("zero-incident doc generates bytes", False, f"ERROR: {e}")
        for _ in range(3):
            test("(skipped — generation error)", False)

    # =========================================================================
    # TEST GROUP 3: seasonal_word_adapter — Worker with incidents
    # =========================================================================
    print("\n--- TEST GROUP 3: seasonal_word_adapter worker with incidents ---")

    try:
        seasonal_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(payload_with)
        test("seasonal generates bytes", isinstance(seasonal_bytes, bytes) and len(seasonal_bytes) > 0,
             f"size={len(seasonal_bytes)} bytes")

        doc_seasonal = Document(io.BytesIO(seasonal_bytes))
        tables = doc_seasonal.tables

        # Find metrics table — should have 13 rows now
        seasonal_metrics_found = False
        for t in tables:
            if len(t.columns) == 2 and len(t.rows) >= 13:
                first_cell = t.cell(0, 0).text
                if 'Total Incidents' in first_cell:
                    seasonal_metrics_found = True
                    test("seasonal metrics: High Severity row",
                         'High Severity' in t.cell(1, 0).text)
                    test("seasonal metrics: Medium Severity row",
                         'Medium Severity' in t.cell(2, 0).text)
                    test("seasonal metrics: Low Severity row",
                         'Low Severity' in t.cell(3, 0).text)
                    test("seasonal metrics: Good Feedback row",
                         'Good Feedback' in t.cell(4, 0).text)
                    test("seasonal metrics: Bad Feedback row",
                         'Bad Feedback' in t.cell(5, 0).text)
                    test("seasonal metrics: Neutral Feedback row",
                         'Neutral Feedback' in t.cell(6, 0).text)
                    break

        test("seasonal 13-row metrics table found", seasonal_metrics_found)

        # Find incidents detail table (7 cols)
        seasonal_incidents_found = False
        for t in tables:
            if len(t.columns) == 7 and len(t.rows) > 1:
                header = t.cell(0, 0).text
                if 'Date' in header:
                    seasonal_incidents_found = True
                    test("seasonal incidents table has Classification column",
                         'Classification' in t.cell(0, 4).text)
                    test("seasonal incidents data rows count",
                         len(t.rows) - 1 == 3, f"data_rows={len(t.rows) - 1}")
                    break

        test("seasonal incidents detail table found", seasonal_incidents_found)

    except Exception as e:
        test("seasonal generates bytes", False, f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        for _ in range(9):
            test("(skipped — generation error)", False)

    # =========================================================================
    # TEST GROUP 4: seasonal_word_adapter — Worker zero incidents
    # =========================================================================
    print("\n--- TEST GROUP 4: seasonal_word_adapter zero incidents ---")

    try:
        seasonal_zero_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(payload_zero)
        test("seasonal zero generates bytes", isinstance(seasonal_zero_bytes, bytes) and len(seasonal_zero_bytes) > 0)

        doc_zero_seasonal = Document(io.BytesIO(seasonal_zero_bytes))
        all_text = '\n'.join([p.text for p in doc_zero_seasonal.paragraphs])
        test("seasonal zero: has elegant message",
             'No incidents recorded' in all_text or 'Clean record' in all_text,
             "found elegant message")
        test("seasonal zero: has praise",
             'Excellent' in all_text or 'Clean record' in all_text)

    except Exception as e:
        test("seasonal zero generates bytes", False, f"ERROR: {e}")
        for _ in range(2):
            test("(skipped — generation error)", False)

    # =========================================================================
    # TEST GROUP 5: Doctor report (person_report_word_adapter)
    # =========================================================================
    print("\n--- TEST GROUP 5: Doctor report through person_report_word_adapter ---")

    doctor_payload = {
        'doctor_identity': {
            'id': 100,
            'name': 'Dr. Test',
            'specialty': 'Cardiology'
        },
        'period': {
            'start': date(2025, 1, 1),
            'end': date(2025, 6, 30)
        },
        'metrics': {
            'total_incidents': 2,
            'high_severity': 1,
            'medium_severity': 0,
            'low_severity': 1,
            'good_feedback_count': 1,
            'bad_feedback_count': 1,
            'neutral_feedback_count': 0
        },
        'incidents_details': [
            {
                'id': 500,
                'date': '2025-03-01',
                'patient_name': 'Patient A',
                'category': 'Medical',
                'severity': 'High',
                'status': 'Closed',
                'classification': 'bad'
            },
            {
                'id': 501,
                'date': '2025-04-15',
                'patient_name': 'Patient B',
                'category': 'Communication',
                'severity': 'Low',
                'status': 'Open',
                'classification': 'good'
            }
        ]
    }

    try:
        doc_bytes_dr = generate_person_seasonal_word_report(
            person_type="doctor",
            payload=doctor_payload
        )
        test("doctor report generates bytes", isinstance(doc_bytes_dr, bytes) and len(doc_bytes_dr) > 0)

        doc_dr = Document(io.BytesIO(doc_bytes_dr))
        # Verify document has expected structure (3 tables: identity, metrics, incidents)
        test("doctor report has correct table count",
             len(doc_dr.tables) >= 3,
             f"tables={len(doc_dr.tables)}")

    except Exception as e:
        test("doctor report generates bytes", False, f"ERROR: {e}")
        test("(skipped)", False)

    # =========================================================================
    # FINAL RESULTS
    # =========================================================================
    print(f"\n{'='*70}")
    print(f"PHASE 4 RESULTS: {PASS} passed, {FAIL} failed, {PASS+FAIL} total")
    print(f"{'='*70}")

    if FAIL == 0:
        print("🎉 ALL TESTS PASSED! Phase 4 complete.")
    else:
        print(f"⚠ {FAIL} test(s) failed. Review output above.")


if __name__ == "__main__":
    run_all_tests()
