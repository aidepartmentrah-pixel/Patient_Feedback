"""
COMPREHENSIVE VERIFICATION - ALL THREE WORD REPORT TYPES
Tests and validates all implemented Word exports
"""

from docx import Document
import os
import glob

def verify_all_reports():
    """Run comprehensive verification of all three report types"""
    
    print("\n" + "="*70)
    print("COMPREHENSIVE WORD REPORTS VERIFICATION")
    print("="*70)
    print("\nThis verifies all three Word export implementations:")
    print("1. Patient History Export (JSON → Word)")
    print("2. Single Doctor Seasonal Report (with detailed incident table)")
    print("3. Aggregate Doctors Comparison Report (filtered & sorted)")
    print("="*70)
    
    results = {}
    
    # ========== 1. PATIENT WORD EXPORT ==========
    print("\n" + "="*70)
    print("1. PATIENT HISTORY WORD EXPORT")
    print("="*70)
    
    patient_files = glob.glob("test_patient_*_export_*.docx") + glob.glob("patient_*_report_*.docx")
    if patient_files:
        latest = max(patient_files, key=os.path.getmtime)
        doc = Document(latest)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"File: {latest}")
        print(f"Size: {os.path.getsize(latest):,} bytes")
        print(f"Tables: {len(doc.tables)}")
        
        checks = {
            'Arabic title': 'تقرير تاريخ المريض' in text,
            'Patient info table': len(doc.tables) >= 1,
            'Proper message for no complaints': 'لا توجد شكاوى' in text or len([t for t in doc.tables if len(t.columns) == 6]) > 0
        }
        
        for check, passed in checks.items():
            print(f"  {'✓' if passed else '✗'} {check}")
        
        results['patient'] = all(checks.values())
    else:
        print("✗ No patient Word files found")
        results['patient'] = False
    
    # ========== 2. SINGLE DOCTOR SEASONAL ==========
    print("\n" + "="*70)
    print("2. SINGLE DOCTOR SEASONAL REPORT")
    print("="*70)
    
    doctor_files = glob.glob("doctor_*_seasonal_*.docx")
    if doctor_files:
        latest = max(doctor_files, key=os.path.getmtime)
        doc = Document(latest)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"File: {latest}")
        print(f"Size: {os.path.getsize(latest):,} bytes")
        print(f"Tables: {len(doc.tables)}")
        
        # Count tables by columns
        tables_by_cols = {}
        for table in doc.tables:
            cols = len(table.columns) if table.rows else 0
            tables_by_cols[cols] = tables_by_cols.get(cols, 0) + 1
        
        checks = {
            'Arabic title': 'التقرير الموسمي' in text or 'طبيب' in text,
            'Person info table (2 cols)': 2 in tables_by_cols,
            'Summary metrics table (2 cols)': tables_by_cols.get(2, 0) >= 2,
            'Detailed incidents table (6 cols)': 6 in tables_by_cols,
            'Total tables (3 expected)': len(doc.tables) == 3
        }
        
        for check, passed in checks.items():
            print(f"  {'✓' if passed else '✗'} {check}")
        
        # Show table structure
        print("\n  Table breakdown:")
        for cols, count in sorted(tables_by_cols.items()):
            print(f"    {count}x tables with {cols} columns")
        
        results['single_doctor'] = all(checks.values())
    else:
        print("✗ No single doctor seasonal files found")
        results['single_doctor'] = False
    
    # ========== 3. AGGREGATE DOCTORS COMPARISON ==========
    print("\n" + "="*70)
    print("3. AGGREGATE DOCTORS COMPARISON REPORT")
    print("="*70)
    
    agg_files = glob.glob("aggregate_doctors_*.docx")
    if agg_files:
        latest = max(agg_files, key=os.path.getmtime)
        doc = Document(latest)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"File: {latest}")
        print(f"Size: {os.path.getsize(latest):,} bytes")
        print(f"Tables: {len(doc.tables)}")
        
        # Find comparison table (7 columns)
        comparison_table = None
        for table in doc.tables:
            if len(table.columns) == 7:
                comparison_table = table
                break
        
        # Check sorting if table exists
        sorted_correctly = False
        doctor_count = 0
        if comparison_table:
            doctor_count = len(comparison_table.rows) - 1  # Exclude header
            totals = []
            for ridx in range(1, len(comparison_table.rows)):
                try:
                    total = int(comparison_table.rows[ridx].cells[3].text.strip())
                    totals.append(total)
                except:
                    pass
            sorted_correctly = totals == sorted(totals, reverse=True)
        
        checks = {
            'Arabic comparison title': 'تقرير مقارنة الأطباء' in text,
            'Summary statistics section': 'إحصائيات' in text or 'إجمالي' in text,
            'Comparison table (7 cols)': comparison_table is not None,
            'Has doctors with incidents': doctor_count > 0,
            'Sorted by total cases desc': sorted_correctly
        }
        
        for check, passed in checks.items():
            print(f"  {'✓' if passed else '✗'} {check}")
        
        if comparison_table:
            print(f"\n  Doctors with incidents: {doctor_count}")
            print(f"  Comparison table: {len(comparison_table.rows)} rows × {len(comparison_table.columns)} cols")
        
        results['aggregate_doctors'] = all(checks.values())
    else:
        print("✗ No aggregate doctors files found")
        results['aggregate_doctors'] = False
    
    # ========== FINAL SUMMARY ==========
    print("\n" + "="*70)
    print("FINAL RESULTS")
    print("="*70)
    print(f"\n1. Patient Word Export:           {'✅ PASS' if results.get('patient') else '❌ FAIL'}")
    print(f"2. Single Doctor Seasonal:        {'✅ PASS' if results.get('single_doctor') else '❌ FAIL'}")
    print(f"3. Aggregate Doctors Comparison:  {'✅ PASS' if results.get('aggregate_doctors') else '❌ FAIL'}")
    
    all_passed = all(results.values())
    print("\n" + "="*70)
    if all_passed:
        print("🎉 ALL THREE WORD REPORT TYPES VERIFIED SUCCESSFULLY!")
        print("="*70)
        print("\nImplementation Summary:")
        print("✅ Patient export: Changed from JSON to Word with Arabic template")
        print("✅ Doctor seasonal: Added detailed 6-column incident table")
        print("✅ Aggregate doctors: Filtered to doctors with cases, sorted by total")
        print("\n✨ NO frontend changes required - same endpoints & parameters")
    else:
        print("❌ SOME REPORTS FAILED VERIFICATION")
        print("="*70)
    
    return all_passed

if __name__ == "__main__":
    verify_all_reports()
