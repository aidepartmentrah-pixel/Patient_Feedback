"""
Debug patient Word file - show all content
"""

from docx import Document
import sys

filename = "test_patient_100022_export_20260212_092917.docx"

doc = Document(filename)

print(f"\n{'='*70}")
print(f"PATIENT WORD FILE CONTENT")
print(f"{'='*70}")
print(f"\nFile: {filename}")
print(f"Tables: {len(doc.tables)}")
print(f"Paragraphs: {len(doc.paragraphs)}")

print(f"\n{'='*70}")
print("ALL PARAGRAPHS:")
print(f"{'='*70}")
for idx, para in enumerate(doc.paragraphs, 1):
    text = para.text.strip()
    if text:
        print(f"{idx:2}. {text}")

print(f"\n{'='*70}")
print("ALL TABLES:")
print(f"{'='*70}")
for tidx, table in enumerate(doc.tables, 1):
    rows = len(table.rows)
    cols = len(table.columns) if table.rows else 0
    print(f"\nTable {tidx}: {rows} rows × {cols} cols")
    
    # Show all content
    for ridx, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        print(f"  Row {ridx}: {row_data}")
