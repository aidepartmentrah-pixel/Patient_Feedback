"""
Endpoint Test: All Reporting Endpoints

Tests all report endpoints via HTTP to verify they work end-to-end through the API.
Covers:
1. Worker profile endpoint (GET /api/workers/{id}/profile)
2. Worker seasonal Word export (GET /api/person-reports/worker/{id}/seasonal-word)
3. Doctor statistics endpoint (GET /api/doctors/{id}/statistics)
4. Doctor incidents endpoint (GET /api/doctors/{id}/incidents)
5. Doctor seasonal Word export (GET /api/person-reports/doctor/{id}/seasonal-word)
6. Doctor seasonal via seasonal_export_router (GET /api/doctors/{id}/seasonal-report)
7. Worker seasonal via seasonal_export_router (GET /api/workers/{id}/seasonal-report)
8. Validates new fields (severity, intent classification) are in responses
"""
import sys
import os
sys.path.insert(0, '.')
sys.path.insert(0, os.path.join('.', 'backend'))

import requests
import json
from datetime import date
from docx import Document
import io

BASE_URL = "http://localhost:8000"

PASS = 0
FAIL = 0
SKIP = 0

def test(name, condition, detail=""):
    global PASS, FAIL
    if condition:
        PASS += 1
        icon = "✅"
    else:
        FAIL += 1
        icon = "❌"
    print(f"  {icon} {name}" + (f" — {detail}" if detail else ""))

def skip(name, reason=""):
    global SKIP
    SKIP += 1
    print(f"  ⏭ {name}" + (f" — {reason}" if reason else ""))


def authenticate(session):
    """Try to authenticate with known test credentials."""
    credentials = [
        ("software_admin", "admin123"),
        ("section_admin", "section123"),
        ("department_admin", "dept123"),
        ("administration_admin", "adminis123"),
    ]
    for username, password in credentials:
        try:
            resp = session.post(
                f"{BASE_URL}/api/auth/login",
                json={"username": username, "password": password},
                timeout=5
            )
            if resp.status_code == 200:
                print(f"   Authenticated as: {username}")
                return True
        except Exception:
            pass
    return False


def find_test_data():
    """Find valid employee/doctor IDs from the database."""
    from backend.core.database import get_connection
    conn = get_connection()
    cursor = conn.cursor()
    
    # Employee with incidents
    cursor.execute("""
        SELECT TOP 1 ice.EmployeeID
        FROM dbo.APP_IncidentCaseEmployee ice
        ORDER BY ice.EmployeeID
    """)
    row = cursor.fetchone()
    emp_with = row.EmployeeID if row else None
    
    # Employee without incidents
    cursor.execute("""
        SELECT TOP 1 EmployeeID FROM dbo.APP_VIEWTABLE_HR_EMPLOYEES
        WHERE EmployeeID NOT IN (SELECT EmployeeID FROM dbo.APP_IncidentCaseEmployee)
        ORDER BY EmployeeID
    """)
    row = cursor.fetchone()
    emp_zero = row.EmployeeID if row else None
    
    # Doctor with incidents
    cursor.execute("""
        SELECT TOP 1 icd.DoctorID
        FROM dbo.APP_IncidentCaseDoctor icd
        ORDER BY icd.DoctorID
    """)
    row = cursor.fetchone()
    doc_id = row.DoctorID if row else None
    
    conn.close()
    return emp_with, emp_zero, doc_id


def run_all_tests():
    global PASS, FAIL, SKIP
    
    print("=" * 70)
    print("ENDPOINT TEST SUITE: All Reporting Endpoints")
    print("=" * 70)
    
    # Check if server is running
    print("\n--- Pre-check: Server connectivity ---")
    try:
        resp = requests.get(f"{BASE_URL}/docs", timeout=5)
        test("Server is reachable", resp.status_code == 200, f"status={resp.status_code}")
    except requests.ConnectionError:
        print("  ❌ Server is NOT running at localhost:8000")
        print("     Start the server first: cd backend && uvicorn main:app --reload --port 8000")
        print(f"\n{'='*70}")
        print("CANNOT PROCEED — server not running")
        print(f"{'='*70}")
        return
    
    # Get test data
    emp_with, emp_zero, doc_id = find_test_data()
    print(f"\n--- Test Data ---")
    print(f"   Employee with incidents: {emp_with}")
    print(f"   Employee zero incidents: {emp_zero}")
    print(f"   Doctor with incidents: {doc_id}")
    
    # Authenticate
    print(f"\n--- Authentication ---")
    session = requests.Session()
    authenticated = authenticate(session)
    if not authenticated:
        print("  ⚠ Could not authenticate — will try endpoints without auth")
        print("    (endpoints requiring auth will fail with 401/403)")
    
    # =========================================================================
    # TEST GROUP 1: Worker Profile Endpoint
    # =========================================================================
    print(f"\n--- TEST GROUP 1: GET /api/workers/{{id}}/profile ---")
    
    if emp_with:
        resp = session.get(f"{BASE_URL}/api/workers/{emp_with}/profile", timeout=10)
        test("worker profile returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            data = resp.json()
            
            # Identity block
            test("response has 'worker' block", 'worker' in data)
            test("worker has employee_id", data.get('worker', {}).get('employee_id') == emp_with)
            test("worker has full_name", bool(data.get('worker', {}).get('full_name')))
            
            # Metrics block — old fields
            metrics = data.get('metrics', {})
            test("metrics has total_incidents", 'total_incidents' in metrics)
            test("metrics has total_action_items", 'total_action_items' in metrics)
            
            # Metrics block — NEW severity fields
            test("metrics has high_severity", 'high_severity' in metrics,
                 f"value={metrics.get('high_severity')}")
            test("metrics has medium_severity", 'medium_severity' in metrics,
                 f"value={metrics.get('medium_severity')}")
            test("metrics has low_severity", 'low_severity' in metrics,
                 f"value={metrics.get('low_severity')}")
            
            # Metrics block — NEW intent classification fields
            test("metrics has good_feedback_count", 'good_feedback_count' in metrics,
                 f"value={metrics.get('good_feedback_count')}")
            test("metrics has bad_feedback_count", 'bad_feedback_count' in metrics,
                 f"value={metrics.get('bad_feedback_count')}")
            test("metrics has neutral_feedback_count", 'neutral_feedback_count' in metrics,
                 f"value={metrics.get('neutral_feedback_count')}")
            
            # NEW incidents list
            test("response has 'incidents' list", 'incidents' in data,
                 f"count={len(data.get('incidents', []))}")
            
            incidents = data.get('incidents', [])
            if incidents:
                first = incidents[0]
                test("incident has classification", 'classification' in first,
                     f"value={first.get('classification')}")
                test("incident has intent_type_ar", 'intent_type_ar' in first)
                test("incident has severity", 'severity' in first)
            
            # Severity sum check
            sev_sum = metrics.get('high_severity', 0) + metrics.get('medium_severity', 0) + metrics.get('low_severity', 0)
            test("severity sum == total_incidents",
                 sev_sum == metrics.get('total_incidents', -1),
                 f"sum={sev_sum}, total={metrics.get('total_incidents')}")
            
            # Intent sum check
            intent_sum = metrics.get('good_feedback_count', 0) + metrics.get('bad_feedback_count', 0) + metrics.get('neutral_feedback_count', 0)
            test("intent sum == total_incidents",
                 intent_sum == metrics.get('total_incidents', -1),
                 f"sum={intent_sum}, total={metrics.get('total_incidents')}")
        else:
            detail = ""
            try:
                detail = resp.json().get('detail', resp.text[:100])
            except:
                detail = resp.text[:100]
            print(f"    Response: {detail}")
    else:
        skip("worker profile", "no employee with incidents found")
    
    # =========================================================================
    # TEST GROUP 2: Worker Profile with Date Range
    # =========================================================================
    print(f"\n--- TEST GROUP 2: Worker profile with date range ---")
    
    if emp_with:
        resp = session.get(
            f"{BASE_URL}/api/workers/{emp_with}/profile",
            params={"date_from": "2020-01-01", "date_to": "2030-12-31"},
            timeout=10
        )
        test("worker profile with dates returns 200",
             resp.status_code == 200, f"status={resp.status_code}")
        
        if resp.status_code == 200:
            data = resp.json()
            test("period_from is set", data.get('period_from') == "2020-01-01")
            test("period_to is set", data.get('period_to') == "2030-12-31")
    
    # =========================================================================
    # TEST GROUP 3: Worker Profile — Zero Incidents
    # =========================================================================
    print(f"\n--- TEST GROUP 3: Worker profile zero incidents ---")
    
    if emp_zero:
        resp = session.get(f"{BASE_URL}/api/workers/{emp_zero}/profile", timeout=10)
        test("zero worker profile returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            data = resp.json()
            test("zero total_incidents", data.get('metrics', {}).get('total_incidents') == 0)
            test("zero incidents list empty", len(data.get('incidents', [])) == 0)
    else:
        skip("zero worker profile", "no zero-incident employee found")
    
    # =========================================================================
    # TEST GROUP 4: Worker Seasonal Word Export (person-reports)
    # =========================================================================
    print(f"\n--- TEST GROUP 4: GET /api/person-reports/worker/{{id}}/seasonal-word ---")
    
    if emp_with:
        resp = session.get(
            f"{BASE_URL}/api/person-reports/worker/{emp_with}/seasonal-word",
            params={"season_start": "2020-01-01", "season_end": "2030-12-31"},
            timeout=15
        )
        test("worker seasonal Word returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            test("content-type is docx",
                 'wordprocessingml' in resp.headers.get('content-type', ''))
            test("content-disposition has filename",
                 'attachment' in resp.headers.get('content-disposition', ''))
            test("body size > 1KB", len(resp.content) > 1000,
                 f"size={len(resp.content)}")
            
            # Parse the Word doc
            try:
                doc = Document(io.BytesIO(resp.content))
                test("Word doc is parseable", len(doc.paragraphs) > 0,
                     f"paragraphs={len(doc.paragraphs)}")
                
                # Check for 7-col incidents table
                has_7col = any(len(t.columns) == 7 for t in doc.tables if len(t.rows) > 1)
                test("Word doc has 7-col incidents table (with classification)", has_7col)
            except Exception as e:
                test("Word doc is parseable", False, f"ERROR: {e}")
        else:
            detail = ""
            try:
                detail = resp.json().get('detail', resp.text[:200])
            except:
                detail = resp.text[:200]
            print(f"    Response: {detail}")
    
    # =========================================================================
    # TEST GROUP 5: Doctor Statistics Endpoint
    # =========================================================================
    print(f"\n--- TEST GROUP 5: GET /api/doctors/{{id}}/statistics ---")
    
    if doc_id:
        resp = session.get(
            f"{BASE_URL}/api/doctors/{doc_id}/statistics",
            params={"from_date": "2020-01-01", "to_date": "2030-12-31"},
            timeout=10
        )
        test("doctor statistics returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            data = resp.json()
            stats = data.get('statistics', data)
            test("stats has total", 'total' in stats or 'totalIncidents' in stats,
                 f"keys={list(stats.keys())}")
    else:
        skip("doctor statistics", "no doctor with incidents found")
    
    # =========================================================================
    # TEST GROUP 6: Doctor Incidents Endpoint (check classification fields)
    # =========================================================================
    print(f"\n--- TEST GROUP 6: GET /api/doctors/{{id}}/incidents ---")
    
    if doc_id:
        resp = session.get(
            f"{BASE_URL}/api/doctors/{doc_id}/incidents",
            params={"from_date": "2020-01-01", "to_date": "2030-12-31"},
            timeout=10
        )
        test("doctor incidents returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            data = resp.json()
            incidents = data.get('incidents', [])
            test("incidents returned", len(incidents) > 0, f"count={len(incidents)}")
            
            if incidents:
                first = incidents[0]
                test("incident has 'classification'", 'classification' in first,
                     f"value={first.get('classification')}")
                test("incident has 'intent_type_ar'", 'intent_type_ar' in first)
                test("incident has 'patient_name'", 'patient_name' in first)
    else:
        skip("doctor incidents", "no doctor found")
    
    # =========================================================================
    # TEST GROUP 7: Doctor Seasonal Word Export (person-reports)
    # =========================================================================
    print(f"\n--- TEST GROUP 7: GET /api/person-reports/doctor/{{id}}/seasonal-word ---")
    
    if doc_id:
        resp = session.get(
            f"{BASE_URL}/api/person-reports/doctor/{doc_id}/seasonal-word",
            params={"season_start": "2020-01-01", "season_end": "2030-12-31"},
            timeout=15
        )
        test("doctor seasonal Word returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            test("content-type is docx",
                 'wordprocessingml' in resp.headers.get('content-type', ''))
            test("body size > 1KB", len(resp.content) > 1000,
                 f"size={len(resp.content)}")
            
            try:
                doc = Document(io.BytesIO(resp.content))
                test("doctor Word doc is parseable", len(doc.paragraphs) > 0)
            except Exception as e:
                test("doctor Word doc is parseable", False, f"ERROR: {e}")
        else:
            detail = ""
            try:
                detail = resp.json().get('detail', resp.text[:200])
            except:
                detail = resp.text[:200]
            print(f"    Response: {detail}")
    
    # =========================================================================
    # TEST GROUP 8: Seasonal Export Router — Doctor (GET /api/doctors/{id}/seasonal-report)
    # =========================================================================
    print(f"\n--- TEST GROUP 8: GET /api/doctors/{{id}}/seasonal-report ---")
    
    if doc_id:
        resp = session.get(
            f"{BASE_URL}/api/doctors/{doc_id}/seasonal-report",
            params={"season_start": "2020-01-01", "season_end": "2030-12-31"},
            timeout=15
        )
        test("doctor seasonal-report returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            test("content-type is docx",
                 'wordprocessingml' in resp.headers.get('content-type', ''))
            test("body size > 1KB", len(resp.content) > 1000)
        else:
            detail = ""
            try:
                detail = resp.json().get('detail', resp.text[:200])
            except:
                detail = resp.text[:200]
            print(f"    Response: {detail}")
    
    # =========================================================================
    # TEST GROUP 9: Seasonal Export Router — Worker (GET /api/workers/{id}/seasonal-report)
    # =========================================================================
    print(f"\n--- TEST GROUP 9: GET /api/workers/{{id}}/seasonal-report ---")
    
    if emp_with:
        resp = session.get(
            f"{BASE_URL}/api/workers/{emp_with}/seasonal-report",
            params={"season_start": "2020-01-01", "season_end": "2030-12-31"},
            timeout=15
        )
        test("worker seasonal-report returns 200", resp.status_code == 200,
             f"status={resp.status_code}")
        
        if resp.status_code == 200:
            test("content-type is docx",
                 'wordprocessingml' in resp.headers.get('content-type', ''))
            test("body size > 1KB", len(resp.content) > 1000)
        else:
            detail = ""
            try:
                detail = resp.json().get('detail', resp.text[:200])
            except:
                detail = resp.text[:200]
            print(f"    Response: {detail}")
    
    # =========================================================================
    # TEST GROUP 10: Non-existent Employee — should return 404
    # =========================================================================
    print(f"\n--- TEST GROUP 10: Edge case — non-existent employee ---")
    
    resp = session.get(f"{BASE_URL}/api/workers/999999/profile", timeout=10)
    test("non-existent worker returns 404", resp.status_code == 404,
         f"status={resp.status_code}")
    
    # =========================================================================
    # FINAL RESULTS
    # =========================================================================
    print(f"\n{'='*70}")
    print(f"ENDPOINT TEST RESULTS: {PASS} passed, {FAIL} failed, {SKIP} skipped, {PASS+FAIL+SKIP} total")
    print(f"{'='*70}")
    
    if FAIL == 0:
        print("🎉 ALL ENDPOINT TESTS PASSED!")
    else:
        print(f"⚠ {FAIL} test(s) failed. Review output above.")


if __name__ == "__main__":
    run_all_tests()
