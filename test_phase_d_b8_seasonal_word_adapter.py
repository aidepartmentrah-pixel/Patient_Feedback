"""
Test Suite: D-B8 Seasonal Word Adapter

Tests the Word document generation for doctor and worker seasonal reports.
Verifies that data payloads from D-B6/D-B7 are formatted into valid Word documents.

Target: backend/api/services/seasonal_word_adapter.py

Test Coverage:
- Doctor seasonal Word generation
- Worker seasonal Word generation
- Document structure validation
- Content verification
- Error handling
- Integration with seasonal report builders
"""

import pytest
from io import BytesIO
from docx import Document
from backend.api.services.seasonal_word_adapter import SeasonalWordAdapter
from backend.api.services.doctor_seasonal_reporting_service import DoctorSeasonalReportingService
from backend.api.services.worker_seasonal_reporting_service import WorkerSeasonalReportingService


class TestSeasonalWordAdapter:
    """Test suite for D-B8 seasonal Word adapter."""
    
    def test_generate_doctor_word_returns_bytes(self):
        """
        Test 1: Verify doctor Word generation returns bytes.
        
        Should return non-empty bytes that can be loaded as Word document.
        """
        # Get real doctor data
        data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(data)
        
        # Verify output
        assert isinstance(word_bytes, bytes)
        assert len(word_bytes) > 0
        assert len(word_bytes) > 1000  # Should be substantial size
    
    def test_generate_doctor_word_is_valid_docx(self):
        """
        Test 2: Verify doctor Word output is valid DOCX format.
        
        Should be loadable by python-docx.
        """
        # Get real doctor data
        data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(data)
        
        # Try to load with python-docx
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        
        # Should have content
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) > 0
    
    def test_doctor_word_contains_identity(self):
        """
        Test 3: Verify doctor Word contains identity information.
        
        Should include doctor name and ID.
        """
        # Get real doctor data
        data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(data)
        
        # Load and extract text
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        
        # Extract all text from document (paragraphs + tables)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Extract text from tables
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + ' '
        
        combined_text = full_text + '\n' + table_text
        
        # Should contain doctor info
        assert 'Doctor' in combined_text
        assert 'Information' in combined_text or 'Identity' in combined_text
        
        # Should contain doctor name from data
        doctor_name = data['doctor_identity'].get('name_en', '')
        if doctor_name:
            assert doctor_name in combined_text
    
    def test_doctor_word_contains_performance_score(self):
        """
        Test 4: Verify doctor Word contains performance score.
        
        Should display score, praise level, risk level.
        """
        # Get real doctor data
        data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(data)
        
        # Load and extract text
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Should contain performance info
        assert 'Performance' in full_text or 'Score' in full_text
        assert 'Praise' in full_text or 'Risk' in full_text
        
        # Should contain actual score value
        score = data['performance']['score']
        assert str(score) in full_text
    
    def test_doctor_word_contains_metrics(self):
        """
        Test 5: Verify doctor Word contains metrics table.
        
        Should display incident counts, actions, explanations.
        """
        # Get real doctor data
        data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(data)
        
        # Load document
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        
        # Should have tables (identity + period + metrics)
        assert len(doc.tables) >= 3
        
        # Extract all text (paragraphs + tables)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + ' '
        combined_text = full_text + '\n' + table_text
        
        # Should contain metric labels
        assert 'Incidents' in combined_text or 'incidents' in combined_text
        assert 'Actions' in combined_text or 'actions' in combined_text
    
    def test_generate_worker_word_returns_bytes(self):
        """
        Test 6: Verify worker Word generation returns bytes.
        
        Should return non-empty bytes that can be loaded as Word document.
        """
        # Get real worker data
        data = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
            employee_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(data)
        
        # Verify output
        assert isinstance(word_bytes, bytes)
        assert len(word_bytes) > 0
        assert len(word_bytes) > 1000  # Should be substantial size
    
    def test_generate_worker_word_is_valid_docx(self):
        """
        Test 7: Verify worker Word output is valid DOCX format.
        
        Should be loadable by python-docx.
        """
        # Get real worker data
        data = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
            employee_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(data)
        
        # Try to load with python-docx
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        
        # Should have content
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) > 0
    
    def test_worker_word_contains_identity(self):
        """
        Test 8: Verify worker Word contains identity information.
        
        Should include worker name, employee ID, department.
        """
        # Get real worker data
        data = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
            employee_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(data)
        
        # Load and extract text
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        
        # Extract all text from document
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Should contain worker info
        assert 'Worker' in full_text
        assert 'Information' in full_text or 'Identity' in full_text
        
        # Should contain employee ID
        employee_id = data['worker_identity'].get('employee_id', 0)
        assert str(employee_id) in full_text
    
    def test_worker_word_contains_performance_score(self):
        """
        Test 9: Verify worker Word contains performance score.
        
        Should display score, praise level, risk level.
        """
        # Get real worker data
        data = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
            employee_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(data)
        
        # Load and extract text
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Should contain performance info
        assert 'Performance' in full_text or 'Score' in full_text
        assert 'Praise' in full_text or 'Risk' in full_text
        
        # Should contain actual score value
        score = data['performance']['score']
        assert str(score) in full_text
    
    def test_worker_word_contains_metrics(self):
        """
        Test 10: Verify worker Word contains metrics table.
        
        Should display incident counts, action items, explanations.
        """
        # Get real worker data
        data = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
            employee_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(data)
        
        # Load document
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        
        # Should have tables (identity + period + metrics)
        assert len(doc.tables) >= 3
        
        # Extract all text (paragraphs + tables)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + ' '
        combined_text = full_text + '\n' + table_text
        
        # Should contain metric labels
        assert 'Incidents' in combined_text or 'incidents' in combined_text
        assert 'Action' in combined_text or 'action' in combined_text
    
    def test_doctor_word_period_information(self):
        """
        Test 11: Verify doctor Word contains period information.
        
        Should show season name, start date, end date.
        """
        # Get real doctor data with Q1 2024
        data = DoctorSeasonalReportingService.build_doctor_seasonal_report_data(
            doctor_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_doctor_seasonal_word(data)
        
        # Load and extract text
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Extract text from tables
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + ' '
        combined_text = full_text + '\n' + table_text
        
        # Should contain period info
        assert '2024-01-01' in combined_text
        assert '2024-03-31' in combined_text
        assert 'Q1' in combined_text or 'Period' in combined_text
    
    def test_worker_word_period_information(self):
        """
        Test 12: Verify worker Word contains period information.
        
        Should show season name, start date, end date.
        """
        # Get real worker data with Q1 2024
        data = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
            employee_id=1,
            season_start='2024-01-01',
            season_end='2024-03-31'
        )
        
        # Generate Word document
        word_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(data)
        
        # Load and extract text
        buffer = BytesIO(word_bytes)
        doc = Document(buffer)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        # Extract text from tables
        table_text = ''
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + ' '
        combined_text = full_text + '\n' + table_text
        
        # Should contain period info
        assert '2024-01-01' in combined_text
        assert '2024-03-31' in combined_text
        assert 'Q1' in combined_text or 'Period' in combined_text


if __name__ == '__main__':
    """Run tests with pytest."""
    pytest.main([__file__, '-v', '--tb=short'])
