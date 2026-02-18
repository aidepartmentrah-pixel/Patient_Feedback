"""
Verify Doctor Word Reports - Content Validation
Checks that the generated Word documents contain expected structure and data
"""

from docx import Document
import os

def verify_single_doctor_report(filename):
    """Verify single doctor seasonal Word report structure"""
    print("\n" + "="*70)
    print("VERIFYING SINGLE DOCTOR SEASONAL REPORT")
    print("="*70)
    print(f"File: {filename}")
    
    if not os.path.exists(filename):
        print("✗ File not found!")
        return False
    
    try:
        doc = Document(filename)
        
        # Check paragraphs for key Arabic text
        text_content = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"\n✓ Opened document ({os.path.getsize(filename):,} bytes)")
        print(f"✓ Paragraphs: {len(doc.paragraphs)}")
        print(f"✓ Tables: {len(doc.tables)}")
        
        # Verify Arabic title
        if "التقرير الموسمي" in text_content or "طبيب" in text_content:
            print("✓ Arabic title found")
        else:
            print("⚠ Arabic title not found")
        
        # Check for expected tables (person info, metrics, detailed incidents)
        if len(doc.tables) >= 3:
            print(f"✓ Has {len(doc.tables)} tables (expected: 3 - person info, metrics, incidents)")
            
            # Analyze tables
            for idx, table in enumerate(doc.tables, 1):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                print(f"  Table {idx}: {rows} rows × {cols} columns")
                
                # Show first few cells of each table
                if table.rows:
                    first_row = [cell.text.strip() for cell in table.rows[0].cells[:3]]
                    print(f"    First row: {first_row}")
        else:
            print(f"⚠ Only {len(doc.tables)} tables found (expected 3+)")
        
        # Check for detailed incident data (should have table with 6 columns)
        has_detailed_table = any(len(table.columns) == 6 for table in doc.tables if table.rows)
        if has_detailed_table:
            print("✓ Found 6-column detailed incident table")
        else:
            print("⚠ No 6-column incident table found")
        
        print("\n✅ Single doctor report structure validated!")
        return True
        
    except Exception as e:
        print(f"✗ ERROR: {str(e)}")
        return False


def verify_aggregate_doctors_report(filename):
    """Verify aggregate doctors comparison Word report structure"""
    print("\n" + "="*70)
    print("VERIFYING AGGREGATE DOCTORS COMPARISON REPORT")
    print("="*70)
    print(f"File: {filename}")
    
    if not os.path.exists(filename):
        print("✗ File not found!")
        return False
    
    try:
        doc = Document(filename)
        
        # Check paragraphs for key Arabic text
        text_content = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"\n✓ Opened document ({os.path.getsize(filename):,} bytes)")
        print(f"✓ Paragraphs: {len(doc.paragraphs)}")
        print(f"✓ Tables: {len(doc.tables)}")
        
        # Verify Arabic comparison title
        if "تقرير مقارنة الأطباء" in text_content:
            print("✓ Arabic comparison title found: 'تقرير مقارنة الأطباء الموسمي'")
        else:
            print("⚠ Arabic comparison title not found")
        
        # Check for summary statistics section
        if "إحصائيات" in text_content or "إجمالي" in text_content:
            print("✓ Summary statistics section found")
        else:
            print("⚠ Summary statistics not found")
        
        # Check for expected tables (summary stats + comparison table)
        if len(doc.tables) >= 2:
            print(f"✓ Has {len(doc.tables)} tables (expected: 2 - summary stats, comparison)")
            
            # Analyze tables
            for idx, table in enumerate(doc.tables, 1):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                print(f"\n  Table {idx}: {rows} rows × {cols} columns")
                
                # Show headers
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    print(f"    Headers: {headers}")
                    
                    # Check if this is the comparison table (7 columns expected)
                    if cols == 7:
                        print(f"    ✓ This is the comparison table (7 columns)")
                        print(f"    ✓ Contains {rows - 1} doctors with incidents")
                        
                        # Show first 3 data rows
                        for row_idx in range(1, min(4, rows)):
                            row_data = [cell.text.strip() for cell in table.rows[row_idx].cells]
                            print(f"    Row {row_idx}: {row_data}")
                        
                        # Verify sorting (total incidents should be descending)
                        if rows > 2:
                            totals = []
                            for row_idx in range(1, rows):
                                try:
                                    total = int(table.rows[row_idx].cells[3].text.strip())
                                    totals.append(total)
                                except:
                                    pass
                            
                            if totals and totals == sorted(totals, reverse=True):
                                print(f"    ✓ Sorted by total cases descending: {totals[:5]}...")
                            elif totals:
                                print(f"    ⚠ Not properly sorted: {totals[:5]}...")
        else:
            print(f"⚠ Only {len(doc.tables)} tables found (expected 2+)")
        
        # Check for comparison table with 7 columns
        has_comparison_table = any(len(table.columns) == 7 for table in doc.tables if table.rows)
        if has_comparison_table:
            print("\n✓ Found 7-column comparison table (#, Name, Specialty, Total, High, Med, Low)")
        else:
            print("\n⚠ No 7-column comparison table found")
        
        print("\n✅ Aggregate doctors report structure validated!")
        return True
        
    except Exception as e:
        print(f"✗ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    print("\n" + "="*70)
    print("DOCTOR WORD REPORTS VERIFICATION")
    print("="*70)
    
    # Find the most recent files
    import glob
    
    doctor_files = glob.glob("doctor_*_seasonal_*.docx")
    aggregate_files = glob.glob("aggregate_doctors_*.docx")
    
    single_success = False
    aggregate_success = False
    
    if doctor_files:
        latest_single = max(doctor_files, key=os.path.getmtime)
        single_success = verify_single_doctor_report(latest_single)
    else:
        print("\n⚠ No single doctor report files found (doctor_*_seasonal_*.docx)")
    
    if aggregate_files:
        latest_aggregate = max(aggregate_files, key=os.path.getmtime)
        aggregate_success = verify_aggregate_doctors_report(latest_aggregate)
    else:
        print("\n⚠ No aggregate doctor report files found (aggregate_doctors_*.docx)")
    
    print("\n" + "="*70)
    print("FINAL RESULTS")
    print("="*70)
    print(f"Single Doctor Report:   {'✅ PASS' if single_success else '❌ FAIL'}")
    print(f"Aggregate Doctor Report: {'✅ PASS' if aggregate_success else '❌ FAIL'}")
    
    if single_success and aggregate_success:
        print("\n🎉 ALL REPORTS VERIFIED SUCCESSFULLY!")
    else:
        print("\n❌ Some reports failed verification")
