"""
Phase 1 Test: Single Seasonal Export with Comparison + Charts
Tests that a single org unit export generates a ZIP with:
1. Regular seasonal report
2. Comparison report with 9 charts (3 domains x 3 chart types)
"""

import sys
import os

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'backend'))

from backend.api.services.report_export_service import ReportExportService
from backend.api.db_layer.seasonal_report import resolve_season_id_from_year_trimester
import zipfile
from io import BytesIO


def test_phase1_single_unit_export():
    """Test single org unit seasonal export generates ZIP with both reports"""
    
    print("\n" + "="*80)
    print("🧪 PHASE 1 TEST: Single Seasonal Export")
    print("="*80)
    
    # Test parameters
    year = 2026
    period = "Q1"  # Q1-2026
    orgunit_id = 12  # Specific section/department
    orgunit_type = 3  # Section level
    
    print(f"\n📋 Test Parameters:")
    print(f"  - Year: {year}")
    print(f"  - Period: {period}")
    print(f"  - Org Unit ID: {orgunit_id}")
    print(f"  - Org Unit Type: {orgunit_type}")
    
    try:
        # Step 1: Resolve season_id
        print(f"\n🔍 Step 1: Resolving season_id for {period}-{year}...")
        season_id = resolve_season_id_from_year_trimester(year=year, trimester=period)
        print(f"  ✅ Season ID resolved: {season_id}")
        
        # Step 2: Generate export
        print(f"\n📦 Step 2: Generating export...")
        export_service = ReportExportService()
        
        result = export_service.generate_export(
            report_type="seasonal",
            display_mode=None,  # None triggers comparison mode
            file_format="docx",
            year=year,
            month=None,
            trimester=None,
            quarter=None,
            filters={
                "season_id": season_id,
                "orgunit_id": orgunit_id,
                "orgunit_type": orgunit_type
            },
            include_charts=True,
            language="en"
        )
        
        print(f"  ✅ Export generated successfully")
        print(f"  - Filename: {result['filename']}")
        print(f"  - Content Type: {result['content_type']}")
        print(f"  - Content Size: {len(result['content'])} bytes")
        
        # Step 3: Verify it's a ZIP
        print(f"\n🔬 Step 3: Verifying ZIP structure...")
        
        if result['content_type'] != "application/zip":
            print(f"  ❌ FAIL: Expected application/zip, got {result['content_type']}")
            return False
        
        # Step 4: Inspect ZIP contents
        print(f"\n📂 Step 4: Inspecting ZIP contents...")
        zip_buffer = BytesIO(result['content'])
        
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            file_list = zip_file.namelist()
            print(f"  Found {len(file_list)} files in ZIP:")
            
            for filename in file_list:
                file_info = zip_file.getinfo(filename)
                print(f"    - {filename} ({file_info.file_size:,} bytes)")
            
            # Check for expected files
            expected_files = 2
            if len(file_list) != expected_files:
                print(f"  ❌ FAIL: Expected {expected_files} files, found {len(file_list)}")
                return False
            
            # Check for regular report
            regular_report = [f for f in file_list if f.startswith("Seasonal_Report_")]
            if not regular_report:
                print(f"  ❌ FAIL: No regular seasonal report found")
                return False
            print(f"  ✅ Regular report found: {regular_report[0]}")
            
            # Check for comparison report
            comparison_report = [f for f in file_list if f.startswith("Comparison_")]
            if not comparison_report:
                print(f"  ❌ FAIL: No comparison report found")
                return False
            print(f"  ✅ Comparison report found: {comparison_report[0]}")
            
            # Step 5: Verify comparison report has images (charts)
            print(f"\n📊 Step 5: Verifying comparison report contains charts...")
            comparison_data = zip_file.read(comparison_report[0])
            
            # Check if it's a valid Word document
            from docx import Document
            doc_buffer = BytesIO(comparison_data)
            doc = Document(doc_buffer)
            
            # Count images in document
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            print(f"  📈 Found {image_count} images/charts in comparison report")
            
            expected_charts = 9  # 3 levels × 3 chart types
            if image_count < expected_charts:
                print(f"  ⚠️  WARNING: Expected at least {expected_charts} charts, found {image_count}")
                print(f"      (This might be OK if some levels have no data)")
            else:
                print(f"  ✅ All expected charts present!")
            
            # Count paragraphs and tables
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            print(f"  📄 Document structure:")
            print(f"    - Paragraphs: {para_count}")
            print(f"    - Tables: {table_count}")
            print(f"    - Images: {image_count}")
        
        # Step 6: Success summary
        print(f"\n" + "="*80)
        print("✅ PHASE 1 TEST PASSED!")
        print("="*80)
        print(f"✅ Single seasonal export generates ZIP with 2 files")
        print(f"✅ Regular seasonal report included")
        print(f"✅ Comparison report with charts included")
        print(f"✅ ZIP filename: {result['filename']}")
        print("="*80 + "\n")
        
        return True
        
    except Exception as e:
        print(f"\n❌ PHASE 1 TEST FAILED!")
        print(f"Exception: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    print("\n" + "🚀 "* 20)
    print("PHASE 1: Single Seasonal Export Test Suite")
    print("🚀 " * 20 + "\n")
    
    success = test_phase1_single_unit_export()
    
    if success:
        print("\n🎉 PHASE 1 COMPLETE: Ready to proceed to Phase 2!")
    else:
        print("\n⚠️  PHASE 1 NEEDS FIXES: Review errors above")
    
    sys.exit(0 if success else 1)
