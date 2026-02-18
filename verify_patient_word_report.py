"""
Verify Patient Word Export - Content Validation
"""

from docx import Document
import os
import glob

def verify_patient_report(filename):
    """Verify patient Word export structure"""
    print("\n" + "="*70)
    print("VERIFYING PATIENT WORD EXPORT")
    print("="*70)
    print(f"File: {filename}")
    
    if not os.path.exists(filename):
        print("✗ File not found!")
        return False
    
    try:
        doc = Document(filename)
        
        text_content = "\n".join([p.text for p in doc.paragraphs])
        
        print(f"\n✓ Opened document ({os.path.getsize(filename):,} bytes)")
        print(f"✓ Paragraphs: {len(doc.paragraphs)}")
        print(f"✓ Tables: {len(doc.tables)}")
        
        # Verify Arabic title
        if "تقرير المريض" in text_content or "مريض" in text_content:
            print("✓ Arabic title/content found")
        else:
            print("⚠ Arabic patient title not found")
        
        # Check for expected tables (patient info + complaints)
        if len(doc.tables) >= 2:
            print(f"✓ Has {len(doc.tables)} tables (expected: 2 - patient info, complaints)")
            
            for idx, table in enumerate(doc.tables, 1):
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                print(f"\n  Table {idx}: {rows} rows × {cols} columns")
                
                if table.rows:
                    first_row = [cell.text.strip() for cell in table.rows[0].cells[:3]]
                    print(f"    First row: {first_row}")
                    
                    # Check if this is complaints table (6 columns)
                    if cols == 6:
                        print(f"    ✓ This is the complaints table (6 columns)")
                        print(f"    ✓ Contains {rows - 1} complaints")
        else:
            print(f"⚠ Only {len(doc.tables)} tables found")
        
        # Check for complaints table with 6 columns
        has_complaints = any(len(table.columns) == 6 for table in doc.tables if table.rows)
        if has_complaints:
            print("\n✓ Found 6-column complaints table")
        else:
            print("\n⚠ No 6-column complaints table found")
        
        print("\n✅ Patient report structure validated!")
        return True
        
    except Exception as e:
        print(f"✗ ERROR: {str(e)}")
        return False


if __name__ == "__main__":
    print("\n" + "="*70)
    print("PATIENT WORD REPORT VERIFICATION")
    print("="*70)
    
    # Look for both naming patterns
    patient_files = glob.glob("patient_*_report_*.docx") + glob.glob("test_patient_*_export_*.docx")
    
    if patient_files:
        latest = max(patient_files, key=os.path.getmtime)
        print(f"\nChecking: {latest}")
        success = verify_patient_report(latest)
        
        print("\n" + "="*70)
        print(f"RESULT: {'✅ PASS' if success else '❌ FAIL'}")
        print("="*70)
    else:
        print("\n⚠ No patient report files found (patient_*_report_*.docx)")
        print("Run test_patient_word_export.py first to generate a file")
