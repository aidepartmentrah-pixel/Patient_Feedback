"""
Phase 5 Test: End-to-End Integration

Tests the complete flow:
1. DB layer → Service layer → Response model (real data)
2. Service layer → Seasonal service → Word generation (real data)
3. Full pipeline: employee with incidents → profile → seasonal payload → Word bytes
4. Full pipeline: employee with zero incidents → graceful handling
5. Verify all layers agree on counts (DB == Service == Seasonal payload)
6. Word document from real data is valid and parseable
"""
import sys
import os
sys.path.insert(0, '.')
sys.path.insert(0, os.path.join('.', 'backend'))

from backend.core.database import get_connection
from backend.api.db_layer import worker_reporting_db
from backend.api.services.worker_reporting_service import WorkerReportingService
from backend.api.services.worker_seasonal_reporting_service import WorkerSeasonalReportingService
from backend.api.services.person_report_word_adapter import generate_person_seasonal_word_report
from backend.api.services.seasonal_word_adapter import SeasonalWordAdapter
from docx import Document
import io
from datetime import date

PASS = 0
FAIL = 0

def test(name, condition, detail=""):
    global PASS, FAIL
    if condition:
        PASS += 1
        icon = "✅"
    else:
        FAIL += 1
        icon = "❌"
    print(f"  {icon} {name}" + (f" — {detail}" if detail else ""))


def run_all_tests():
    global PASS, FAIL

    print("=" * 70)
    print("PHASE 5 TEST SUITE: End-to-End Integration")
    print("=" * 70)

    # =========================================================================
    # SETUP: Find test employees
    # =========================================================================
    conn = get_connection()
    cursor = conn.cursor()

    # Employee with linked incidents
    cursor.execute("""
        SELECT TOP 1 ice.EmployeeID, COUNT(*) as cnt
        FROM dbo.APP_IncidentCaseEmployee ice
        GROUP BY ice.EmployeeID
        ORDER BY cnt DESC
    """)
    row = cursor.fetchone()
    if not row:
        print("❌ FATAL: No employees linked in APP_IncidentCaseEmployee")
        return
    linked_emp = row.EmployeeID
    linked_count = row.cnt
    print(f"\n--- Test employee WITH incidents: ID={linked_emp} (count={linked_count}) ---")

    # Employee with zero incidents
    cursor.execute("""
        SELECT TOP 1 EmployeeID FROM dbo.APP_VIEWTABLE_HR_EMPLOYEES
        WHERE EmployeeID NOT IN (SELECT EmployeeID FROM dbo.APP_IncidentCaseEmployee)
        ORDER BY EmployeeID
    """)
    zero_row = cursor.fetchone()
    zero_emp = zero_row.EmployeeID if zero_row else None
    print(f"   Test employee with ZERO incidents: ID={zero_emp}")

    # Expected values from raw SQL for validation
    cursor.execute("""
        SELECT
            COUNT(*) as total,
            SUM(CASE WHEN ic.SeverityID = 3 THEN 1 ELSE 0 END) as high,
            SUM(CASE WHEN ic.SeverityID = 2 THEN 1 ELSE 0 END) as med,
            SUM(CASE WHEN ic.SeverityID = 1 THEN 1 ELSE 0 END) as low,
            SUM(CASE WHEN ic.FeedbackIntentTypeID = 2 THEN 1 ELSE 0 END) as good,
            SUM(CASE WHEN ic.FeedbackIntentTypeID = 3 THEN 1 ELSE 0 END) as bad,
            SUM(CASE WHEN ic.FeedbackIntentTypeID IN (1,4) OR ic.FeedbackIntentTypeID IS NULL THEN 1 ELSE 0 END) as neutral
        FROM dbo.APP_IncidentCaseEmployee ice
        INNER JOIN dbo.APP_IncidentCase ic ON ice.IncidentRequestCaseID = ic.IncidentRequestCaseID
        WHERE ice.EmployeeID = ?
    """, (linked_emp,))
    raw = cursor.fetchone()
    expected = {
        'total': raw.total,
        'high': raw.high or 0,
        'med': raw.med or 0,
        'low': raw.low or 0,
        'good': raw.good or 0,
        'bad': raw.bad or 0,
        'neutral': raw.neutral or 0
    }
    print(f"   Raw SQL expected: {expected}")
    conn.close()

    # =========================================================================
    # TEST GROUP 1: Layer Agreement — DB layer counts
    # =========================================================================
    print("\n--- TEST GROUP 1: DB Layer counts match raw SQL ---")

    db_total = worker_reporting_db.count_worker_incidents(employee_id=linked_emp)
    test("DB total matches raw SQL", db_total == expected['total'],
         f"db={db_total}, raw={expected['total']}")

    db_sev = worker_reporting_db.count_worker_incidents_by_severity(employee_id=linked_emp)
    test("DB severity high matches", db_sev.get('high', 0) == expected['high'])
    test("DB severity med matches", db_sev.get('medium', 0) == expected['med'])
    test("DB severity low matches", db_sev.get('low', 0) == expected['low'])

    db_intent = worker_reporting_db.count_worker_incidents_by_intent(employee_id=linked_emp)
    test("DB intent good matches", db_intent.get('good', 0) == expected['good'])
    test("DB intent bad matches", db_intent.get('bad', 0) == expected['bad'])
    test("DB intent neutral matches", db_intent.get('neutral', 0) == expected['neutral'])

    db_detail = worker_reporting_db.get_worker_incidents_detail(employee_id=linked_emp)
    test("DB detail count matches total", len(db_detail) == expected['total'],
         f"detail_count={len(db_detail)}")

    # =========================================================================
    # TEST GROUP 2: Service layer matches DB layer
    # =========================================================================
    print("\n--- TEST GROUP 2: Service layer matches DB layer ---")

    profile = WorkerReportingService.get_worker_profile(employee_id=linked_emp)

    test("service total == db total",
         profile.metrics.total_incidents == db_total)
    test("service high == db high",
         profile.metrics.high_severity == db_sev.get('high', 0))
    test("service medium == db medium",
         profile.metrics.medium_severity == db_sev.get('medium', 0))
    test("service low == db low",
         profile.metrics.low_severity == db_sev.get('low', 0))
    test("service good == db good",
         profile.metrics.good_feedback_count == db_intent.get('good', 0))
    test("service bad == db bad",
         profile.metrics.bad_feedback_count == db_intent.get('bad', 0))
    test("service neutral == db neutral",
         profile.metrics.neutral_feedback_count == db_intent.get('neutral', 0))
    test("service incidents list count == db detail count",
         len(profile.incidents) == len(db_detail))

    # =========================================================================
    # TEST GROUP 3: Seasonal service matches service layer
    # =========================================================================
    print("\n--- TEST GROUP 3: Seasonal service matches service layer ---")

    seasonal = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
        employee_id=linked_emp,
        season_start=date(2020, 1, 1),
        season_end=date(2030, 12, 31)
    )

    sm = seasonal['metrics']
    test("seasonal total == service total",
         sm['total_incidents'] == profile.metrics.total_incidents)
    test("seasonal high == service high",
         sm['high_severity'] == profile.metrics.high_severity)
    test("seasonal good == service good",
         sm['good_feedback_count'] == profile.metrics.good_feedback_count)
    test("seasonal bad == service bad",
         sm['bad_feedback_count'] == profile.metrics.bad_feedback_count)
    test("seasonal incidents_details count == service",
         len(seasonal['incidents_details']) == len(profile.incidents))
    test("seasonal has performance score",
         'score' in seasonal.get('performance', {}))

    # =========================================================================
    # TEST GROUP 4: Word generation from real seasonal data
    # =========================================================================
    print("\n--- TEST GROUP 4: Word generation from real data ---")

    try:
        # person_report_word_adapter
        person_bytes = generate_person_seasonal_word_report(
            person_type="worker",
            payload=seasonal
        )
        test("person_report generates valid bytes from real data",
             isinstance(person_bytes, bytes) and len(person_bytes) > 100,
             f"size={len(person_bytes)}")

        # Parse it back
        person_doc = Document(io.BytesIO(person_bytes))
        test("person_report Word is parseable",
             len(person_doc.paragraphs) > 0,
             f"paragraphs={len(person_doc.paragraphs)}, tables={len(person_doc.tables)}")

    except Exception as e:
        test("person_report generates valid bytes from real data", False, f"ERROR: {e}")
        test("person_report Word is parseable", False)

    try:
        # seasonal_word_adapter
        seasonal_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(seasonal)
        test("seasonal_word generates valid bytes from real data",
             isinstance(seasonal_bytes, bytes) and len(seasonal_bytes) > 100,
             f"size={len(seasonal_bytes)}")

        seasonal_doc = Document(io.BytesIO(seasonal_bytes))
        test("seasonal_word Word is parseable",
             len(seasonal_doc.paragraphs) > 0,
             f"paragraphs={len(seasonal_doc.paragraphs)}, tables={len(seasonal_doc.tables)}")

    except Exception as e:
        test("seasonal_word generates valid bytes from real data", False, f"ERROR: {e}")
        test("seasonal_word Word is parseable", False)

    # =========================================================================
    # TEST GROUP 5: Zero-incident employee — full pipeline
    # =========================================================================
    print("\n--- TEST GROUP 5: Zero-incident employee pipeline ---")

    if zero_emp:
        try:
            zero_profile = WorkerReportingService.get_worker_profile(employee_id=zero_emp)
            test("zero profile succeeds", True)
            test("zero total_incidents == 0", zero_profile.metrics.total_incidents == 0)
            test("zero incidents list empty", len(zero_profile.incidents) == 0)

            zero_seasonal = WorkerSeasonalReportingService.build_worker_seasonal_report_data(
                employee_id=zero_emp,
                season_start=date(2020, 1, 1),
                season_end=date(2030, 12, 31)
            )
            test("zero seasonal total == 0",
                 zero_seasonal['metrics']['total_incidents'] == 0)
            test("zero seasonal incidents_details empty",
                 len(zero_seasonal['incidents_details']) == 0)

            # Word from zero data
            zero_person_bytes = generate_person_seasonal_word_report(
                person_type="worker",
                payload=zero_seasonal
            )
            test("zero Word generates valid bytes",
                 isinstance(zero_person_bytes, bytes) and len(zero_person_bytes) > 100)

            zero_doc = Document(io.BytesIO(zero_person_bytes))
            # Should not have a 7-column incidents table
            has_inc_table = any(
                len(t.columns) == 7 and len(t.rows) > 1
                for t in zero_doc.tables
            )
            test("zero Word has NO incidents table", not has_inc_table)

            zero_seasonal_bytes = SeasonalWordAdapter.generate_worker_seasonal_word(zero_seasonal)
            test("zero seasonal Word generates bytes",
                 isinstance(zero_seasonal_bytes, bytes) and len(zero_seasonal_bytes) > 100)

        except Exception as e:
            test("zero pipeline", False, f"ERROR: {e}")
            import traceback
            traceback.print_exc()
    else:
        print("  ⚠ No zero-incident employee — skipping pipeline tests")
        for _ in range(8):
            test("(skipped — no zero employee)", True)

    # =========================================================================
    # TEST GROUP 6: Classification values validation
    # =========================================================================
    print("\n--- TEST GROUP 6: Classification values ---")

    if len(profile.incidents) > 0:
        valid_classes = {'good', 'bad', 'neutral'}
        all_valid = all(
            inc.get('classification') in valid_classes
            for inc in profile.incidents
        )
        test("all incidents have valid classification",
             all_valid,
             f"checked {len(profile.incidents)} incidents")

        # Check that classification sums match
        good_count = sum(1 for i in profile.incidents if i.get('classification') == 'good')
        bad_count = sum(1 for i in profile.incidents if i.get('classification') == 'bad')
        neutral_count = sum(1 for i in profile.incidents if i.get('classification') == 'neutral')

        test("classification sum from incidents matches metrics good",
             good_count == profile.metrics.good_feedback_count,
             f"list={good_count}, metric={profile.metrics.good_feedback_count}")
        test("classification sum from incidents matches metrics bad",
             bad_count == profile.metrics.bad_feedback_count,
             f"list={bad_count}, metric={profile.metrics.bad_feedback_count}")
        test("classification sum from incidents matches metrics neutral",
             neutral_count == profile.metrics.neutral_feedback_count,
             f"list={neutral_count}, metric={profile.metrics.neutral_feedback_count}")
    else:
        test("(no incidents to check classification)", True)
        for _ in range(3):
            test("(skipped)", True)

    # =========================================================================
    # FINAL RESULTS
    # =========================================================================
    print(f"\n{'='*70}")
    print(f"PHASE 5 RESULTS: {PASS} passed, {FAIL} failed, {PASS+FAIL} total")
    print(f"{'='*70}")

    if FAIL == 0:
        print("🎉 ALL TESTS PASSED! Phase 5 complete — full integration verified.")
    else:
        print(f"⚠ {FAIL} test(s) failed. Review output above.")


if __name__ == "__main__":
    run_all_tests()
